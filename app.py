from flask import Flask, render_template, request, redirect, url_for, session, render_template_string, escape  
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import os
import pandas as pd
from datetime import date, datetime as dt
from datetime import datetime
from itsdangerous import URLSafeTimedSerializer, BadSignature
from sqlalchemy import or_
from sqlalchemy import and_ 
# from flask import Flask, render_template, request, redirect_url
import psycopg2
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
import ast
import datetime
import os
import json

from flask_cors import CORS
import re
# import spacy
from flask_mail import Mail, Message
from flask import render_template, redirect, url_for, flash
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer
# from spacy.matcher import Matcher
from flask import Flask, request, render_template
from sqlalchemy import ForeignKey
from sqlalchemy.orm import relationship
from itsdangerous import URLSafeTimedSerializer
from flask import request, render_template, flash, redirect, url_for
import secrets
import secrets
from urllib.parse import quote_plus
from flask_migrate import Migrate
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from datetime import datetime
import pytz 
from sqlalchemy import case, desc
from sqlalchemy.orm import aliased
import hashlib
import random
import string

app = Flask(__name__)
cors = CORS(app)
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
# app.config['MAIL_SERVER'] = 'smtp.office365.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
# app.config['MAIL_USERNAME'] = 'pavan.k@makonissoft.com'
# app.config['MAIL_PASSWORD'] = 'Roh64272'
# app.config['MAIL_USERNAME'] = 'ganesh.s@makonissoft.com'
# app.config['MAIL_PASSWORD'] = 'Fol98135'
app.config['MAIL_USERNAME'] = 'kanuparthisaiganesh582@gmail.com'
# app.config['MAIL_PASSWORD'] = 'Ganesh@2022'
app.config['MAIL_PASSWORD'] = 'cdxfkuefixpigwae'
mail = Mail(app)

app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get("DATABASE_URL")

app.config['SECRET_KEY'] = secrets.token_hex(16)
engine = create_engine(app.config['SQLALCHEMY_DATABASE_URI'])
Session = sessionmaker(bind=engine)
app.config['SECRET_KEY'] = secrets.token_hex(16)
# Specify the folder where uploaded resumes will be stored
UPLOAD_FOLDER = 'C:/Users/Makonis/PycharmProjects/login/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# cors = CORS(app)
# Specify the allowed resume file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
db = SQLAlchemy(app)
migrate = Migrate(app, db)
from datetime import timedelta
#hello

# Specify the folder where uploaded resumes will be stored
# UPLOAD_FOLDER = 'static/'
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# cors = CORS(app)
# Specify the allowed resume file extensions
# ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# db = SQLAlchemy(app)
# migrate = Migrate(app, db)

def generate_verification_token(user_id):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    return serializer.dumps(user_id)

class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True,nullable=False)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(50), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    user_type = db.Column(db.String(20), nullable=False)
    client = db.Column(db.String(100))
    candidate = relationship("Candidate", back_populates="user", uselist=False)
    is_active = db.Column(db.Boolean, default=True)
    is_verified = db.Column(db.Boolean, default=False)
    created_by = db.Column(db.String(50))
    otp = db.Column(db.String(6), default=False)
    registration_completed = db.Column(db.String(50))
    filename = db.Column(db.String(100))
    # image_file = db.Column(db.String(1000))
    image_file=db.Column(db.LargeBinary)
    image_deleted=db.Column(db.Boolean, default=False)
    def serialize(self):
        return {
            'id': self.id,
            'username': self.username,
            'name': self.name,
            'email': self.email,
            'user_type': self.user_type,
            'client': self.client,
            'is_active': self.is_active,
            'is_verified': self.is_verified,
            'created_by': self.created_by,
            'otp': self.otp,
            'registration_completed': self.registration_completed
        }
        
class Candidate(db.Model):
    __tablename__ = 'candidates'
    id = db.Column(db.Integer, primary_key=True)
    job_id = db.Column(db.Integer)
    name = db.Column(db.String(100), nullable=False)
    mobile = db.Column(db.String(20), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    client = db.Column(db.String(100), nullable=False)
    current_company = db.Column(db.String(100))
    position = db.Column(db.String(100))
    profile = db.Column(db.String(200))
    current_job_location = db.Column(db.String(100))
    preferred_job_location = db.Column(db.String(100))
    # resume = db.Column(db.String(1000))
    resume = db.Column(db.LargeBinary)
    skills = db.Column(db.String(500))
    qualifications = db.Column(db.String(200))
    experience = db.Column(db.String(200))
    relevant_experience = db.Column(db.String(200))
    current_ctc = db.Column(db.String(200))
    expected_ctc = db.Column(db.String(200))
    notice_period = db.Column(db.String(20))
    last_working_date = db.Column(db.Date)
    buyout = db.Column(db.Boolean, default=False)
    holding_offer = db.Column(db.String(20))
    total = db.Column(db.Integer)
    package_in_lpa = db.Column(db.Float)
    recruiter = db.Column(db.String(100))
    management = db.Column(db.String(100))
    status = db.Column(db.String(100))
    reason_for_job_change = db.Column(db.String(200))
    remarks = db.Column(db.String(200))
    screening_done = db.Column(db.Boolean, default=False)
    rejected_at_screening = db.Column(db.Boolean, default=False)
    l1_cleared = db.Column(db.Boolean, default=False)
    rejected_at_l1 = db.Column(db.Boolean, default=False)
    dropped_after_clearing_l1 = db.Column(db.Boolean, default=False)
    l2_cleared = db.Column(db.Boolean, default=False)
    rejected_at_l2 = db.Column(db.Boolean, default=False)
    dropped_after_clearing_l2 = db.Column(db.Boolean, default=False)
    onboarded = db.Column(db.Boolean, default=False)
    dropped_after_onboarding = db.Column(db.Boolean, default=False)
    date_created = db.Column(db.Date, default=date.today)
    time_created = db.Column(db.Time, default=datetime.now().time)
    comments = db.Column(db.String(1000))
    linkedin_url = db.Column(db.String(1000))
    user_id = db.Column(db.Integer, ForeignKey('users.id'))
    serving_notice_period = db.Column(db.String(200))
    period_of_notice = db.Column(db.String(1000))
    user = relationship("User", back_populates="candidate")
    reference = db.Column(db.String(200))
    reference_name = db.Column(db.String(200))
    reference_position = db.Column(db.String(200))
    reference_information = db.Column(db.String(200))
    data_updated_date = db.Column(db.Date)
    data_updated_time = db.Column(db.Time)
    resume_present = db.Column(db.Boolean, default=False)
    # resume_present = db.Column(db.Boolean, default=True)
    def serialize(self):
        return {
            'id': self.id,
            'job_id': self.job_id,
            'name': self.name,
            'mobile': self.mobile,
            'email': self.email,
            'client': self.client,
            'current_company': self.current_company,
            'position': self.position,
            'profile': self.profile,
            'current_job_location': self.current_job_location,
            'preferred_job_location': self.preferred_job_location,
            'resume': self.resume,
            'skills': self.skills,
            'qualifications': self.qualifications,
            'experience': self.experience,
            'relevant_experience': self.relevant_experience,
            'current_ctc': self.current_ctc,
            'expected_ctc': self.expected_ctc,
            'notice_period': self.notice_period,
            'last_working_date': self.last_working_date.strftime('%Y-%m-%d') if self.last_working_date else None,
            'buyout': self.buyout,
            'holding_offer': self.holding_offer,
            'total': self.total,
            'package_in_lpa': self.package_in_lpa,
            'recruiter': self.recruiter,
            'management': self.management,
            'status': self.status,
            'reason_for_job_change': self.reason_for_job_change,
            'remarks': self.remarks,
            'screening_done': self.screening_done,
            'rejected_at_screening': self.rejected_at_screening,
            'l1_cleared': self.l1_cleared,
            'rejected_at_l1': self.rejected_at_l1,
            'dropped_after_clearing_l1': self.dropped_after_clearing_l1,
            'l2_cleared': self.l2_cleared,
            'rejected_at_l2': self.rejected_at_l2,
            'dropped_after_clearing_l2': self.dropped_after_clearing_l2,
            'onboarded': self.onboarded,
            'dropped_after_onboarding': self.dropped_after_onboarding,
            'date_created': self.date_created.strftime('%Y-%m-%d'),
            'time_created': self.time_created.strftime('%H:%M:%S'),
            'comments': self.comments,
            'linkedin_url': self.linkedin_url,
            'user_id': self.user_id,
            'period_of_notice': self.period_of_notice,
            'reference': self.reference,
            'reference_name': self.reference_name,
            'reference_position': self.reference_position,
            'reference_information': self.reference_information,
            
        }
        
class Career_user(db.Model):
    __tablename__ = 'career_users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True, nullable=False)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(50), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    user_type = db.Column(db.String(50), default="career_visitor")


class Career_notification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    recruiter_name = db.Column(db.String(100), nullable=False)
    notification_status = db.Column(db.Boolean, default=False)

    def __init__(self, recruiter_name, notification_status=False):
        self.recruiter_name = recruiter_name
        self.notification_status = notification_status


class JobPost(db.Model):
    __tablename__ = 'job_posts'

    id = db.Column(db.Integer, primary_key=True)
    client = db.Column(db.String(100))
    # experience_min = db.Column(db.Integer)
    # experience_max = db.Column(db.Integer)
    experience_min = db.Column(db.String(100))
    experience_max = db.Column(db.String(100))
    budget_min = db.Column(db.String(300))
    budget_max = db.Column(db.String(300))
    location = db.Column(db.String(100))
    shift_timings = db.Column(db.String(100))
    notice_period = db.Column(db.String(100))
    role = db.Column(db.String(100))
    detailed_jd = db.Column(db.Text)
    # jd_pdf =  db.Column(db.String(1000))
    jd_pdf =  db.Column(db.LargeBinary)
    mode = db.Column(db.String(100))
    recruiter = db.Column(db.String(1000))
    management = db.Column(db.String(100))
    date_created = db.Column(db.Date)
    time_created = db.Column(db.Time)
    job_status = db.Column(db.String(20))
    job_type = db.Column(db.String(100))
    contract_in_months = db.Column(db.String(100))
    # contract_in_months = db.Column(db.Integer, nullable=True)
    skills = db.Column(db.String(500))
    notification = db.Column(db.String(20))
    data_updated_date = db.Column(db.Date)
    data_updated_time = db.Column(db.Time)
    jd_pdf_present = db.Column(db.Boolean, default=False)
    # jd_pdf_present = db.Column(db.Boolean, default=True)
    def __init__(self, client, experience_min, experience_max, budget_min, budget_max, location, shift_timings, notice_period, role, detailed_jd, mode, recruiter, management, job_status, job_type, skills, jd_pdf, jd_pdf_present,contract_in_months):
        self.client = client
        self.experience_min = experience_min
        self.experience_max = experience_max
        self.budget_min = budget_min
        self.budget_max = budget_max
        self.location = location
        self.shift_timings = shift_timings
        self.notice_period = notice_period
        self.role = role
        self.detailed_jd = detailed_jd
        self.mode = mode
        self.recruiter = recruiter
        self.management = management
        self.job_status = job_status
        self.job_type = job_type
        self.skills = skills
        self.jd_pdf = jd_pdf
        self.contract_in_months = contract_in_months
        self.jd_pdf_present = jd_pdf_present

class Deletedcandidate(db.Model):
    _tablename_ = 'deletedcandidate'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), nullable=False)
    candidate_name = db.Column(db.String(100), nullable=False)
    candidate_email = db.Column(db.String(100), nullable=False)
    client = db.Column(db.String(100), nullable=False)
    profile = db.Column(db.String(100), nullable=False)
    status = db.Column(db.String(100), nullable=False)

# class Notification(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     recruiter_name = db.Column(db.String(100), nullable=False)
#     notification_status = db.Column(db.Boolean, default=False)

#     def __init__(self, recruiter_name, notification_status=False):
#         self.recruiter_name = recruiter_name
#         self.notification_status = notification_status
class Notification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_post_id = db.Column(db.Integer, db.ForeignKey('job_posts.id'))
    recruiter_name = db.Column(db.String(100), nullable=False)
    notification_status = db.Column(db.Boolean, default=False)
    num_notification = db.Column(db.Integer, default=0)  # New column added

    def __init__(self, job_post_id, recruiter_name, notification_status=False):
        self.job_post_id = job_post_id
        self.recruiter_name = recruiter_name
        self.notification_status = notification_status
        self.num_notification = 0  # Default value for num_notification


@app.route('/check_candidate', methods=['POST'])
def check_candidate():
    data = request.json
    clients = []
    profiles = []
    dates = []
    job_ids = []
    status = []

    email = data.get('email')
    mobile = data.get('mobile')

    # Query the database to check for an existing candidate with the provided mobile or email
    existing_candidates = Candidate.query.filter(or_(Candidate.mobile == mobile, Candidate.email == email)).all()
    
    for candidate in existing_candidates:
        clients.append(candidate.client)
        profiles.append(candidate.profile)
        dates.append(candidate.date_created.strftime('%Y-%m-%d'))
        job_ids.append(candidate.job_id)
        status.append(candidate.status)

    if existing_candidates:
        response = {
            'message': "Candidate with this mobile or email already exists.",
            'clients': clients,
            'profiles': profiles,
            'dates': dates,
            'jobIds': job_ids,
            'status': status
        }
    else:
        response = {
            'message': "Mobile and email not available.",
            'clients': None,
            'profiles': None,
            'dates': None,
            'jobIds': None,
            'status': None
        }

    return jsonify(response)


# @app.route('/check_candidate', methods=['POST'])
# def check_candidate():
#     clients = []
#     profiles = []
#     dates=[]
#     job_ids=[]
#     status=[]
#     field = request.json['field']
#     value = request.json['value']

#     # Query the database to check for an existing candidate with the provided mobile or email
#     existing_candidate = Candidate.query.filter(or_(Candidate.mobile == value, Candidate.email == value)).all()
#     for i in existing_candidate:
#         clients.append(" " + i.client + " ")
#         profiles.append(" " + i.profile + " " )
#         dates.append(i.date_created.strftime('%Y-%m-%d'))
#         job_ids.append(i.job_id)
#         status.append(i.status)

#     # candidate = Candidate.query.filter_by(mobile=existing_candidate.mobile).first()
#     if existing_candidate:
#         response = {
#             'message' : f"Candidate with this {field} already exists.",
#             'client' : clients,
#             'profile' : profiles,
#             'dates':dates,
#             'jobId':job_ids,
#             'status':status
#         }

#     else:
#         response = {
#             'message': f"{field.capitalize()} is available.",
#             'client': None,
#             'profile': None,
#             'dates':None,
#             'jobId':None,
#             'status':None
#         }
#     return json.dumps(response)

@app.route('/recruiter')
def recruiter_index():
    return render_template('recruiter_index.html')

@app.route('/')
def index():
    session_timeout_msg = request.args.get("session_timeout_msg")
    reset_message = request.args.get("reset_message")
    signup_message = request.args.get('signup_message')
    password_message = request.args.get('password_message')
    return render_template('index.html',reset_message=reset_message,session_timeout_msg=session_timeout_msg,signup_message=signup_message,password_message=password_message)

@app.route('/forgot_password')
def forgot_password():
    return render_template('forgot_password.html')

def generate_6otp():
    digits = "0123456789"
    otp = "".join(random.choice(digits) for _ in range(6))
    return otp


@app.route('/generate_otp', methods=['POST'])
def generate_otp():
    if request.method == 'POST':
        data = request.json
        username = data.get('username')
        email = data.get('email')
        user = User.query.filter_by(username=username, email=email).first()
        if user:
            otp = generate_6otp()
            user.otp = otp
            db.session.commit()
            msg = Message('New OTP Generated', sender='kanuparthisaiganesh582@gmail.com', recipients=[email])
            msg.html = f'''
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        background-color: #f4f4f4;
                        color: #333;
                        margin: 0;
                        padding: 20px;
                    }}
                    .container {{
                        background-color: #ffffff;
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                        border: 1px solid #dddddd;
                        border-radius: 8px;
                        box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                    }}
                    .header {{
                        font-size: 24px;
                        font-weight: bold;
                        margin-bottom: 20px;
                        color: #4CAF50;
                    }}
                    .content p {{
                        font-size: 16px;
                        line-height: 1.6;
                        margin: 10px 0;
                    }}
                    .otp-label {{
                        display: block;
                        margin-bottom: 10px;
                        font-weight: bold;
                    }}
                    .otp-input {{
                        width: 50%;
                        padding: 10px;
                       font-size: 20px;
                        font-weight: bold;
                        border: 1px solid #eeeeee;
                        border-radius: 5px;
                        background-color: #f9f9f9;
                        box-sizing: border-box;
                        outline: none; /* Remove the default outline */
                    }}
                    .footer {{
                        font-size: 12px;
                        color: #999;
                        margin-top: 20px;
                        text-align: center;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">New OTP Generated</div>
                    <div class="content">
                        <p>Hi {user.name},</p>
                        <p>OTP for resetting your password:</p>
                        <p><input type="text" class="otp-input" id="otp" value="{otp}" readonly></p>
                    </div>
                    <div class="footer">
                        <p>If you did not request this change, please contact our support team immediately.</p>
                        <p><b>Makonis Talent Track Pro Team</b></p>
                    </div>
                </div>
            </body>
            </html>
            '''
            mail.send(msg)
            return jsonify({'status': 'success', 'message': 'OTP has been sent to your email.'})
        else:
            return jsonify({'status': 'error', 'message': 'User does not exist.'})
    else:
        return jsonify({'status': 'error', 'message': 'Invalid request method.'})



# @app.route('/generate_otp', methods=['POST'])
# def generate_otp():
#     if request.method == 'POST':
#         data = request.json
#         username = data.get('username')
#         email = data.get('email')
#         user = User.query.filter_by(username=username, email=email).first()
#         if user:
#             otp = generate_6otp()
#             user.otp = otp
#             db.session.commit()
#             msg = Message('New OTP Notification', sender='kanuparthisaiganesh582@gmail.com', recipients=[email])
#             msg.html = f'''
#             <!DOCTYPE html>
#             <html lang="en">
#             <head>
#                 <meta charset="UTF-8">
#                 <meta name="viewport" content="width=device-width, initial-scale=1.0">
#                 <style>
#                     body {{
#                         font-family: Arial, sans-serif;
#                         background-color: #f4f4f4;
#                         color: #333;
#                         margin: 0;
#                         padding: 20px;
#                     }}
#                     .container {{
#                         background-color: #ffffff;
#                         max-width: 600px;
#                         margin: 0 auto;
#                         padding: 20px;
#                         border: 1px solid #dddddd;
#                         border-radius: 8px;
#                         box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#                     }}
#                     .header {{
#                         font-size: 24px;
#                         font-weight: bold;
#                         margin-bottom: 20px;
#                         color: #4CAF50;
#                     }}
#                     .content {{
#                         font-size: 16px;
#                         line-height: 1.6;
#                     }}
#                     .otp-container {{
#                         margin-top: 10px;
#                         position: relative;
#                     }}
#                     .otp {{
#                         background-color: #f9f9f9;
#                         padding: 10px;
#                         border: 1px solid #eeeeee;
#                         border-radius: 5px;
#                         margin-top: 10px;
#                         font-size: 20px;
#                         font-weight: bold;
#                         text-align: center;
#                         position: relative;
#                         width: 100%;
#                         box-sizing: border-box; /* Ensure padding and border are included in width */
#                     }}
#                     .footer {{
#                         font-size: 12px;
#                         color: #999;
#                         margin-top: 20px;
#                         text-align: center;
#                     }}
#                 </style>
#             </head>
#             <body>
#                 <div class="container">
#                     <div class="header">New OTP Generated</div>
#                     <div class="content">
#                         <p>Hi {user.name},</p>
#                         <p>OTP for resetting your password:</p>
#                         <div class="otp-container">
#                             <input type="text" class="otp" id="otp" value="{otp}" readonly>
#                         </div>
#                     </div>
#                     <div class="footer">
#                         <p><b>Makonis Talent Track Pro Team</b></p>
#                     </div>
#                 </div>
#             </body>
#             </html>
#             '''
#             mail.send(msg)
#             return jsonify({'status': 'success', 'message': 'OTP has been sent to your email.'})
#         else:
#             return jsonify({'status': 'error', 'message': 'User does not exist.'})
#     else:
#         return jsonify({'status': 'error', 'message': 'Invalid request method.'})

# @app.route('/generate_otp', methods=['POST'])
# def generate_otp():
#     if request.method == 'POST':
#         username = request.json.get('username')
#         email = request.json.get('email')
#         user = User.query.filter_by(username=username, email=email).first()
#         if user:
#             otp = generate_6otp()
#             user.otp = otp
#             db.session.commit()
#             msg = Message('Account Verification', sender='kanuparthisaiganesh582@gmail.com', recipients=[email])
#             msg.body = f'Hi {user.name},\n\n OTP for resetting your password {otp}.'
#             mail.send(msg)
#             return jsonify({'status': 'success', 'message': 'OTP has been sent to your email.'})
#         else:
#             return jsonify({'status': 'error', 'message': 'User does not exist.'})
#     else:
#         return jsonify({'status': 'error', 'message': 'Invalid request method.'})
    



@app.route('/reset_password', methods=['POST'])
def reset_password():
    if request.method == 'POST':
        data = request.json
        otp = data['otp']
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
        new_password_hashed = hashlib.sha256(new_password.encode()).hexdigest()

        user = User.query.filter_by(otp=otp).first()

        if user and user.otp == otp and new_password == confirm_password:
            # Check if the new password is different from the old password
            if new_password_hashed != user.password:  # comparing hashes
                user.password = new_password_hashed
                db.session.commit()
                # Send the updated password to the user's email
                msg = Message('Password Changed', sender='your-email@gmail.com', recipients=[user.email])
                msg.html = f'''
                <!DOCTYPE html>
                <html lang="en">
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <style>
                        body {{
                            font-family: Arial, sans-serif;
                            background-color: #f4f4f4;
                            color: #333;
                            margin: 0;
                            padding: 20px;
                        }}
                        .container {{
                            background-color: #ffffff;
                            max-width: 600px;
                            margin: 0 auto;
                            padding: 20px;
                            border: 1px solid #dddddd;
                            border-radius: 8px;
                            box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                        }}
                        .header {{
                            font-size: 24px;
                            font-weight: bold;
                            margin-bottom: 20px;
                            color: #4CAF50;
                        }}
                        .content {{
                            font-size: 16px;
                            line-height: 1.6;
                        }}
                        .credentials {{
                            background-color: #f9f9f9;
                            padding: 10px;
                            border: 1px solid #eeeeee;
                            border-radius: 5px;
                            margin-top: 10px;
                        }}
                        .footer {{
                            font-size: 12px;
                            color: #999;
                            margin-top: 20px;
                            text-align: center;
                        }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <div class="header">Password Changed</div>
                        <div class="content">
                            <p>Hello {user.name},</p>
                            <p>Your password has been successfully changed. </p>
                            <p>Here are your updated credentials:,</p>
                            <div class="credentials">
                                <p><strong>Username:</strong> {user.username}</p>
                                <p><strong>Password:</strong> {new_password}</p>
                            </div>
                        </div>
                        <div class="footer">
                            <p>If you did not request this change, please contact our support team immediately.</p>
                            <p><b>Makonis Talent Track Pro Team</b></p>
                        </div>
                    </div>
                </body>
                </html>
                '''
                mail.send(msg)

                return jsonify({'status': 'success', 'message': 'Password changed successfully.'})
            else:
                return jsonify({'status': 'error', 'message': 'New password is the same as the old password'})
        else:
            return jsonify({'status': 'error', 'message': 'Invalid OTP or password confirmation. Please try again.'})

    return jsonify({'status': 'error', 'message': 'Invalid request method.'})


# @app.route('/reset_password', methods=['POST'])
# def reset_password():
#     if request.method == 'POST':
#         otp = request.json['otp']
#         new_password = request.json.get('new_password')
#         confirm_password = request.json.get('confirm_password')
#         new_password_hashed = hashlib.sha256(new_password.encode()).hexdigest()

#         user = User.query.filter_by(otp=otp).first()

#         if user and user.otp == otp and new_password == confirm_password:
#             # Check if the new password is different from the old password
#             if new_password_hashed != user.password:  # comparing hashes
#                 user.password = new_password_hashed
#                 db.session.commit()
#                 # Send the updated password to the user's email
#                 msg = Message('Password Changed', sender='kanuparthisaiganesh582@gmail.com', recipients=[user.email])
#                 msg.body = f'Hello {user.name},\n\nYour password has been successfully changed. Here are your updated credentials:\n\nUsername: {user.username}\nPassword: {new_password}'
#                 mail.send(msg)

#                 return jsonify({'status': 'success', 'message': 'Password changed successfully.'})
#             else:
#                 return jsonify({'status': 'error', 'message': 'New password is the same as the old password'})
#         else:
#             return jsonify({'status': 'error', 'message': 'Invalid OTP or password confirmation. Please try again.'})

#     return jsonify({'status': 'error', 'message': 'Invalid request method.'})



def generate_html_message(message, redirect_url=None):
    html_message = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Message</title>
            <!-- Add CSS styles for your message -->
            <style>
                /* Example CSS styles */
                body {{
                    font-family: Arial, sans-serif;
                    background-color: #f4f4f4;
                    text-align: center;
                }}
                .message {{
                    margin-top: 50px;
                    background-color: #fff;
                    border-radius: 10px;
                    padding: 20px;
                    box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                    display: inline-block;
                }}
                .message p {{
                    font-size: 20px;
                }}
                .button {{
                    display: inline-block;
                    padding: 10px 20px;
                    background-color: #007bff;
                    color: #fff;
                    text-decoration: none;
                    border-radius: 5px;
                    transition: background-color 0.3s;
                    margin-top: 20px;
                }}
                .button:hover {{
                    background-color: #0056b3;
                }}
            </style>
        </head>
        <body>
            <div class="message">
                <p>{message}</p>
                {'' if redirect_url is None else f'<a href="{redirect_url}" class="button">Login</a>'}
            </div>
        </body>
        </html>
    """
    return html_message

@app.route('/verify/<token>')
def verify(token):
    user_id = verify_token(token)
    if user_id:
        user = User.query.get(user_id)
        if user.is_verified:
            if user.user_type == 'management':
                message = 'Your Management Account is already Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/ManagementLogin')
            elif user.user_type == 'recruiter':
                message = 'Your Recruiter Account is already Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/RecruitmentLogin')
        else:
            user.is_verified = True
            db.session.commit()
            if user.user_type == 'management':
                message = 'Your Management Account has been Successfully Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/ManagementLogin')
            elif user.user_type == 'recruiter':
                message = 'Your Recruiter Account has been Successfully Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/RecruitmentLogin')
    else:
        message = 'Your verification link has expired. Please contact management to activate your account.'
        return generate_html_message(message)


# @app.route('/verify/<token>')
# def verify(token):
#     user_id = verify_token(token)
#     if user_id:
#         user = User.query.get(user_id)
#         user.is_verified = True
#         db.session.commit()
#         if user.user_type == 'management':
#             message = 'Your Management Account has been Successfully Verified. Please Login!'
#             return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/ManagementLogin')
#         elif user.user_type == 'recruiter':
#             message = 'Your Recruiter Account has been Successfully Verified!'
#             return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/RecruitmentLogin')
#     else:
#         message = 'Your verification link has expired. Please contact management to activate your account.'
#         return generate_html_message(message)


# @app.route('/verify/<token>')
# def verify(token):
#     user_id = verify_token(token)
#     if user_id:
#         user = User.query.get(user_id)
#         user.is_verified = True
#         db.session.commit()
#         if user.user_type == 'management':
#             return jsonify({'status': 'success', 'message': 'Account verified successfully!', 'redirect': url_for('management_login', verification_msg_manager='Your Account has been Successfully Verified. Please Login.')})
#         elif user.user_type == 'recruiter':
#             return jsonify({'status': 'success', 'message': 'Account verified successfully!', 'redirect': url_for('recruiter_index')})
#     else:
#         return jsonify({'status': 'error', 'message': 'Your verification link has expired. Please contact management to activate your account.'})
#     return jsonify({'status': 'error', 'message': 'An error occurred while verifying your account.'})



# Function to generate a random password
def generate_random_password(length=8):
    digits = string.digits
    password = ''.join(random.choice(digits) for _ in range(length - 3))
    return "Mak" + password


@app.route('/signup', methods=['POST'])
def signup():
    data = request.json
    user_id = data.get('user_id')  # Using get method to avoid KeyError
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error', 'message': 'Invalid user ID or user does not exist.'})

    user_type = user.user_type
    user_name = user.username

    if user_type == 'management':
        username = data.get('username')
        name = data.get('name')
        email = data.get('email')
        user_type = data.get('user_type')

        # Check if required fields are provided
        if not all([username, name, email, user_type]):
            return jsonify({'status': 'error', 'message': 'All fields are required'})

        # Generate a random password
        password = generate_random_password()
        hashed_password = hashlib.sha256(password.encode()).hexdigest()

        created_by = user_name

        existing_user = User.query.filter(or_(User.username == username, User.email == email, User.name == name)).first()

        if existing_user:
            return jsonify({'status': 'error', 'message': 'Account with the same Username, Email, or Name already exists.'})

        new_user = User(username=username, password=hashed_password, name=name, email=email, user_type=user_type, created_by=created_by)
        
        db.session.add(new_user)
        db.session.commit()

        # Generate a verification token
        verification_token = generate_verification_token(new_user.id)

        # Create the verification link
        verification_link = url_for('verify', token=verification_token, _external=True)

        # Send the verification email
        msg = Message('Account Verification', sender='kanuparthisaiganesh582@gmail.com', recipients=[new_user.email])
        
        msg.html = f'''
        <html>
        <head>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                }}
                .container {{
                    max-width: 600px;
                    margin: auto;
                    padding: 20px;
                    border: 1px solid #ddd;
                    border-radius: 5px;
                    box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                }}
                h2 {{
                    color: #333;
                }}
                p {{
                    color: #555;
                }}
                ul {{
                    list-style-type: none;
                    padding: 0;
                }}
                ul li {{
                    background: #f9f9f9;
                    margin: 5px 0;
                    padding: 10px;
                    border: 1px solid #ddd;
                    border-radius: 3px;
                }}
                a {{
                    color: #1a73e8;
                    text-decoration: none;
                }}
                a:hover {{
                    text-decoration: underline;
                }}
                .footer {{
                    margin-top: 20px;
                    font-size: 0.9em;
                    color: #888;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h2>Hello {new_user.name},</h2>
                <p>We are pleased to inform you that your account has been successfully created for the <strong>ATS Makonis Talent Track Pro</strong>. Here are your login credentials:</p>
                <ul>
                    <li><strong>Username:</strong> {new_user.username}</li>
                    <li><strong>Password:</strong> {password}</li>
                </ul>
                <p>Please note that the verification link will expire after 24 hours.</p>
                <p>To verify your account, please click on the following link:</p>
                <p><a href="{verification_link}">Verify Your Account</a></p>
                <p>After successfully verifying your account, you can access the application using the following link:</p>
                <p><a href="https://ats-makonis.netlify.app/">Application Link (Post Verification)</a></p>
                <p>If you have any questions or need assistance, please feel free to reach out.</p>
                <p>Best regards,</p>
                <p><strong>ATS Makonis Talent Track Pro Team</strong></p>
                <div class="footer">
                    <p>This is an automated message, please do not reply.</p>
                </div>
            </div>
        </body>
        </html>
        '''
        mail.send(msg)

        return jsonify({'status': 'success',
            'message': 'A verification email has been sent to your email address. Please check your inbox.',
            'success_message': 'Account created successfully'
        })
    else:
        return jsonify({'message': 'You do not have permission to create recruiter accounts.'})



# @app.route('/signup', methods=['POST'])
# def signup():
#     data = request.json
#     user_id = data.get('user_id')  # Using get method to avoid KeyError
#     user = User.query.filter_by(id=user_id).first()

#     if not user:
#         return jsonify({'status': 'error', 'message': 'Invalid user ID or user does not exist.'})

#     user_type = user.user_type
#     user_name = user.username

#     if user_type == 'management':
#         username = data.get('username')
#         name = data.get('name')
#         email = data.get('email')
#         user_type = data.get('user_type')

#         # Check if required fields are provided
#         if not all([username, name, email, user_type]):
#             return jsonify({'status': 'error', 'message': 'All fields are required'})

#         # Generate a random password
#         password = generate_random_password()
#         hashed_password = hashlib.sha256(password.encode()).hexdigest()

#         created_by = user_name

#         existing_user = User.query.filter(or_(User.username == username, User.email == email, User.name == name)).first()

#         if existing_user:
#             return jsonify({'status': 'error', 'message': 'Account with the same Username, Email, or Name already exists.'})

#         new_user = User(username=username, password=hashed_password, name=name, email=email, user_type=user_type, created_by=created_by)
        
#         db.session.add(new_user)
#         db.session.commit()

#         # Generate a verification token
#         verification_token = generate_verification_token(new_user.id)

#         # Create the verification link
#         verification_link = url_for('verify', token=verification_token, _external=True)

#         # Send the verification email
#         msg = Message('Account Verification', sender='kanuparthisaiganesh582@gmail.com', recipients=[new_user.email])
#         msg.body = f'Hello {new_user.name},\n\n We are pleased to inform you that your account has been successfully created for the ATS Makonis Talent Track Pro. Here are your login credentials:\n\nUsername: {new_user.username}\nPassword: {password}\n\n Please note that the verification link will expire after 24 hours. \n\n After successfully verifying your account, you can access the application using the following link : \n\n Application Link (Post Verification): https://ats-makonis.netlify.app/ \n\n To verify your account, please click on the following link: {verification_link} \n\n If you have any questions or need assistance, please feel free to reach out. \n\n Best regards, '
#         mail.send(msg)

#         # return jsonify({'message': 'A verification email has been sent to your email address. Please check your inbox.'})
#         return jsonify({'status': 'success',
#             'message': 'A verification email has been sent to your email address. Please check your inbox.',
#             'success_message': 'Account created successfully'
#             })
#     else:
#         return jsonify({'message': 'You do not have permission to create recruiter accounts.'})


import hashlib

@app.route('/signup-onetime', methods=['POST'])
def signup_onetime():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    name = data.get('name')
    email = data.get('email')
    user_type = 'management'
    registration_completed = 'one_time'

    # Hash the password using SHA-256
    hashed_password = hashlib.sha256(password.encode()).hexdigest()

    user_onetime = User.query.filter_by(registration_completed='one_time').first()
    if user_onetime:
        return jsonify({'status': 'error','message': 'The one-time registration for this application has already been completed.'})

    new_user = User(username=username, password=hashed_password, name=name,
                    email=email, user_type=user_type, registration_completed=registration_completed)

    db.session.add(new_user)
    db.session.commit()

    # Generate a verification token
    verification_token = generate_verification_token(new_user.id)

    # Create the verification link
    verification_link = url_for('verify', token=verification_token, _external=True)

    # Send the verification email
    msg = Message('Account Verification', sender='kanuparthisaiganesh582@gmail.com', recipients=[new_user.email])
    
    msg.html = f'''
    <html>
    <head>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                background-color: #f7f7f7;
                margin: 0;
                padding: 0;
            }}
            .container {{
                max-width: 600px;
                margin: 20px auto;
                background-color: #ffffff;
                padding: 20px;
                border: 1px solid #e0e0e0;
                border-radius: 10px;
            }}
            h2 {{
                color: #333333;
                font-size: 24px;
                margin-bottom: 20px;
            }}
            p {{
                color: #666666;
                line-height: 1.6;
                margin: 10px 0;
            }}
            .credentials {{
                background-color: #f9f9f9;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #dddddd;
            }}
            .credentials li {{
                margin: 5px 0;
                padding: 5px;
            }}
            a {{
                color: #1a73e8;
                text-decoration: none;
            }}
            a:hover {{
                text-decoration: underline;
            }}
            .button {{
                display: inline-block;
                padding: 10px 20px;
                margin: 20px 0;
                font-size: 16px;
                color: #ffffff;
                background-color: #333333;
                border-radius: 5px;
                text-align: center;
                text-decoration: none;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 0.9em;
                color: #999999;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h2>Account Verification</h2>
            <p>Hello {new_user.name},</p>
            <p>We are pleased to inform you that your account has been successfully created for the <strong>ATS Makonis Talent Track Pro</strong>. Here are your login credentials:</p>
            <div class="credentials">
                <ul>
                    <li><strong>Username:</strong> {new_user.username}</li>
                    <li><strong>Password:</strong> {password}</li>
                </ul>
            </div>
            <p>Please note that the verification link will expire in 24 hours.</p>
            <p>To verify your account, please click the button below:</p>
            <p><a href="{verification_link}" class="button">Verify Your Account</a></p>
            <p>After verifying your account, you can access the application using the following link:</p>
            <p><a href="https://ats-makonis.netlify.app/" class="button">Go to ATS Makonis Talent Track Pro</a></p> 
            <p>If you have any questions or need assistance, please feel free to reach out to our support team.</p>
            <p>Best regards,</p>
            <p><strong>ATS Makonis Talent Track Pro Team</strong></p>
            <div class="footer">
                <p>This is an automated message, please do not reply.</p>
            </div>
        </div>
    </body>
    </html>
    '''
    mail.send(msg)

    return jsonify({'status': 'success',
        'message': 'A verification email has been sent to your email address. Please check your inbox.',
        'success_message': 'Account created successfully'
    }), 200


# @app.route('/signup-onetime', methods=['POST'])
# def signup_onetime():
#     if request.method == 'POST':
#         username = request.json.get('username')
#         password = request.json.get('password')
#         name = request.json.get('name')
#         email = request.json.get('email')
#         user_type = 'management'
#         registration_completed = 'one_time'

#         # Hash the password using SHA-256
#         hashed_password = hashlib.sha256(password.encode()).hexdigest()

#         user_onetime = User.query.filter_by(registration_completed='one_time').first()
#         if user_onetime:
#             return jsonify({'message': 'The one-time registration for this application has already been completed.'}),400

#         new_user = User(username=username, password=hashed_password, name=name,
#                         email=email, user_type=user_type, registration_completed=registration_completed)

#         db.session.add(new_user)
#         db.session.commit()

#         # Generate a verification token
#         verification_token = generate_verification_token(new_user.id)

#         # Create the verification link
#         verification_link = url_for('verify', token=verification_token, _external=True)

#         # Construct the message body with username and plaintext password
#         message_body = f'Hello {new_user.name},\n\nWe are pleased to inform you that your account has been successfully created for the ATS Makonis Talent Track Pro.\n\nYour login credentials:\n\nUsername: {new_user.username}\nPassword: {password}\n\nTo complete the account setup, kindly click on the verification link below:\n{verification_link}\n\nPlease note that the verification link will expire after 24 hours.\n\nAfter successfully verifying your account, you can access the application using the following link:\n\nApplication Link (Post Verification): https://ats-makonis.netlify.app/\n\nIf you have any questions or need assistance, please feel free to reach out.\n\nBest regards,'

#         # Send the verification email
#         msg = Message('Account Verification', sender='kanuparthisaiganesh582@gmail.com', recipients=[new_user.email])
#         msg.body = message_body
#         mail.send(msg)

#         # return jsonify({'message': 'A verification email has been sent to your email address. Please check your inbox.'})
#         return jsonify({
#             'message': 'A verification email has been sent to your email address. Please check your inbox.',
#             'success_message': 'Account created successfully'
#             }),200

#     return jsonify({'message': 'Invalid request method.'}),400

# @app.route('/login/recruiter', methods=['POST'])
# def recruiter_login():
#     verification_msg = request.args.get('verification_msg')
#     reset_message = request.args.get('reset_message')
#     session_timeout_msg = request.args.get("session_timeout_msg")
#     password_message = request.args.get('password_message')

#     if request.method == 'POST':
#         username = request.json.get('username')
#         password = request.json.get('password')

#         # Hash the entered password
#         hashed_password = hashlib.sha256(password.encode()).hexdigest()

#         # Check if the user exists and the password is correct
#         user = User.query.filter_by(username=username, password=hashed_password, user_type='recruiter').first()

#         if user:
#             if user.is_active:  # Check if the user is active
#                 if user.is_verified:
#                     # Set the user session variables
#                     session['user_id'] = user.id
#                     session['user_type'] = user.user_type
#                     session['username'] = user.username
#                     session['user_name'] = user.name
#                     session['JWT Token'] = secrets.token_hex(16)
#                     return jsonify({'status': 'success', 'redirect': url_for('dashboard'),'user_id': user.id})
#                 else:
#                     error = 'Your account is not verified yet. Please check your email for the verification link.'
#             else:
#                 error = 'Your account is not active. Please contact the administrator.'
#         else:
#             error = 'Invalid username or password'

#         return jsonify({'status': 'error', 'error': error})

#     # For GET requests, return necessary data
#     return jsonify({
#         'status': 'success',
#         'verification_msg': verification_msg,
#         'reset_message': reset_message,
#         'session_timeout_msg': session_timeout_msg,
#         'password_message': password_message
#     })


@app.route('/login/recruiter', methods=['POST'])
def recruiter_login():
    verification_msg = request.args.get('verification_msg')
    reset_message = request.args.get('reset_message')
    session_timeout_msg = request.args.get("session_timeout_msg")
    password_message = request.args.get('password_message')

    if request.method == 'POST':
        username = request.json.get('username')
        password = request.json.get('password')

        # Hash the entered password
        hashed_password = hashlib.sha256(password.encode()).hexdigest()

        # Check if the user exists and the password is correct
        user = User.query.filter_by(username=username, password=hashed_password, user_type='recruiter').first()

        if user:
            if user.is_active:  # Check if the user is active
                if user.is_verified:
                    # Set the user session variables
                    session['user_id'] = user.id
                    session['user_type'] = user.user_type
                    session['username'] = user.username
                    session['user_name'] = user.name
                    session['JWT Token'] = secrets.token_hex(16)
                    return jsonify({'status': 'success', 'redirect': url_for('dashboard'),'user_id': user.id})
                else:
                    message = 'Your account is not verified yet. Please check your email for the verification link.'
            else:
                message = 'Your account is not active. Please contact the administrator.'
        else:
            message= 'Invalid username or password'

        return jsonify({'status': 'error', 'message': message})

    # For GET requests, return necessary data
    return jsonify({
        'status': 'success',
        'verification_msg': verification_msg,
        'reset_message': reset_message,
        'session_timeout_msg': session_timeout_msg,
        'password_message': password_message
    })


import hashlib

import hashlib

@app.route('/login/management', methods=['POST'])
def management_login():
    username = request.json.get('username')
    password = request.json.get('password')
    verification_msg_manager = request.args.get('verification_msg_manager')
    
    # Check if the user exists
    user = User.query.filter_by(username=username, user_type='management').first()
    
    if user:
        # Hash the provided password using the same hash function and parameters used to hash the passwords in the database
        hashed_password = hashlib.sha256(password.encode()).hexdigest()
        
        # Compare the hashed password with the hashed password stored in the database
        if hashed_password == user.password:
            if user.is_active:  # Check if the user is active
                if user.is_verified:
                    # Set the user session variables
                    session['user_id'] = user.id
                    session['user_type'] = user.user_type
                    session['username'] = user.username
                    session['user_name'] = user.name
                    session['JWT Token'] = secrets.token_hex(16)
                    return jsonify({'status': 'success', 'redirect': url_for('dashboard'),'user_id':user.id})
                else:
                    message = 'Your account is not verified yet. Please check your email for the verification link.'
            else:
                message = 'Your account is not active. Please contact the administrator.'
        else:
            message = 'Invalid username or password'
    else:
        message = 'Invalid username or password'

    return jsonify({'status': 'error', 'message': message, 'verification_msg_manager': verification_msg_manager})

# @app.route('/get_recruiters', methods=['GET'])   
# def get_recruiters_list():
#     recruiters = User.query.filter_by(user_type='recruiter').all()
    
#     # Assuming you want to return a list of dictionaries containing user details
#     recruiters_list = []
#     for recruiter in recruiters:
#         recruiter_dict = {
#             'id': recruiter.id,
#             'username': recruiter.username,
#             'user_type': recruiter.user_type
#             # Add more fields if needed
#         }
#         recruiters_list.append(recruiter_dict)
    
#     return jsonify(recruiters_list)


@app.route('/get_recruiters', methods=['GET'])
def get_recruiters_list():
    recruiters = User.query.filter_by(user_type='recruiter').all()
    management = User.query.filter_by(user_type='management').all()

    # Extracting only usernames
    recruiter_usernames = [recruiter.username for recruiter in recruiters]
    management_usernames = [manager.username for manager in management]

    return jsonify({
        'recruiters': recruiter_usernames,
        'management': management_usernames
    })


@app.route('/get_recruiters_candidate', methods=['POST'])
def get_recruiters_candidate():
    data = request.json
    
    if not data or 'user_name' not in data:
        return jsonify({'status': 'error', 'message': 'Invalid input'}), 400
    
    username = data['user_name']
    
    # Find the user with the given username who is either a recruiter or in management
    user = User.query.filter((User.username == username) & (User.user_type.in_(['recruiter', 'management']))).first()
    
    if user:
        # If the user is a recruiter, find all candidates linked with the user's username
        if user.user_type == 'recruiter':
            candidates = Candidate.query.filter_by(recruiter=username).all()
        # If the user is in management, find all candidates where recruiter matches the username
        elif user.user_type == 'management':
            candidates = Candidate.query.filter(
                (Candidate.recruiter == username) | (Candidate.management == username)
            ).all()
        else:
            return jsonify({'status': 'error', 'message': 'User type not authorized'}), 403
        
        # Prepare response data
        candidates_list = [
            {
                'id': candidate.id,
                'username': candidate.name,
                'status': candidate.status,
                'client':candidate.client,
                'profile': candidate.profile,
                'recruiter': candidate.recruiter,
                'management': candidate.management
            } 
            for candidate in candidates
        ]
        return jsonify(candidates_list)
    else:
        return jsonify({'status': 'error', 'message': 'User not found or not authorized'}), 404


from flask_mail import Message


def assign_candidates_notification(recruiter_email, new_recruiter_name, candidates_data):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                Candidate Assignment Notification
            </div>
            <p>Hi {new_recruiter_name},</p>
            <p>Candidate data has been transferred to your ATS account.</p>
            <p>Please find the details below:</p>
            <table>
                <tr>
                    <th>Job ID</th>
                    <th>Client Name</th>
                    <th>Role/Profile</th>
                    <th>Candidate Name</th>
                    <th>Previous Recruiter</th>
                </tr>
                {candidates_data}
            </table>
            <p>Please check <b>ATS Dashboard</b> page for more details.</p>
            <p>Best Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        'Candidate Assignment Notification',
        sender='kanuparthisaiganesh582@gmail.com',
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)


# def assign_candidates_notification(recruiter_email, new_recruiter_name, candidates_data):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 color: #333;
#                 line-height: 1.6;
#                 background-color: #f4f4f4;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #ffffff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #4CAF50;
#                 color: white;
#                 padding: 10px;
#                 text-align: center;
#                 font-size: 20px;
#                 border-radius: 8px 8px 0 0;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 10px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 8px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #4CAF50;
#                 color: white;
#             }}
            # # th.job-id-column {{
            # #     width: 120px; /* Increase the width for the Job ID column */
            # # }}
            #  th.clientName-column {{
            #     width: 100px; /* Increase the width for the Job ID column */
            # }}
#             tr:nth-child(even) {{
#                 background-color: #f9f9f9;
#             }}
#             p {{
#                 margin: 10px 0;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
    # <body>
    #     <div class="container">
    #         <div class="header">
    #             Candidate Assignment Notification
    #         </div>
    #         <p>Hi {new_recruiter_name},</p>
    #         <p>Candidate data has been transferred to your ATS account.</p>
    #         <p> Please find the details below:</p>
    #         <table>
    #             <tr>
    #                 <th class="job-id-column">Job ID</th>
    #                 <th class="clientName-column">Client Name</th>
    #                 <th>Role/Profile</th>
    #                 <th>Candidate Name</th>
    #                 <th>Previous Recruiter</th>
    #             </tr>
    #             {candidates_data}
    #         </table>
    #         <p>Please check <b>ATS Dashboard</b> page for more details.</p>
    #         <p>Best Regards,</p>
    #         <p><b>Makonis Talent Track Pro Team</b></p>
    #     </div>
    # </body>
#     </html>
#     """

    # msg = Message(
    #     # 'Candidate Assignment Notification',
    #     f'Candidate Data Transferred',
    #     sender='kanuparthisaiganesh582@gmail.com',
    #     recipients=[recruiter_email]
    # )
    # msg.html = html_body
    # mail.send(msg)


# def assign_candidates_notification(recruiter_email, new_recruiter_name, candidates_data):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 color: #333;
#                 line-height: 1.6;
#                 background-color: #f4f4f4;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #ffffff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #4CAF50;
#                 color: white;
#                 padding: 10px;
#                 text-align: center;
#                 font-size: 20px;
#                 border-radius: 8px 8px 0 0;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 10px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 8px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #4CAF50;
#                 color: white;
#             }}
#             th.job-id-column {{
#                 width: 120px; /* Increase the width for the Job ID column */
#             }}
#             tr:nth-child(even) {{
#                 background-color: #f9f9f9;
#             }}
#             p {{
#                 margin: 10px 0;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
#     <body>
#         <div class="container">
#             <div class="header">
#                 Candidate Assignment Notification
#             </div>
#             <p>Dear {new_recruiter_name},</p>
#             <p>The following candidates have been assigned to you:</p>
#             <table>
#                 <tr>
#                     <th class="job-id-column">Job ID</th>
#                     <th>Client</th>
#                     <th>Profile</th>
#                     <th>Candidate Name</th>
#                 </tr>
#                 {candidates_data}
#             </table>
#             <p>Check your dashboard for more details.</p>
#             <p>Regards,</p>
#             <p>Your Company</p>
#         </div>
#     </body>
#     </html>
#     """

#     msg = Message(
#         'Candidate Assignment Notification',
#         sender='ganesh.s@makonissoft.com',
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     mail.send(msg)

###########################################################################

# def assign_candidates_notification(recruiter_email, new_recruiter_name, candidates_data):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
#                 color: #444;
#                 line-height: 1.6;
#                 background-color: #f8f8f8;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #fff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #007BFF;
#                 color: #fff;
#                 padding: 20px;
#                 text-align: center;
#                 font-size: 24px;
#                 border-top-left-radius: 8px;
#                 border-top-right-radius: 8px;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 20px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 12px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #007BFF;
#                 color: #fff;
#             }}
#             tr:nth-child(even) {{
#                 background-color: #f2f2f2;
#             }}
#             p {{
#                 margin: 10px 0;
#                 text-align: justify;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
#     <body>
#         <div class="container">
#             <div class="header">
#                 Candidate Assignment Notification
#             </div>
#             <p>Dear {new_recruiter_name},</p>
#             <p>The following candidates have been assigned to you:</p>
#             <table>
#                 <tr>
#                     <th>Job ID</th>
#                     <th>Candidate ID</th>
#                     <th>Client</th>
#                     <th>Profile</th>
#                     <th>Candidate Name</th>
#                 </tr>
#                 {candidates_data}
#             </table>
#             <p>Check your dashboard for more details.</p>
#             <p>Regards,</p>
#             <p>Your Company</p>
#             # <div class="footer">
#             #     &copy; 2024 Your Company. All rights reserved.
#             # </div>
#         </div>
#     </body>
#     </html>
#     """

#     msg = Message(
#         'Candidate Assignment Notification',
#         sender='ganesh.s@makonissoft.com',
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     mail.send(msg)


# def assign_candidates_notification(recruiter_email, candidate_data):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 color: #333;
#                 line-height: 1.6;
#                 background-color: #f4f4f4;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #ffffff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #4CAF50;
#                 color: white;
#                 padding: 10px;
#                 text-align: center;
#                 font-size: 20px;
#                 border-radius: 8px 8px 0 0;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 10px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 8px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #4CAF50;
#                 color: white;
#             }}
#             tr:nth-child(even) {{
#                 background-color: #f9f9f9;
#             }}
#             p {{
#                 margin: 10px 0;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
#     <body>
#         <div class="container">
#             <div class="header">
#                 Candidate Assignment Notification
#             </div>
#             <p>Dear Recruiter,</p>
#             <p>The following candidates have been assigned to you:</p>
#             <table>
#                 <tr>
#                     <th>Candidate ID</th>
#                     <th>Name</th>
#                 </tr>
#                 {candidate_data}
#             </table>
#             <p>Check your dashboard for more details.</p>
#             <p>Regards,</p>
#             <p>Your Company</p>
#             <div class="footer">
#                 &copy; 2024 Your Company. All rights reserved.
#             </div>
#         </div>
#     </body>
#     </html>
#     """

#     msg = Message(
#         'Candidate Assignment Notification',
#         sender='ganesh.s@makonissoft.com',
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     mail.send(msg)


# def assign_candidates_notification(recruiter_email, candidate_data):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 color: #333;
#                 line-height: 1.6;
#                 background-color: #f8f8f8;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #ffffff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #007BFF;
#                 color: white;
#                 padding: 10px;
#                 text-align: center;
#                 font-size: 24px;
#                 border-radius: 8px 8px 0 0;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 20px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 12px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #007BFF;
#                 color: white;
#             }}
#             tr:nth-child(even) {{
#                 background-color: #f9f9f9;
#             }}
#             p {{
#                 margin: 10px 0;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
#     <body>
#         <div class="container">
#             <div class="header">
#                 Candidate Assignment Notification
#             </div>
#             <p>Dear Recruiter,</p>
#             <p>The following candidates have been assigned to you:</p>
#             <table>
#                 <tr>
#                     <th>Candidate ID</th>
#                     <th>Name</th>
#                 </tr>
#                 {candidate_data}
#             </table>
#             <p>Check your dashboard for more details.</p>
#             <p>Regards,</p>
#             <p>Your Company</p>
#             <div class="footer">
#                 &copy; 2024 Your Company. All rights reserved.
#             </div>
#         </div>
#     </body>
#     </html>
#     """

#     msg = Message(
#         'Candidate Assignment Notification',
#         sender='ganesh.s@makonissoft.com',
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     mail.send(msg)



# def assign_candidates_notification(recruiter_email, candidate_data):
#     html_body = """
#     <html>
#     <head></head>
#     <body>
#         <p>Dear Recruiter,</p>
#         <p>The following candidates have been assigned to you:</p>
#         <table border="1">
#             <tr>
#                 <th>Candidate ID</th>
#                 <th>Name</th>
#             </tr>
#             {}
#         </table>
#         <p>Check your dashboard for more details.</p>
#         <p>Regards,</p>
#         <p>Your Company</p>
#     </body>
#     </html>
#     """.format(candidate_data)

#     msg = Message(
#         'Candidate Assignment Notification',
#         sender='ganesh.s@makonissoft.com',
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     mail.send(msg)


def job_transfered_to_new_recruiter_notification(recruiter_email, new_recruiter_name, job_data):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                New Job Requirement Assigned 
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>A new requirement has been assigned while transfering candidates to you.</p>
            <p> Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 30%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check in Job Listing page for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        f'New Job Requirement Assigned',
        sender='kanuparthisaiganesh582@gmail.com',
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)

@app.route('/assign_candidate_new_recuriter', methods=['POST'])
def assign_candidate_to_a_new_recruiter():
    data = request.json

    try:
        candidates_data = ""
        current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        new_recruiter = None
        new_recruiter_name = ""

        for candidate_data in data['candidates']:
            candidate_id = candidate_data.get('candidate_id')
            new_recruiter_username = candidate_data.get('new_recruiter')
            current_recruiter_username = candidate_data.get('current_recruiter')

            if not candidate_id or not new_recruiter_username or not current_recruiter_username:
                return jsonify({"error": "Candidate ID, new recruiter username, or current recruiter username not provided"}), 400

            # Fetch candidate details from the database
            candidate = Candidate.query.filter(
                Candidate.id == candidate_id,
                or_(
                    Candidate.recruiter == current_recruiter_username,
                    Candidate.management == current_recruiter_username
                )
            ).first()

            if candidate is None:
                return jsonify({"error": f"Candidate with ID {candidate_id} not found or not assigned to current recruiter/management {current_recruiter_username}"}), 404

            # Fetch job_id associated with the candidate
            job_id = candidate.job_id

            # Update the recruiter for the candidate
            candidate.recruiter = new_recruiter_username
            candidate.data_updated_date = current_datetime.date()
            candidate.data_updated_time = current_datetime.time()

            # Remove the current recruiter from the management field if it matches
            if candidate.management == current_recruiter_username:
                candidate.management = None

            # Append candidate details to the candidates_data string
            candidates_data += f"<tr><td>{candidate.job_id}</td><td>{candidate.client}</td><td>{candidate.profile}</td><td>{candidate.name}</td><td>{escape(current_recruiter_username)}</td></tr>"

            # Update recruiter column in job_post table
            job_post = JobPost.query.filter_by(id=job_id).first()
            if job_post:
                recruiters_list = job_post.recruiter.split(", ") if job_post.recruiter else []

                if current_recruiter_username in recruiters_list:
                    # Check if there are any candidates still linked with this job post and the current recruiter
                    linked_candidates = Candidate.query.filter(
                        Candidate.job_id == job_id,
                        Candidate.recruiter == current_recruiter_username
                    ).count()
                    if linked_candidates == 0:
                        recruiters_list.remove(current_recruiter_username)

                if new_recruiter_username not in recruiters_list:
                    recruiters_list.append(new_recruiter_username)
                    job_post.recruiter = ", ".join(recruiters_list)

                    # Add new notification for the new recruiter
                    new_notification = Notification(job_post_id=job_id, recruiter_name=new_recruiter_username)
                    db.session.add(new_notification)
                    new_notification.num_notification = 1

                    # Send new job post notification to the new recruiter
                    new_recruiter = User.query.filter_by(username=new_recruiter_username).first()
                    new_recruiter_name = new_recruiter.username if new_recruiter else "New Recruiter"
                    job_data = f"""
                    <tr>
                        <td>{job_post.id}</td>
                        <td>{job_post.client}</td>
                        <td>{job_post.role}</td>
                        <td>{job_post.location}</td>
                    </tr>
                    """
                    job_transfered_to_new_recruiter_notification(new_recruiter.email, new_recruiter_name, job_data)

        # Commit changes to the database
        db.session.commit()

        # Send notification email to the new recruiter for candidate assignment
        if candidates_data and new_recruiter:
            assign_candidates_notification(new_recruiter.email, new_recruiter_name, candidates_data)

        return jsonify({'status': 'success', "message": "Candidates assigned successfully."})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', "error": f"Error assigning candidates: {str(e)}"}), 500



# @app.route('/assign_candidate_new_recuriter', methods=['POST'])
# def assign_candidate_to_a_new_recruiter():
#     data = request.json

#     try:
#         candidates_data = ""
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))

#         for candidate_data in data['candidates']:
#             candidate_id = candidate_data.get('candidate_id')
#             new_recruiter_username = candidate_data.get('new_recruiter')
#             current_recruiter_username = candidate_data.get('current_recruiter')

#             if not candidate_id or not new_recruiter_username or not current_recruiter_username:
#                 return jsonify({"error": "Candidate ID, new recruiter username, or current recruiter username not provided"}), 400

#             # Fetch candidate details from the database
#             candidate = Candidate.query.filter(
#                 Candidate.id == candidate_id,
#                 or_(
#                     Candidate.recruiter == current_recruiter_username,
#                     Candidate.management == current_recruiter_username
#                 )
#             ).first()

#             if candidate is None:
#                 return jsonify({"error": f"Candidate with ID {candidate_id} not found or not assigned to current recruiter/management {current_recruiter_username}"}), 404

#             # Fetch job_id associated with the candidate
#             job_id = candidate.job_id

#             # Update the recruiter for the candidate
#             candidate.recruiter = new_recruiter_username
#             candidate.data_updated_date = current_datetime.date()
#             candidate.data_updated_time = current_datetime.time()

#             # Append candidate details to the candidates_data string
#             candidates_data += f"<tr><td>{candidate.job_id}</td><td>{candidate.client}</td><td>{candidate.profile}</td><td>{candidate.name}</td><td>{escape(current_recruiter_username)}</td></tr>"

#             # Update recruiter column in job_post table
#             job_post = JobPost.query.filter_by(id=job_id).first()
#             if job_post:
#                 recruiters_list = job_post.recruiter.split(", ") if job_post.recruiter else []

#                 if current_recruiter_username in recruiters_list:
#                     # Check if there are any candidates still linked with this job post and the current recruiter
#                     linked_candidates = Candidate.query.filter(
#                         Candidate.job_id == job_id,
#                         Candidate.recruiter == current_recruiter_username
#                     ).count()
#                     if linked_candidates == 0:
#                         recruiters_list.remove(current_recruiter_username)

#                 if new_recruiter_username not in recruiters_list:
#                     recruiters_list.append(new_recruiter_username)
#                     job_post.recruiter = ", ".join(recruiters_list)

#                     # # Add new notification for the new recruiter
#                     # new_notification = Notification(job_post_id=job_id, recruiter_name=new_recruiter_username, num_notification=1)
#                     # db.session.add(new_notification)
                    
#                     # Add new notification for the new recruiter
#                     new_notification = Notification(job_post_id=job_id, recruiter_name=new_recruiter_username)
#                     db.session.add(new_notification)
#                     new_notification.num_notification = 1

#                     # Send new job post notification to the new recruiter
#                     new_recruiter = User.query.filter_by(username=new_recruiter_username).first()
#                     new_recruiter_name = new_recruiter.username if new_recruiter else "New Recruiter"
#                     job_data = f"""
#                     <tr>
#                         <td>{job_post.id}</td>
#                         <td>{job_post.client}</td>
#                         <td>{job_post.role}</td>
#                         <td>{job_post.location}</td>
#                     </tr>
#                     """
#                     job_transfered_to_new_recruiter_notification(new_recruiter.email, new_recruiter_name, job_data)

#         # Commit changes to the database
#         db.session.commit()

#         # Send notification email to the new recruiter for candidate assignment
#         if candidates_data:
#             assign_candidates_notification(new_recruiter.email, new_recruiter_name, candidates_data)

#         return jsonify({'status': 'success', "message": "Candidates assigned successfully."})
#     except Exception as e:
#         db.session.rollback()
#         return jsonify({'status': 'error', "error": f"Error assigning candidates: {str(e)}"}), 500
        





from flask import jsonify

@app.route('/candidate_details/<int:candidate_id>/<user_type>/<int:page_no>', methods=['GET'])
def candidate_details(candidate_id, user_type, page_no):
    user_name = session.get('user_name')
    count_notification_no = Notification.query.filter(Notification.notification_status == 'false', Notification.recruiter_name == user_name).count()
    candidate = Candidate.query.get(candidate_id)
    if candidate:
        # Return JSON response with candidate details
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "mobile": candidate.mobile,
            "email": candidate.email,
            "client": candidate.client,
            "current_company": candidate.current_company,
            "position": candidate.position,
            "profile": candidate.profile,
            "current_job_location": candidate.current_job_location,
            "preferred_job_location": candidate.preferred_job_location,
            "resume": candidate.resume,
            "skills": candidate.skills,
            "qualifications": candidate.qualifications,
            "experience": candidate.experience,
            "relevant_experience": candidate.relevant_experience,
            "current_ctc": candidate.current_ctc,
            "expected_ctc": candidate.expected_ctc,
            "linkedin_url": candidate.linkedin_url,
            "notice_period": candidate.notice_period,
            "holding_offer": candidate.holding_offer,
            "user_type": user_type,
            "user_name": user_name,
            "count_notification_no": count_notification_no,
            "page_no": page_no
        })
    else:
        # Return JSON response with error message
        return jsonify({"error_message": "Candidate not found"}), 404


from flask import Flask, jsonify, request, Response
from datetime import date
import json
from sqlalchemy import and_


def date_handler(obj):
    if isinstance(obj, datetime):
        return obj.strftime('%H:%M:%S')  # Format time without microseconds
    elif isinstance(obj, date):
        return obj.isoformat()
    else:
        return None
        
@app.route('/dashboard', methods=['POST'])
def dashboard():
    data = request.json
    print(data)  # Just to verify if data is received properly
    edit_candidate_message = data.get('edit_candidate_message')
    page_no = data.get('page_no')
    candidate_message = data.get('candidate_message')
    signup_message = data.get('signup_message')
    job_message = data.get('job_message')
    update_candidate_message = data.get('update_candidate_message')
    delete_message = data.get("delete_message")

    user_id = data.get('user_id')
    if user_id is None:
        return jsonify({"message": "User ID missing"}), 400

    user = User.query.filter_by(id=user_id).first()
    if user is None:
        return jsonify({"message": "User not found"}), 404

    user_type = user.user_type
    user_name = user.username

    response_data = {}

    # Define case statements for conditional ordering
    conditional_order_date = case(
        (Candidate.data_updated_date != None, Candidate.data_updated_date),
        (Candidate.date_created != None, Candidate.date_created),
        else_=Candidate.date_created
    )

    conditional_order_time = case(
        (Candidate.data_updated_time != None, Candidate.data_updated_time),
        (Candidate.time_created != None, Candidate.time_created),
        else_=Candidate.time_created
    )

    if user_type == 'recruiter':
        recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
        if recruiter is None:
            return jsonify({"message": "Recruiter not found"}), 404

        user_name = recruiter.username
        recruiters = user_name.split(',')  # Splitting the recruiter usernames separated by commas

        print("Recruiter usernames:", recruiters)  # Debugging statement to check the recruiter usernames

        candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.username, Candidate.reference.is_(None)))\
            .order_by(
                desc(conditional_order_date),
                desc(conditional_order_time),
                desc(Candidate.id)  # Ensure newer candidates appear first if dates are equal
            )\
            .all()

        for candidate in candidates:
            print(f"Candidate ID: {candidate.id}, Time Created: {candidate.time_created}")

        jobs_query = JobPost.query.filter(
            or_(*[JobPost.recruiter.like(f"%{recruiter}%") for recruiter in recruiters])
         )
        jobs = jobs_query.all()

        print("Jobs retrieved:", jobs)  # Debugging statement to check the jobs retrieved

        response_data = {
            'user': {
                'id': recruiter.id,
                'name': recruiter.username,
                'user_type': recruiter.user_type,
                'email': recruiter.email
            },
            'user_type': user_type,
            'user_name': user_name,
            'candidates': [{
                'id': candidate.id,
                'job_id': candidate.job_id,
                'name': candidate.name,
                'mobile': candidate.mobile,
                'email': candidate.email,
                'client': candidate.client,
                'current_company': candidate.current_company,
                'position': candidate.position,
                'profile': candidate.profile,
                'current_job_location': candidate.current_job_location,
                'preferred_job_location': candidate.preferred_job_location,
                'qualifications': candidate.qualifications,
                'experience': candidate.experience,
                'relevant_experience': candidate.relevant_experience,
                'current_ctc': candidate.current_ctc,
                'expected_ctc': candidate.expected_ctc,
                'notice_period': candidate.notice_period,
                'linkedin': candidate.linkedin_url,
                'reason_for_job_change':candidate.reason_for_job_change,
                'holding_offer': candidate.holding_offer,
                'recruiter': candidate.recruiter,
                'management': candidate.management,
                'status': candidate.status,
                'remarks': candidate.remarks,
                'skills': candidate.skills,
                'resume': candidate.resume if candidate.resume is not None else "",
                'serving_notice_period': candidate.notice_period,
                'period_of_notice': candidate.period_of_notice,
                'last_working_date': candidate.last_working_date,
                'total_offers': candidate.total,
                'highest_package_in_lpa': candidate.package_in_lpa,
                'buyout': candidate.buyout,
                'date_created': candidate.date_created.isoformat() if candidate.date_created else None,
                'time_created': candidate.time_created.strftime('%H:%M:%S') if candidate.time_created else None,
                'data_updated_date': candidate.data_updated_date.isoformat() if candidate.data_updated_date else None,
                'data_updated_time': candidate.data_updated_time.strftime('%H:%M:%S') if candidate.data_updated_time else None,
                'resume_present': candidate.resume_present
            } for candidate in candidates],
            'jobs': [{
                'id': job.id,
                'client': job.client,
                'experience_min': job.experience_min,
                'experience_max': job.experience_max,
                'budget_min': job.budget_min,
                'budget_max': job.budget_max,
                'location': job.location,
                'shift_timings': job.shift_timings,
                'notice_period': job.notice_period,
                'role': job.role,
                'detailed_jd': job.detailed_jd,
                'jd_pdf': job.jd_pdf,
                'mode': job.mode,
                'recruiter': job.recruiter,
                'management': job.management,
                'date_created': job.date_created.isoformat() if job.date_created else None,
                'time_created': job.time_created.strftime('%H:%M:%S') if job.time_created else None,
                'job_status': job.job_status,
                'job_type': job.job_type,
                'contract_in_months': job.contract_in_months,
                'skills': job.skills,
                'notification': job.notification,
                'data_updated_date': job.data_updated_date.isoformat() if job.data_updated_date else None,
                'data_updated_time': job.data_updated_time.strftime('%H:%M:%S') if job.data_updated_time else None,
                'jd_pdf_present': job.jd_pdf_present
            } for job in jobs],
            'edit_candidate_message': edit_candidate_message,
            'page_no': page_no,
        }
        
    elif user_type == 'management':
        # Define case statements for conditional ordering
        conditional_order_date = case(
        (Candidate.date_created != None, Candidate.date_created),
        else_=Candidate.date_created
        )

        conditional_order_time = case(
        (Candidate.time_created != None, Candidate.time_created),
        else_=Candidate.time_created
        )
        users = User.query.all()
        
        candidates = Candidate.query.filter(Candidate.reference.is_(None))\
            .order_by(
                desc(conditional_order_date),
                desc(conditional_order_time),
                desc(Candidate.id)
            )\
            .all()

        for candidate in candidates:
            print(f"Candidate ID: {candidate.id}, Time Created: {candidate.time_created}")

        jobs = JobPost.query.all()
        
        response_data = {
            'users': [{
                'id': user.id,
                'name': user.name,
                'user_type': user.user_type,
                'email': user.email
            } for user in users],
            'user_type': user_type,
            'user_name': user_name,
            'candidates': [{
                'id': candidate.id,
                'job_id': candidate.job_id,
                'name': candidate.name,
                'mobile': candidate.mobile,
                'email': candidate.email,
                'client': candidate.client,
                'current_company': candidate.current_company,
                'position': candidate.position,
                'profile': candidate.profile,
                'current_job_location': candidate.current_job_location,
                'preferred_job_location': candidate.preferred_job_location,
                'qualifications': candidate.qualifications,
                'experience': candidate.experience,
                'relevant_experience': candidate.relevant_experience,
                'current_ctc': candidate.current_ctc,
                'expected_ctc': candidate.expected_ctc,
                'notice_period': candidate.notice_period,
                'reason_for_job_change':candidate.reason_for_job_change,
                'linkedin': candidate.linkedin_url,
                'holding_offer': candidate.holding_offer,
                'recruiter': candidate.recruiter,
                'management': candidate.management,
                'status': candidate.status,
                'remarks': candidate.remarks,
                'skills': candidate.skills,
                'resume': candidate.resume if candidate.resume is not None else "",
                'serving_notice_period': candidate.notice_period,
                'period_of_notice': candidate.period_of_notice,
                'last_working_date': candidate.last_working_date,
                'total_offers': candidate.total,
                'highest_package_in_lpa': candidate.package_in_lpa,
                'buyout': candidate.buyout,
                'date_created': candidate.date_created.isoformat() if candidate.date_created else None,
                'time_created': candidate.time_created.strftime('%H:%M:%S') if candidate.time_created else None,
                'data_updated_date': candidate.data_updated_date.isoformat() if candidate.data_updated_date else None,
                'data_updated_time': candidate.data_updated_time.strftime('%H:%M:%S') if candidate.data_updated_time else None,
                'resume_present': candidate.resume_present
            } for candidate in candidates],
            'jobs': [{
                'id': job.id,
                'client': job.client,
                'experience_min': job.experience_min,
                'experience_max': job.experience_max,
                'budget_min': job.budget_min,
                'budget_max': job.budget_max,
                'location': job.location,
                'shift_timings': job.shift_timings,
                'notice_period': job.notice_period,
                'role': job.role,
                'detailed_jd': job.detailed_jd,
                'jd_pdf': job.jd_pdf,
                'mode': job.mode,
                'recruiter': job.recruiter,
                'management': job.management,
                'date_created': job.date_created.isoformat() if job.date_created else None,
                'time_created': job.time_created.strftime('%H:%M:%S') if job.time_created else None,
                'job_status': job.job_status,
                'job_type': job.job_type,
                'contract_in_months': job.contract_in_months,
                'skills': job.skills,
                'notification': job.notification,
                'data_updated_date': job.data_updated_date.isoformat() if job.data_updated_date else None,
                'data_updated_time': job.data_updated_time.strftime('%H:%M:%S') if job.data_updated_time else None,
                'jd_pdf_present': job.jd_pdf_present
            } for job in jobs],
            'signup_message': signup_message,
            'job_message': job_message,
            'page_no': page_no,
            'edit_candidate_message': edit_candidate_message
        }
        
    else:
        candidates = Candidate.query.filter_by(recruiter=user.name)\
            .order_by(
                desc(conditional_order_date),
                desc(conditional_order_time),
                desc(Candidate.id)
            )\
            .all()

        for candidate in candidates:
            print(f"Candidate ID: {candidate.id}, Time Created: {candidate.time_created}")
        
        response_data = {
            'user': {
                'id': user.id,
                'name': user.name,
                'user_type': user.user_type,
                'email': user.email
            },
            'user_type': user_type,
            'user_name': user_name,
            'candidates': [{
                'id': candidate.id,
                'job_id': candidate.job_id,
                'name': candidate.name,
                'mobile': candidate.mobile,
                'email': candidate.email,
                'client': candidate.client,
                'current_company': candidate.current_company,
                'position': candidate.position,
                'profile': candidate.profile,
                'current_job_location': candidate.current_job_location,
                'preferred_job_location': candidate.preferred_job_location,
                'qualifications': candidate.qualifications,
                'experience': candidate.experience,
                'relevant_experience': candidate.relevant_experience,
                'current_ctc': candidate.current_ctc,
                'expected_ctc': candidate.expected_ctc,
                'notice_period': candidate.notice_period,
                'linkedin': candidate.linkedin_url,
                'reason_for_job_change':candidate.reason_for_job_change,
                'holding_offer': candidate.holding_offer,
                'recruiter': candidate.recruiter,
                'management': candidate.management,
                'status': candidate.status,
                'remarks': candidate.remarks,
                'skills': candidate.skills,
                'resume': candidate.resume if candidate.resume is not None else "",
                'serving_notice_period': candidate.notice_period,
                'period_of_notice': candidate.period_of_notice,
                'last_working_date': candidate.last_working_date,
                'buyout': candidate.buyout,
                'total_offers': candidate.total,
                'highest_package_in_lpa': candidate.package_in_lpa,
                'date_created': candidate.date_created.isoformat() if candidate.date_created else None,
                'time_created': candidate.time_created.strftime('%H:%M:%S') if candidate.time_created else None,
                'data_updated_date': candidate.data_updated_date.isoformat() if candidate.data_updated_date else None,
                'data_updated_time': candidate.data_updated_time.strftime('%H:%M:%S') if candidate.data_updated_time else None,
                'resume_present': candidate.resume_present
            } for candidate in candidates],
        }
    
    # Convert response_data to JSON string
    response_json = json.dumps(response_data, default=date_handler)

    # Create the response
    return response_json

# def date_handler(obj):
#     if isinstance(obj, date):
#         return obj.isoformat()
#     else:
#         return None

# @app.route('/dashboard', methods=['POST'])
# def dashboard():
#     data = request.json
#     print(data)  # Just to verify if data is received properly
#     edit_candidate_message = data.get('edit_candidate_message')
#     page_no = data.get('page_no')
#     candidate_message = data.get('candidate_message')
#     signup_message = data.get('signup_message')
#     job_message = data.get('job_message')
#     update_candidate_message = data.get('update_candidate_message')
#     delete_message = data.get("delete_message")
    
#     user_id = data.get('user_id')
#     if user_id is None:
#         return jsonify({"message": "User ID missing"}), 400
    
#     user = User.query.filter_by(id=user_id).first()
#     if user is None:
#         return jsonify({"message": "User not found"}), 404
    
#     user_type = user.user_type
#     user_name = user.username
    
#     response_data = {}

#     # Define case statements for conditional ordering
#     conditional_order_date = case(
#         (Candidate.data_updated_date != None, Candidate.data_updated_date),
#         (Candidate.date_created != None, Candidate.date_created),
#         else_=Candidate.date_created
#     )

#     conditional_order_time = case(
#         (Candidate.data_updated_time != None, Candidate.data_updated_time),
#         (Candidate.time_created != None, Candidate.time_created),
#         else_=Candidate.time_created
#     )

#     if user_type == 'recruiter':
#         recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
#         # username = recruiter.username
#         if recruiter is None:
#             return jsonify({"message": "Recruiter not found"}), 404
            
#         # recruiters = recruiter.username.split(',')  # Splitting the recruiter usernames separated by commas
#         user_name = recruiter.username
#         recruiters = user_name.split(',')  # Splitting the recruiter usernames separated by commas

#         print("Recruiter usernames:", recruiters)  # Debugging statement to check the recruiter usernames
        
#         candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.username, Candidate.reference.is_(None)))\
#             .order_by(
#                 desc(conditional_order_date),
#                 desc(conditional_order_time),
#                 desc(Candidate.id)  # Ensure newer candidates appear first if dates are equal
#             )\
#             .all()
#         jobs_query = JobPost.query.filter(
#             or_(*[JobPost.recruiter.like(f"%{recruiter}%") for recruiter in recruiters])
#          )
#         jobs = jobs_query.all()

#         print("Jobs retrieved:", jobs)  # Debugging statement to check the jobs retrieved


#         # jobs = JobPost.query.filter_by(recruiter=user_name).all()
#         # jobs = JobPost.query.filter(JobPost.recruiter.in_(recruiters)).all()
        
#         response_data = {
#             'user': {
#                 'id': recruiter.id,
#                 'name': recruiter.username,
#                 'user_type': recruiter.user_type,
#                 'email': recruiter.email
#                 # Add more attributes as needed
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                 'id': candidate.id,
#                 'job_id': candidate.job_id,
#                 'name': candidate.name,
#                 'mobile': candidate.mobile,
#                 'email': candidate.email,
#                 'client': candidate.client,
#                 'current_company': candidate.current_company,
#                 'position': candidate.position,
#                 'profile': candidate.profile,
#                 'current_job_location': candidate.current_job_location,
#                 'preferred_job_location': candidate.preferred_job_location,
#                 'qualifications': candidate.qualifications,
#                 'experience': candidate.experience,
#                 'relevant_experience': candidate.relevant_experience,
#                 'current_ctc': candidate.current_ctc,
#                 'expected_ctc': candidate.expected_ctc,
#                 'notice_period': candidate.notice_period,
#                 'linkedin_url': candidate.linkedin_url,
#                 'holding_offer': candidate.holding_offer,
#                 'recruiter': candidate.recruiter,
#                 'management': candidate.management,
#                 'status': candidate.status,
#                 'remarks': candidate.remarks,
#                 'skills': candidate.skills,
#                 'resume': candidate.resume if candidate.resume is not None else "",
#                 # 'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                 # 'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                 'serving_notice_period' :candidate.notice_period,
#                 'period_of_notice': candidate.period_of_notice,
#                 'last_working_date': candidate.last_working_date,
#                 'total_offers': candidate.total,
#                 'highest_package_in_lpa' : candidate.package_in_lpa,
#                 'buyout': candidate.buyout,
#                 'date_created': candidate.date_created,
#                 'time_created': candidate.time_created,
#                 'data_updated_date': candidate.data_updated_date,
#                 'data_updated_time': candidate.data_updated_time,
#                 'resume_present':candidate.resume_present
#                 # Add more attributes as needed
#             } for candidate in candidates],
#             'jobs': [{
#                 'id': job.id,
#                 'client': job.client,
#                 'experience_min': job.experience_min,
#                 'experience_max': job.experience_max,
#                 'budget_min': job.budget_min,
#                 'budget_max': job.budget_max,
#                 'location': job.location,
#                 'shift_timings': job.shift_timings,
#                 'notice_period': job.notice_period,
#                 'role': job.role,
#                 'detailed_jd': job.detailed_jd,
#                 'jd_pdf': job.jd_pdf,
#                 'mode': job.mode,
#                 'recruiter': job.recruiter,
#                 'management': job.management,
#                 'date_created': job.date_created,
#                 'time_created': job.time_created,
#                 'job_status': job.job_status,
#                 'job_type': job.job_type,
#                 'contract_in_months': job.contract_in_months,
#                 'skills': job.skills,
#                 'notification': job.notification,
#                 'date_created': job.date_created,
#                 'time_created': job.time_created,
#                 'data_updated_date': job.data_updated_date,
#                 'data_updated_time': job.data_updated_time,
#                 'jd_pdf_present':job.jd_pdf_present
#                 # Add more attributes as needed
#             } for job in jobs],
#             'edit_candidate_message': edit_candidate_message,
#             'page_no': page_no,
#         }
        
#     elif user_type == 'management':
#         # Define case statements for conditional ordering
#         conditional_order_date = case(
#         (Candidate.date_created != None, Candidate.date_created),
#         else_=Candidate.date_created
#         )

#         conditional_order_time = case(
#         (Candidate.time_created != None, Candidate.time_created),
#         else_=Candidate.time_created
#         )

#         users = User.query.all()
        
#         candidates = Candidate.query.filter(Candidate.reference.is_(None))\
#             .order_by(
#                 desc(conditional_order_date),
#                 desc(conditional_order_time),
#                 desc(Candidate.id)  # Ensure newer candidates appear first if dates are equal
#             )\
#             .all()

#         jobs = JobPost.query.all()
        
#         response_data = {
#             'users': [{
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#                 # Add more attributes as needed
#             } for user in users],
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                 'id': candidate.id,
#                 'job_id': candidate.job_id,
#                 'name': candidate.name,
#                 'mobile': candidate.mobile,
#                 'email': candidate.email,
#                 'client': candidate.client,
#                 'current_company': candidate.current_company,
#                 'position': candidate.position,
#                 'profile': candidate.profile,
#                 'current_job_location': candidate.current_job_location,
#                 'preferred_job_location': candidate.preferred_job_location,
#                 'qualifications': candidate.qualifications,
#                 'experience': candidate.experience,
#                 'relevant_experience': candidate.relevant_experience,
#                 'current_ctc': candidate.current_ctc,
#                 'expected_ctc': candidate.expected_ctc,
#                 'notice_period': candidate.notice_period,
#                 'linkedin_url': candidate.linkedin_url,
#                 'holding_offer': candidate.holding_offer,
#                 'recruiter': candidate.recruiter,
#                 'management': candidate.management,
#                 'status': candidate.status,
#                 'remarks': candidate.remarks,
#                 'skills': candidate.skills,
#                 'resume': candidate.resume if candidate.resume is not None else "",
#                 # 'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                 # 'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                 'serving_notice_period' :candidate.notice_period,
#                 'period_of_notice': candidate.period_of_notice,
#                 'last_working_date': candidate.last_working_date,
#                 'total_offers': candidate.total,
#                 'highest_package_in_lpa' :candidate.package_in_lpa,
#                 'buyout': candidate.buyout,
#                 'date_created': candidate.date_created,
#                 'time_created': candidate.time_created,
#                 'data_updated_date': candidate.data_updated_date,
#                 'data_updated_time': candidate.data_updated_time,
#                 'resume_present':candidate.resume_present
#                 # Add more attributes as needed
#             } for candidate in candidates],
#             'jobs': [{
#                 'id': job.id,
#                 'client': job.client,
#                 'experience_min': job.experience_min,
#                 'experience_max': job.experience_max,
#                 'budget_min': job.budget_min,
#                 'budget_max': job.budget_max,
#                 'location': job.location,
#                 'shift_timings': job.shift_timings,
#                 'notice_period': job.notice_period,
#                 'role': job.role,
#                 'detailed_jd': job.detailed_jd,
#                 'jd_pdf': job.jd_pdf,
#                 'mode': job.mode,
#                 'recruiter': job.recruiter,
#                 'management': job.management,
#                 'date_created': job.date_created,
#                 'time_created': job.time_created,
#                 'job_status': job.job_status,
#                 'job_type': job.job_type,
#                 'contract_in_months': job.contract_in_months,
#                 'skills': job.skills,
#                 'notification': job.notification,
#                 'date_created': job.date_created,
#                 'time_created': job.time_created,
#                 'data_updated_date': job.data_updated_date,
#                 'data_updated_time': job.data_updated_time,
#                 'jd_pdf_present':job.jd_pdf_present
#                 # Add more attributes as needed
#             } for job in jobs],
#             'signup_message': signup_message,
#             'job_message': job_message,
#             'page_no': page_no,
#             'edit_candidate_message': edit_candidate_message
#         }
        
#     else:
#         candidates = Candidate.query.filter_by(recruiter=user.name)\
#             .order_by(
#                 desc(conditional_order_date),
#                 desc(conditional_order_time),
#                 desc(Candidate.id)  # Ensure newer candidates appear first if dates are equal
#             )\
#             .all()
        
#         response_data = {
#             'user': {
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#                 # Add more attributes as needed
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                 'id': candidate.id,
#                 'job_id': candidate.job_id,
#                 'name': candidate.name,
#                 'mobile': candidate.mobile,
#                 'email': candidate.email,
#                 'client': candidate.client,
#                 'current_company': candidate.current_company,
#                 'position': candidate.position,
#                 'profile': candidate.profile,
#                 'current_job_location': candidate.current_job_location,
#                 'preferred_job_location': candidate.preferred_job_location,
#                 'qualifications': candidate.qualifications,
#                 'experience': candidate.experience,
#                 'relevant_experience': candidate.relevant_experience,
#                 'current_ctc': candidate.current_ctc,
#                 'expected_ctc': candidate.expected_ctc,
#                 'notice_period': candidate.notice_period,
#                 'linkedin_url': candidate.linkedin_url,
#                 'holding_offer': candidate.holding_offer,
#                 'recruiter': candidate.recruiter,
#                 'management': candidate.management,
#                 'status': candidate.status,
#                 'remarks': candidate.remarks,
#                 'skills': candidate.skills,
#                 'resume': candidate.resume if candidate.resume is not None else "",
#                 # 'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                 # 'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                 'serving_notice_period' :candidate.notice_period,
#                 'period_of_notice': candidate.period_of_notice,
#                 'last_working_date': candidate.last_working_date,
#                 'buyout': candidate.buyout,
#                 'total_offers': candidate.total,
#                 'highest_package_in_lpa' : candidate.package_in_lpa,
#                 'date_created': candidate.date_created,
#                 'time_created': candidate.time_created,
#                 'data_updated_date': candidate.data_updated_date,
#                 'data_updated_time': candidate.data_updated_time,
#                 'resume_present':candidate.resume_present
#                 # Add more attributes as needed
#             } for candidate in candidates],
#         }
        
#     # Convert response_data to JSON string
#     response_json = json.dumps(response_data, default=date_handler)

#     # Create the response
#     return response_json
############################################

# def date_handler(obj):
#     if isinstance(obj, date):
#         return obj.isoformat()
#     else:
#         return None

# @app.route('/dashboard', methods=['POST'])
# def dashboard():
#     data = request.json
#     print(data)  # Just to verify if data is received properly
#     edit_candidate_message = data.get('edit_candidate_message')
#     page_no = data.get('page_no')
#     candidate_message = data.get('candidate_message')
#     signup_message = data.get('signup_message')
#     job_message = data.get('job_message')
#     update_candidate_message = data.get('update_candidate_message')
#     delete_message = data.get("delete_message")
    
#     user_id = data.get('user_id')
#     if user_id is None:
#         return jsonify({"message": "User ID missing"}), 400
    
#     user = User.query.filter_by(id=user_id).first()
#     if user is None:
#         return jsonify({"message": "User not found"}), 404
    
#     user_type = user.user_type
#     user_name = user.username
    
#     response_data = {}
    
#     if user_type == 'recruiter':
#         recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
#         if recruiter is None:
#             return jsonify({"message": "Recruiter not found"}), 404
        
#         candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None))).all()  
#         candidates = sorted(candidates, key=lambda candidate: candidate.id)
#         jobs = JobPost.query.filter_by(recruiter=user_name).all()
        
#         response_data = {
#             'user': {
#                 'id': recruiter.id,
#                 'name': recruiter.name,
#                 'user_type': recruiter.user_type,
#                 'email': recruiter.email
#                 # Add more attributes as needed
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         # 'resume': candidate.resume if candidate.resume is not None else "",
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created': candidate.date_created,
#                         'time_created': candidate.time_created,
#                         'data_updated_date': candidate.data_updated_date,
#                         'data_updated_time': candidate.data_updated_time
#                 # Add more attributes as needed
#             } for candidate in candidates],
#             'jobs': [{
#                         'id': job.id,
#                         'client': job.client,
#                         'experience_min': job.experience_min,
#                         'experience_max': job.experience_max,
#                         'budget_min': job.budget_min,
#                         'budget_max': job.budget_max,
#                         'location': job.location,
#                         'shift_timings': job.shift_timings,
#                         'notice_period': job.notice_period,
#                         'role': job.role,
#                         'detailed_jd': job.detailed_jd,
#                         'jd_pdf': job.jd_pdf,
#                         'mode': job.mode,
#                         'recruiter': job.recruiter,
#                         'management': job.management,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'job_status': job.job_status,
#                         'job_type': job.job_type,
#                         'skills': job.skills,
#                         'notification': job.notification,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'data_updated_date': job.data_updated_date,
#                         'data_updated_time': job.data_updated_time

#                 # Add more attributes as needed
#             } for job in jobs],
#             'edit_candidate_message': edit_candidate_message,
#             'page_no': page_no,
#         }
        
#     elif user_type == 'management':
#         users = User.query.all()
#         candidates = Candidate.query.filter(Candidate.reference.is_(None)).all()
#         candidates = sorted(candidates, key=lambda candidate: candidate.id)
#         jobs = JobPost.query.all()
        
#         response_data = {
#             'users': [{
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#                 # Add more attributes as needed
#             } for user in users],
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         # 'resume': candidate.resume if candidate.resume is not None else "",
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created': candidate.date_created,
#                         'time_created': candidate.time_created,
#                         'data_updated_date': candidate.data_updated_date,
#                         'data_updated_time': candidate.data_updated_time
#                 # Add more attributes as needed
#             } for candidate in candidates],
#             'jobs': [{
#                         'id': job.id,
#                         'client': job.client,
#                         'experience_min': job.experience_min,
#                         'experience_max': job.experience_max,
#                         'budget_min': job.budget_min,
#                         'budget_max': job.budget_max,
#                         'location': job.location,
#                         'shift_timings': job.shift_timings,
#                         'notice_period': job.notice_period,
#                         'role': job.role,
#                         'detailed_jd': job.detailed_jd,
#                         'jd_pdf': job.jd_pdf,
#                         'mode': job.mode,
#                         'recruiter': job.recruiter,
#                         'management': job.management,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'job_status': job.job_status,
#                         'job_type': job.job_type,
#                         'skills': job.skills,
#                         'notification': job.notification,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'data_updated_date': job.data_updated_date,
#                         'data_updated_time': job.data_updated_time
#                 # Add more attributes as needed
#             } for job in jobs],
#             'signup_message': signup_message,
#             'job_message': job_message,
#             'page_no': page_no,
#             'edit_candidate_message': edit_candidate_message
#         }
        
#     else:
#         candidates = Candidate.query.filter_by(recruiter=user.name).all()  
        
#         response_data = {
#             'user': {
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#                 # Add more attributes as needed
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume if candidate.resume is not None else "",
#                         # 'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created': candidate.date_created,
#                         'time_created': candidate.time_created,
#                         'data_updated_date': candidate.data_updated_date,
#                         'data_updated_time': candidate.data_updated_time
#                 # Add more attributes as needed
#             } for candidate in candidates],
            
#         }
        
#     # Convert response_data to JSON string
#     response_json = json.dumps(response_data, default=date_handler)

#     # Create the response
#     return response_json

# @app.route('/dashboard', methods=['POST'])
# def dashboard():
#     data = request.json
#     print(data)  # Just to verify if data is received properly
#     edit_candidate_message = data.get('edit_candidate_message')
#     page_no = data.get('page_no')
#     candidate_message = data.get('candidate_message')
#     signup_message = data.get('signup_message')
#     job_message = data.get('job_message')
#     update_candidate_message = data.get('update_candidate_message')
#     delete_message = data.get("delete_message")
    
#     user_id = data.get('user_id')
#     if user_id is None:
#         return jsonify({"message": "User ID missing"}), 400
    
#     user = User.query.filter_by(id=user_id).first()
#     if user is None:
#         return jsonify({"message": "User not found"}), 404
    
#     user_type = user.user_type
#     user_name = user.username
    
#     response_data = {}
    
#     if user_type == 'recruiter':
#         recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
#         if recruiter is None:
#             return jsonify({"message": "Recruiter not found"}), 404
        
#         candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None))).all()  
#         candidates = sorted(candidates, key=lambda candidate: candidate.id)
#         jobs = JobPost.query.filter_by(recruiter=user_name).all()
        
#         response_data = {
#             'user': {
#                 'id': recruiter.id,
#                 'name': recruiter.name,
#                 'user_type': recruiter.user_type,
#                 'email': recruiter.email
#                 # Add more attributes as needed
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created':candidate.date_created,
#                         'time_created':candidate.time_created,
#                         'data_updated_date':candidate.data_updated_date,
#                         'data_updated_time':candidate.data_updated_time
#                 # Add more attributes as needed
#             } for candidate in candidates],
#             'jobs': [{
#                         'id': job.id,
#                         'client': job.client,
#                         'experience_min': job.experience_min,
#                         'experience_max': job.experience_max,
#                         'budget_min': job.budget_min,
#                         'budget_max': job.budget_max,
#                         'location': job.location,
#                         'shift_timings': job.shift_timings,
#                         'notice_period': job.notice_period,
#                         'role': job.role,
#                         'detailed_jd': job.detailed_jd,
#                         'jd_pdf': job.jd_pdf,
#                         'mode': job.mode,
#                         'recruiter': job.recruiter,
#                         'management': job.management,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'job_status': job.job_status,
#                         'job_type': job.job_type,
#                         'skills': job.skills,
#                         'notification': job.notification,
#                         'date_created':job.date_created,
#                         'time_created':job.time_created,
#                         'data_updated_date':job.data_updated_date,
#                         'data_updated_time':job.data_updated_time

#                 # Add more attributes as needed
#             } for job in jobs],
#             'edit_candidate_message': edit_candidate_message,
#             'page_no': page_no,
#         }
        
#     elif user_type == 'management':
#         users = User.query.all()
#         candidates = Candidate.query.filter(Candidate.reference.is_(None)).all()
#         candidates = sorted(candidates, key=lambda candidate: candidate.id)
#         jobs = JobPost.query.all()
        
#         response_data = {
#             'users': [{
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#                 # Add more attributes as needed
#             } for user in users],
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created':candidate.date_created,
#                         'time_created':candidate.time_created,
#                         'data_updated_date':candidate.data_updated_date,
#                         'data_updated_time':candidate.data_updated_time
#                 # Add more attributes as needed
#             } for candidate in candidates],
#             'jobs': [{
#                         'id': job.id,
#                         'client': job.client,
#                         'experience_min': job.experience_min,
#                         'experience_max': job.experience_max,
#                         'budget_min': job.budget_min,
#                         'budget_max': job.budget_max,
#                         'location': job.location,
#                         'shift_timings': job.shift_timings,
#                         'notice_period': job.notice_period,
#                         'role': job.role,
#                         'detailed_jd': job.detailed_jd,
#                         'jd_pdf': job.jd_pdf,
#                         'mode': job.mode,
#                         'recruiter': job.recruiter,
#                         'management': job.management,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'job_status': job.job_status,
#                         'job_type': job.job_type,
#                         'skills': job.skills,
#                         'notification': job.notification,
#                         'date_created':job.date_created,
#                         'time_created':job.time_created,
#                         'data_updated_date':job.data_updated_date,
#                         'data_updated_time':job.data_updated_time

#                 # Add more attributes as needed
#             } for job in jobs],
#             'signup_message': signup_message,
#             'job_message': job_message,
#             'page_no': page_no,
#             'edit_candidate_message': edit_candidate_message
#         }
        
#     else:
#         candidates = Candidate.query.filter_by(recruiter=user.name).all()  
        
#         response_data = {
#             'user': {
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#                 # Add more attributes as needed
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created':candidate.date_created,
#                         'time_created':candidate.time_created,
#                         'data_updated_date':candidate.data_updated_date,
#                         'data_updated_time':candidate.data_updated_time
#                 # Add more attributes as needed
#             } for candidate in candidates],
#         }
        
#     # Convert response_data to JSON string
#     response_json = json.dumps(response_data)

#     # Create the response
#     return jsonify(response_data)

# @app.route('/dashboard', methods=['POST'])
# def dashboard():
#     data = request.json
#     print(data)  # Just to verify if data is received properly
#     edit_candidate_message = data.get('edit_candidate_message')
#     page_no = data.get('page_no')
#     candidate_message = data.get('candidate_message')
#     signup_message = data.get('signup_message')
#     job_message = data.get('job_message')
#     update_candidate_message = data.get('update_candidate_message')
#     delete_message = data.get("delete_message")
#     # data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_type = user.user_type
#     user_name = user.username
#     response_data = {}
#     if user_id and user_type:
#         if user_type == 'recruiter':
#             recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
#             if recruiter:
#                 candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None))).all()  # Filter candidates by recruiter's name
#                 candidates = sorted(candidates, key=lambda candidate: candidate.id)
#                 jobs = JobPost.query.filter_by(recruiter=user_name).all()  # Filter jobs by recruiter's name
#                 count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                                   Notification.recruiter_name == user_name).count()
#                 career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
#                                                                   Career_notification.recruiter_name == user_name).count()
#                 response_data = {
#                     'user': {
#                         'id': recruiter.id,
#                         'name': recruiter.name,
#                         'user_type': recruiter.user_type,
#                         'email': recruiter.email
#                         # Add more attributes as needed
#                     },
#                     'user_type': user_type,
#                     'user_name': user_name,
#                     'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created':candidate.date_created,
#                         'time_created':candidate.time_created
            
#                         # Add more attributes as needed
#                     } for candidate in candidates],
#                     'jobs': [{
#                         'id': job.id,
#                         'client': job.client,
#                         'experience_min': job.experience_min,
#                         'experience_max': job.experience_max,
#                         'budget_min': job.budget_min,
#                         'budget_max': job.budget_max,
#                         'location': job.location,
#                         'shift_timings': job.shift_timings,
#                         'notice_period': job.notice_period,
#                         'role': job.role,
#                         'detailed_jd': job.detailed_jd,
#                         'jd_pdf': job.jd_pdf,
#                         'mode': job.mode,
#                         'recruiter': job.recruiter,
#                         'management': job.management,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'job_status': job.job_status,
#                         'job_type': job.job_type,
#                         'skills': job.skills,
#                         'notification': job.notification
#                         # Add more attributes as needed
#                     } for job in jobs],
#                     'edit_candidate_message': edit_candidate_message,
#                     'page_no': page_no,
#                     'count_notification_no': count_notification_no,
#                     'career_count_notification_no': career_count_notification_no
#                 }
#         elif user_type == 'management':
#             users = User.query.all()
#             candidates = Candidate.query.filter(Candidate.reference.is_(None)).all()
#             candidates = sorted(candidates, key=lambda candidate: candidate.id)
#             jobs = JobPost.query.all()
#             response_data = {
#                 'users': [{
#                     'id': user.id,
#                     'name': user.name,
#                     'user_type': user.user_type,
#                     'email': user.email
                     
#                     # Add more attributes as needed
#                 } for user in users],
#                 'user_type': user_type,
#                 'user_name': user_name,
#                 'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created':candidate.date_created,
#                         'time_created':candidate.time_created
#                     # Add more attributes as needed
#                 } for candidate in candidates],
#                 'jobs': [{
#                     'id': job.id,
#                     'client': job.client,
#                     'experience_min': job.experience_min,
#                     'experience_max': job.experience_max,
#                     'budget_min': job.budget_min,
#                     'budget_max': job.budget_max,
#                     'location': job.location,
#                     'shift_timings': job.shift_timings,
#                     'notice_period': job.notice_period,
#                     'role': job.role,
#                     'detailed_jd': job.detailed_jd,
#                     'jd_pdf': job.jd_pdf,
#                     'mode': job.mode,
#                     'recruiter': job.recruiter,
#                     'management': job.management,
#                     'date_created': job.date_created,
#                     'time_created': job.time_created,
#                     'job_status': job.job_status,
#                     'job_type': job.job_type,
#                     'skills': job.skills,
#                     'notification': job.notification
#                     # Add more attributes as needed
#                 } for job in jobs],
#                 'signup_message': signup_message,
#                 'job_message': job_message,
#                 'page_no': page_no,
#                 'edit_candidate_message': edit_candidate_message
#             }
#         else:
#             user = User.query.filter_by(id=user_id).first()
#             if user:
#                 candidates = Candidate.query.filter_by(recruiter=user.name).all()  # Filter candidates by user's name
#                 response_data = {
#                     'user': {
#                         'id': user.id,
#                         'name': user.name,
#                         'user_type': user.user_type,
#                         'email': user.email
#                         # Add more attributes as needed
#                     },
#                     'user_type': user_type,
#                     'user_name': user_name,
#                     'candidates': [{
#                         'id': candidate.id,
#                         'job_id':candidate.job_id,
#                         'name': candidate.name,
#                         'email': candidate.email,
#                         'mobile': candidate.mobile,
#                         'client':candidate.client,
#                         'skills':candidate.skills,
#                         "profile": candidate.profile, 
#                         'recruiter':candidate.recruiter,
#                         "management":candidate.management,
#                         'resume': candidate.resume,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications':candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience':candidate.relevant_experience,
#                         'current_ctc':candidate.current_ctc,
#                         'experted_ctc': candidate.expected_ctc,
#                         "total":candidate.total,
#                         'package_in_lpa':candidate.package_in_lpa,
#                         'holding_offer':candidate.holding_offer,
#                         'status': candidate.status,
#                         'reason_for_job_change':candidate.reason_for_job_change,
#                         'remarks':candidate.remarks,
#                         'screening_done': candidate.screening_done,
#                         'rejected_at_screening': candidate.rejected_at_screening,
#                         'l1_cleared':candidate.l1_cleared,
#                         'rejected_at_l1':candidate.rejected_at_l1,
#                         "dropped_after_clearing_l1": candidate.dropped_after_clearing_l1,
#                         'l2_cleared':candidate.l1_cleared,
#                         'rejected_at_l2':candidate.rejected_at_l1,
#                         "dropped_after_clearing_l2": candidate.dropped_after_clearing_l1,
#                         'onboarded': candidate.onboarded,
#                         'dropped_after_onboarding': candidate.dropped_after_onboarding,
#                         'linkedin_url': candidate.linkedin_url,
#                         'period_of_notice': candidate.period_of_notice,
#                         'reference': candidate.reference,
#                         'reference_name': candidate.reference_name,
#                         'reference_position': candidate.reference_position,
#                         'reference_information': candidate.reference_information,
#                         'comments':candidate.comments,
#                         "time_created":str(candidate.time_created),
#                         "date_created": str(candidate.date_created)
#                         # Add more attributes as needed
#                     } for candidate in candidates],
#                 }
#     else:
#         response_data = {"message": "User ID or User Type missing"}
#     # Convert date objects to string representations before returning the response
#     for job in response_data.get('jobs', []):
#         if job.get('date_created'):
#             job['date_created'] = job['date_created'].isoformat()
#     # Convert response_data to JSON string
#     response_json = json.dumps(response_data)

#     # Create the response
#     return Response(response_json, content_type='application/json')
#     # return Response(json.dumps(response_data, default=str), content_type='application/json')

# @app.route('/dashboard', methods=['POST'])
# def dashboard():
#     data = request.json
#     print(data)  # Just to verify if data is received properly

#     edit_candidate_message = data.get('edit_candidate_message')
#     page_no = data.get('page_no')
#     candidate_message = data.get('candidate_message')
#     signup_message = data.get('signup_message')
#     job_message = data.get('job_message')
#     update_candidate_message = data.get('update_candidate_message')
#     delete_message = data.get("delete_message")

#     # data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     print("user :",user)
#     user_type = user.user_type
#     print("user_type :",user_type)
#     user_name = user.username

#     response_data = {}

#     if user_id and user_type:
#         if user_type == 'recruiter':
#             recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
#             if recruiter:
#                 candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None))).all()  # Filter candidates by recruiter's name
#                 candidates = sorted(candidates, key=lambda candidate: candidate.id)
#                 jobs = JobPost.query.filter_by(recruiter=user_name).all()  # Filter jobs by recruiter's name
#                 count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                                   Notification.recruiter_name == user_name).count()
#                 career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
#                                                                   Career_notification.recruiter_name == user_name).count()
#                 response_data = {
#                     'user': {
#                         'id': recruiter.id,
#                         'name': recruiter.name,
#                         'user_type': recruiter.user_type,
#                         'email': recruiter.email
#                         # Add more attributes as needed
#                     },
#                     'user_type': user_type,
#                     'user_name': user_name,
#                     'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created':candidate.date_created,
#                         'time_created':candidate.time_created

            
#                         # Add more attributes as needed
#                     } for candidate in candidates],
#                     'jobs': [{
#                         'id': job.id,
#                         'client': job.client,
#                         'experience_min': job.experience_min,
#                         'experience_max': job.experience_max,
#                         'budget_min': job.budget_min,
#                         'budget_max': job.budget_max,
#                         'location': job.location,
#                         'shift_timings': job.shift_timings,
#                         'notice_period': job.notice_period,
#                         'role': job.role,
#                         'detailed_jd': job.detailed_jd,
#                         'jd_pdf': job.jd_pdf,
#                         'mode': job.mode,
#                         'recruiter': job.recruiter,
#                         'management': job.management,
#                         'date_created': job.date_created,
#                         'time_created': job.time_created,
#                         'job_status': job.job_status,
#                         'job_type': job.job_type,
#                         'skills': job.skills,
#                         'notification': job.notification
#                         # Add more attributes as needed
#                     } for job in jobs],
#                     'edit_candidate_message': edit_candidate_message,
#                     'page_no': page_no,
#                     'count_notification_no': count_notification_no,
#                     'career_count_notification_no': career_count_notification_no
#                 }
#         elif user_type == 'management':
#             users = User.query.all()
#             candidates = Candidate.query.filter(Candidate.reference.is_(None)).all()
#             candidates = sorted(candidates, key=lambda candidate: candidate.id)
#             jobs = JobPost.query.all()
#             response_data = {
#                 'users': [{
#                     'id': user.id,
#                     'name': user.name,
#                     'user_type': user.user_type,
#                     'email': user.email
                     
#                     # Add more attributes as needed
#                 } for user in users],
#                 'user_type': user_type,
#                 'user_name': user_name,
#                 'candidates': [{
#                         'id': candidate.id,
#                         'job_id': candidate.job_id,
#                         'name': candidate.name,
#                         'mobile': candidate.mobile,
#                         'email': candidate.email,
#                         'client': candidate.client,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'profile': candidate.profile,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications': candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience': candidate.relevant_experience,
#                         'current_ctc': candidate.current_ctc,
#                         'expected_ctc': candidate.expected_ctc,
#                         'notice_period': candidate.notice_period,
#                         'linkedin_url': candidate.linkedin_url,
#                         'holding_offer': candidate.holding_offer,
#                         'recruiter': candidate.recruiter,
#                         'management': candidate.management,
#                         'status': candidate.status,
#                         'remarks': candidate.remarks,
#                         'skills': candidate.skills,
#                         'resume': candidate.resume,
#                         'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
#                         'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
#                         'buyout': candidate.buyout,
#                         'date_created':candidate.date_created,
#                         'time_created':candidate.time_created

#                     # Add more attributes as needed
#                 } for candidate in candidates],
#                 'jobs': [{
#                     'id': job.id,
#                     'client': job.client,
#                     'experience_min': job.experience_min,
#                     'experience_max': job.experience_max,
#                     'budget_min': job.budget_min,
#                     'budget_max': job.budget_max,
#                     'location': job.location,
#                     'shift_timings': job.shift_timings,
#                     'notice_period': job.notice_period,
#                     'role': job.role,
#                     'detailed_jd': job.detailed_jd,
#                     'jd_pdf': job.jd_pdf,
#                     'mode': job.mode,
#                     'recruiter': job.recruiter,
#                     'management': job.management,
#                     'date_created': job.date_created,
#                     'time_created': job.time_created,
#                     'job_status': job.job_status,
#                     'job_type': job.job_type,
#                     'skills': job.skills,
#                     'notification': job.notification
#                     # Add more attributes as needed
#                 } for job in jobs],
#                 'signup_message': signup_message,
#                 'job_message': job_message,
#                 'page_no': page_no,
#                 'edit_candidate_message': edit_candidate_message
#             }
#         else:
#             user = User.query.filter_by(id=user_id).first()
#             if user:
#                 candidates = Candidate.query.filter_by(recruiter=user.name).all()  # Filter candidates by user's name
#                 response_data = {
#                     'user': {
#                         'id': user.id,
#                         'name': user.name,
#                         'user_type': user.user_type,
#                         'email': user.email
#                         # Add more attributes as needed
#                     },
#                     'user_type': user_type,
#                     'user_name': user_name,
#                     'candidates': [{
#                         'id': candidate.id,
#                         'job_id':candidate.job_id,
#                         'name': candidate.name,
#                         'email': candidate.email,
#                         'mobile': candidate.mobile,
#                         'client':candidate.client,
#                         'skills':candidate.skills,
#                         "profile": candidate.profile, 
#                         'recruiter':candidate.recruiter,
#                         "management":candidate.management,
#                         'resume': candidate.resume,
#                         'current_company': candidate.current_company,
#                         'position': candidate.position,
#                         'current_job_location': candidate.current_job_location,
#                         'preferred_job_location': candidate.preferred_job_location,
#                         'qualifications':candidate.qualifications,
#                         'experience': candidate.experience,
#                         'relevant_experience':candidate.relevant_experience,
#                         'current_ctc':candidate.current_ctc,
#                         'experted_ctc': candidate.expected_ctc,
#                         "total":candidate.total,
#                         'package_in_lpa':candidate.package_in_lpa,
#                         'holding_offer':candidate.holding_offer,
#                         'status': candidate.status,
#                         'reason_for_job_change':candidate.reason_for_job_change,
#                         'remarks':candidate.remarks,
#                         'screening_done': candidate.screening_done,
#                         'rejected_at_screening': candidate.rejected_at_screening,
#                         'l1_cleared':candidate.l1_cleared,
#                         'rejected_at_l1':candidate.rejected_at_l1,
#                         "dropped_after_clearing_l1": candidate.dropped_after_clearing_l1,
#                         'l2_cleared':candidate.l1_cleared,
#                         'rejected_at_l2':candidate.rejected_at_l1,
#                         "dropped_after_clearing_l2": candidate.dropped_after_clearing_l1,
#                         'onboarded': candidate.onboarded,
#                         'dropped_after_onboarding': candidate.dropped_after_onboarding,
#                         'linkedin_url': candidate.linkedin_url,
#                         'period_of_notice': candidate.period_of_notice,
#                         'reference': candidate.reference,
#                         'reference_name': candidate.reference_name,
#                         'reference_position': candidate.reference_position,
#                         'reference_information': candidate.reference_information,
#                         'comments':candidate.comments,
#                         "time_created":str(candidate.time_created),
#                         "date_created": str(candidate.date_created)
#                         # Add more attributes as needed
#                     } for candidate in candidates],
#                 }
#     else:
#         response_data = {"message": "User ID or User Type missing"}

#     # # Convert date objects to string representations before returning the response
#     # for job in response_data.get('jobs', []):
#     #     job['date_created'] = job['date_created'].isoformat()

#     return Response(json.dumps(response_data, default=str), content_type='application/json')


# Mocked function for demonstration
# Mocked function for demonstration
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'doc', 'docx'}


import binascii  

import binascii  

@app.route('/add_candidate', methods=['POST'])
def add_candidate():
    try:
        
        # Retrieve request data from JSON
        data = request.json
        user_id = data['user_id']
        user = User.query.filter_by(id=user_id).first()
        user_type = user.user_type
        user_name = user.username

        job_id = data.get('job_id')
        client = data.get('client')
        name = data.get('name')
        mobile = data.get('mobile')
        email = data.get('email')
        profile = data.get('profile')
        skills = data.get('skills')
        current_company = data.get('current_company')
        position = data.get('position')
        current_job_location = data.get('current_job_location')
        preferred_job_location = data.get('preferred_job_location')
        
        # notice_period = data.get('notice_period')  # yes no completed
        # period_of_notice = data.get('period_of_notice')
        
        qualifications = data.get('qualifications')
        experience = data.get('experience')
        experience_months=data.get('experience')
        relevant_experience = data.get('relevant_experience')
        relevant_experience_months=data.get('relevant_experience_months')
        reason_for_job_change=data.get('reason_for_job_change')
        current_ctc = data.get('current_ctc')
        expected_ctc = data.get('expected_ctc')
        linkedin = data.get('linkedin')
        
        resume = data.get('resume')
        resume_binary = base64.b64decode(resume)
        print("Resume : ",type(resume_binary))

        # Set jd_pdf_present based on the presence of jd_pdf
        if resume_binary is not None:
            resume_present = True
        else:
            resume_present = False

        notice_period = data.get('serving_notice_period')
        last_working_date = None
        buyout = False
        period_of_notice = None
        if notice_period == 'yes':
            last_working_date=data.get('last_working_date')
            buyout=data.get('buyout')
        elif notice_period == 'no':
            period_of_notice = data.get('period_of_notice')
            buyout=data.get('buyout')
        # elif notice_period == 'completed':
        #     last_working_date=data.get('last_working_date')

        holding_offer = data.get('holding_offer')
        if holding_offer == 'yes':
            total_offers=data.get('total_offers')
            if total_offers == '':
                total_offers = 0
            else:
                total_offers=data.get('total_offers')
            highest_package_lpa=data.get('highest_package')
            if highest_package_lpa == '':
                highest_package_lpa = 0
            else:
                highest_package_lpa=data.get('highest_package')
        else:
            total_offers = None
            highest_package_lpa = None

        # # Check if the user is logged in
        if request.method == 'POST':
            # Retrieve the recruiter and management names based on user type
            if user_type == 'recruiter':
                # recruiter = User.query.get(user_id).name
                recruiter=user_name
                management = None
            elif user_type == 'management':
                recruiter = None
                # management = User.query.get(user_id).name
                management=user_name
            else:
                recruiter = None
                management = None

            # Check if the job_id is provided and job is active
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
            if not matching_job_post:
                return jsonify({'status': 'error',"message": "Job on hold"})

            # Create new candidate object
            new_candidate = Candidate(
                user_id=user_id,
                job_id=job_id,
                name=name,
                mobile=mobile,
                email=email,
                client=client,
                current_company=current_company,
                position=position,
                profile=profile,
                current_job_location=current_job_location,
                preferred_job_location=preferred_job_location,
                qualifications=qualifications,
                experience=experience,
                relevant_experience=relevant_experience,
                current_ctc=current_ctc,
                expected_ctc=expected_ctc,
                linkedin_url=linkedin,
                holding_offer=holding_offer,
                recruiter=recruiter,
                management=management,
                status='SCREENING',
                remarks=data.get('remarks'),
                skills=skills,
                resume=resume_binary,
                # serving_notice_period=serving_notice_period,
                notice_period=notice_period,
                period_of_notice=period_of_notice,
                # last_working_date=data.get('last_working_date') if notice_period in {'yes', 'completed'} else None,
                last_working_date=last_working_date,
                buyout=buyout,
                package_in_lpa=highest_package_lpa,
                total=total_offers,
                reason_for_job_change=reason_for_job_change,
                resume_present=resume_present
                # buyout='buyout' in data
            )

            print("Hello !!")
            
            # new_candidate.date_created = date.today()
            # new_candidate.time_created = datetime.now().time()
    
            # Created data and time
            current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
            new_candidate.date_created = current_datetime.date()
            new_candidate.time_created = current_datetime.time()


            db.session.add(new_candidate)
            db.session.commit()

            return jsonify({'status': 'success',"message": "Candidate Added Successfully", "candidate_id": new_candidate.id})

        return jsonify({"error_message": "Method not found"})

    except Exception as e:
        print(e)
        return jsonify({'status': 'error',"message": "Candidate unable to add"})
        

# @app.route('/add_candidate', methods=['POST'])
# def add_candidate():
#     try:
        
#         # Retrieve request data from JSON
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
#         user_type = user.user_type
#         user_name = user.username

#         job_id = data.get('job_id')
#         client = data.get('client')
#         name = data.get('name')
#         mobile = data.get('mobile')
#         email = data.get('email')
#         profile = data.get('profile')
#         skills = data.get('skills')
#         current_company = data.get('current_company')
#         position = data.get('position')
#         current_job_location = data.get('current_job_location')
#         preferred_job_location = data.get('preferred_job_location')
#         qualifications = data.get('qualifications')
#         experience = data.get('experience')
#         experience_months=data.get('experience')
#         relevant_experience = data.get('relevant_experience')
#         relevant_experience_months=data.get('relevant_experience_months')
#         reason_for_job_change=data.get('reason_for_job_change')
#         current_ctc = data.get('current_ctc')
#         expected_ctc = data.get('expected_ctc')
#         linkedin = data.get('linkedin')
#         serving_notice_period = data.get('serving_notice_period')
#         notice_period = data.get('notice_period')
#         holding_offer = data.get('holding_offer')
#         buyout=data.get('buyout')
#         last_working_date=data.get('last_working_date')
#         total_offers=data.get('total_offers')
#         highest_package_lpa=data.get('highest_package')
#         resume = data.get('resume')
#         resume_binary = base64.b64decode(resume)
#         print("Resume : ",type(resume_binary))

#         # Set jd_pdf_present based on the presence of jd_pdf
#         if resume_binary is not None:
#             resume_present = True
#         else:
#             resume_present = False

#         # # Check if the user is logged in
#         if request.method == 'POST':
#             # Retrieve the recruiter and management names based on user type
#             if user_type == 'recruiter':
#                 # recruiter = User.query.get(user_id).name
#                 recruiter=user_name
#                 management = None
#             elif user_type == 'management':
#                 recruiter = None
#                 # management = User.query.get(user_id).name
#                 management=user_name
#             else:
#                 recruiter = None
#                 management = None

#             # Check if the job_id is provided and job is active
#             matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
#             if not matching_job_post:
#                 return jsonify({'status': 'error',"message": "Job on hold"})

#             # Create new candidate object
#             new_candidate = Candidate(
#                 user_id=user_id,
#                 job_id=job_id,
#                 name=name,
#                 mobile=mobile,
#                 email=email,
#                 client=client,
#                 current_company=current_company,
#                 position=position,
#                 profile=profile,
#                 current_job_location=current_job_location,
#                 preferred_job_location=preferred_job_location,
#                 qualifications=qualifications,
#                 experience=experience,
#                 relevant_experience=relevant_experience,
#                 current_ctc=current_ctc,
#                 expected_ctc=expected_ctc,
#                 linkedin_url=linkedin,
#                 holding_offer=holding_offer,
#                 recruiter=recruiter,
#                 management=management,
#                 status='SCREENING',
#                 remarks=data.get('remarks'),
#                 skills=skills,
#                 resume=resume_binary,
#                 serving_notice_period=serving_notice_period,
#                 period_of_notice=notice_period,
#                 # last_working_date=data.get('last_working_date') if notice_period in {'yes', 'completed'} else None,
#                 last_working_date=last_working_date,
#                 buyout=buyout,
#                 package_in_lpa=highest_package_lpa,
#                 total=total_offers,
#                 resume_present=resume_present
#                 # buyout='buyout' in data
#             )

#             print("Hello !!")
            
#             # new_candidate.date_created = date.today()
#             # new_candidate.time_created = datetime.now().time()
    
#             # Created data and time
#             current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#             new_candidate.date_created = current_datetime.date()
#             new_candidate.time_created = current_datetime.time()


#             db.session.add(new_candidate)
#             db.session.commit()

#             return jsonify({'status': 'success',"message": "Candidate Added Successfully", "candidate_id": new_candidate.id})

#         return jsonify({"error_message": "Method not found"})

#     except Exception as e:
#         print(e)
#         return jsonify({'status': 'error',"message": "Candidate unable to add"})
        
        
        
# from flask import jsonify
# @app.route('/add_candidate', methods=['POST'])
# def add_candidate():
#     try:
        
#         # Retrieve request data from JSON
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
#         user_type = user.user_type
#         user_name = user.username

#         job_id = data.get('job_id')
#         client = data.get('client')
#         name = data.get('name')
#         mobile = data.get('mobile')
#         email = data.get('email')
#         profile = data.get('profile')
#         skills = data.get('skills')
#         current_company = data.get('current_company')
#         position = data.get('position')
#         current_job_location = data.get('current_job_location')
#         preferred_job_location = data.get('preferred_job_location')
#         qualifications = data.get('qualifications')
#         experience = data.get('experience')
#         experience_months=data.get('experience')
#         relevant_experience = data.get('relevant_experience')
#         relevant_experience_months=data.get('relevant_experience_months')
#         reason_for_job_change=data.get('reason_for_job_change')
#         current_ctc = data.get('current_ctc')
#         expected_ctc = data.get('expected_ctc')
#         linkedin = data.get('linkedin')
#         notice_period = data.get('notice_period')
#         holding_offer = data.get('holding_offer')
#         resume = data.get('resume')
#         print("Resume : ",type(resume))


        

#         # # Check if the user is logged in
#         if request.method == 'POST':
              
#             # Retrieve the recruiter and management names based on user type
#             if user_type == 'recruiter':
#                 recruiter = User.query.get(user_id).name
#                 management = None
#             elif user_type == 'management':
#                 recruiter = None
#                 management = User.query.get(user_id).name
#             else:
#                 recruiter = None
#                 management = None

#             # Check if the job_id is provided and job is active
#             matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
#             if not matching_job_post:
#                 return jsonify({"error_message": "Job on hold"})

#             # Create new candidate object
#             new_candidate = Candidate(
#                 user_id=user_id,
#                 job_id=job_id,
#                 name=name,
#                 mobile=mobile,
#                 email=email,
#                 client=client,
#                 current_company=current_company,
#                 position=position,
#                 profile=profile,
#                 current_job_location=current_job_location,
#                 preferred_job_location=preferred_job_location,
#                 qualifications=qualifications,
#                 experience=experience,
#                 relevant_experience=relevant_experience,
#                 current_ctc=current_ctc,
#                 expected_ctc=expected_ctc,
#                 notice_period=notice_period,
#                 linkedin_url=linkedin,
#                 holding_offer=holding_offer,
#                 recruiter=recruiter,
#                 management=management,
#                 status='None',
#                 remarks=data.get('remarks'),
#                 skills=skills,
#                 resume=resume,
#                 period_of_notice=data.get('months') if notice_period == 'no' else None,
#                 last_working_date=data.get('last_working_date') if notice_period in {'yes', 'completed'} else None,
#                 buyout='buyout' in data
#             )

#             new_candidate.date_created = date.today()
#             new_candidate.time_created = datetime.now().time()

#             db.session.add(new_candidate)
#             db.session.commit()

#             return jsonify({"message": "Candidate Added Successfully", "candidate_id": new_candidate.id})

#         return jsonify({"error_message": "Method not found"})

#     except Exception as e:
#         return jsonify({"error_message": str(e)}),500

@app.route('/get_job_role', methods=['GET'])
def get_job_role():
    job_id = request.args.get('job_id')

    job_post = JobPost.query.filter_by(id=job_id).first()
    if job_post:
        return jsonify({"role": job_post.role})
    else:
        return jsonify({"role": ""})

@app.route('/delete_candidate/<int:candidate_id>', methods=["POST"])
def delete_candidate(candidate_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()
    user_type = user.user_type
    username = user.username

    if user_type == 'management':
        candidate = Candidate.query.filter_by(id=candidate_id).first()

        if candidate:
            if request.method == "POST":
                # Check for null values
                if None in [candidate.name, candidate.email, candidate.client, candidate.profile, candidate.status]:
                    # If any value is null, delete the candidate
                    Candidate.query.filter_by(id=candidate_id).delete()
                    db.session.commit()
                    return jsonify({'status': 'success', "message": "Candidate details deleted successfully"})

                # Save deletion details before deleting the candidate
                deleted_candidate = Deletedcandidate(
                    username=username,
                    candidate_name=candidate.name,
                    candidate_email=candidate.email,
                    client=candidate.client,
                    profile=candidate.profile,
                    status=candidate.status
                )
                db.session.add(deleted_candidate)
                db.session.commit()

                # Delete the candidate
                Candidate.query.filter_by(id=candidate_id).delete()
                db.session.commit()

                return jsonify({'status': 'success', "message": "Candidate details deleted successfully"})

            return jsonify({
                "candidate": {
                    "id": candidate.id,
                    "name": candidate.name,
                    "email": candidate.email,
                    "client": candidate.client,
                    "profile": candidate.profile,
                    "status": candidate.status
                },
                "user_name": username
            })

        else:
            return jsonify({'status': 'error', "message": "Candidate not found"})

    return jsonify({'status': 'error', "message": "Unauthorized: Only management can delete candidates"})


# @app.route('/delete_candidate/<int:candidate_id>', methods=["POST"])
# def delete_candidate(candidate_id):
#     data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_type = user.user_type
#     username=user.username
    
#     if user_type == 'management':
#         candidate = Candidate.query.filter_by(id=candidate_id).first()

#         if candidate:
#             if request.method == "POST":
#                 # Save deletion details before deleting the candidate
#                 deleted_candidate = Deletedcandidate(
#                     username=username,
#                     candidate_name=candidate.name,
#                     candidate_email=candidate.email,
#                     client=candidate.client,
#                     profile=candidate.profile,
#                     status=candidate.status
#                 )
#                 db.session.add(deleted_candidate)
#                 db.session.commit()

#                 # Delete the candidate
#                 Candidate.query.filter_by(id=candidate_id).delete()
#                 db.session.commit()

#                 return jsonify({'status': 'success',"message": "Candidate details deleted successfully"})

#             return jsonify({
#                 "candidate": {
#                     "id": candidate.id,
#                     "name": candidate.name,
#                     "email": candidate.email,
#                     "client": candidate.client,
#                     "profile": candidate.profile,
#                     "status": candidate.status
#                 },
#                 "user_name": username
#             })

#         else:
#             return jsonify({'status': 'error',"message": "Candidate not found"}), 404

#     return jsonify({'status': 'error',"message": "Unauthorized: Only management can delete candidates"}), 401


@app.route('/delete_candidate_recruiter/<int:candidate_id>', methods=["GET", "POST"])
def delete_candidate_recruiter(candidate_id):
    if 'user_id' in session and 'user_type' in session:
        user_type = session['user_type']
        user_name = session['user_name']

        if user_type == 'management':
            candidate = Candidate.query.filter_by(id=candidate_id,recruiter=user_name)

            if request.method == "POST":
                # Save deletion details before deleting the candidate
                deleted_candidate = Deletedcandidate(
                    username=user_name,
                    candidate_name=candidate.name,
                    candidate_email=candidate.email,
                    client=candidate.client,
                    profile=candidate.profile,
                    status=candidate.status
                )
                db.session.add(deleted_candidate)
                db.session.commit()

                # Delete the candidate
                Candidate.query.filter_by(id=candidate_id).delete()
                db.session.commit()

                return redirect(url_for('dashboard', delete_message="Candidate details deleted successfully"))

            return render_template('delete_candidate.html', candidate=candidate, user_name=user_name)

        return "Unauthorized: Only management can delete candidates", 401

    return "Unauthorized: You must log in to access this page", 401


def verify_token(token):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        user_id = serializer.loads(token, max_age=86400)  
        return user_id
    except BadSignature:
        return None  
    except Exception as e:
        return None

# Search String Changed
# @app.route('/update_candidate/<int:candidate_id>/<page_no>/<search_string>', methods=['GET', 'POST'])
from flask import session, jsonify

@app.route('/update_candidate/<int:candidate_id>', methods=['POST'])
def update_candidate(candidate_id):
    data = request.json

    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()
    user_type = user.user_type
    user_name = user.username
    count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                      Notification.recruiter_name == user_name).count()
    career_count_notification_no = Career_notification.query.filter(
        Career_notification.notification_status == 'false',
        Career_notification.recruiter_name == user_name).count()
    
    if user_type == 'recruiter':
        recruiter = User.query.get(user_id).name
        management = None
    elif user_type == 'management':
        recruiter = None
        management = User.query.get(user_id).name
    else:
        recruiter = None
        management = None

    if user_type == 'recruiter':
        user_email = User.query.get(user_id).email
        management_email = None
    elif user_type == 'management':
        user_email = None
        management_email = User.query.get(user_id).email
    else:
        user_email = None
        management_email = None

    candidate = Candidate.query.filter_by(id=candidate_id).first()
    if not candidate:
        return jsonify({'status': 'error',"message": "Candidate not found"})
    
    previous_status = candidate.status

    candidate_status = request.json.get('candidate_status')
    candidate_comment = request.json.get('comments')

    candidate.status = candidate_status
    candidate.comments = candidate_comment

    # Update data_updated_date and data_updated_time
    current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
    candidate.data_updated_date = current_datetime.date()
    candidate.data_updated_time = current_datetime.time()

    db.session.commit()

    if candidate_status in [
            "SCREENING", "SCREEN REJECTED", "NO SHOW", "DROP", "CANDIDATE HOLD", "OFFERED - DECLINED", "DUPLICATE", "SCREENING SELECTED",
            "L1-SCHEDULE", "L1-FEEDBACK", "L1-SELECTED", "L1-REJECTED", "CANDIDATE RESCHEDULE", "PANEL RESCHEDULE", "L2-SCHEDULE", 
            "L2-FEEDBACK", "L2-SELECTED", "L2-REJECTED", "HR-ROUND", "MANAGERIAL ROUND", "NEGOTIATION", "SELECTED", "OFFER-REJECTED",
            "OFFER-DECLINED", "ON-BOARDED", "HOLD", "CANDIDATE NO-SHOW"
            ]:
        candidate_name = candidate.name
        candidate_position = candidate.position
        candidate_email = candidate.email

        if candidate_position:
            candidate_position = candidate_position.upper()
        else:
            candidate_position = ""

        if candidate.client:
            client = candidate.client.upper()
        else:
            client = ""

        if candidate_status in ["SCREENING", "SCREEN REJECTED"]:
            message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
        else:
            message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Previous Status : "{previous_status}"\n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
    else:
        message = ""
        candidate_name = ""
        candidate_position = ""
        candidate_email = ""

    return jsonify({
        'status': 'success',
        "message": "Candidate Status Updated Successfully",
        "user_id": user_id,
        "user_type": user_type,
        "user_name": user_name,
        "count_notification_no": count_notification_no,
        "career_count_notification_no": career_count_notification_no,
        "recruiter": recruiter,
        "management": management,
        "recruiter_email": user_email,
        "management_email": management_email,
        "candidate_name": candidate_name,
        "candidate_position": candidate_position,
        "candidate_email": candidate_email,
        "message_body": message 
    })

# @app.route('/update_candidate/<int:candidate_id>', methods=['POST'])
# def update_candidate(candidate_id):
#     data = request.json

#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_type = user.user_type
#     user_name = user.username
#     count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                       Notification.recruiter_name == user_name).count()
#     career_count_notification_no = Career_notification.query.filter(
#         Career_notification.notification_status == 'false',
#         Career_notification.recruiter_name == user_name).count()
    
#     if user_type == 'recruiter':
#         recruiter = User.query.get(user_id).name
#         management = None
#     elif user_type == 'management':
#         recruiter = None
#         management = User.query.get(user_id).name
#     else:
#         recruiter = None
#         management = None

#     if user_type == 'recruiter':
#         user_email = User.query.get(user_id).email
#         management_email = None
#     elif user_type == 'management':
#         user_email = None
#         management_email = User.query.get(user_id).email
#     else:
#         user_email = None
#         management_email = None

#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return jsonify({"error_message": "Candidate not found"}), 500
    
#     previous_status = candidate.status

#     candidate_status = request.json.get('candidate_status')
#     candidate_comment = request.json.get('comments')

#     candidate.status = candidate_status
#     candidate.comments = candidate_comment

#     # Update data_updated_date and data_updated_time
#     current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#     candidate.data_updated_date = current_datetime.date()
#     candidate.data_updated_time = current_datetime.time()

#     db.session.commit()

#     if candidate_status in [
#             "SCREENING", "SCREEN REJECTED", "NO SHOW", "DROP", "CANDIDATE HOLD", "OFFERED - DECLINED", "DUPLICATE", "SCREENING SELECTED",
#             "L1-SCHEDULE", "L1-FEEDBACK", "L1-SELECTED", "L1-REJECTED", "CANDIDATE RESCHEDULE", "PANEL RESCHEDULE", "L2-SCHEDULE", 
#             "L2-FEEDBACK", "L2-SELECTED", "L2-REJECTED", "HR-ROUND", "MANAGERIAL ROUND", "NEGOTIATION", "SELECTED", "OFFER-REJECTED",
#             "OFFER-DECLINED", "ON-BOARDED", "HOLD", "CANDIDATE NO-SHOW"
#             ]:
#         candidate_name = candidate.name
#         candidate_position = candidate.position
#         candidate_email = candidate.email

#         if candidate_position:
#             candidate_position = candidate_position.upper()
#         else:
#             candidate_position = ""

#         if candidate.client:
#             client = candidate.client.upper()
#         else:
#             client = ""

#         if candidate_status in ["SCREENING", "SCREEN REJECTED"]:
#             message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#         else:
#             message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Previous Status : "{previous_status}"\n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#     else:
#         message = ""
#         candidate_name = ""
#         candidate_position = ""
#         candidate_email = ""

#     return jsonify({
#         "message": "Candidate Status Updated Successfully",
#         "user_id": user_id,
#         "user_type": user_type,
#         "user_name": user_name,
#         "count_notification_no": count_notification_no,
#         "career_count_notification_no": career_count_notification_no,
#         "recruiter": recruiter,
#         "management": management,
#         "recruiter_email": user_email,
#         "management_email": management_email,
#         "candidate_name": candidate_name,
#         "candidate_position": candidate_position,
#         "candidate_email": candidate_email,
#         "message_body": message 
#     })

# @app.route('/update_candidate/<int:candidate_id>', methods=['POST'])
# def update_candidate(candidate_id):
#     data = request.json

#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_type = user.user_type
#     user_name = user.username
#     count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                       Notification.recruiter_name == user_name).count()
#     career_count_notification_no = Career_notification.query.filter(
#         Career_notification.notification_status == 'false',
#         Career_notification.recruiter_name == user_name).count()
#     if request.method == 'POST':
#         if user_type == 'recruiter':
#             recruiter = User.query.get(user_id).name
#             management = None
#         elif user_type == 'management':
#             recruiter = None
#             management = User.query.get(user_id).name
#         else:
#             recruiter = None
#             management = None

#         if user_type == 'recruiter':
#             user_email = User.query.get(user_id).email
#             management_email = None
#         elif user_type == 'management':
#             user_email = None
#             management_email = User.query.get(user_id).email
#         else:
#             user_email = None
#             management_email = None

#         candidate = Candidate.query.filter_by(id=candidate_id).first()
#         print(candidate)
        
#         previous_status = candidate.status

#         candidate_status = request.json.get('candidate_status')
#         candidate_comment = request.json.get('comments')

#         candidate.status = candidate_status
#         candidate.comments = candidate_comment

#         db.session.commit()

#         if candidate_status in [
#                 "SCREENING", "SCREEN REJECTED", "NO SHOW", "DROP", "CANDIDATE HOLD", "OFFERED - DECLINED", "DUPLICATE", "SCREENING SELECTED",
#                 "L1-SCHEDULE", "L1-FEEDBACK", "L1-SELECTED", "L1-REJECTED", "CANDIDATE RESCHEDULE", "PANEL RESCHEDULE", "L2-SCHEDULE", 
#                 "L2-FEEDBACK", "L2-SELECTED", "L2-REJECTED", "HR-ROUND", "MANAGERIAL ROUND", "NEGOTIATION", "SELECTED", "OFFER-REJECTED",
#                 "OFFER-DECLINED", "ON-BOARDED", "HOLD", "CANDIDATE NO-SHOW"
#                 ]:
#             candidate_name = candidate.name
#             candidate_position = candidate.position
#             candidate_email = candidate.email

#             if candidate_position:
#                 candidate_position = candidate_position.upper()
#             else:
#                 candidate_position = ""

#             if candidate.client:
#                 client = candidate.client.upper()
#             else:
#                 client = ""

#             if candidate_status in ["SCREENING", "SCREEN REJECTED"]:
#                 message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#             else:
#                 message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Previous Status : "{previous_status}"\n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#         else:
#             message = ""
#             candidate_name = ""
#             candidate_position = ""
#             candidate_email = ""

#         return jsonify({
#         "message": "Candidate Status Updated Successfully",
#         "user_id": user_id,
#         "user_type": user_type,
#         "user_name": user_name,
#         "count_notification_no": count_notification_no,
#         "career_count_notification_no": career_count_notification_no,
#         "recruiter": recruiter,
#         "management": management,
#         "recruiter_email": user_email,
#         "management_email": management_email,
#         "candidate_name": candidate_name,
#         "candidate_position": candidate_position,
#         "candidate_email": candidate_email,
#         # "message": message
#         "message_body": message 
#     })

@app.route('/update_candidate_careers/<int:candidate_id>/<page_no>/<search_string>', methods=['GET', 'POST'])
@app.route('/update_candidate_careers/<int:candidate_id>/<page_no>', methods=['GET', 'POST'])
def update_candidate_careers(candidate_id, page_no):
    if 'user_id' in session and 'user_type' in session:
        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']
        count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                          Notification.recruiter_name == user_name).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name).count()
        if request.method == 'POST':
            # Retrieve the logged-in user's ID and user type from the session
            user_id = session['user_id']
            user_type = session['user_type']

            # Retrieve the recruiter and management names based on user type
            if user_type == 'recruiter':
                recruiter = User.query.get(user_id).name
                management = None
            elif user_type == 'management':
                recruiter = None
                management = User.query.get(user_id).name
            else:
                recruiter = None
                management = None

            if user_type == 'recruiter':
                recruiter_email = User.query.get(user_id).email
                management = None
            elif user_type == 'management':
                recruiter = None
                recruiter_email = User.query.get(user_id).email
            else:
                recruiter_email = None
                management_email = None

            # Retrieve the form data for the candidate
            candidate = Candidate.query.get(candidate_id)
            previous_status = candidate.status
            # candidate.recruiter = recruiter
            # candidate.management = management

            # Get the selected candidate status from the form
            candidate_status = request.form.get('candidate_status')
            candidate_comment = request.form.get('comments')

            # Update the candidate status field
            candidate.status = candidate_status

            candidate.comments = candidate_comment

            db.session.commit()

            if candidate_status == "SCREENING" or candidate_status == "SCREEN REJECTED":
                candidate_name = candidate.name
                candidate_position = candidate.position

                # Retrieve the candidate's email
                candidate_email = candidate.email

                # Determine if the logged-in user is a recruiter or management
                user_type = session.get('user_type')

                if user_type == 'recruiter' or user_type == 'management':
                    # Retrieve the corresponding user's email
                    user_email = User.query.get(session.get('user_id')).email

                    message = Message(f'Job Application Status - {candidate_position}',
                                      sender='kanuparthisaiganesh582@gmail.com', recipients=[candidate_email])

                    if user_type == 'management':
                        management_email = user_email
                        message.cc = [management_email]
                    elif user_type == 'recruiter':
                        recruiter_email = user_email
                        message.cc = [recruiter_email]
                    message.body = f'''Dear {candidate.name}, 

Greetings! 

We hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate.position.upper()} position and participating in the recruitment process. 

We are writing to inform you about the latest update we received from our client {candidate.client.upper()} regarding your interview. 

        Current Status :  "{candidate.status}"

Thank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. 

If you have any questions or need further information, please feel free to reach out to us. 

Thanks, 
                            '''
                #mail.send(message)
                pass
            elif candidate_status == "NO SHOW" or candidate_status == "DROP" or candidate_status == "CANDIDATE HOLD" or candidate_status == "OFFERED - DECLINED" or candidate_status == "DUPLICATE":
                pass
            else:
                candidate_name = candidate.name
                candidate_position = candidate.position
                candidate_email = candidate.email

                user_type = session.get('user_type')

                if user_type == 'recruiter' or user_type == 'management':
                    user_email = User.query.get(session.get('user_id')).email

                    message = Message(f'Job Application Status - {candidate_position}',
                                      sender='kanuparthisaiganesh582@gmail.com', recipients=[candidate_email])

                    if user_type == 'management':
                        management_email = user_email
                        message.cc = [management_email]
                    elif user_type == 'recruiter':
                        recruiter_email = user_email
                        message.cc = [recruiter_email]
                    message.body = f'''Dear {candidate.name}, 

Greetings! 

We hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate.position.upper()} position and participating in the recruitment process. 

We are writing to inform you about the latest update we received from our client {candidate.client.upper()} regarding your interview. 

        Previous Status : "{previous_status}"

        Current Status :  "{candidate.status}"

Thank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. 

If you have any questions or need further information, please feel free to reach out to us. 

Thanks, 
                            '''
                #mail.send(message)
                pass

            return redirect(
                url_for('career_dashboard', update_candidate_message='Candidate Status Updated Sucessfully', page_no=page_no))

        candidate = Candidate.query.get(candidate_id)
        candidate_data = {
            'id': candidate.id,
            'name': candidate.name,
            'mobile': candidate.mobile,
            'email': candidate.email,
            'client': candidate.client,
            'current_company': candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location': candidate.current_job_location,
            'preferred_job_location': candidate.preferred_job_location,
            'resume': candidate.resume,
            'qualifications': candidate.qualifications,
            'experience': candidate.experience,
            'relevant_experience': candidate.relevant_experience,
            'current_ctc': candidate.current_ctc,
            'expected_ctc': candidate.expected_ctc,
            'notice_period': candidate.notice_period,
            'last_working_date': candidate.last_working_date,
            'buyout': candidate.buyout,
            'holding_offer': candidate.holding_offer,
            'total': candidate.total,
            'package_in_lpa': candidate.package_in_lpa,
            'reason_for_job_change': candidate.reason_for_job_change,
            'remarks': candidate.remarks,
            'candidate_status': candidate.status,
        }

        return render_template('update_candidate.html', candidate_data=candidate_data, user_id=user_id,
                               user_type=user_type, user_name=user_name, candidate=candidate,
                               count_notification_no=count_notification_no,
                               career_count_notification_no=career_count_notification_no)

    return redirect(url_for('career_dashboard'))


from flask import jsonify

@app.route('/logout', methods=['POST'])
def logout():
    data = request.json
    
    if data:
        user_id = data.get('user_id')
        
        if user_id:
            user = User.query.filter_by(id=user_id).first()
            
            if user:
                user_type = user.user_type
                user_name = user.username
                
                return jsonify({"message": "Logged out successfully"}), 200
            
            return jsonify({"message": "User not found"}), 404
        else:
            return jsonify({"message": "'user_id' not provided in JSON data"}), 400
    
    return jsonify({"message": "No JSON data provided"}), 400



from datetime import datetime

@app.route('/edit_candidate/<int:candidate_id>', methods=['POST'])
def edit_candidate(candidate_id):
    try:
        data = request.json

        user_id = data.get('user_id')
        if not user_id:
            return jsonify({"error_message": "User ID is required"}), 400

        user = User.query.filter_by(id=user_id).first()
        if not user:
            return jsonify({"error_message": "User not found"}), 404

        user_name = user.username
        count_notification_no = Notification.query.filter(
            Notification.notification_status == 'false',
            Notification.recruiter_name == user_name
        ).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name
        ).count()

        # Retrieve the candidate object
        candidate = Candidate.query.get(candidate_id)
        if not candidate:
            return jsonify({"error_message": "Candidate not found"}), 404

        # Update the candidate fields with the new data
        candidate.name = data.get('name')
        candidate.mobile = data.get('mobile')
        candidate.email = data.get('email')
        candidate.client = data.get('client')
        candidate.current_company = data.get('current_company')
        candidate.position = data.get('position')
        candidate.profile = data.get('profile')
        candidate.current_job_location = data.get('current_job_location')
        candidate.preferred_job_location = data.get('preferred_job_location')
        candidate.qualifications = data.get('qualifications')
        candidate.experience = data.get('experience')
        candidate.relevant_experience = data.get('relevant_experience')
        candidate.current_ctc = data.get('current_ctc')
        candidate.expected_ctc = data.get('expected_ctc')    
        candidate.reason_for_job_change = data.get('reason_for_job_change')
        candidate.linkedin_url = data.get('linkedin')
        candidate.remarks = data.get('remarks')
        candidate.skills = data.get('skills')
        # candidate.holding_offer = data.get('holding_offer')
        candidate.total = data.get('total_offers')
        candidate.package_in_lpa = data.get('highest_package')
        
        
        # Handle resume decoding
        resume_data = data.get('resume')
        if resume_data is not None:
            try:
                resume_binary = base64.b64decode(resume_data)
                candidate.resume = resume_binary
                candidate.resume_present = True  # Update resume_present to True if resume is provided and valid
            except (binascii.Error, TypeError) as e:
                return jsonify({"error_message": "Invalid resume format"}), 500

        # Serving notice period logic
        notice_period = data.get('serving_notice_period')
        candidate.notice_period = notice_period
        if notice_period == 'yes':
            candidate.last_working_date = data.get('last_working_date')
            candidate.buyout = data.get('buyout')
        elif notice_period == 'no':
            candidate.period_of_notice = data.get('period_of_notice')
            candidate.buyout = data.get('buyout')
        # elif notice_period == 'completed':
        #     candidate.last_working_date = data.get('last_working_date')

        # Holding offer logic
        holding_offer = data.get('holding_offer')
        candidate.holding_offer = holding_offer
        if holding_offer == 'yes':
            total_offers = data.get('total_offers')
            candidate.total = 0 if total_offers == '' else total_offers
            highest_package = data.get('highest_package')
            candidate.package_in_lpa = 0 if highest_package == '' else highest_package
        else:
            candidate.total = None
            candidate.package_in_lpa = None

        # Update data_updated_date and data_updated_time
        # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        # candidate.data_updated_date = current_datetime.date()
        # candidate.data_updated_time = current_datetime.time()

        db.session.commit()
        return jsonify({'status': 'success', "message": "Candidate Details Edited Successfully"})

    except Exception as e:
        print(e)
        return jsonify({'status': 'error', "message": "Candidate Details not Edited "})


# @app.route('/edit_candidate/<int:candidate_id>', methods=['POST'])
# def edit_candidate(candidate_id):
#     try:
#         data = request.json

#         user_id = data.get('user_id')
#         if not user_id:
#             return jsonify({"error_message": "User ID is required"}), 400

#         user = User.query.filter_by(id=user_id).first()
#         if not user:
#             return jsonify({"error_message": "User not found"}), 404

#         user_name = user.username
#         count_notification_no = Notification.query.filter(
#             Notification.notification_status == 'false',
#             Notification.recruiter_name == user_name
#         ).count()
#         career_count_notification_no = Career_notification.query.filter(
#             Career_notification.notification_status == 'false',
#             Career_notification.recruiter_name == user_name
#         ).count()

#         # Retrieve the candidate object
#         candidate = Candidate.query.get(candidate_id)
#         if not candidate:
#             return jsonify({"error_message": "Candidate not found"}), 404

#         # Update the candidate fields with the new data
#         candidate.name = data.get('name')
#         candidate.mobile = data.get('mobile')
#         candidate.email = data.get('email')
#         candidate.client = data.get('client')
#         candidate.current_company = data.get('current_company')
#         candidate.position = data.get('position')
#         candidate.profile = data.get('profile')
#         candidate.current_job_location = data.get('current_job_location')
#         candidate.preferred_job_location = data.get('preferred_job_location')
#         candidate.qualifications = data.get('qualifications')
#         candidate.experience = data.get('experience')
#         candidate.relevant_experience = data.get('relevant_experience')
#         candidate.current_ctc = data.get('current_ctc')
#         candidate.expected_ctc = data.get('expected_ctc')
#         # candidate.serving_notice_period = data.get('serving_notice_period')
#         candidate.notice_period = data.get('notice_period')
#         candidate.period_of_notice = data.get('period_of_notice')
#         candidate.reason_for_job_change = data.get('reason_for_job_change')
#         candidate.linkedin_url = data.get('linkedin')
#         candidate.remarks = data.get('remarks')
#         candidate.skills = data.get('skills')
#         candidate.holding_offer = data.get('holding_offer')
#         candidate.total = data.get('total')
#         candidate.package_in_lpa = data.get('package_in_lpa')
#         candidate.buyout = data.get('buyout')
#         candidate.total = data.get('total_offers')
#         candidate.package_in_lpa =data.get('highest_package')
#         candidate.last_working_date = data.get('last_working_date')

#         # Handle resume decoding
#         resume_data = data.get('resume')
#         if resume_data is not None:
#             try:
#                 resume_binary = base64.b64decode(resume_data)
#                 candidate.resume = resume_binary
#                 candidate.resume_present = True  # Update resume_present to True if resume is provided and valid
#             except (binascii.Error, TypeError) as e:
#                 return jsonify({"error_message": "Invalid resume format"}), 500

#         # Update data_updated_date and data_updated_time
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         candidate.data_updated_date = current_datetime.date()
#         candidate.data_updated_time = current_datetime.time()

#         db.session.commit()
#         return jsonify({'status': 'success',"message": "Candidate Details Edited Successfully"})

#     except Exception as e:
#         return jsonify({'status': 'error',"message": "Candidate Details not Edited Successfully"})


# @app.route('/edit_candidate/<int:candidate_id>', methods=['POST'])
# def edit_candidate(candidate_id):
#     try:
#         data = request.json

#         user_id = data.get('user_id')
#         if not user_id:
#             return jsonify({"error_message": "User ID is required"}), 400

#         user = User.query.filter_by(id=user_id).first()
#         if not user:
#             return jsonify({"error_message": "User not found"}), 404

#         user_name = user.username
#         count_notification_no = Notification.query.filter(
#             Notification.notification_status == 'false',
#             Notification.recruiter_name == user_name
#         ).count()
#         career_count_notification_no = Career_notification.query.filter(
#             Career_notification.notification_status == 'false',
#             Career_notification.recruiter_name == user_name
#         ).count()

#         # Retrieve the candidate object
#         candidate = Candidate.query.get(candidate_id)
#         if not candidate:
#             return jsonify({"error_message": "Candidate not found"}), 404

#         # Update the candidate fields with the new data
#         candidate.name = data.get('name')
#         candidate.mobile = data.get('mobile')
#         candidate.email = data.get('email')
#         candidate.client = data.get('client')
#         candidate.current_company = data.get('current_company')
#         candidate.position = data.get('position')
#         candidate.profile = data.get('profile')
#         candidate.current_job_location = data.get('current_job_location')
#         candidate.preferred_job_location = data.get('preferred_job_location')
#         candidate.qualifications = data.get('qualifications')
#         candidate.experience = data.get('experience')
#         candidate.relevant_experience = data.get('relevant_experience')
#         candidate.current_ctc = data.get('current_ctc')
#         candidate.expected_ctc = data.get('expected_ctc')
#         candidate.notice_period = data.get('notice_period')
#         candidate.reason_for_job_change = data.get('reason_for_job_change')
#         candidate.linkedin_url = data.get('linkedin')
#         candidate.remarks = data.get('remarks')
#         candidate.skills = data.get('skills')
#         candidate.holding_offer = data.get('holding_offer')
#         candidate.total = data.get('total')
#         candidate.package_in_lpa = data.get('package_in_lpa')
#         candidate.buyout = data.get('buyout')
#         candidate.last_working_date=data.get('last_working_date')
        
#         # # Handle resume decoding
#         # resume_data = data.get('resume')
#         # if resume_data:
#         #     try:
#         #         resume_binary = base64.b64decode(resume_data)
#         #         candidate.resume = resume_binary
#         #     except (base64.binascii.Error, TypeError) as e:
#         #         return jsonify({"error_message": "Invalid resume format"}), 400
                
#         # if resume_binary is not None:
#         #     resume_present = True
#         # else:
#         #     resume_present = False

#         # resume_data = data.get('resume')
#         # resume_binary = None  # Initialize resume_binary
#         # resume_present = False  # Initialize resume_present
#         # if resume_data:
#         #     try:
#         #         resume_binary = base64.b64decode(resume_data)
#         #         resume_present = True if resume_binary is not None
#         #         candidate.resume = resume_binary
#         #     except (binascii.Error, TypeError) as e:
#         #         return jsonify({"error_message": "Invalid resume format"}), 400

                
                
#         # Update resume_present status
#         candidate.resume_present = resume_present
            
#         # Update data_updated_date and data_updated_time
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         candidate.data_updated_date = current_datetime.date()
#         candidate.data_updated_time = current_datetime.time()

#         db.session.commit()
#         return jsonify({"message": "Candidate Details Edited Successfully"})

#     except Exception as e:
#         return jsonify({"error_message": str(e)}), 500

# @app.route('/edit_candidate/<int:candidate_id>', methods=['POST'])
# def edit_candidate(candidate_id):
#     data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_name = user.username
#     count_notification_no = Notification.query.filter(
#         Notification.notification_status == 'false',
#         Notification.recruiter_name == user_name
#     ).count()
#     career_count_notification_no = Career_notification.query.filter(
#         Career_notification.notification_status == 'false',
#         Career_notification.recruiter_name == user_name
#     ).count()

#     if request.method == 'POST':
#         # Retrieve the form data for the candidate from JSON payload
#         data = request.json

#         # Retrieve the candidate object
#         candidate = Candidate.query.get(candidate_id)
#         if candidate:
#             # Update the candidate fields with the new data
#             candidate.name = data.get('name')
#             candidate.mobile = data.get('mobile')
#             candidate.email = data.get('email')
#             candidate.client = data.get('client')
#             candidate.current_company = data.get('current_company')
#             candidate.position = data.get('position')
#             candidate.profile = data.get('profile')
#             candidate.current_job_location = data.get('current_job_location')
#             candidate.preferred_job_location = data.get('preferred_job_location')
#             candidate.qualifications = data.get('qualifications')
#             candidate.experience = data.get('experience')
#             candidate.relevant_experience=data.get('relevant_experience')
#             candidate.current_ctc=data.get('current_ctc')
#             candidate.expected_ctc=data.get('expected_ctc')
#             candidate.notice_period = data.get('notice_period')
#             candidate.reason_for_job_change = data.get('reason_for_job_change')
#             candidate.linkedin_url = data.get('linkedin')
#             candidate.remarks = data.get('remarks')
#             candidate.skills = data.get('skills')
#             candidate.holding_offer = data.get('holding_offer')
#             candidate.total = data.get('total')
#             candidate.package_in_lpa = data.get('package_in_lpa')
#             candidate.buyout=data.get('buyout')
#             # candidate.resume=data.get('resume')
            
#             # Handle resume decoding
#             resume_data =data.get('resume')
#             if resume_data:
#                 try:
#                     resume_binary = base64.b64decode(resume_data)
#                     candidate.resume = resume_binary
#                 except base64.binascii.Error as e:
#                     return jsonify({"error_message": "Invalid resume format"}), 400
            

#             # Update data_updated_date and data_updated_time
#             current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#             candidate.data_updated_date = current_datetime.date()
#             candidate.data_updated_time = current_datetime.time()

#             db.session.commit()
#             return jsonify({"message": "Candidate Details Edited Successfully"})
#         else:
#             return jsonify({"error_message": "Candidate not found"}), 500

# # Search String Changed
# # @app.route('/edit_candidate/<int:candidate_id>/<int:page_no>/<search_string>', methods=['GET', 'POST'])
# @app.route('/edit_candidate/<int:candidate_id>', methods=['POST'])
# def edit_candidate(candidate_id):
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
#         user_name = user.username
#         count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                           Notification.recruiter_name == user_name).count()
#         career_count_notification_no = Career_notification.query.filter(
#             Career_notification.notification_status == 'false',
#             Career_notification.recruiter_name == user_name).count()

#         if request.method == 'POST':
#             # Retrieve the form data for the candidate from JSON payload
#             data = request.json

#             # Retrieve the candidate object
#             candidate = Candidate.query.get(candidate_id)
#             if candidate:
#                 # Update the candidate fields with the new data
#                 candidate.name = data.get('name')
#                 candidate.mobile = data.get('mobile')
#                 candidate.email = data.get('email')
#                 candidate.client = data.get('client')
#                 candidate.current_company = data.get('current_company')
#                 candidate.position = data.get('position')
#                 candidate.profile = data.get('profile')
#                 candidate.current_job_location = data.get('current_job_location')
#                 candidate.preferred_job_location = data.get('preferred_job_location')
#                 candidate.qualifications = data.get('qualifications')
#                 candidate.experience = data.get('experience')
#                 candidate.notice_period = data.get('notice_period')
#                 candidate.reason_for_job_change = data.get('reason_for_job_change')
#                 candidate.linkedin_url = data.get('linkedin')
#                 candidate.remarks = data.get('remarks')
#                 candidate.skills = data.get('skills')
#                 candidate.holding_offer = data.get('holding_offer')
#                 candidate.total = data.get('total')
#                 candidate.package_in_lpa = data.get('package_in_lpa')
#                 candidate.period_of_notice = data.get('period_of_notice')

#                 db.session.commit()
#                 return jsonify({"message": "Candidate Details Edited Successfully"})
#             else:
#                 return jsonify({"error_message": "Candidate not found"}), 500


@app.route('/edit_candidate_careers/<int:candidate_id>/<int:page_no>/<search_string>', methods=['GET', 'POST'])
@app.route('/edit_candidate_careers/<int:candidate_id>/<int:page_no>', methods=['GET', 'POST'])
def edit_candidate_careers(candidate_id, page_no):
    if 'user_id' in session and 'user_type' in session:
        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']
        count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                          Notification.recruiter_name == user_name).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name).count()

        if request.method == 'POST':
            # Retrieve the logged-in user's ID and user type from the session
            user_id = session['user_id']
            user_type = session['user_type']

            # Retrieve the form data for the candidate
            candidate = Candidate.query.get(candidate_id)

            # Update the candidate information based on user type
            if user_type == 'recruiter':
                candidate.recruiter = User.query.get(user_id).name
            elif user_type == 'management':
                candidate.management = User.query.get(user_id).name

            # Update the candidate fields with the new form data
            candidate.name = request.form.get('name')
            candidate.mobile = request.form.get('mobile')
            candidate.email = request.form.get('email')
            candidate.client = request.form.get('client')
            candidate.current_company = request.form.get('current_company')
            candidate.position = request.form.get('position')
            candidate.profile = request.form.get('profile')
            candidate.current_job_location = request.form.get('current_job_location')
            candidate.preferred_job_location = request.form.get('preferred_job_location')
            candidate.qualifications = request.form.get('qualifications')
            experience = request.form.get('experience')
            exp_months = request.form.get('exp_months')
            candidate.experience = experience +'.'+exp_months
            relevant_experience = request.form.get('relevant_experience')
            relevant_exp_months = request.form.get('relevant_exp_months')
            candidate.relevant_experience = relevant_experience + '.' + relevant_exp_months
            candidate.current_ctc = request.form.get('current_ctc')
            candidate.expected_ctc = request.form.get('expected_ctc')
            currency_type_current = request.form['currency_type_current']
            currency_type_except = request.form['currency_type_except']
            candidate.current_ctc = currency_type_current + " " + request.form['current_ctc']
            candidate.expected_ctc = currency_type_except + " " + request.form['expected_ctc']
            candidate.notice_period = request.form.get('notice_period')
            candidate.reason_for_job_change = request.form.get('reason_for_job_change')
            candidate.linkedin_url = request.form.get('linkedin')
            candidate.remarks = request.form.get('remarks')
            candidate.skills = request.form.get('skills')
            candidate.holding_offer = request.form.get('holding_offer')
            candidate.total = request.form.get('total')
            candidate.package_in_lpa = request.form.get('package_in_lpa')
            candidate.period_of_notice = request.form.get('period_of_notice')

            # Handle the resume file upload
            resume_file = request.files['resume']
            if resume_file.filename != '':
                # Save the new resume to the candidate's resume field as bytes
                candidate.resume = resume_file.read()

            holding_offer = request.form.get('holding_offer')
            if holding_offer == 'yes':
                total = request.form.get('total')
                package_in_lpa = request.form.get('package_in_lpa')

                candidate.total = total
                candidate.package_in_lpa = package_in_lpa
            elif holding_offer in ['no', 'pipeline']:
                candidate.total = None
                candidate.package_in_lpa = None

            notice_period = request.form.get('notice_period')
            if notice_period == 'yes':
                last_working_date = request.form['last_working_date']
                buyout = 'buyout' in request.form
                candidate.last_working_date = last_working_date
                candidate.buyout = buyout
            elif notice_period == 'no':
                period_of_notice = request.form['months']
                buyout = 'buyout' in request.form
                candidate.period_of_notice = period_of_notice
                candidate.buyout = buyout
            elif notice_period == 'completed':
                last_working_date = request.form['last_working_date']
                candidate.last_working_date = last_working_date

            db.session.commit()

            return redirect(
                url_for('career_dashboard', page_no=page_no, edit_candidate_message='Candidate Details Edited Successfully'))

        candidate = Candidate.query.get(candidate_id)
        candidate_data = {
            'id': candidate.id,
            'name': candidate.name,
            'mobile': candidate.mobile,
            'email': candidate.email,
            'client': candidate.client,
            'current_company': candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location': candidate.current_job_location,
            'preferred_job_location': candidate.preferred_job_location,
            'qualifications': candidate.qualifications,
            'experience': candidate.experience,
            'relevant_experience': candidate.relevant_experience,
            'current_ctc': candidate.current_ctc,
            'expected_ctc': candidate.expected_ctc,
            'notice_period': candidate.notice_period,
            'reason_for_job_change': candidate.reason_for_job_change,
            'remarks': candidate.remarks,
            'candidate_status': candidate.status,
            'linkedin_url': candidate.linkedin_url,
            'skills': candidate.skills,
            'resume': candidate.resume,
            'holding_offer': candidate.holding_offer,
            'total': candidate.total,
            'package_in_lpa': candidate.package_in_lpa,
            'last_working_date': candidate.last_working_date,
            'buyout': candidate.buyout,
            'period_of_notice': candidate.period_of_notice,
        }

        return render_template('edit_candidate_careers.html', candidate_data=candidate_data, user_id=user_id,
                               user_type=user_type, user_name=user_name, count_notification_no=count_notification_no,
                               page_no=page_no,career_count_notification_no=career_count_notification_no)

    return redirect(url_for('career_dashboard'))


# Function to check if a filename has an allowed extension
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


from flask import send_file

import io
import os
import base64
from flask import redirect, url_for, send_file

@app.route('/download_resume/<int:candidate_id>', methods=['GET','POST'])
def download_resume(candidate_id):
    candidate = Candidate.query.get(candidate_id)

    if candidate is None or candidate.resume is None:
        return redirect(url_for('dashboard'))

    # Decode the base64 encoded resume
    resume_data = candidate.resume.split(',')[1]  # Get the data part after the comma
    resume_bytes = base64.b64decode(resume_data)

    # Determine the file extension
    if candidate.resume.startswith("data:application/pdf"):
        resume_filename = f"{candidate.name}_resume.pdf"
    else:
        resume_filename = f"{candidate.name}_resume.docx"

    # Send the resume data for download
    return send_file(io.BytesIO(resume_bytes),
                     attachment_filename=resume_filename,
                     as_attachment=True)



# def send_notification(recruiter_email):
#     msg = Message('New Job Posted', sender='kanuparthisaiganesh582@gmail.com', recipients=[recruiter_email])
#     msg.body = 'A new job has been posted. Check your dashboard for more details.'
#     mail.send(msg)

def post_job_send_notification(recruiter_email, new_recruiter_name, job_data):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                New Job Posted
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>A new requirement has been assigned to you.</p>
            <p> Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 30%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check in Job Listing page for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        'New Requirement Assigned',
        sender='kanuparthisaiganesh582@gmail.com',
        recipients=[recruiter_email]
    )
    msg.html = html_body
    try:
        mail.send(msg)
    except Exception as e:
        print("mail error", str(e))
        return f'Failed to send mail: {str(e)}'
    return None

# def post_job_send_notification(recruiter_email, new_recruiter_name, job_data):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 color: #333;
#                 line-height: 1.6;
#                 background-color: #f4f4f4;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #ffffff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #4CAF50;
#                 color: white;
#                 padding: 10px;
#                 text-align: center;
#                 font-size: 20px;
#                 border-radius: 8px 8px 0 0;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 10px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 8px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #4CAF50;
#                 color: white;
#             }}
#             tr:nth-child(even) {{
#                 background-color: #f9f9f9;
#             }}
#             p {{
#                 margin: 10px 0;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
#     <body>
#         <div class="container">
#             <div class="header">
#                 New Job Posted
#             </div>
#             <p>Dear {new_recruiter_name},</p>
#             <p>A new requirement has been assigned to you.</p>
#             <p> Please find the details below:</p>
#             <table>
#                 <tr>
#                     <th style="width: 20%;">Job ID</th>
#                     <th style="width: 30%;">Client</th>
#                     <th style="width: 30%;">Role/Profile</th>
#                     <th style="width: 30%;">Location</th>
#                 </tr>
#                 {job_data}
#             </table>
#             <p>Please check in Job Listing page for more details.</p>
#             <p>Regards,</p>
#             <p><b>Makonis Talent Track Pro Team</b></p>
#         </div>
#     </body>
#     </html>
#     """

#     msg = Message(
#         # 'New Job Notification',
#         f'New Requirement Assigned',
#         sender='kanuparthisaiganesh582@gmail.com',
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     try:
#         mail.send(msg)
#     except Exception as e:
#         print("mail error",str(e))
#         return jsonify({'status': 'error', 'message': f'Failed to send mail: {str(e)}'}),500


# def post_job_send_notification(recruiter_email, new_recruiter_name, job_data):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 color: #333;
#                 line-height: 1.6;
#                 background-color: #f4f4f4;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #ffffff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #4CAF50;
#                 color: white;
#                 padding: 10px;
#                 text-align: center;
#                 font-size: 20px;
#                 border-radius: 8px 8px 0 0;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 10px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 8px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #4CAF50;
#                 color: white;
#             }}
#             tr:nth-child(even) {{
#                 background-color: #f9f9f9;
#             }}
#             p {{
#                 margin: 10px 0;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
#     <body>
#         <div class="container">
#             <div class="header">
#                 New Job Posted
#             </div>
#             <p>Dear {new_recruiter_name},</p>
#             <p>The following job have been assigned to you:</p>
#             <table>
#                 <tr>
#                     <th style="width: 20%;">Job ID</th>
#                     <th style="width: 30%;">Client</th>
#                     <th style="width: 30%;">Role/Profile</th>
#                     <th style="width: 30%;">Location</th>
#                 </tr>
#                 {job_data}
#             </table>
#             <p>Check your dashboard for more details.</p>
#             <p>Regards,</p>
#             <p>Your Company</p>
#         </div>
#     </body>
#     </html>
#     """

#     msg = Message(
#         'New Job Notification',
#         sender='ganesh.s@makonissoft.com',
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     mail.send(msg)

import re

def is_valid_email(email):
    """ Validate email format using regex. """
    regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(regex, email) is not None

@app.route('/post_job', methods=['POST'])
def post_job():
    data = request.json

    user_id = data.get('user_id')
    if not user_id:
        return jsonify({'status': 'error', 'message': 'user_id is required'}), 400

    user = User.query.filter_by(id=user_id).first()
    if not user:
        return jsonify({'status': 'error', 'message': 'User not found'}), 404

    if user.user_type != 'management':
        return jsonify({'status': 'error', 'message': 'You do not have permission to post a job'}), 401

    try:
        job_details = {
            'client': data['client'],
            'experience_min': data['experience_min'],
            'experience_max': data['experience_max'],
            'budget_min': f"{data['currency_type_min']} {data['budget_min']}",
            'budget_max': f"{data['currency_type_max']} {data['budget_max']}",
            'location': data['location'],
            'shift_timings': data['shift_timings'],
            'notice_period': data['notice_period'],
            'role': data['role'],
            'detailed_jd': data['detailed_jd'],
            'mode': data['mode'],
            'job_status': data['job_status'],
            'skills': data['skills'],
            'job_type': data['Job_Type'],
            'contract_in_months': data['Job_Type_details'] if data['Job_Type'] == 'Contract' else None
        }
    except KeyError as e:
        return jsonify({'status': 'error', 'message': f'Missing required field: {e}'}), 400

    jd_pdf = data.get('jd_pdf')
    jd_binary = None
    jd_pdf_present = False

    if jd_pdf:
        try:
            jd_binary = base64.b64decode(jd_pdf)
            jd_pdf_present = bool(jd_binary)
        except Exception:
            return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 400

    new_job_post = JobPost(
        client=job_details['client'],
        experience_min=job_details['experience_min'],
        experience_max=job_details['experience_max'],
        budget_min=job_details['budget_min'],
        budget_max=job_details['budget_max'],
        location=job_details['location'],
        shift_timings=job_details['shift_timings'],
        notice_period=job_details['notice_period'],
        role=job_details['role'],
        detailed_jd=job_details['detailed_jd'],
        recruiter=', '.join(data.get('recruiter', [])),
        management=user.username,
        mode=job_details['mode'],
        job_status=job_details['job_status'],
        job_type=job_details['job_type'],
        skills=job_details['skills'],
        contract_in_months=job_details['contract_in_months'],
        jd_pdf=jd_binary,
        jd_pdf_present=jd_pdf_present
    )

    try:
        current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        new_job_post.date_created = current_datetime.date()
        new_job_post.time_created = current_datetime.time()

        db.session.add(new_job_post)
        db.session.commit()

        job_post_id = new_job_post.id

        return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200
    except Exception as e:
        db.session.rollback()
        print(e)
        return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500

@app.route('/send_notifications', methods=['POST'])
def send_notifications():
    data = request.json
    job_post_id = data.get('job_post_id')

    if not job_post_id:
        return jsonify({'status': 'error', 'message': 'job_post_id is required'}), 400

    job_post = JobPost.query.filter_by(id=job_post_id).first()
    if not job_post:
        return jsonify({'status': 'error', 'message': 'Job post not found'}), 404

    job_data = f"<tr><td>{job_post.id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

    invalid_emails = []
    for recruiter_name in job_post.recruiter.split(','):
        recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
        if recruiter:
            if is_valid_email(recruiter.email):
                error_msg = post_job_send_notification(recruiter.email, recruiter.username, job_data)
                if error_msg:
                    return jsonify({'status': 'error', 'message': error_msg}), 500
            else:
                invalid_emails.append(recruiter.email)

    if invalid_emails:
        print("Invalid Emails:", invalid_emails)  # Print invalid_emails list
        return jsonify({'status': 'error', 'message': f'Invalid email format for: {", ".join(invalid_emails)}'}), 400

    return jsonify({'status': 'success', 'message': 'Notifications sent successfully'}), 200



# def is_valid_email(email):
#     """ Validate email format using regex. """
#     regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
#     return re.match(regex, email) is not None
    
# @app.route('/post_job', methods=['POST'])
# def post_job():
#     data = request.json

#     user_id = data.get('user_id')
#     if not user_id:
#         return jsonify({'status': 'error', 'message': 'user_id is required'}), 400

#     user = User.query.filter_by(id=user_id).first()
#     if not user:
#         return jsonify({'status': 'error', 'message': 'User not found'}), 404

#     if user.user_type != 'management':
#         return jsonify({'status': 'error', 'message': 'You do not have permission to post a job'}), 401

#     try:
#         job_details = {
#             'client': data['client'],
#             'experience_min': data['experience_min'],
#             'experience_max': data['experience_max'],
#             'budget_min': f"{data['currency_type_min']} {data['budget_min']}",
#             'budget_max': f"{data['currency_type_max']} {data['budget_max']}",
#             'location': data['location'],
#             'shift_timings': data['shift_timings'],
#             'notice_period': data['notice_period'],
#             'role': data['role'],
#             'detailed_jd': data['detailed_jd'],
#             'mode': data['mode'],
#             'job_status': data['job_status'],
#             'skills': data['skills'],
#             'job_type': data['Job_Type'],
#             'contract_in_months': data['Job_Type_details'] if data['Job_Type'] == 'Contract' else None
#         }
#     except KeyError as e:
#         return jsonify({'status': 'error', 'message': f'Missing required field: {e}'}), 400

#     jd_pdf = data.get('jd_pdf')
#     jd_binary = None
#     jd_pdf_present = False

#     if jd_pdf:
#         try:
#             jd_binary = base64.b64decode(jd_pdf)
#             jd_pdf_present = bool(jd_binary)
#         except Exception:
#             return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 400

#     new_job_post = JobPost(
#         client=job_details['client'],
#         experience_min=job_details['experience_min'],
#         experience_max=job_details['experience_max'],
#         budget_min=job_details['budget_min'],
#         budget_max=job_details['budget_max'],
#         location=job_details['location'],
#         shift_timings=job_details['shift_timings'],
#         notice_period=job_details['notice_period'],
#         role=job_details['role'],
#         detailed_jd=job_details['detailed_jd'],
#         recruiter=', '.join(data.get('recruiter', [])),
#         management=user.username,
#         mode=job_details['mode'],
#         job_status=job_details['job_status'],
#         job_type=job_details['job_type'],
#         skills=job_details['skills'],
#         contract_in_months=job_details['contract_in_months'],
#         jd_pdf=jd_binary,
#         jd_pdf_present=jd_pdf_present
#     )

#     try:
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         db.session.add(new_job_post)
#         db.session.commit()

#         job_post_id = new_job_post.id
#         job_data = f"<tr><td>{job_post_id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"
        
#         invalid_emails = []
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 if is_valid_email(recruiter.email):
#                     error_msg = post_job_send_notification(recruiter.email, recruiter.username, job_data)
#                     # pass
#                     if error_msg:
#                         return jsonify({'status': 'error', 'message': error_msg}), 500
#                         # pass
#                 else:
#                     invalid_emails.append(recruiter.email)

#         if invalid_emails:
#             print("Invalid Emails:", invalid_emails)  # Print invalid_emails list
#             return jsonify({'status': 'error', 'message': f'Invalid email format for: {", ".join(invalid_emails)}'}), 400

#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200
#     except Exception as e:
#         db.session.rollback()
#         print(e)
#         return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500


# @app.route('/post_job', methods=['POST'])
# def post_job():
#     data = request.json

#     user_id = data.get('user_id')
#     if not user_id:
#         return jsonify({'status': 'error', 'message': 'user_id is required'}), 400

#     user = User.query.filter_by(id=user_id).first()
#     if not user:
#         return jsonify({'status': 'error', 'message': 'User not found'}), 404

#     if user.user_type != 'management':
#         return jsonify({'status': 'error', 'message': 'You do not have permission to post a job'}), 401

#     try:
#         job_details = {
#             'client': data['client'],
#             'experience_min': data['experience_min'],
#             'experience_max': data['experience_max'],
#             'budget_min': f"{data['currency_type_min']} {data['budget_min']}",
#             'budget_max': f"{data['currency_type_max']} {data['budget_max']}",
#             'location': data['location'],
#             'shift_timings': data['shift_timings'],
#             'notice_period': data['notice_period'],
#             'role': data['role'],
#             'detailed_jd': data['detailed_jd'],
#             'mode': data['mode'],
#             'job_status': data['job_status'],
#             'skills': data['skills'],
#             'job_type': data['Job_Type'],
#             'contract_in_months': data['Job_Type_details'] if data['Job_Type'] == 'Contract' else None
#         }
#     except KeyError as e:
#         return jsonify({'status': 'error', 'message': f'Missing required field: {e}'}), 400

#     jd_pdf = data.get('jd_pdf')
#     jd_binary = None
#     jd_pdf_present = False

#     if jd_pdf:
#         try:
#             jd_binary = base64.b64decode(jd_pdf)
#             jd_pdf_present = bool(jd_binary)
#         except Exception:
#             return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 400

#     new_job_post = JobPost(
#         client=job_details['client'],
#         experience_min=job_details['experience_min'],
#         experience_max=job_details['experience_max'],
#         budget_min=job_details['budget_min'],
#         budget_max=job_details['budget_max'],
#         location=job_details['location'],
#         shift_timings=job_details['shift_timings'],
#         notice_period=job_details['notice_period'],
#         role=job_details['role'],
#         detailed_jd=job_details['detailed_jd'],
#         recruiter=', '.join(data.get('recruiter', [])),
#         management=user.username,
#         mode=job_details['mode'],
#         job_status=job_details['job_status'],
#         job_type=job_details['job_type'],
#         skills=job_details['skills'],
#         contract_in_months=job_details['contract_in_months'],
#         jd_pdf=jd_binary,
#         jd_pdf_present=jd_pdf_present
#     )

#     try:
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         db.session.add(new_job_post)
#         db.session.commit()

#         job_post_id = new_job_post.id

#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)

#         db.session.commit()

#         job_data = f"<tr><td>{job_post_id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"
        
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 error_msg = post_job_send_notification(recruiter.email, recruiter.username, job_data)
#                 if error_msg:
#                     return jsonify({'status': 'error', 'message': error_msg}), 500

#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200
#     except Exception as e:
#         db.session.rollback()
#         print(e)
#         return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500

# @app.route('/post_job', methods=['POST'])
# def post_job():
#     data = request.json
    
#     user_id = data.get('user_id')
#     if not user_id:
#         return jsonify({'status': 'error', 'message': 'user_id is required'}), 400

#     user = User.query.filter_by(id=user_id).first()
#     if not user:
#         return jsonify({'status': 'error', 'message': 'User not found'}), 404

#     if user.user_type != 'management':
#         return jsonify({'status': 'error', 'message': 'You do not have permission to post a job'}), 401

#     try:
#         job_details = {
#             'client': data['client'],
#             'experience_min': data['experience_min'],
#             'experience_max': data['experience_max'],
#             'budget_min': f"{data['currency_type_min']} {data['budget_min']}",
#             'budget_max': f"{data['currency_type_max']} {data['budget_max']}",
#             'location': data['location'],
#             'shift_timings': data['shift_timings'],
#             'notice_period': data['notice_period'],
#             'role': data['role'],
#             'detailed_jd': data['detailed_jd'],
#             'mode': data['mode'],
#             'job_status': data['job_status'],
#             'skills': data['skills'],
#             'job_type': data['Job_Type'],
#             'contract_in_months': data['Job_Type_details'] if data['Job_Type'] == 'Contract' else None
#         }
#     except KeyError as e:
#         return jsonify({'status': 'error', 'message': f'Missing required field: {e}'}), 400
#     except ValueError as e:
#         return jsonify({'status': 'error', 'message': f'Invalid data type for job details: {e}'}), 400
#     except Exception as e:
#         print("job gettnng error", e)
#         return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500

#     jd_pdf = data.get('jd_pdf')
#     jd_binary = None
#     jd_pdf_present = False

#     if jd_pdf:
#         try:
#             jd_binary = base64.b64decode(jd_pdf)
#             jd_pdf_present = bool(jd_binary)
#         except Exception:
#             return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 400

#     new_job_post = JobPost(
#         client=job_details['client'],
#         experience_min=job_details['experience_min'],
#         experience_max=job_details['experience_max'],
#         budget_min=job_details['budget_min'],
#         budget_max=job_details['budget_max'],
#         location=job_details['location'],
#         shift_timings=job_details['shift_timings'],
#         notice_period=job_details['notice_period'],
#         role=job_details['role'],
#         detailed_jd=job_details['detailed_jd'],
#         recruiter=', '.join(data.get('recruiter', [])),
#         management=user.username,
#         mode=job_details['mode'],
#         job_status=job_details['job_status'],
#         job_type=job_details['job_type'],
#         skills=job_details['skills'],
#         contract_in_months=job_details['contract_in_months'],
#         jd_pdf=jd_binary,
#         jd_pdf_present=jd_pdf_present
#     )

#     try:
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         db.session.add(new_job_post)
#         db.session.commit()

#         job_post_id = new_job_post.id

#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)

#         db.session.commit()

#         job_data = f"<tr><td>{job_post_id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"
#         print('recruiter', data.get('recruiter', []))
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 post_job_send_notification(recruiter.email, recruiter.username, job_data)

#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200
#     except Exception as e:
#         db.session.rollback()
#         print("final exception ", e)
#         return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500


# @app.route('/post_job', methods=['POST'])
# def post_job():
#     data = request.json
    
#     user_id = data.get('user_id')
#     if not user_id:
#         return jsonify({'status': 'error', 'message': 'user_id is required'}), 500

#     user = User.query.filter_by(id=user_id).first()
#     if not user:
#         return jsonify({'status': 'error', 'message': 'User not found'}), 500

#     if user.user_type != 'management':
#         return jsonify({'status': 'error', 'message': 'You do not have permission to post a job'}), 500

#     try:
#         job_details = {
#             'client': data['client'],
#             'experience_min': data['experience_min'],
#             'experience_max': data['experience_max'],
#             'budget_min': f"{data['currency_type_min']} {data['budget_min']}",
#             'budget_max': f"{data['currency_type_max']} {data['budget_max']}",
#             'location': data['location'],
#             'shift_timings': data['shift_timings'],
#             'notice_period': data['notice_period'],
#             'role': data['role'],
#             'detailed_jd': data['detailed_jd'],
#             'mode': data['mode'],
#             'job_status': data['job_status'],
#             'skills': data['skills'],
#             'job_type': data['Job_Type'],
#             'contract_in_months': data['Job_Type_details'] if data['Job_Type'] == 'Contract' else None
#         }
#     except KeyError as e:
#         return jsonify({'status': 'error', 'message': f'Missing required field: {e}'}),500
#     except ValueError as e:
#         return jsonify({'status': 'error', 'message': f'Invalid data type for job details: {e}'}),500
#     except Exception as e:
#         print(e)
#         return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}),500

#     jd_pdf = data.get('jd_pdf')
#     jd_binary = None
#     jd_pdf_present = False

#     if jd_pdf:
#         try:
#             jd_binary = base64.b64decode(jd_pdf)
#             jd_pdf_present = bool(jd_binary)
#         except Exception:
#             return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 500

#     new_job_post = JobPost(
#         client=job_details['client'],
#         experience_min=job_details['experience_min'],
#         experience_max=job_details['experience_max'],
#         budget_min=job_details['budget_min'],
#         budget_max=job_details['budget_max'],
#         location=job_details['location'],
#         shift_timings=job_details['shift_timings'],
#         notice_period=job_details['notice_period'],
#         role=job_details['role'],
#         detailed_jd=job_details['detailed_jd'],
#         recruiter=', '.join(data.get('recruiter', [])),
#         management=user.username,
#         mode=job_details['mode'],
#         job_status=job_details['job_status'],
#         job_type=job_details['job_type'],
#         skills=job_details['skills'],
#         contract_in_months=job_details['contract_in_months'],
#         jd_pdf=jd_binary,
#         jd_pdf_present=jd_pdf_present
#     )

#     try:
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         db.session.add(new_job_post)
#         db.session.commit()

#         job_post_id = new_job_post.id

#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)

#         db.session.commit()

#         job_data = f"<tr><td>{job_post_id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"
#         print('reecruier',  data.get('recruiter', []))
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             print('recruiter 1', recruiter.username)
#             if recruiter:
#                 print('recruiter 2', recruiter.username)
#                 post_job_send_notification(recruiter.email, recruiter.username, job_data)

#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200
#     except Exception as e:
#         db.session.rollback()
#         print(e)
#         return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500


# @app.route('/post_job', methods=['POST'])
# def post_job():
#     data = request.json
#     user_id = data.get('user_id')
    
#     if user_id is None:
#         return jsonify({'status': 'error', 'message': 'user_id is required'}), 400

#     user = User.query.filter_by(id=user_id).first()

#     if not user:
#         return jsonify({'status': 'error', 'message': 'User not found'}), 400

#     if user.user_type != 'management':
#         return jsonify({'status': 'error', 'message': 'Job post not added successfully'}), 400

#     job_details = {
#         'client': data.get('client'),
#         'experience_min': data.get('experience_min'),
#         'experience_max': data.get('experience_max'),
#         'budget_min': f"{data.get('currency_type_min')} {data.get('budget_min')}",
#         'budget_max': f"{data.get('currency_type_max')} {data.get('budget_max')}",
#         'location': data.get('location'),
#         'shift_timings': data.get('shift_timings'),
#         'notice_period': data.get('notice_period'),
#         'role': data.get('role'),
#         'detailed_jd': data.get('detailed_jd'),
#         'mode': data.get('mode'),
#         'job_status': data.get('job_status'),
#         'skills': data.get('skills'),
#         'job_type': data.get('Job_Type'),
#         'contract_in_months': data.get('Job_Type_details') if data.get('Job_Type') == 'Contract' else None
#     }

#     jd_pdf = data.get('jd_pdf')
#     jd_binary = None
#     jd_pdf_present = False

#     if jd_pdf:
#         try:
#             jd_binary = base64.b64decode(jd_pdf)
#             jd_pdf_present = bool(jd_binary)
#         except Exception:
#             return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 400

#     new_job_post = JobPost(
#         client=job_details['client'],
#         experience_min=job_details['experience_min'],
#         experience_max=job_details['experience_max'],
#         budget_min=job_details['budget_min'],
#         budget_max=job_details['budget_max'],
#         location=job_details['location'],
#         shift_timings=job_details['shift_timings'],
#         notice_period=job_details['notice_period'],
#         role=job_details['role'],
#         detailed_jd=job_details['detailed_jd'],
#         recruiter=', '.join(data.get('recruiter', [])),
#         management=user.username,
#         mode=job_details['mode'],
#         job_status=job_details['job_status'],
#         job_type=job_details['job_type'],
#         skills=job_details['skills'],
#         contract_in_months=job_details['contract_in_months'],
#         jd_pdf=jd_binary,
#         jd_pdf_present=jd_pdf_present
#     )

#     current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#     new_job_post.date_created = current_datetime.date()
#     new_job_post.time_created = current_datetime.time()

#     db.session.add(new_job_post)
#     db.session.commit()

#     job_post_id = new_job_post.id

#     for recruiter_name in data.get('recruiter', []):
#         notification = Notification(
#             job_post_id=job_post_id,
#             recruiter_name=recruiter_name.strip(),
#             notification_status=False
#         )
#         db.session.add(notification)

#     db.session.commit()

#     job_data = f"<tr><td>{job_post_id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"

#     for recruiter_name in data.get('recruiter', []):
#         recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#         if recruiter:
#             post_job_send_notification(recruiter.email, recruiter.username, job_data)

#     return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200


# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()

#         if not user:
#             return jsonify({'status': 'error', 'message': 'User not found'}), 400

#         if user.user_type != 'management':
#             return jsonify({'status': 'error', 'message': 'Job post not added successfully'}), 400

#         job_details = {
#             'client': data.get('client'),
#             'experience_min': data.get('experience_min'),
#             'experience_max': data.get('experience_max'),
#             'budget_min': f"{data.get('currency_type_min')} {data.get('budget_min')}",
#             'budget_max': f"{data.get('currency_type_max')} {data.get('budget_max')}",
#             'location': data.get('location'),
#             'shift_timings': data.get('shift_timings'),
#             'notice_period': data.get('notice_period'),
#             'role': data.get('role'),
#             'detailed_jd': data.get('detailed_jd'),
#             'mode': data.get('mode'),
#             'job_status': data.get('job_status'),
#             'skills': data.get('skills'),
#             'job_type': data.get('Job_Type'),
#             'contract_in_months': data.get('Job_Type_details') if data.get('Job_Type') == 'Contract' else None
#         }

#         jd_pdf = data.get('jd_pdf')
#         jd_binary = None
#         jd_pdf_present = False

#         if jd_pdf:
#             try:
#                 jd_binary = base64.b64decode(jd_pdf)
#                 jd_pdf_present = bool(jd_binary)
#             except Exception as e:
#                 return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file', 'details': str(e)}), 400

#         new_job_post = JobPost(
#             client=job_details['client'],
#             experience_min=job_details['experience_min'],
#             experience_max=job_details['experience_max'],
#             budget_min=job_details['budget_min'],
#             budget_max=job_details['budget_max'],
#             location=job_details['location'],
#             shift_timings=job_details['shift_timings'],
#             notice_period=job_details['notice_period'],
#             role=job_details['role'],
#             detailed_jd=job_details['detailed_jd'],
#             recruiter=', '.join(data.get('recruiter', [])),
#             management=user.username,
#             mode=job_details['mode'],
#             job_status=job_details['job_status'],
#             job_type=job_details['job_type'],
#             skills=job_details['skills'],
#             contract_in_months=job_details['contract_in_months'],
#             jd_pdf=jd_binary,
#             jd_pdf_present=jd_pdf_present
#         )

#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         db.session.add(new_job_post)
#         db.session.commit()

#         job_post_id = new_job_post.id

#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)

#         db.session.commit()

#         job_data = f"<tr><td>{job_post_id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"

#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 post_job_send_notification(recruiter.email, recruiter.username, job_data)

#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200

#     except Exception as e:
#         print(e)
#         return jsonify({"status": "error", "message": str(e)}), 500



# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()

#         if not user:
#             return jsonify({'status': 'error', 'message': 'User not found'}), 400

#         if user.user_type != 'management':
#             return jsonify({'status': 'error', 'message': 'Job post not added successfully'}), 400

#         job_details = {
#             'client': data.get('client'),
#             'experience_min': data.get('experience_min'),
#             'experience_max': data.get('experience_max'),
#             'budget_min': f"{data.get('currency_type_min')} {data.get('budget_min')}",
#             'budget_max': f"{data.get('currency_type_max')} {data.get('budget_max')}",
#             'location': data.get('location'),
#             'shift_timings': data.get('shift_timings'),
#             'notice_period': data.get('notice_period'),
#             'role': data.get('role'),
#             'detailed_jd': data.get('detailed_jd'),
#             'mode': data.get('mode'),
#             'job_status': data.get('job_status'),
#             'skills': data.get('skills'),
#             'job_type': data.get('Job_Type'),
#             'contract_in_months': data.get('Job_Type_details') if data.get('Job_Type') == 'Contract' else None
#         }

#         jd_pdf = data.get('jd_pdf')
#         jd_binary = None
#         jd_pdf_present = False

#         if jd_pdf:
#             try:
#                 jd_binary = base64.b64decode(jd_pdf)
#                 jd_pdf_present = bool(jd_binary)
#             except Exception as e:
#                 return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file', 'details': str(e)}), 400

#         new_job_post = JobPost(
#             client=job_details['client'],
#             experience_min=job_details['experience_min'],
#             experience_max=job_details['experience_max'],
#             budget_min=job_details['budget_min'],
#             budget_max=job_details['budget_max'],
#             location=job_details['location'],
#             shift_timings=job_details['shift_timings'],
#             notice_period=job_details['notice_period'],
#             role=job_details['role'],
#             detailed_jd=job_details['detailed_jd'],
#             recruiter=', '.join(data.get('recruiter', [])),
#             management=user.username,
#             mode=job_details['mode'],
#             job_status=job_details['job_status'],
#             job_type=job_details['job_type'],
#             skills=job_details['skills'],
#             contract_in_months=job_details['contract_in_months'],
#             jd_pdf=jd_binary,
#             jd_pdf_present=jd_pdf_present
#         )

#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         db.session.add(new_job_post)
#         db.session.commit()

#         job_post_id = new_job_post.id

#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)

#         db.session.commit()

#         job_data = f"<tr><td>{new_job_post.id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"

#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 post_job_send_notification(recruiter.email, recruiter.username, job_data)

#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_id': job_post_id}), 200

#     except KeyError as e:
#         return jsonify({"status": "error", "message": f"KeyError: {e}"}), 500

#     except Exception as e:
#         return jsonify({"status": "error", "message": str(e)}), 500

# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()

#         # Check if the user exists
#         if not user:
#             return jsonify({'status': 'error', 'message': 'User not found'}), 400

#         user_type = user.user_type
#         if user_type != 'management':
#             return jsonify({'status': 'error', 'message': 'Job post not added successfully'}), 400

#         # Extracting job details from the request
#         job_details = {
#             'client': data.get('client'),
#             'experience_min': data.get('experience_min'),
#             'experience_max': data.get('experience_max'),
#             'budget_min': f"{data.get('currency_type_min')} {data.get('budget_min')}",
#             'budget_max': f"{data.get('currency_type_max')} {data.get('budget_max')}",
#             'location': data.get('location'),
#             'shift_timings': data.get('shift_timings'),
#             'notice_period': data.get('notice_period'),
#             'role': data.get('role'),
#             'detailed_jd': data.get('detailed_jd'),
#             'mode': data.get('mode'),
#             'job_status': data.get('job_status'),
#             'skills': data.get('skills'),
#             'job_type': data.get('Job_Type'),
#             'contract_in_months': data.get('Job_Type_details') if data.get('Job_Type') == 'Contract' else None
#         }

#         # Decode the base64 encoded PDF file
#         jd_pdf = data.get('jd_pdf')
#         jd_binary = None
#         jd_pdf_present = False  # Default value
#         if jd_pdf:
#             try:
#                 jd_binary = base64.b64decode(jd_pdf)
#                 jd_pdf_present = bool(jd_binary)  # If jd_binary is not None, set jd_pdf_present to True
#             except Exception as e:
#                 return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file', 'details': str(e)}), 400

#         # Create a new job post instance
#         new_job_post = JobPost(
#             client=job_details['client'],
#             experience_min=job_details['experience_min'],
#             experience_max=job_details['experience_max'],
#             budget_min=job_details['budget_min'],
#             budget_max=job_details['budget_max'],
#             location=job_details['location'],
#             shift_timings=job_details['shift_timings'],
#             notice_period=job_details['notice_period'],
#             role=job_details['role'],
#             detailed_jd=job_details['detailed_jd'],
#             recruiter=', '.join(data.get('recruiter', [])),
#             management=user.username,
#             mode=job_details['mode'],
#             job_status=job_details['job_status'],
#             job_type=job_details['job_type'],
#             skills=job_details['skills'],
#             contract_in_months=job_details['contract_in_months'],
#             jd_pdf=jd_binary,  # Store the binary data in the database
#             jd_pdf_present=jd_pdf_present  # Store whether PDF is present
#         )

#         # Set created date and time
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         # Add the new job post to the session and commit
#         db.session.add(new_job_post)
#         db.session.commit()

#         # Generate job_post_id after committing the new_job_post
#         job_post_id = new_job_post.id

#         # Add notifications for each recruiter
#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)

#         db.session.commit()

#         # Create job data for the email notification
#         job_data = f"<tr><td>{new_job_post.id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"

#         # Send notifications to recruiters
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 post_job_send_notification(recruiter.email, recruiter.username, job_data)

#         # Return the job_id along with the success message
#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_id': job_post_id}), 200

#     except KeyError as e:
#         return jsonify({"status": "error", "message": f"KeyError: {e}"}), 400

#     except Exception as e:
#         return jsonify({"status": "error", "message": str(e)}), 500

############################################################################################
# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()

#         # Check if the user exists
#         if not user:
#             return jsonify({'status': 'error', 'message': 'User not found'}), 400

#         user_type = user.user_type
#         if user_type != 'management':
#             return jsonify({'status': 'error', 'message': 'Job post not added successfully'}), 400

#         # Extracting job details from the request
#         job_details = {
#             'client': data.get('client'),
#             'experience_min': data.get('experience_min'),
#             'experience_max': data.get('experience_max'),
#             'budget_min': f"{data.get('currency_type_min')} {data.get('budget_min')}",
#             'budget_max': f"{data.get('currency_type_max')} {data.get('budget_max')}",
#             'location': data.get('location'),
#             'shift_timings': data.get('shift_timings'),
#             'notice_period': data.get('notice_period'),
#             'role': data.get('role'),
#             'detailed_jd': data.get('detailed_jd'),
#             'mode': data.get('mode'),
#             'job_status': data.get('job_status'),
#             'skills': data.get('skills'),
#             'job_type': data.get('Job_Type'),
#             'contract_in_months': data.get('Job_Type_details') if data.get('Job_Type') == 'Contract' else None
#         }

#         # Decode the base64 encoded PDF file
#         jd_pdf = data.get('jd_pdf')
#         jd_binary = None
#         jd_pdf_present = False  # Default value
#         if jd_pdf:
#             try:
#                 jd_binary = base64.b64decode(jd_pdf)
#                 jd_pdf_present = bool(jd_binary)  # If jd_binary is not None, set jd_pdf_present to True
#             except Exception as e:
#                 return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file', 'details': str(e)}), 400

#         # Create a new job post instance
#         new_job_post = JobPost(
#             client=job_details['client'],
#             experience_min=job_details['experience_min'],
#             experience_max=job_details['experience_max'],
#             budget_min=job_details['budget_min'],
#             budget_max=job_details['budget_max'],
#             location=job_details['location'],
#             shift_timings=job_details['shift_timings'],
#             notice_period=job_details['notice_period'],
#             role=job_details['role'],
#             detailed_jd=job_details['detailed_jd'],
#             recruiter=', '.join(data.get('recruiter', [])),
#             management=user.username,
#             mode=job_details['mode'],
#             job_status=job_details['job_status'],
#             job_type=job_details['job_type'],
#             skills=job_details['skills'],
#             contract_in_months=job_details['contract_in_months'],
#             jd_pdf=jd_binary,  # Store the binary data in the database
#             jd_pdf_present=jd_pdf_present  # Store whether PDF is present
#         )

#         # Set created date and time
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         # Add the new job post to the session and commit
#         db.session.add(new_job_post)
#         db.session.commit()

#         # Generate job_post_id after committing the new_job_post
#         job_post_id = new_job_post.id

#         # Add notifications for each recruiter
#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)

#         db.session.commit()

#         # Send notifications to recruiters
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 send_notification(recruiter.email)

#         # Return the job_id along with the success message
#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_id': job_post_id}), 200

#     except KeyError as e:
#         return jsonify({"status": "error", "message": f"KeyError: {e}"})

#     except Exception as e:
#         return jsonify({"status": "error", "message": str(e)}), 500

################################################################################################

# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists
#         if not user:
#             return jsonify({'status': 'error', 'message': 'User not found'}), 400
        
#         user_type = user.user_type
#         if user_type != 'management':
#             return jsonify({'status': 'error', 'message': 'Job post not added successfully'})
        
#         # Extracting job details from the request
#         job_details = {
#             'client': data.get('client'),
#             'experience_min': data.get('experience_min'),
#             'experience_max': data.get('experience_max'),
#             'budget_min': f"{data.get('currency_type_min')} {data.get('budget_min')}",
#             'budget_max': f"{data.get('currency_type_max')} {data.get('budget_max')}",
#             'location': data.get('location'),
#             'shift_timings': data.get('shift_timings'),
#             'notice_period': data.get('notice_period'),
#             'role': data.get('role'),
#             'detailed_jd': data.get('detailed_jd'),
#             'mode': data.get('mode'),
#             'job_status': data.get('job_status'),
#             'skills': data.get('skills'),
#             'job_type': data.get('Job_Type'),
#             'contract_in_months': data.get('Job_Type_details') if data.get('Job_Type') == 'Contract' else None
#         }
        
#         # Decode the base64 encoded PDF file
#         jd_pdf = data.get('jd_pdf')
#         jd_binary = None
#         jd_pdf_present = False  # Default value
#         if jd_pdf:
#             try:
#                 jd_binary = base64.b64decode(jd_pdf)
#                 jd_pdf_present = bool(jd_binary)  # If jd_binary is not None, set jd_pdf_present to True
#             except Exception as e:
#                 return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file', 'details': str(e)})
        
#         # Create a new job post instance
#         new_job_post = JobPost(
#             client=job_details['client'],
#             experience_min=job_details['experience_min'],
#             experience_max=job_details['experience_max'],
#             budget_min=job_details['budget_min'],
#             budget_max=job_details['budget_max'],
#             location=job_details['location'],
#             shift_timings=job_details['shift_timings'],
#             notice_period=job_details['notice_period'],
#             role=job_details['role'],
#             detailed_jd=job_details['detailed_jd'],
#             recruiter=', '.join(data.get('recruiter', [])),
#             management=user.username,
#             mode=job_details['mode'],
#             job_status=job_details['job_status'],
#             job_type=job_details['job_type'],
#             skills=job_details['skills'],
#             contract_in_months=job_details['contract_in_months'],
#             jd_pdf=jd_binary,  # Store the binary data in the database
#             jd_pdf_present=jd_pdf_present  # Store whether PDF is present
#         )

#         # Set created date and time
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()
        
#         # Add the new job post to the session and commit
#         db.session.add(new_job_post)
#         db.session.commit()
        
#         # Generate job_post_id after committing the new_job_post
#         job_post_id = new_job_post.id
        
#         # Add notifications for each recruiter
#         for recruiter_name in data.get('recruiter', []):
#             notification = Notification(
#                 job_post_id=job_post_id,
#                 recruiter_name=recruiter_name.strip(),
#                 notification_status=False
#             )
#             db.session.add(notification)
        
#         db.session.commit()

#         # Send notifications to recruiters
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 send_notification(recruiter.email)
        
#         # Return the job_id along with the success message
#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_id': job_post_id}), 200

#     except KeyError as e:
#         return jsonify({"status": "error", "message": f"KeyError: {e}"})

#     except Exception as e:
#         return jsonify({"status": "error", "message": str(e)}), 500

# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists
#         if user:
#             user_type = user.user_type

#             if user_type == 'management':
#                 # Extract data from the request
#                 client = data.get('client')
#                 experience_min = data.get('experience_min')
#                 experience_max = data.get('experience_max')
#                 budget_min = data.get('budget_min')
#                 budget_max = data.get('budget_max')
#                 currency_type_min = data.get('currency_type_min')
#                 currency_type_max = data.get('currency_type_max')
#                 budget_min = currency_type_min + ' ' + budget_min
#                 budget_max = currency_type_max + ' ' + budget_max
#                 location = data.get('location')
#                 shift_timings = data.get('shift_timings')
#                 notice_period = data.get('notice_period')
#                 role = data.get('role')
#                 detailed_jd = data.get('detailed_jd')
#                 mode = data.get('mode')
#                 job_status = data.get('job_status')
#                 skills = data.get('skills')
#                 jd_pdf = data.get('jd_pdf')

#                 job_type = data.get('job_type')
#                 contract_in_months = None  # Initialize the variable here
#                 if job_type == 'Contract':
#                     contract_in_months = data.get('Job_Type_details')
#                     # job_type = job_type + '(' + Job_Type_details + ' Months )'
#                 else:
#                     pass

#                 # Decode the base64 encoded PDF file
#                 jd_binary = None
#                 jd_pdf_present = False  # Default value
#                 if jd_pdf:
#                     try:
#                         jd_binary = base64.b64decode(jd_pdf)
#                         jd_pdf_present = bool(jd_binary)  # If jd_binary is not None, set jd_pdf_present to True
#                     except Exception as e:
#                         return jsonify({'status': 'error',"message": "Error decoding base64 PDF file", "details": str(e)})
                        
#                 recruiter_names = data.get('recruiter', [])
#                 joined_recruiters = ', '.join(recruiter_names)

#                 # Create a new job post instance
#                 new_job_post = JobPost(
#                     client=client,
#                     experience_min=experience_min,
#                     experience_max=experience_max,
#                     budget_min=budget_min,
#                     budget_max=budget_max,
#                     location=location,
#                     shift_timings=shift_timings,
#                     notice_period=notice_period,
#                     role=role,
#                     detailed_jd=detailed_jd,
#                     recruiter=joined_recruiters,
#                     management=user.username,
#                     mode=mode,
#                     job_status=job_status,
#                     job_type=job_type,
#                     skills=skills,
#                     contract_in_months = contract_in_months,
#                     jd_pdf=jd_binary,  # Store the binary data in the database
#                     jd_pdf_present=jd_pdf_present  # Store whether PDF is present
#                 )

#                 # Created data and time
#                 current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#                 new_job_post.date_created = current_datetime.date()
#                 new_job_post.time_created  = current_datetime.time()

                       
#                 # Add the new job post to the session and commit
#                 db.session.add(new_job_post)
#                 db.session.commit()

#                 # Generate job_post_id after committing the new_job_post
#                 job_post_id = new_job_post.id

#                 # Add notifications for each recruiter
#                 recruiter_names = data.get('recruiter', [])
#                 for recruiter_name in recruiter_names:
#                     notification = Notification(
#                         job_post_id=job_post_id,
#                         recruiter_name=recruiter_name.strip(),
#                         notification_status=False
#                     )
#                     db.session.add(notification)

#                 db.session.commit()

#                 # Return the job_id along with the success message
#                 return jsonify({'status': 'success',"message": "Job posted successfully", "job_id": job_post_id}), 200
#             else:
#                 return jsonify({'status': 'error',"message": "Job post are not add successfully"})
#         else:
#             return jsonify({'status': 'error',"message": "User not found"}), 400

#     except KeyError as e:
#         return jsonify({"error": f"KeyError: {e}"}), 400

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500



# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
#         print("user :",user)
#         user_name = user.username
#         print("user_name:",user_name)
#         # Check if the "user_name" field exists
#         if user_name:
#             user_type = user.user_type

#             if user_type == 'management':
#                 client = data.get('client')
#                 experience_min = data.get('experience_min')
#                 experience_max = data.get('experience_max')
#                 budget_min = data.get('budget_min')
#                 budget_max = data.get('budget_max')
#                 currency_type_min = data.get('currency_type_min')
#                 currency_type_max = data.get('currency_type_max')
#                 budget_min = currency_type_min + ' ' + budget_min
#                 budget_max = currency_type_max + ' ' + budget_max
#                 location = data.get('location')
#                 shift_timings = data.get('shift_timings')
#                 notice_period = data.get('notice_period')
#                 role = data.get('role')
#                 detailed_jd = data.get('detailed_jd')
#                 mode = data.get('mode')
#                 job_status = data.get('job_status')
#                 job_type = data.get('job_type')
#                 skills = data.get('skills')
#                 jd_pdf = data.get('jd_pdf')
#                 # jd_binary = base64.b64decode(jd_pdf)
#                 # # Set jd_pdf_present based on the presence of jd_pdf
#                 # jd_pdf_present = bool(jd_binary)  # If jd_binary is not None, set jd_pdf_present to True

                
#                 # # Set jd_pdf_present based on the presence of jd_pdf
#                 # if jd_binary is not None:
#                 #     jd_pdf_present = True
#                 # else:
#                 #     jd_pdf_present = False
                    
                
#                 # Job_Type_details=data.get('Job_Type_details')

#                 if job_type == 'Contract':
#                     Job_Type_details = data.get('Job_Type_details')
#                     job_type = job_type + '(' + Job_Type_details + ' Months )'

                # recruiter_names = data.get('recruiter', [])
                # joined_recruiters = ', '.join(recruiter_names)

#                 new_job_post = JobPost(
#                     client=client,
#                     experience_min=experience_min,
#                     experience_max=experience_max,
#                     budget_min=budget_min,
#                     budget_max=budget_max,
#                     location=location,
#                     shift_timings=shift_timings,
#                     notice_period=notice_period,
#                     role=role,
#                     detailed_jd=detailed_jd,
#                     mode=mode,
#                     recruiter=joined_recruiters,
#                     management=user.username,
#                     job_status=job_status,
#                     job_type=job_type,
#                     skills=skills,
#                     jd_pdf=jd_pdf
#                     # jd_pdf_present=jd_pdf_present
#                 )

#                 new_job_post.notification = 'no'
#                 # new_job_post.date_created = date.today()
#                 # new_job_post.time_created = datetime.now().time()
            
                # # Created data and time
                # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
                # new_job_post.date_created = current_datetime.date()
                # new_job_post.time_created  = current_datetime.time()

#                 # Add the new_job_post to the session and commit to generate the job_post_id
#                 db.session.add(new_job_post)
#                 db.session.commit()

#                 # Generate job_post_id after committing the new_job_post
#                 job_post_id = new_job_post.id

#                 # Define an empty list to hold Notification instances
#                 notifications = []

#                 for recruiter_name in joined_recruiters.split(','):
#                     notification_status = False
#                     notification = Notification(
#                         job_post_id=job_post_id,  # Add job_post_id to Notification
#                         recruiter_name=recruiter_name.strip(),
#                         notification_status=notification_status
#                     )
#                     # Append each Notification instance to the notifications list
#                     notifications.append(notification)

#                 # Add the notifications to the session and commit
#                 db.session.add_all(notifications)
#                 db.session.commit()
#                 notifications = Notification.query.filter_by(job_post_id=job_post_id).all()
#                 for notification in notifications:
#                     notification.num_notification += 1
#                 db.session.commit()

                # # Retrieve the email addresses of the recruiters
                # recruiter_emails = [recruiter.email for recruiter in User.query.filter(User.username.in_(recruiter_names),
                #                                                                          User.user_type == 'recruiter',
                #                                                                          User.is_active == True,
                #                                                                          User.is_verified == True)]
                # for email in recruiter_emails:
                #     send_notification(email)

#                 # Return the job_id along with the success message
#                 return jsonify({"message": "Job posted successfully", "job_id": job_post_id}), 200
#             else:
#                 return jsonify({"error": "Invalid user type"}), 400
#         else:
#             return jsonify({"error": "Missing 'user_name' field in the request"}), 400

#     except KeyError as e:
#         return jsonify({"error": f"KeyError: {e}"}), 400

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
#         user_name = user.username
#         # Check if the "user_name" field exists
#         if user_name:
#             user_type = user.user_type

#             if user_type == 'management':
#                 client = data.get('client')
#                 experience_min = data.get('experience_min')
#                 experience_max = data.get('experience_max')
#                 budget_min = data.get('budget_min')
#                 budget_max = data.get('budget_max')
#                 currency_type_min = data.get('currency_type_min')
#                 currency_type_max = data.get('currency_type_max')
#                 budget_min = currency_type_min + ' ' + budget_min if currency_type_min and budget_min else budget_min
#                 budget_max = currency_type_max + ' ' + budget_max if currency_type_max and budget_max else budget_max
#                 location = data.get('location')
#                 shift_timings = data.get('shift_timings')
#                 notice_period = data.get('notice_period')
#                 role = data.get('role')
#                 detailed_jd = data.get('detailed_jd')
#                 mode = data.get('mode')
#                 job_status = data.get('job_status')
#                 job_type = data.get('job_type')
#                 skills = data.get('skills')
#                 jd_pdf = data.get('jd_pdf')

#                 # Debug print statements
#                 print(f"Received job_type: {job_type}")
#                 print(f"Received data: {data}")

#                 if job_type == 'Contract':
#                     Job_Type_details = data.get('Job_Type_details')
#                     job_type = f'{job_type} ({Job_Type_details} Months)' if Job_Type_details else job_type

#                 recruiter_names = data.get('recruiter', [])
#                 joined_recruiters = ', '.join(recruiter_names)

                # new_job_post = JobPost(
                #     client=client,
                #     experience_min=experience_min,
                #     experience_max=experience_max,
                #     budget_min=budget_min,
                #     budget_max=budget_max,
                #     location=location,
                #     shift_timings=shift_timings,
                #     notice_period=notice_period,
                #     role=role,
                #     detailed_jd=detailed_jd,
                #     mode=mode,
                #     recruiter=joined_recruiters,
                #     management=user.username,
                #     job_status=job_status,
                #     job_type=job_type,
                #     skills=skills,
                #     jd_pdf=jd_pdf
                # )

#                 new_job_post.notification = 'no'
#                 new_job_post.date_created = date.today()
#                 new_job_post.time_created = datetime.now().time()

#                 # Define an empty list to hold Notification instances
#                 notifications = []

#                 if ',' in joined_recruiters:
#                     recruiter_names_lst = joined_recruiters.split(',')
#                     for recruiter_name in recruiter_names_lst:
#                         notification_status = False
#                         notification = Notification(
#                             recruiter_name=recruiter_name.strip(),
#                             notification_status=notification_status
#                         )
#                         # Append each Notification instance to the notifications list
#                         notifications.append(notification)
#                 else:
#                     recruiter_name = joined_recruiters
#                     notification_status = False
#                     notification = Notification(
#                         recruiter_name=recruiter_name,
#                         notification_status=notification_status
#                     )
#                     # Append each Notification instance to the notifications list
#                     notifications.append(notification)

#                 # Add the new_job_post and all associated notifications to the session
#                 db.session.add(new_job_post)
#                 db.session.add_all(notifications)
#                 db.session.commit()

#                 # Retrieve the email addresses of the recruiters
#                 recruiter_emails = [recruiter.email for recruiter in User.query.filter(User.username.in_(recruiter_names),
#                                                                                        User.user_type == 'recruiter',
#                                                                                        User.is_active == True,
#                                                                                        User.is_verified == True)]
#                 for email in recruiter_emails:
#                     send_notification(email)

#                 # Return the job_id along with the success message
#                 return jsonify({"message": "Job posted successfully", "job_id": new_job_post.id}), 200
#             else:
#                 return jsonify({"error": "Invalid user type"}), 400
#         else:
#             return jsonify({"error": "Missing 'user_name' field in the request"}), 400

#     except KeyError as e:
#         return jsonify({"error": f"KeyError: {e}"}), 400

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/recruiter_job_posts/<int:user_id>', methods=['GET'])
# def recruiter_job_posts(user_id):
    
#     if not user_id:
#         return jsonify({"error": "Missing user_id parameter"})

#     # Validate user existence
#     recruiter = User.query.get(user_id)
#     if not recruiter:
#         return jsonify({"error": "Recruiter not found"})

#     recruiter_name = recruiter.name

#     # Filter unread notifications efficiently using the recruiter's ID
#     unread_notifications = Notification.query.filter(
#         Notification.id == recruiter.id,
#         Notification.notification_status == False
#     ).all()

#     # Filter active and on-hold job posts
#     active_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Active'
#     ).order_by(JobPost.id).all()

#     on_hold_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Hold'
#     ).order_by(JobPost.id).all()

#     # Update notification statuses after retrieving them
#     for notification in unread_notifications:
#         notification.notification_status = True
#     db.session.commit()

#     # Construct JSON response with relevant data
#     response_data = {
#         "count_notification_no": len(unread_notifications),
#         "user_name": recruiter_name,
#         "job_posts": [
#             {  # Include only necessary job post fields
#                 "id": job_post.id,
#                 "title": job_post.title,
#                 "description": job_post.description,
#                 "created_at": job_post.created_at.isoformat()  # Example for date formatting
#             }
#             for job_post in active_job_posts
#         ],
#         "job_posts_hold": [
#             {  # Include only necessary job post fields (optional)
#                 "id": job_post.id,
#                 "title": job_post.title,
#                 "description": job_post.description,
#                 "created_at": job_post.created_at.isoformat()  # Example for date formatting
#             }
#             for job_post in on_hold_job_posts
#         ],
#         "career_count_notification_no": 0  # Placeholder, implement career notification logic
#     }

#     # # Include optional parameters conditionally
#     # if url_for('add_candidate'):
#     #     response_data["redirect_url"] = url_for('add_candidate')
#     # if request.args.get('no_doc_message'):
#     #     response_data["no_doc_message"] = request.args.get('no_doc_message')

#     return jsonify(response_data)


# @app.route('/recruiter_job_posts/<int:user_id>', methods=['GET'])
# def recruiter_job_posts(user_id):
   
#     if not user_id:
#         return jsonify({"error": "Missing user_id parameter"})

#     # Validate user existence
#     recruiter = User.query.get(user_id)
#     if not recruiter:
#         return jsonify({"error": "Recruiter not found"})

#     recruiter_name = recruiter.name

#     # Filter unread notifications efficiently using the recruiter's ID
#     unread_notifications = Notification.query.filter(
#         Notification.recruiter_id == recruiter.id,
#         Notification.notification_status == False
#     ).all()

#     # Filter active and on-hold job posts
#     active_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Active'
#     ).order_by(JobPost.id).all()

#     on_hold_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Hold'
#     ).order_by(JobPost.id).all()

#     # Update notification statuses after retrieving them
#     for notification in unread_notifications:
#         notification.notification_status = True
#     db.session.commit()

#     # Construct JSON response with serialized job post data
#     response_data = {
#         "count_notification_no": len(unread_notifications),
#         "job_posts": [job_post.serialize() for job_post in active_job_posts],
#         "user_name": recruiter_name,
#         "job_posts_hold": [job_post.serialize() for job_post in on_hold_job_posts],
#         "redirect_url": url_for('add_candidate'),  # Optional, include if needed
#         "no_doc_message": request.args.get('no_doc_message'),  # Optional, include if needed
#         "career_count_notification_no": 0  # Placeholder, implement career notification logic
#     }

#     return jsonify(response_data)

# @app.route('/recruiter_job_posts', methods=['POST'])
# def recruiter_job_posts():
#     data = request.json
#     user_id = data.get('user_id')  # Using get() to avoid KeyError if 'user_id' is missing
#     if not user_id:
#         return jsonify({"error": "User ID is missing"}), 400

#     # Validate user existence
#     recruiter = User.query.get(user_id)
#     if not recruiter:
#         return jsonify({"error": "Recruiter not found"}), 404

#     recruiter_name = recruiter.name

#     # Filter unread notifications based on recruiter name
#     unread_notifications = Career_notification.query.filter(
#         Career_notification.recruiter_name == recruiter_name,
#         Career_notification.notification_status == False
#     ).all()

#     # Filter active and on-hold job posts
#     active_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
#         JobPost.job_status == 'Active'
#     ).order_by(JobPost.id).all()

#     on_hold_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
#         JobPost.job_status == 'Hold'
#     ).order_by(JobPost.id).all()

#     # Update notification statuses after retrieving them
#     for notification in unread_notifications:
#         notification.notification_status = True
#     db.session.commit()

#     # Construct JSON response
#     response_data = {
#         "count_notification_no": len(unread_notifications),
#         "job_posts": [job_post_to_dict(job_post) for job_post in active_job_posts],
#         "user_name": recruiter_name,
#         "job_posts_hold": [job_post_to_dict(job_post) for job_post in on_hold_job_posts],
#         "redirect_url": url_for('add_candidate'),  # Optional, include if needed
#         "no_doc_message": request.args.get('no_doc_message'),  # Optional, include if needed
#         "career_count_notification_no": 0  # Placeholder, implement career notification logic
#     }

#     return jsonify(response_data)

@app.route('/recruiter_job_posts', methods=['POST'])
def recruiter_job_posts():
    data = request.json
    user_id = data.get('user_id')  # Using get() to avoid KeyError if 'user_id' is missing
    if not user_id:
        return jsonify({"error": "User ID is missing"}), 400

    # Validate user existence
    recruiter = User.query.get(user_id)
    if not recruiter:
        return jsonify({"error": "Recruiter not found"}), 404

    recruiter_name = recruiter.name

    # Filter unread notifications based on recruiter name
    unread_notifications = Career_notification.query.filter(
        Career_notification.recruiter_name == recruiter_name,
        Career_notification.notification_status == False
    ).all()

    # Filter active and on-hold job posts
    active_job_posts = JobPost.query.filter(
        JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
        JobPost.job_status == 'Active'
    ).order_by(JobPost.id).all()

    on_hold_job_posts = JobPost.query.filter(
        JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
        JobPost.job_status == 'Hold'
    ).order_by(JobPost.id).all()

    # Update notification statuses after retrieving them
    for notification in unread_notifications:
        notification.notification_status = True
    db.session.commit()

    # Construct JSON response
    response_data = {
        "count_notification_no": len(unread_notifications),
        "job_posts": [job_post_to_dict(job_post) for job_post in active_job_posts],
        "user_name": recruiter_name,
        "job_posts_hold": [job_post_to_dict(job_post) for job_post in on_hold_job_posts],
        "redirect_url": url_for('add_candidate'),  # Optional, include if needed
        "no_doc_message": request.args.get('no_doc_message'),  # Optional, include if needed
        "career_count_notification_no": 0  # Placeholder, implement career notification logic
    }

    return jsonify(response_data)

# Helper function to convert JobPost object to dictionary
def job_post_to_dict(job_post):
    data_updated_date_str = job_post.data_updated_date.strftime('%Y-%m-%d') if job_post.data_updated_date else None
    data_updated_time_str = job_post.data_updated_time.strftime('%H:%M:%S') if job_post.data_updated_time else None

    return {
        "id": job_post.id,
        "client": job_post.client,
        "experience_min": job_post.experience_min,
        "experience_max": job_post.experience_max,
        "budget_min": job_post.budget_min,
        "budget_max": job_post.budget_max,
        "location": job_post.location,
        "shift_timings": job_post.shift_timings,
        "notice_period": job_post.notice_period,
        "role": job_post.role,
        "detailed_jd": job_post.detailed_jd,
        "mode": job_post.mode,
        "recruiter": job_post.recruiter,
        "management": job_post.management,
        "date_created": job_post.date_created.strftime('%Y-%m-%d'),
        "time_created": job_post.time_created.strftime('%H:%M:%S'),
        "job_status": job_post.job_status,
        "job_type": job_post.job_type,
        "skills": job_post.skills,
        "notification": job_post.notification,
        "data_updated_date": data_updated_date_str,
        "data_updated_time": data_updated_time_str
    }


from flask import jsonify

@app.route('/update_job_status/<int:job_id>', methods=['POST'])
def update_job_status(job_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error', "message": "User not found"}), 404

    user_type = user.user_type
    username = user.username

    # Retrieve the job post from the database based on the provided job_id
    job_post = JobPost.query.get(job_id)

    if job_post:
        try:
            # Extract the new job status from the form data
            new_job_status = data['new_job_status']

            # Update the job status
            job_post.job_status = new_job_status
            
            # Update data_updated_date and data_updated_time
            # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
            # job_post.data_updated_date = current_datetime.date()
            # job_post.data_updated_time = current_datetime.time()
            
            # Commit the changes to the database
            db.session.commit()

            # Return a JSON response indicating success
            return jsonify({'status': 'success', "message": "Job status updated successfully"})
        
        except KeyError:
            # If 'new_job_status' key is missing in form data
            return jsonify({'status': 'error', "message": "Missing 'new_job_status' in form data"})
        
        except Exception as e:
            # Handle other exceptions
            db.session.rollback()  # Rollback any changes made to the session
            return jsonify({'status': 'error', "message": str(e)})

    # If job_post is None (job not found)
    return jsonify({'status': 'error', "message": "Job post not found"})




import base64

@app.route('/view_all_jobs', methods=['POST'])
def view_all_jobs():
    data = request.json
    user_name = data['username']

    # Define case statements for conditional ordering
    conditional_order_date = case(
        (JobPost.data_updated_date != None, JobPost.data_updated_date),
        (JobPost.date_created != None, JobPost.date_created),
        else_=JobPost.date_created
    )

    conditional_order_time = case(
        (JobPost.data_updated_time != None, JobPost.data_updated_time),
        (JobPost.time_created != None, JobPost.time_created),
        else_=JobPost.time_created
    )

    # Retrieve all job posts with conditional ordering
    job_posts_active = JobPost.query.filter_by(job_status='Active')\
        .order_by(
            desc(conditional_order_date),
            desc(conditional_order_time),
            desc(JobPost.id)  # Ensure newer posts appear first if dates are equal
        )\
        .all()

    job_posts_hold = JobPost.query.filter_by(job_status='Hold')\
        .order_by(
            desc(conditional_order_date),
            desc(conditional_order_time),
            desc(JobPost.id)  # Ensure newer posts appear first if dates are equal
        )\
        .all()

    # Construct JSON response
    response_data = {
        "user_name": user_name,
        "job_posts_active": [
            {
                "id": job_post.id,
                "client": job_post.client,
                "role": job_post.role,
                "experience_min": job_post.experience_min,
                "experience_max": job_post.experience_max,
                "budget_min": job_post.budget_min,
                "budget_max": job_post.budget_max,
                "location": job_post.location,
                "shift_timings": job_post.shift_timings,
                "notice_period": job_post.notice_period,
                "detailed_jd": job_post.detailed_jd,
                "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
                "mode": job_post.mode,
                "recruiter": job_post.recruiter,
                "management": job_post.management,
                "job_status": job_post.job_status,
                "job_type": job_post.job_type,
                "contract_in_months": job_post.contract_in_months,
                "skills": job_post.skills,
                # "date_created": str(job_post.date_created),
                # "time_created": str(job_post.time_created),
                # "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
                # "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None,
                "date_created": job_post.date_created.isoformat() if job_post.date_created else None,
                "time_created": job_post.time_created.strftime('%H:%M:%S') if job_post.time_created else None,
                "data_updated_date": job_post.data_updated_date.isoformat() if job_post.data_updated_date else None,
                "data_updated_time": job_post.data_updated_time.strftime('%H:%M:%S') if job_post.data_updated_time else None,
                "jd_pdf_present": job_post.jd_pdf_present  # Added line
            }
            for job_post in job_posts_active
        ],
        "job_posts_hold": [
            {
                "id": job_post.id,
                "client": job_post.client,
                "role": job_post.role,
                "experience_min": job_post.experience_min,
                "experience_max": job_post.experience_max,
                "budget_min": job_post.budget_min,
                "budget_max": job_post.budget_max,
                "location": job_post.location,
                "shift_timings": job_post.shift_timings,
                "notice_period": job_post.notice_period,
                "detailed_jd": job_post.detailed_jd,
                "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
                "mode": job_post.mode,
                "recruiter": job_post.recruiter,
                "management": job_post.management,
                "job_status": job_post.job_status,
                "job_type": job_post.job_type,
                "contract_in_months": job_post.contract_in_months,
                "skills": job_post.skills,
                # "date_created": str(job_post.date_created),
                # "time_created": str(job_post.time_created),
                # "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
                # "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None,
                "date_created": job_post.date_created.isoformat() if job_post.date_created else None,
                "time_created": job_post.time_created.strftime('%H:%M:%S') if job_post.time_created else None,
                "data_updated_date": job_post.data_updated_date.isoformat() if job_post.data_updated_date else None,
                "data_updated_time": job_post.data_updated_time.strftime('%H:%M:%S') if job_post.data_updated_time else None,
                "jd_pdf_present": job_post.jd_pdf_present  # Added line
            }
            for job_post in job_posts_hold
        ]
    }

    # Return JSON response
    return jsonify(response_data)
    
# @app.route('/view_all_jobs', methods=['POST'])
# def view_all_jobs():
#     data = request.json
#     user_name = data['username']

#     # Define case statements for conditional ordering
#     conditional_order_date = case(
#         (JobPost.data_updated_date != None, JobPost.data_updated_date),
#         (JobPost.date_created != None, JobPost.date_created),
#         else_=JobPost.date_created
#     )

#     conditional_order_time = case(
#         (JobPost.data_updated_time != None, JobPost.data_updated_time),
#         (JobPost.time_created != None, JobPost.time_created),
#         else_=JobPost.time_created
#     )

#     # Retrieve all job posts with conditional ordering
#     job_posts_active = JobPost.query.filter_by(job_status='Active')\
#         .order_by(desc(conditional_order_date),
#                   desc(conditional_order_time),
#                   desc(JobPost.date_created),  # Ensure newer posts appear first
#                   desc(JobPost.time_created),
#                   desc(JobPost.id))\
#         .all()

#     job_posts_hold = JobPost.query.filter_by(job_status='Hold')\
#         .order_by(desc(conditional_order_date),
#                   desc(conditional_order_time),
#                   desc(JobPost.date_created),  # Ensure newer posts appear first
#                   desc(JobPost.time_created),
#                   desc(JobPost.id))\
#         .all()
#     # Construct JSON response
#     response_data = {
#         "user_name": user_name,
#         "job_posts_active": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
#                 "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None
#             }
#             for job_post in job_posts_active
#         ],
#         "job_posts_hold": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
#                 "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None
#             }
#             for job_post in job_posts_hold
#         ]
#     }

#     # Return JSON response
#     return jsonify(response_data)

# @app.route('/view_all_jobs', methods=['POST'])
# def view_all_jobs():
#     data = request.json
#     user_name = data['username']

#     # Define case statements for conditional ordering
#     conditional_order_date = case(
#         (JobPost.data_updated_date != None, JobPost.data_updated_date),
#         else_=JobPost.date_created
#     )

#     conditional_order_time = case(
#         (JobPost.data_updated_time != None, JobPost.data_updated_time),
#         else_=JobPost.time_created
#     )

#     # Retrieve all job posts with conditional ordering
#     job_posts_active = JobPost.query.filter_by(job_status='Active')\
#         .order_by(desc(conditional_order_date), desc(conditional_order_time), desc(JobPost.id))\
#         .all()

#     job_posts_hold = JobPost.query.filter_by(job_status='Hold')\
#         .order_by(desc(conditional_order_date), desc(conditional_order_time), desc(JobPost.id))\
#         .all()

#     # Construct JSON response
#     response_data = {
#         "user_name": user_name,
#         "job_posts_active": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
#                 "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None
#             }
#             for job_post in job_posts_active
#         ],
#         "job_posts_hold": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
#                 "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None
#             }
#             for job_post in job_posts_hold
#         ]
#     }

#     # Return JSON response
#     return jsonify(response_data)

# @app.route('/view_all_jobs', methods=['POST'])
# def view_all_jobs():
#     # Get data from JSON request
#     data = request.json

#     # Extract any parameters you need from the JSON data
#     user_name = data['username']

#     # Retrieve all job posts from the database
#     job_posts_active = JobPost.query.filter_by(job_status='Active').order_by(JobPost.id).all()
#     job_posts_hold = JobPost.query.filter_by(job_status='Hold').order_by(JobPost.id).all()

#     # Construct JSON response
#     response_data = {
#         "user_name": user_name,
#         "job_posts_active": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 # "jd_pdf": job_post.jd_pdf ,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date":str(job_post.data_updated_date),
#                 "data_updated_time":str(job_post.data_updated_time)
#                 # Include other attributes as needed
#             }
#             for job_post in job_posts_active
#         ],
#         "job_posts_hold": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 # "jd_pdf": job_post.jd_pdf ,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date":str(job_post.data_updated_date),
#                 "data_updated_time":str(job_post.data_updated_time)
#                 # Include other attributes as needed
#             }
#             for job_post in job_posts_hold
#         ]
#     }

#     # Return JSON response
#     return jsonify(response_data)



# def send_notification(recruiter_email):
#     msg = Message('New Job Posted', sender='ganesh.s@makonissoft.com', recipients=[recruiter_email])
#     msg.body = 'A new job has been posted. Check your dashboard for more details.'
#     mail.send(msg)

@app.route('/other_job_posts', methods=['GET'])
def other_job_posts():
    if 'user_id' in session and 'user_type' in session:
        if session['user_type'] == 'recruiter':
            # Retrieve the logged-in user's ID from the session
            user_id = session['user_id']

            # Retrieve the recruiter's name based on user ID
            recruiter_name = User.query.get(user_id).name

            job_posts = JobPost.query.filter(JobPost.recruiter != recruiter_name).distinct(JobPost.client).all()

            return render_template('other_job_posts.html', job_posts=job_posts)

    # Redirect or render an appropriate page if the conditions are not met
    return redirect(url_for('login'))

# @app.route('/recruiter_job_posts', methods=['GET'])
# def recruiter_job_posts():
#     no_doc_message = request.args.get('no_doc_message')
#     if 'user_id' in session and 'user_type' in session:
#         if session['user_type'] == 'recruiter':
#             # Retrieve the logged-in user's ID from the session
#             user_id = session['user_id']
#             user_name = session['user_name']
#             count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                               Notification.recruiter_name == user_name).count()
#             career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
#                                                               Career_notification.recruiter_name == user_name).count()
#             recruiter_name = User.query.get(user_id).name

#             job_posts = JobPost.query.filter(JobPost.recruiter.contains(recruiter_name),
#                                              JobPost.job_status == 'Active').order_by(JobPost.id).all()
#             job_posts_hold = JobPost.query.filter(JobPost.recruiter.contains(recruiter_name),
#                                                   JobPost.job_status == 'Hold').order_by(JobPost.id).all()

#             notifications = Notification.query.filter(Notification.recruiter_name.contains(recruiter_name)).all()

#             for notification in notifications:
#                 if notification.notification_status == False:
#                     notification.notification_status = True
#                     db.session.commit()

#             # for job_post in job_posts:
#             #     if job_post.notification == 'no':
#             #         job_post.notification = 'yes'
#             #         db.session.commit()

#             return render_template('recruiter_job_posts.html', count_notification_no=count_notification_no,
#                                    job_posts=job_posts, user_name=user_name, job_posts_hold=job_posts_hold,
#                                    redirect_url=url_for('add_candidate'), recruiter_job_posts=recruiter_job_posts,
#                                    no_doc_message=no_doc_message, career_count_notification_no=career_count_notification_no)

#     return redirect(url_for('login'))
    

import base64
import io
import magic
from flask import Flask, send_file, request

@app.route('/view_resume/<int:candidate_id>', methods=['GET'])
def view_resume(candidate_id):
    # Retrieve the resume data from the database using SQLAlchemy
    candidate = Candidate.query.filter_by(id=candidate_id).first()
    if not candidate:
        return 'Candidate not found', 404

    if candidate.resume is None:
        return 'Resume not found for this candidate', 404

    try:
        # If the resume data is a base64 encoded string, decode it
        if isinstance(candidate.resume, str):
            resume_binary = base64.b64decode(candidate.resume)
        elif isinstance(candidate.resume, bytes):
            resume_binary = candidate.resume
        else:
            return 'Invalid resume format', 400

        # Determine the mimetype based on the file content
        is_pdf = resume_binary.startswith(b"%PDF")
        mimetype = 'application/pdf' if is_pdf else 'application/msword'

        # Send the file as a response
        return send_file(
            io.BytesIO(resume_binary),
            mimetype=mimetype,
            as_attachment=False
        )
    except Exception as e:
        return f'Error processing resume: {str(e)}', 500

# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found'

#     if candidate.resume is None:
#         return 'Resume not found for this candidate', 404

#     if isinstance(candidate.resume, bytes):
#         # If the resume data is already in bytes format
#         resume_binary = candidate.resume
#     else:
#         # If the resume data is base64 encoded, decode it
#         decoded_resume = base64.b64decode(candidate.resume)
#         resume_binary = decoded_resume

#     # Determine the mimetype based on the file content
#     is_pdf = resume_binary.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'

#     if "==" not in str(candidate.resume):  # assuming you are checking for base64 here
#         # If the resume data is not base64 encoded, send it directly
#         return send_file(
#             io.BytesIO(resume_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         # If the resume data is base64 encoded, create a file-like object from the decoded data and send it
#         resume_file = io.BytesIO(resume_binary)
#         return send_file(
#             resume_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )

# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found', 404

#     print("Candidate resume:", candidate.resume)  # Check the resume data
#     print("Resume type:", type(candidate.resume))  # Check the type of resume data

#     if isinstance(candidate.resume, bytes):
#         # If the resume data is already in bytes format
#         resume_binary = candidate.resume
#     else:
#         # If the resume data is base64 encoded, decode it
#         resume_binary = base64.b64decode(candidate.resume)

#     print("Resume binary:", resume_binary)  # Check the binary data

#     # Determine the mimetype based on the file content
#     is_pdf = resume_binary.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'

#     if "==" not in str(candidate.resume):  # assuming you are checking for base64 here
#         # If the resume data is not base64 encoded, send it directly
#         return send_file(
#             io.BytesIO(resume_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         # If the resume data is base64 encoded, create a file-like object from the decoded data and send it
#         resume_file = io.BytesIO(resume_binary)
#         return send_file(
#             resume_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )


# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found', 404

#     if isinstance(candidate.resume, bytes):
#         # If the resume data is already in bytes format
#         resume_binary = candidate.resume
#     else:
#         # If the resume data is base64 encoded, decode it
#         resume_binary = base64.b64decode(candidate.resume)

#     # Determine the mimetype based on the file content
#     is_pdf = resume_binary.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'

#     if "==" not in str(candidate.resume):  # assuming you are checking for base64 here
#         # If the resume data is not base64 encoded, send it directly
#         return send_file(
#             io.BytesIO(resume_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         # If the resume data is base64 encoded, create a file-like object from the decoded data and send it
#         resume_file = io.BytesIO(resume_binary)
#         return send_file(
#             resume_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )

# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found'
#     # Decode the base64 encoded resume data
#     # print("candidate.resume",candidate.resume.tobytes())
#     if "==" not in str(candidate.resume.tobytes()):
#         if request.args.get('decode') == 'base64':
#             # Decode the base64 encoded resume data
#             decoded_resume = base64.b64decode(candidate.resume)
#             resume_binary = decoded_resume
#         else:
#             # Retrieve the resume binary data from the database
#             resume_binary = candidate.resume.tobytes()  # Convert memoryview to bytes

#         # Determine the mimetype based on the file content
#         is_pdf = resume_binary.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             io.BytesIO(resume_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         decoded_resume = base64.b64decode(candidate.resume)
#         # Create a file-like object (BytesIO) from the decoded resume data
#         resume_file = io.BytesIO(decoded_resume)
#         # Determine the mimetype based on the file content
#         is_pdf = decoded_resume.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             resume_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )



# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found', 404

#     # Check if the request specifies to retrieve the resume data directly from the database
#     if request.args.get('decode') == 'base64':
#         # Decode the base64 encoded resume data
#         decoded_resume = base64.b64decode(candidate.resume)
#         resume_binary = decoded_resume
#     else:
#         # Retrieve the resume binary data from the database
#         resume_binary = candidate.resume.tobytes()  # Convert memoryview to bytes

#     # Determine the mimetype based on the file content
#     is_pdf = resume_binary.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'

#     # Send the file as a response
#     return send_file(
#         io.BytesIO(resume_binary),
#         mimetype=mimetype,
#         as_attachment=False
#     )


 
# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found'
#     # Decode the base64 encoded resume data
        
#     decoded_resume = base64.b64decode(candidate.resume)
#     # Create a file-like object (BytesIO) from the decoded resume data
#     resume_file = io.BytesIO(decoded_resume)
#     # Determine the mimetype based on the file content
#     is_pdf = decoded_resume.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'
 
#     # Send the file as a response
#     return send_file(
#         resume_file,
#         mimetype=mimetype,
#         as_attachment=False
#     )
###############################################################################################
# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data=request.json
#     # if not data or 'image' not in data:
#     #     return jsonify({'error': 'No image data provided'}), 400
    
#     filename=data['file_name']
#     image_file=data['image_file']

#     user = User.query.filter_by(id=user_id).first()
#     # if not user:
#     #     return jsonify({'error': 'User not found'}), 400
    
#     user.filename = filename
#     user.image_file=image_file
#     db.session.commit()

#     return jsonify({'message': 'Image updated successfully'}), 200
#######################################################################################


# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data = request.json
#     if not data:
#         return jsonify({'error': 'Invalid JSON data provided'}), 400

#     image_content = data.get('image')
#     file_name = data.get('filename')

#     # if not image_content or not file_name:
#     #     return jsonify({'error': 'Image content or filename missing in the request'}), 400

#     # Find the user by user_id
#     user = User.query.get(user_id)
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     # Update user's filename and image content
#     user.filename = file_name
#     user.image_file = image_content

#     # Commit changes to the database
#     db.session.commit()

#     return jsonify({'message': 'Image updated successfully'}), 200

# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data = request.json
#     print("\n\n\n\n\n")
#     print("Data :",data)
#     image_content = data['image']
#     file_name = data['filename']
#     # Find the user by user_id
#     user = User.query.get(user_id)
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     # Update user's filename and image content
#     user.filename = file_name
#     user.image_file = image_content
#     # db.session.commit()

#     # Commit changes to the database
#     try:
#         db.session.commit()
#     except:
#         print("Failed to Upload !!")

#     return jsonify({'message': 'Image updated successfully'}), 200


# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data = request.form

#     # Extract file name and image content
#     image_content = data['image']
#     filename = data['filename']  # Retrieve file object
     
#     # Find the user by user_id
#     user = User.query.get(user_id)
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     # Update user's filename and image content
#     user.filename = filename
#     user.image_file = image_content

#     # Commit changes to the database
#     db.session.commit()

#     return jsonify({'message': 'Image updated successfully'}), 200

# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     try:
#         # Extract file from request
#         data=request.json
#         image_file = data['image']
#         filename = data['filename']
#         image_delete_status=data['image_delete_status']
        
#         # Find the user by user_id
#         user = User.query.get(user_id)
#         if not user:
#             return jsonify({'error': 'User not found'}), 404

#         # Update user's filename and image content
#         user.filename = filename
#         user.image_file = image_file  # Store image content as binary data
#         user.image_deleted=image_delete_status
#         # Commit changes to the database
#         db.session.commit()

#         return jsonify({'message': 'Image updated successfully'}), 200
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

@app.route('/upload_user_image/<int:user_id>', methods=['POST'])
def upload_user_image(user_id):
    try:
        # Extract data from the request
        data = request.json
        image_base64 = data['image']
        filename = data['filename']
        image_delete_status = data['image_delete_status']

        # Decode the base64 image
        image_binary = base64.b64decode(image_base64)

        # Find the user by user_id
        user = User.query.get(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404

        # Update user's filename and image content
        user.filename = filename
        user.image_file = image_binary  # Store image content as binary data
        user.image_deleted = image_delete_status

        # Commit changes to the database
        db.session.commit()

        return jsonify({'message': 'Image updated successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

#################################################################################################

import base64
import io
# @app.route('/image_status/<int:user_id>', methods=['GET'])
# def image_status(user_id):
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'error': 'Image not found'}), 404

#     return jsonify({'message': user.image_file}), 200


import io
import base64
import mimetypes

# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'message': 'Image not found'}), 400
    
#     # Decode the bytea image data
#     image_data = user.image_file
#     image_file = base64.b64decode(image_data)
    
#     # Determine the MIME type dynamically
#     mime_type, _ = mimetypes.guess_type(user.image_filename)  # Assuming user.image_filename holds the filename
    
#     # Default to 'application/octet-stream' if MIME type couldn't be guessed
#     if not mime_type:
#         mime_type = 'application/octet-stream'
    
#     # Send the file as a response
#     return send_file(
#         io.BytesIO(image_file),
#         mimetype=mime_type,
#         as_attachment=False
#     )

import io
import base64
from PIL import Image
import mimetypes

from flask import send_file, jsonify

# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'message': 'Image not found'}), 400
    
#     # Decode the bytea image data
#     image_data = base64.b64decode(user.image_file)
    
#     # Determine the MIME type
#     image = Image.open(io.BytesIO(image_data))
#     mime_type = Image.MIME.get(image.format)
    
#     # Get the original filename from the user object
#     filename = user.filename
    
#     # Send the file as inline content with filename
#     return send_file(
#         io.BytesIO(image_data),
#         mimetype=mime_type,
#         as_attachment=False,
#         attachment_filename=filename  # Set the filename in the response headers
#     )

# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'message': 'Image not found'}), 400
    
#     # Decode the bytea image data
#     image_data = base64.b64decode(user.image_file)
    
#     # Determine the MIME type
#     image = Image.open(io.BytesIO(image_data))
#     mime_type = Image.MIME.get(image.format)
    
#     # Send the file as inline content
#     return send_file(
#         io.BytesIO(image_data),
#         mimetype=mime_type,
#         as_attachment=False
#     )


# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'message': 'Image not found'}), 400
    
#     # Decode the bytea image data
#     image_data = base64.b64decode(user.image_file)
    
#     # Determine the MIME type
#     image = Image.open(io.BytesIO(image_data))
#     mime_type = Image.MIME.get(image.format)
    
#     # Send the file as a response
#     return send_file(
#         io.BytesIO(image_data),
#         mimetype=mime_type,
#         as_attachment=False
#     )

import io
import base64
import imghdr
@app.route('/user_image/<int:user_id>', methods=['GET'])
def user_image(user_id):
    # Retrieve the user data from the database
    user = User.query.filter_by(id=user_id).first()
    if not user or not user.image_file:
        return jsonify({'message': 'Image not found'}), 404
    
    # Get the image data from the user object
    image_data = user.image_file
    
    # Determine the image format dynamically
    image_format = imghdr.what(None, h=image_data)
    if not image_format:
        return jsonify({'error': 'Unknown image format'}), 500
    
    # Determine the MIME type based on the image format
    if image_format == 'jpeg':
        mimetype = 'image/jpeg'
    elif image_format == 'png':
        mimetype = 'image/png'
    else:
        # Handle other image formats as needed
        return jsonify({'error': 'Unsupported image format'}), 500
    
    # Send the file as a response
    return send_file(
        io.BytesIO(image_data),
        mimetype=mimetype,
        as_attachment=False
    )

# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'message': 'Image not found'}),400
    
#     # Decode the bytea image data
#     image_data = user.image_file
    
#     image_file = base64.b64decode(image_data)
    
#     # Send the file as a response
#     return send_file(
#         io.BytesIO(image_file),
#          mimetype = 'image/jpeg',
#         as_attachment=False
#     )



# import base64
# import io
# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'error': 'Image not found'}), 404
    
#     # Decode the bytea image data
#     image_data = user.image_file

#     # Create a file-like object (BytesIO) from the image data
#     image_file = io.BytesIO(image_data)

#     # Determine the mimetype based on the file content
#     if image_data.startswith(b"\x89PNG"):
#         mimetype = 'image/png'
#     elif image_data.startswith(b"\xff\xd8\xff"):
#         mimetype = 'image/jpeg'
#     elif image_data.startswith(b"\xff\xd8\xff\xe0") and image_data[6:10] in (b"JFIF", b"Exif"):
#         mimetype = 'image/jpeg'
#     else:
#         return jsonify({'error': 'Unsupported image format'}), 400

#     # Send the file as a response
#     return send_file(
#         image_file,
#         mimetype=mimetype,
#         as_attachment=False
#     )

@app.route('/delete_user_image/<int:user_id>', methods=['POST'])
def delete_user_image(user_id):
    data = request.json
    profile_image = data['profileImage']
    image_delete_status=data['image_delete_status']
    if not profile_image:
        return jsonify({"error": "Profile image must be specified"}), 400

    user = User.query.filter_by(id=user_id).first()
    
    if not user:
        return jsonify({"error": "User not found"}), 404

    user.image_file = None
    user.filename = None
    user.image_deleted = image_delete_status
    db.session.commit()

    return jsonify({"message": "Image file deleted successfully"}), 200


# @app.route('/delete_user_image/<int:user_id>', methods=['POST'])
# def delete_user_image(user_id):
#     data = request.json
#     image_file = data.get('image_file')
#     if not image_file:
#         return jsonify({"error": "Image file must be specified"}), 400

#     user = User.query.filter_by(id=user_id,).first()
    
#     if not user:
#         return jsonify({"error": "User not found"}), 400
    
#     # if user.image_file != image_file:
#     #     return jsonify({"error": "Image file does not match the user's image"}), 400

#     user.image_file = None
#     user.filename=None
#     db.session.commit()

#     return jsonify({"message": "Image file deleted successfully"}), 200
    

@app.route('/viewfull_jd/<int:id>')
def viewfull_jd(id):
    user_type = session['user_type']
    job_post = JobPost.query.get(id)
    return render_template('viewfull_jd.html', job_post=job_post,user_type=user_type)

@app.route('/add_candidate_view')
def add_candidate_view():
    user_id = session['user_id']
    user_type = session['user_type']
    user_name = session['user_name']

    if user_type == 'recruiter':
        recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
        if recruiter:
            candidates = Candidate.query.filter_by(
                recruiter=recruiter.name).all()  # Filter candidates by recruiter's name
            # data = json.dumps(candidates, sort_keys=False)
            results = db.session.query(JobPost.client, JobPost.recruiter).filter(
                JobPost.recruiter.contains(user_name)).all()
            client_names = sorted(list(set([result.client for result in results])))
            count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                              Notification.recruiter_name == user_name).count()
            return render_template('add_candidate_view.html', user=recruiter, user_type=user_type, user_name=user_name,
                                   candidates=candidates, count_notification_no=count_notification_no,
                                   client_names=client_names)
    elif user_type == 'management':
        users = User.query.all()
        candidates = Candidate.query.all()
        JobsPosted = JobPost.query.all()
        clients = db.session.query(JobPost.client).all()
        client_names = list(set([client[0] for client in clients]))

        return render_template('add_candidate_view.html', users=users, user_type=user_type, user_name=user_name,
                               JobsPosted=JobsPosted, client_names=client_names)

import os
import shutil
from flask import Flask, request, send_file, redirect, url_for
from zipfile import ZipFile

@app.route('/download_resumes')
def download_resumes():
    candidate_ids = request.args.getlist('candidate_ids')
    
    # Create a temporary directory to store resume files
    temp_dir = 'temp_resumes'
    os.makedirs(temp_dir, exist_ok=True)
    
    resume_paths = []

    for candidate_id in candidate_ids:
        candidate = Candidate.query.get(candidate_id)
        if candidate is None or candidate.resume is None:
            continue
        
        resume_file = io.BytesIO(candidate.resume)
        is_pdf = resume_file.getvalue().startswith(b"%PDF")
        if is_pdf : 
            resume_filename = f"{candidate.name}_resume.pdf" 
            resume_path = os.path.join(temp_dir, resume_filename)
            with open(resume_path, 'wb') as file:
                file.write(candidate.resume)
            
            resume_paths.append(resume_path)
        else:
            resume_filename = f"{candidate.name}_resume.docx" 
            resume_path = os.path.join(temp_dir, resume_filename)
            with open(resume_path, 'wb') as file:
                file.write(candidate.resume)
            
            resume_paths.append(resume_path)

    # Create a zip file containing all resume files
    zip_filename = 'resumes.zip'
    with ZipFile(zip_filename, 'w') as zipf:
        for resume_path in resume_paths:
            zipf.write(resume_path, os.path.basename(resume_path))
    
    # Clean up temporary directory
    shutil.rmtree(temp_dir)
    
    # Send the zip file for download
    return send_file(zip_filename, as_attachment=True)


@app.route('/assign_job/<int:job_id>', methods=['POST'])
def assign_job(job_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error',"message": "User not found"})

    user_type = user.user_type
    username = user.username
    job_post = JobPost.query.get(job_id)  # Retrieve the job post by its ID

    if not job_post:
        return jsonify({'status': 'success',"message": "Job not found"})

    current_recruiters = job_post.recruiter.split(', ') if job_post.recruiter else []

    if request.method == 'POST':
        new_recruiter_names = data.get('recruiters', [])
        
        # Modification: Remove duplicate recruiters by combining lists and converting to a set
        updated_recruiter_names = list(set(current_recruiters + new_recruiter_names))
        
        # Join the recruiter names into a single string
        joined_recruiters = ', '.join(updated_recruiter_names)
        job_post.recruiter = joined_recruiters
        
        # Update data_updated_date and data_updated_time
        # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        # job_post.data_updated_date = current_datetime.date()
        # job_post.data_updated_time = current_datetime.time()

        db.session.commit()

        # Send notification emails to the newly assigned recruiters
        new_recruiter_emails = [recruiter.email for recruiter in
                                User.query.filter(User.name.in_(new_recruiter_names),
                                                  User.user_type == 'recruiter')]
        for email in new_recruiter_emails:
            send_notification(email)

        # Define an empty list to hold Notification instances
        notifications = []

        for recruiter_name in updated_recruiter_names:
            if recruiter_name.strip() in new_recruiter_names:
                notification_status = False  # Set the initial status
                notification = Notification(
                    job_post_id=job_post.id,
                    recruiter_name=recruiter_name.strip(),
                    notification_status=notification_status
                )
                # Append each Notification instance to the notifications list
                notifications.append(notification)

        # Commit the notifications to the database session
        db.session.add_all(notifications)
        db.session.commit()

        return jsonify({'status': 'success',"message": "Job re-assigned successfully"}), 200

    recruiter_names = [recruiter.name for recruiter in User.query.filter_by(user_type='recruiter')]
    return jsonify({
        "user_name": username,
        "job_post": job_post.serialize(),
        "current_recruiters": current_recruiters,
        "recruiters": recruiter_names
    })

# @app.route('/assign_job/<int:job_id>', methods=['POST'])
# def assign_job(job_id):
#     data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()

#     if not user:
#         return jsonify({"error_message": "User not found"}), 404

#     user_type = user.user_type
#     username = user.username
#     job_post = JobPost.query.get(job_id)  # Retrieve the job post by its ID

#     if not job_post:
#         return jsonify({"error_message": "Job not found"}), 404

#     current_recruiters = job_post.recruiter.split(', ') if job_post.recruiter else []

#     if request.method == 'POST':
#         new_recruiter_names = data.get('recruiters', [])
        
#         # Modification: Remove duplicate recruiters by combining lists and converting to a set
#         updated_recruiter_names = list(set(current_recruiters + new_recruiter_names))
        
#         # Join the recruiter names into a single string
#         joined_recruiters = ', '.join(updated_recruiter_names)
#         job_post.recruiter = joined_recruiters
#         db.session.commit()

#         # Send notification emails to the newly assigned recruiters
#         new_recruiter_emails = [recruiter.email for recruiter in
#                                 User.query.filter(User.name.in_(new_recruiter_names),
#                                                   User.user_type == 'recruiter')]
#         for email in new_recruiter_emails:
#             send_notification(email)

#         # Define an empty list to hold Notification instances
#         notifications = []

#         for recruiter_name in updated_recruiter_names:
#             if recruiter_name.strip() in new_recruiter_names:
#                 notification_status = False  # Set the initial status
#                 notification = Notification(
#                     job_post_id=job_post.id,
#                     recruiter_name=recruiter_name.strip(),
#                     notification_status=notification_status
#                 )
#                 # Append each Notification instance to the notifications list
#                 notifications.append(notification)

#         # Commit the notifications to the database session
#         db.session.add_all(notifications)
#         db.session.commit()

#         return jsonify({"message": "Job re-assigned successfully"}), 200

#     recruiter_names = [recruiter.name for recruiter in User.query.filter_by(user_type='recruiter')]
#     return jsonify({
#         "user_name": username,
#         "job_post": job_post.serialize(),
#         "current_recruiters": current_recruiters,
#         "recruiters": recruiter_names
#     })


# @app.route('/assign_job/<int:job_id>', methods=['POST'])
# def assign_job(job_id):
#     data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
    
#     if not user:
#         return jsonify({"error_message": "User not found"}), 404
    
#     user_type = user.user_type
#     username = user.username
#     job_post = JobPost.query.get(job_id)  # Retrieve the job post by its ID

#     if not job_post:
#         return jsonify({"error_message": "Job not found"}), 404

#     current_recruiters = job_post.recruiter.split(', ') if job_post.recruiter else []

#     if request.method == 'POST':
#         new_recruiter_names = data.get('recruiters', [])
#         all_recruiter_names = current_recruiters + new_recruiter_names
#         joined_recruiters = ', '.join(all_recruiter_names)
#         job_post.recruiter = joined_recruiters
#         db.session.commit()

#         # Send notification emails to the newly assigned recruiters
#         new_recruiter_emails = [recruiter.email for recruiter in
#                                 User.query.filter(User.name.in_(new_recruiter_names),
#                                                     User.user_type == 'recruiter')]
#         for email in new_recruiter_emails:
#             send_notification(email)

#         # Define an empty list to hold Notification instances
#         notifications = []

#         if ',' in joined_recruiters:
#             recruiter_names_lst = joined_recruiters.split(',')
#             for recruiter_name in recruiter_names_lst:
#                 if recruiter_name.strip() in new_recruiter_names:
#                     notification_status = False  # Set the initial status
#                     notification = Notification(
#                         recruiter_name=recruiter_name.strip(),
#                         notification_status=notification_status
#                     )
#                     # Append each Notification instance to the notifications list
#                     notifications.append(notification)
#         else:
#             recruiter_name = joined_recruiters
#             if recruiter_name in new_recruiter_names:
#                 notification_status = False  # Set the initial status
#                 notification = Notification(
#                     recruiter_name=recruiter_name,
#                     notification_status=notification_status
#                 )
#                 # Append each Notification instance to the notifications list
#                 notifications.append(notification)

#         # Commit the notifications to the database session
#         db.session.add_all(notifications)
#         db.session.commit()
#         return jsonify({"message": "Job re-assigned successfully"}), 200

#     recruiter_names = [recruiter.name for recruiter in User.query.filter_by(user_type='recruiter')]
#     return jsonify({
#         "user_name": username,
#         "job_post": job_post.serialize(),
#         "current_recruiters": current_recruiters,
#         "recruiters": recruiter_names
#     })



@app.route('/assign_candidate', methods=['POST'])
def assign_candidate():
    assignment_message = request.args.get('assignment_message')
    if 'user_id' in session and 'user_type' in session and session['user_type'] == 'management':
        user_name = session['user_name']
        recruiters = User.query.filter_by(user_type='recruiter').all()

        selected_recruiter_id = request.json.get('selected_recruiter_id')
        selected_candidate_id = request.json.get('selected_candidate_id')
        assign_recruiter_id = request.json.get('assign_recruiter_id')
        assign_candidate_id = request.json.get('assign_candidate_id')

        selected_recruiter = None
        selected_candidate = None
        assigned_recruiter = None

        if selected_recruiter_id:
            selected_recruiter = User.query.get(selected_recruiter_id)

        if selected_candidate_id:
            selected_candidate = Candidate.query.get(selected_candidate_id)

        if assign_recruiter_id:
            assigned_recruiter = User.query.get(assign_recruiter_id)

        if request.method == 'POST':
            selected_candidate_ids = request.json.get('selected_candidate_ids')
            if assigned_recruiter and selected_candidate_ids:
                for candidate_id in selected_candidate_ids:
                    candidate = Candidate.query.get(candidate_id)
                    if candidate:
                        candidate.recruiter = assigned_recruiter.name
                db.session.commit()
                return jsonify({"message": "Candidates Assigned Successfully"}), 200

            if selected_candidate_id:
                selected_candidate = Candidate.query.get(selected_candidate_id)

        candidates = []
        if selected_recruiter:
            candidates = Candidate.query.filter(
                Candidate.recruiter == selected_recruiter.name,
                Candidate.status.in_(['None', "SCREENING","L1 - SCHEDULED" ,"L1 - SELECTED", 'L1 - FEEDBACK', 'L1 - RESCHEDULE',"L2 - SCHEDULED" ,"L2 - SELECTED",
                                      'L2 - FEEDBACK', 'L2 - RESCHEDULE', 'HOLD(POSITION)', 'CANDIDATE HOLD', 'OFFERED',
                                      "L2 - SELECTED"])
            ).all()
        return jsonify({
            "recruiters": [recruiter.serialize() for recruiter in recruiters],
            "candidates": [candidate.serialize() for candidate in candidates],
            "selected_recruiter": selected_recruiter.serialize() if selected_recruiter else None,
            "selected_candidate": selected_candidate.serialize() if selected_candidate else None,
            "assigned_recruiter": assigned_recruiter.serialize() if assigned_recruiter else None,
            "assignment_message": assignment_message,
            "user_name": user_name
        })

    return jsonify({"error_message": "Unauthorized: You must log in as management user to access this page"}), 401



from flask import jsonify

@app.route('/disable_user', methods=['POST'])
def disable_user():
    data = request.json
    user_id = data['user_id']
    user_status = data['user_status']
    user_name = data['user_name']

    if user_id is None or user_status is None or user_name is None:
        return jsonify({'message': 'User ID, user status, and user name are required'}), 400

    # Find the user making the request
    request_user = User.query.get(user_id)

    if request_user is None or request_user.user_type != 'management':
        return jsonify({'message': 'Unauthorized access'}), 403

    # Find the user to be updated
    user = User.query.filter_by(username=user_name).first()

    if user is None:
        return jsonify({'message': 'User not found'}), 404

    # Change verification status for the user
    user.is_verified = user_status

    # If the user is a recruiter, change verification status for management user with the same username
    if user.user_type == 'recruiter':
        management_user = User.query.filter_by(username=user_name, user_type='management').first()
        if management_user:
            management_user.is_verified = user_status

    try:
        db.session.commit()
        # Return different messages based on user_type
        if user.user_type == 'management':
            if user_status:
                return jsonify({'message': 'Verification status updated for management account'}), 200
            else:
                return jsonify({'message': 'Verification status updated to unverified for management account'}), 200
        elif user.user_type == 'recruiter':
            if user_status:
                return jsonify({'message': 'Verification status updated for recruiter account'}), 200
            else:
                return jsonify({'message': 'Verification status updated to unverified for recruiter account'}), 200
    except Exception as e:
        # Log the exception or return an error message
        db.session.rollback()
        return jsonify({'message': 'Failed to update verification status'}), 500

@app.route('/active_users', methods=['POST'])
def update_user_status():
    data = request.json
    username = data.get('user_name')
    new_status = data.get('new_status')

    try:
        user = User.query.filter_by(username=username).first()
        if user:
            # user.is_verified = new_status
            # db.session.commit()

            # Fetch updated active users list
            active_users_manager = User.query.filter_by(user_type='management').all()
            active_users_manager = sorted(active_users_manager, key=lambda user: user.id)
            active_users_recruiter = User.query.filter_by(user_type='recruiter').all()
            active_users_recruiter = sorted(active_users_recruiter, key=lambda user: user.id)

            return jsonify({
                "message": "User status updated successfully",
                "username": username,
                "active_users_manager": [user.serialize() for user in active_users_manager],
                "active_users_recruiter": [user.serialize() for user in active_users_recruiter]
            })
        else:
            return jsonify({"message": "User not found"}), 404
    except Exception as e:
        db.session.rollback()  # Rollback changes in case of error
        return jsonify({"message": "Error updating user status", "error": str(e)}), 500
        

from flask import jsonify

@app.route('/deactivate_user', methods=['POST'])
def deactivate_user():
    data = request.json
    management_user_id = data.get('user_id')
    username = data.get('user_name')
    user_status = data.get('user_status')

    if management_user_id:
        # Find the management user
        management_user = User.query.get(management_user_id)

        if management_user and management_user.user_type == 'management':
            messages = []

            if username:
                # Find the target user by username
                target_user = User.query.filter_by(username=username).first()

                if target_user:
                    if target_user.user_type == 'management' or target_user.user_type == 'recruiter':
                        # Update user account status
                        target_user.is_active = user_status
                        db.session.commit()

                        # Determine the message based on user_status
                        if user_status:
                            messages.append(f'{target_user.user_type.capitalize()} account {username} has been successfully activated.')
                        else:
                            messages.append(f'{target_user.user_type.capitalize()} account {username} has been successfully deactivated.')
                    else:
                        messages.append('User is neither a management nor a recruiter account.')
                else:
                    messages.append('User not found.')

            if messages:
                # Get all user records
                all_users = User.query.all()
                user_data = [{'id': user.id, 'username': user.username, 'is_active': user.is_active} for user in all_users]
                return jsonify({'messages': messages, 'users': user_data})
            else:
                return jsonify({'message': 'No valid username provided.'})
        else:
            return jsonify({'message': 'Management user not found or not a management user.'})
    else:
        return jsonify({'message': 'Management user_id is required.'})


# @app.route('/deactivate_user', methods=['POST'])
# def deactivate_user():
#     data = request.json
#     management_user_id = data.get('user_id')
#     username = data.get('user_name')
#     user_status=data.get('user_status')

#     if management_user_id:
#         # Find the management user
#         management_user = User.query.get(management_user_id)

#         if management_user and management_user.user_type == 'management':
#             messages = []

#             if username:
#                 # Find the target user by username
#                 target_user = User.query.filter_by(username=username).first()

#                 if target_user:
#                     if target_user.user_type == 'management':
#                         # Deactivate management account
#                         target_user.is_active = user_status
#                         db.session.commit()
#                         messages.append(f'Management account {username} has been successfully deactivated.')
#                     elif target_user.user_type == 'recruiter':
#                         # Deactivate recruiter account
#                         target_user.is_active = user_status
#                         db.session.commit()
#                         messages.append(f'Recruiter account {username} has been successfully deactivated.')
#                     else:
#                         messages.append('User is neither a management nor a recruiter account.')
#                 else:
#                     messages.append('User not found.')

#             if messages:
#                 # Get all user records
#                 all_users = User.query.all()
#                 user_data = [{'id': user.id, 'username': user.username, 'is_active': user.is_active} for user in all_users]
#                 return jsonify({'messages': messages, 'users': user_data})
#             else:
#                 return jsonify({'message': 'No valid username provided.'})
#         else:
#             return jsonify({'message': 'Management user not found or not a management user.'})
#     else:
#         return jsonify({'message': 'Management user_id is required.'})

# @app.route('/deactivate_user', methods=['POST'])
# def deactivate_user():
#     data = request.json
#     management_user_id = data.get('user_id')
#     recruiter_username = data.get('user_name')
#     user_status = data.get('user_status')

#     if management_user_id and recruiter_username:  
#         # Find the management user
#         management_user = User.query.get(management_user_id)

#         if management_user and management_user.user_type == 'management':
#             # Find the recruiter user by username
#             recruiter_user = User.query.filter_by(username=recruiter_username, user_type='recruiter').first()

#             if recruiter_user:
#                 # Change active status for the recruiter user
#                 recruiter_user.is_active = user_status
#                 db.session.commit()

#                 # Get all user records
#                 all_users = User.query.all()
                
#                 # Construct response data
#                 user_data = [{'id': user.id, 'username': user.username, 'is_active': user.is_verified} for user in all_users]

#                 if user_status:
#                     return jsonify({'message': f'Recruiter account {recruiter_username} has been successfully activated.', 'users': user_data})
#                 else:
#                     return jsonify({'message': f'Recruiter account {recruiter_username} has been successfully deactivated.', 'users': user_data})
#             else:
#                 return jsonify({'message': 'Recruiter user not found or not a recruiter user'})
#         else:
#             return jsonify({'message': 'Management user not found or not a management user'})
#     else:
#         return jsonify({'message': 'Both management_user_id and recruiter_username are required'})


        
# @app.route('/verify_checkbox', methods=['POST'])
# def verify_checkbox():
#     data = request.json
#     user_id = data.get('userId')
#     checked = data.get('checked')
#     user = User.query.get(user_id)
#     user.is_verified = checked
#     db.session.commit()
#     return redirect(url_for('active_users'))

import hashlib
from flask_mail import Message


@app.route('/change_password', methods=['POST'])
def change_password():
    data = request.json

    if not data:
        return jsonify({'status': 'error', 'message': 'No JSON data provided.'})

    user_id = data.get('user_id')
    username = data.get('username')
    old_password = data.get('old_password')
    new_password = data.get('new_password')
    confirm_password = data.get('confirm_password')

    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error', 'message': 'User not found.'})

    if username != user.username:
        return jsonify({'status': 'error', 'message': 'Logged in user does not match the provided username.'})

    hashed_old_password = hashlib.sha256(old_password.encode()).hexdigest()

    if user.password != hashed_old_password:
        return jsonify({'status': 'error', 'message': 'Invalid old password.'})

    if old_password == new_password:
        return jsonify({'status': 'error', 'message': 'New password cannot be the same as the old password.'})

    if new_password != confirm_password:
        return jsonify({'status': 'error', 'message': 'New password and confirm password is not matching.'})

    hashed_new_password = hashlib.sha256(new_password.encode()).hexdigest()
    user.password = hashed_new_password
    db.session.commit()

    msg = Message('Password Changed', sender='your-email@gmail.com', recipients=[user.email])
    msg.html = f'''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: #f4f4f4;
                color: #333;
                margin: 0;
                padding: 20px;
            }}
            .container {{
                background-color: #ffffff;
                max-width: 600px;
                margin: 0 auto;
                padding: 20px;
                border: 1px solid #dddddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                font-size: 24px;
                font-weight: bold;
                margin-bottom: 20px;
                color: #4CAF50;
            }}
            .content {{
                font-size: 16px;
                line-height: 1.6;
            }}
            .credentials {{
                background-color: #f9f9f9;
                padding: 10px;
                border: 1px solid #eeeeee;
                border-radius: 5px;
                margin-top: 10px;
            }}
            .footer {{
                font-size: 12px;
                color: #999;
                margin-top: 20px;
                text-align: center;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">Password Changed</div>
            <div class="content">
                <p>Hello {user.username},</p>
                <p>Your password has been successfully changed.</p>
                <p>Here are your updated credentials:</p>
                <div class="credentials">
                    <p><strong>Username:</strong> {user.username}</p>
                    <p><strong>Password:</strong> {new_password}</p>
                </div>
            </div>
            <div class="footer">
                <p>If you did not request this change, please contact our support team immediately.</p>
                <p><b>Makonis Talent Track Pro Team</b></p>
            </div>
        </div>
    </body>
    </html>
    '''
    mail.send(msg)

    if user.user_type == 'management':
        return jsonify({'status': 'success', 'message': 'Password changed successfully for management user.'})
    elif user.user_type == 'recruiter':
        return jsonify({'status': 'success', 'message': 'Password changed successfully for recruiter user.'})
    else:
        return jsonify({'status': 'success', 'message': 'Password changed successfully.'})


# @app.route('/change_password', methods=['POST'])
# def change_password():
#     data = request.json

#     if not data:
#         return jsonify({'status': 'error', 'message': 'No JSON data provided.'})

#     user_id = data.get('user_id')
#     username = data.get('username')
#     old_password = data.get('old_password')
#     new_password = data.get('new_password')
#     confirm_password = data.get('confirm_password')

#     user = User.query.filter_by(id=user_id).first()

#     if not user:
#         return jsonify({'status': 'error', 'message': 'User not found.'})

#     if username != user.username:
#         return jsonify({'status': 'error', 'message': 'Logged in user does not match the provided username.'})

#     hashed_old_password = hashlib.sha256(old_password.encode()).hexdigest()

#     if user.password != hashed_old_password:
#         return jsonify({'status': 'error', 'message': 'Invalid old password.'})

#     if old_password == new_password:
#         return jsonify({'status': 'error', 'message': 'New password cannot be the same as the old password.'})

#     if new_password != confirm_password:
#         return jsonify({'status': 'error', 'message': 'New password and confirm password is not matching.'})

#     hashed_new_password = hashlib.sha256(new_password.encode()).hexdigest()
#     user.password = hashed_new_password
#     db.session.commit()

#     msg = Message('Password Changed', sender='kanuparthisaiganesh582@gmail.com', recipients=[user.email])
#     msg.body = f'Hello {user.username},\n\nYour password has been successfully changed. Here are your updated credentials:\n\nUsername: {user.username}\nPassword: {new_password}'
#     mail.send(msg)

#     if user.user_type == 'management':
#         return jsonify({'status': 'success', 'message': 'Password changed successfully for management user.'})
#     elif user.user_type == 'recruiter':
#         return jsonify({'status': 'success', 'message': 'Password changed successfully for recruiter user.'})



@app.route('/delete_job_post_message/<int:job_id>')
def delete_job_post_message(job_id):
    job_post = JobPost.query.get(job_id)
    id = job_post.id
    client = job_post.client
    role = job_post.role
    return redirect(url_for('view_all_jobs',client=client,role=role,id=id))


@app.route('/delete_job_post/<int:job_id>', methods=['POST'])
def delete_job_post(job_id):
    # Fetch the job post
    job_post = JobPost.query.get(job_id)
    
    if not job_post:
        return jsonify({'status': 'error',"message": "Job Post not found"}), 404

    # Fetch all notifications related to the job post
    notifications = Notification.query.filter_by(job_post_id=job_id).all()

    if notifications:
        try:
            # Delete all associated notifications first
            for notification in notifications:
                db.session.delete(notification)
            db.session.commit()
            
            # Now delete the job post
            db.session.delete(job_post)
            db.session.commit()
            
            return jsonify({'status': 'success',"message": "Job Post and Notifications Deleted Successfully"})
        except Exception as e:
            # Handle any potential exceptions
            db.session.rollback()
            return jsonify({'status': 'error',"message": "An error occurred while deleting job post and notifications"})
    else:
        # No notifications found for the job post
        # Delete only the job post
        db.session.delete(job_post)
        db.session.commit()
        return jsonify({'status': 'success',"message": "Job Post and Notifications Deleted Successfully"})
        # return jsonify({'status': 'success',"message": "Job Post Deleted Successfully. No associated notifications found."}), 200


# @app.route('/delete_job_post/<int:job_id>', methods=['POST'])
# def delete_job_post(job_id):
#     # Fetch the job post
#     job_post = JobPost.query.get(job_id)
    
#     if not job_post:
#         return jsonify({"error": "Job Post not found"}), 404

#     # Fetch all notifications related to the job post
#     notifications = Notification.query.filter_by(job_post_id=job_id).all()

#     if notifications:
#         # Delete the job post and all associated notifications
#         db.session.delete(job_post)
#         for notification in notifications:
#             db.session.delete(notification)
#         db.session.commit()
#         return jsonify({"message": "Job Post and Notifications Deleted Successfully"}), 200
#     else:
#         # No notifications found for the job post
#         # Delete only the job post
#         db.session.delete(job_post)
#         db.session.commit()
#         return jsonify({"message": "Job Post Deleted Successfully. No associated notifications found."}), 200

# @app.route('/delete_job_post/<int:job_id>', methods=['POST'])
# def delete_job_post(job_id):
#     # data=request.json
#     # job_id=data['job_id']
#     job_post = JobPost.query.get(job_id)
#     if job_post:
#         JobPost.query.filter_by(id=job_id).delete()
#         db.session.commit()
#         return jsonify({"message": "Job Post Deleted Successfully"}), 200
#     else:
#         return jsonify({"error": "Job Post not found"}), 404

@app.route('/download_jd/<int:job_id>')
def download_jd(job_id):
    jobpost = JobPost.query.get(job_id)
    if jobpost is None or jobpost.jd_pdf is None:
        return redirect(url_for('dashboard'))

    jd_file = io.BytesIO(jobpost.jd_pdf)
    is_pdf = jd_file.getvalue().startswith(b"%PDF")
    if is_pdf : 
        jd_filename = f"{jobpost.client}_jd.pdf"  # Set the filename as desired
        jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_filename)
        with open(jd_path, 'wb') as file:
            file.write(jobpost.jd_pdf)

        # Send the saved resume file for download
        return send_file(jd_path, as_attachment=True)
    else:
        jd_filename = f"{jobpost.client}_jd.docx"  # Set the filename as desired
        jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_filename)
        with open(jd_path, 'wb') as file:
            file.write(jobpost.jd_pdf)

        # Send the saved resume file for download
        return send_file(jd_path, as_attachment=True)


import base64
import io
from flask import send_file


# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status
#     # Decode the base64 encoded resume data
#     print("jobpost.jd_pdf",jobpost.jd_pdf.tobytes())
#     if "==" not in str(jobpost.jd_pdf.tobytes()):
#         if request.args.get('decode') == 'base64':
#             # Decode the base64 encoded resume data
#             decoded_resume = base64.b64decode(jobpost.jd_pdf)
#             resume_binary = decoded_resume
#         else:
#             # Retrieve the resume binary data from the database
#             resume_binary = jobpost.jd_pdf.tobytes()  # Convert memoryview to bytes

#         # Determine the mimetype based on the file content
#         is_pdf = resume_binary.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             io.BytesIO(resume_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         decoded_resume = base64.b64decode(jobpost.jd_pdf)
#         # Create a file-like object (BytesIO) from the decoded resume data
#         resume_file = io.BytesIO(decoded_resume)
#         # Determine the mimetype based on the file content
#         is_pdf = decoded_resume.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             resume_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the job post data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status

#     # Convert the memoryview to bytes
#     jd_pdf_data = jobpost.jd_pdf.tobytes()
#     if request.args.get('decode') == 'base64':
#         # Decode the base64 encoded data if requested
#         try:
#             jd_pdf_binary = base64.b64decode(jd_pdf_data)
#         except Exception as e:
#             return f'Error decoding base64 data: {e}', 400
#     else:
#         jd_pdf_binary = jd_pdf_data

#     # Determine the mimetype based on the file content
#     is_pdf = jd_pdf_binary.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'

#     # Send the file as a response
#     return send_file(
#         io.BytesIO(jd_pdf_binary),
#         mimetype=mimetype,
#         as_attachment=False
#     )

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the job post data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return jsonify({'error': 'Job post not found'}), 404  # Return JSON error with 404 status

#     # Retrieve the job description data (assumed to be in jobpost.jd_pdf)
#     jd_pdf_data = jobpost.jd_pdf.tobytes()  # Convert memoryview to bytes

#     # Check if the request has 'decode=base64' argument
#     if request.args.get('decode') == 'base64':
#         # Decode the base64 encoded job description data
#         jd_pdf_data = base64.b64decode(jd_pdf_data)

#     # Create a file-like object (BytesIO) from the job description data
#     jd_pdf_file = io.BytesIO(jd_pdf_data)

#     # Determine the mimetype based on the file content
#     mimetype = 'application/pdf' if jd_pdf_data.startswith(b"%PDF") else 'application/msword'

#     # Send the file as a response
#     return send_file(jd_pdf_file, mimetype=mimetype, as_attachment=False)

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status
#     # Decode the base64 encoded resume data
#     print("jobpost.jd_pdf",jobpost.jd_pdf.tobytes())
#     if "==" not in str(jobpost.jd_pdf.tobytes()):
#         if request.args.get('decode') == 'base64':
#             # Decode the base64 encoded resume data
#             decoded_jd_pdf = base64.b64decode(jobpost.jd_pdf)
#             jd_pdf_binary = decoded_jd_pdf
#         else:
#             # Retrieve the resume binary data from the database
#             jd_pdf_binary = jobpost.jd_pdf.tobytes()  # Convert memoryview to bytes
#             print("jd_pdf_binary :",jd_pdf_binary)
#         # Determine the mimetype based on the file content
#         is_pdf = jd_pdf_binary.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             io.BytesIO(jd_pdf_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         decoded_jd_pdf = base64.b64decode(jobpost.jd_pdf)
#         # Create a file-like object (BytesIO) from the decoded resume data
#         jd_pdf_file = io.BytesIO(decoded_jd_pdf)
#         # Determine the mimetype based on the file content
#         is_pdf = decoded_jd_pdf.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             jd_pdf_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the job post data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status
    
#     # Check if the job post contains a JD PDF
#     if jobpost.jd_pdf:
#         # Decode the base64 string back to its original binary data
#         jd_binary = base64.b64decode(jobpost.jd_pdf)
 
#         # Create a file-like object (BytesIO) from the decoded binary data
#         jd_file = io.BytesIO(jd_binary)
 
#         # Check if the file is a PDF
#         is_pdf = jd_binary.startswith(b"%PDF")
 
#         # Determine the mimetype based on the file type
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'
 
#         # Return the file as a response
#         return send_file(
#             jd_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         return 'JD PDF not available', 404  # Return 404 Not Found status if JD PDF is not available


@app.route('/view_jd/<int:job_id>', methods=['GET'])
def view_jd(job_id):
    # Retrieve the resume data from the database using SQLAlchemy
    jobpost = JobPost.query.filter_by(id=job_id).first()
    if not jobpost:
        return 'Job post not found', 404  # Return 404 Not Found status

    if jobpost.jd_pdf is None:
        return 'jd_pdf not found for this job', 404

    try:
        # If the resume data is a base64 encoded string, decode it
        if isinstance(jobpost.jd_pdf, str):
            jd_pdf_binary = base64.b64decode(jobpost.jd_pdf)
        elif isinstance(jobpost.jd_pdf, bytes):
            jd_pdf_binary = jobpost.jd_pdf
        else:
            return 'Invalid jd_pdf format', 400

        # Determine the mimetype based on the file content
        is_pdf = jd_pdf_binary.startswith(b"%PDF")
        mimetype = 'application/pdf' if is_pdf else 'application/msword'

        # Send the file as a response
        return send_file(
            io.BytesIO(jd_pdf_binary),
            mimetype=mimetype,
            as_attachment=False
        )
    except Exception as e:
        return f'Error processing resume: {str(e)}', 500

# import magic

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the job post data from the database using SQLAlchemy
    # jobpost = JobPost.query.filter_by(id=job_id).first()
    # if not jobpost:
    #     return 'Job post not found', 404  # Return 404 Not Found status
    
#     # Check if the job post contains a JD PDF
#     if jobpost.jd_pdf:
#         # Decode the base64 string back to its original binary data
#         jd_binary = base64.b64decode(jobpost.jd_pdf)
 
#         # Create a file-like object (BytesIO) from the decoded binary data
#         jd_file = io.BytesIO(jd_binary)
 
#         # Detect the mimetype using more complex logic
#         mimetype = detect_mimetype(jd_binary)
 
#         # Return the file as a response
#         return send_file(
#             jd_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         return 'JD PDF not available', 404  # Return 404 Not Found status if JD PDF is not available

# def detect_mimetype(data):
#     # Check for PDF magic number
#     if data.startswith(b"%PDF"):
#         return 'application/pdf'
#     # Check for MS Word magic number
#     elif data.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
#         return 'application/msword'
#     # Add more checks for other file types if needed
#     else:
#         # If no specific type is detected, fallback to using python-magic
#         mime = magic.Magic(mime=True)
#         return mime.from_buffer(data)


from flask import Flask, request, jsonify
from datetime import datetime, timedelta
# from sqlalchemy import func
from sqlalchemy import func, text,case
from sqlalchemy.sql import text 
from sqlalchemy import text
import pandas as pd  # Import pandas for date_range
import matplotlib.pyplot as plt
from io import BytesIO
import base64

# from sqlalchemy.orm import sessionmaker
# from sqlalchemy import create_engine

# import plotly.express as px
# import plotly.io as pio

import plotly.express as px
import plotly.io as pio
from flask import Flask, request, jsonify
from sqlalchemy import func, extract
from datetime import datetime
import io
import base64

@app.route('/analyze_recruitment', methods=['POST','GET'])
def analyze_recruitment():
    data = request.json

    if not data:
        return jsonify({'error': 'No JSON data provided'})

    recruiter_usernames = data.get('recruiter_names', [])

    if not recruiter_usernames:
        return jsonify({'error': 'Please select any Recruiter'})

    try:
        from_date_str = data.get('from_date')
        to_date_str = data.get('to_date')
        from_date = datetime.strptime(from_date_str, "%d-%m-%Y")
        to_date = datetime.strptime(to_date_str, "%d-%m-%Y")
    except ValueError:
        return jsonify({'status': 'error', 'message': 'Invalid date format. Please use DD-MM-YYYY format.'})

    from_date = from_date.strftime("%Y-%m-%d")
    to_date = to_date.strftime("%Y-%m-%d")

    recruiter_data = {}
    total_candidate_count = 0
    total_selected_candidates = 0
    total_rejected_candidates_count = 0
    total_process_candidates = 0

    for recruiter_username in recruiter_usernames:
        candidates_query = db.session.query(Candidate).filter(
            Candidate.recruiter == recruiter_username,
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        )

        candidates = candidates_query.all()
        recruiter_candidate_count = candidates_query.count()
        total_candidate_count += recruiter_candidate_count

        if recruiter_candidate_count > 0:
            selected_candidates_count = candidates_query.filter(Candidate.status == 'ON-BOARDED').count()
            in_process_candidates_count = candidates_query.filter(Candidate.status.notin_([
                'ON-BOARDED', 'SCREEN REJECTED', 'L1-REJECTED', 'L2-REJECTED', 'L3-REJECTED', 'OFFER-DECLINED', 
                'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO-SHOW'])).count()
            rejected_candidates_count = candidates_query.filter(Candidate.status.in_([
                'SCREEN REJECTED', 'L1-REJECTED', 'L2-REJECTED', 'L3-REJECTED', 'OFFER-DECLINED', 
                'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO-SHOW'])).count()
        else:
            selected_candidates_count = 0
            in_process_candidates_count = 0
            rejected_candidates_count = 0

        total_selected_candidates += selected_candidates_count
        total_process_candidates += in_process_candidates_count
        total_rejected_candidates_count += rejected_candidates_count

        in_process_candidates = recruiter_candidate_count - (selected_candidates_count + rejected_candidates_count)

        # role_industry_location_analysis_result = get_role_industry_location_analysis(
        #     recruiter_username, from_date, to_date)

        conversion_rate = get_conversion_rate(candidates_query)
        # analysis_result = get_role_industry_location_analysis(recruiter_username, from_date, to_date)
        analysis_result = get_role_industry_location_analysis()
        # time_to_close_analysis = get_time_to_close_analysis(recruiter_usernames)
        historical_performance_analysis = calculate_historical_performance_analysis(recruiter_usernames, from_date, to_date)
        client_closure_rates, highest_closure_client, lowest_closure_client, _, _ = get_client_closure_rates(candidates_query)

        percentage_of_selected = (selected_candidates_count / recruiter_candidate_count) * 100 if recruiter_candidate_count > 0 else 0.0

        recruiter_data[recruiter_username] = {
            'submission_counts_daily': get_submission_counts(candidates_query, from_date, to_date, 'daily'),
            'submission_counts_weekly': get_submission_counts(candidates_query, from_date, to_date, 'weekly'),
            'submission_counts_monthly': get_submission_counts(candidates_query, from_date, to_date, 'monthly'),
            'submission_counts_yearly': get_submission_counts(candidates_query, from_date, to_date, 'yearly'),
            'selected_candidates_count': selected_candidates_count,
            'rejected_candidates_count': rejected_candidates_count,
            'in_process_candidates_count': in_process_candidates,
            'conversion_rate': conversion_rate,
            'client_closure_rates': client_closure_rates,
            'highest_closure_client': highest_closure_client,
            'lowest_closure_client': lowest_closure_client,
            'candidate_count': recruiter_candidate_count,
            'percentage_of_selected': percentage_of_selected,
            'candidates': [{
                'candidate_name': candidate.name,
                'job_id': candidate.job_id,
                'client': candidate.client,
                'recruiter': candidate.recruiter,
                'date_created': candidate.date_created.strftime('%Y-%m-%d'),
                'time_created': candidate.date_created.strftime('%H:%M:%S'),
                'profile': candidate.profile,
                'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
                'status': candidate.status
            } for candidate in candidates]
        }

    ranked_recruiters = sorted(recruiter_data.items(), key=lambda x: (x[1]['selected_candidates_count'] / x[1]['candidate_count'] if x[1]['candidate_count'] > 0 else 0), reverse=True)

    rank = 0

    for recruiter, data in ranked_recruiters:
        if data['candidate_count'] > 0:
            if data['selected_candidates_count'] > 0:
                rank += 1
                recruiter_data[recruiter]['ranking'] = rank
            else:
                recruiter_data[recruiter]['ranking'] = 0
        else:
            recruiter_data[recruiter]['ranking'] = 0

    response_data = {
        'status': 'success',
        # 'time_to_close_analysis':time_to_close_analysis,
        'historical_performance_analysis':historical_performance_analysis,
        'Job_Type_Analysis': analysis_result,
        'recruiter_data': recruiter_data,
        'total_candidate_count': total_candidate_count,
        'total_selected_candidates': total_selected_candidates,
        'total_rejected_candidates_count': total_rejected_candidates_count,
        'total_process_candidates_count': total_process_candidates,
        'from_date_str': from_date_str,
        'to_date_str': to_date_str,
        'message': 'Ranking calculations completed successfully',
        'bar_graph_data': generate_bar_graph_data(recruiter_data)
    }

    return jsonify(response_data)
    
def calculate_historical_performance_analysis(recruiter_usernames, from_date_str, to_date_str):
    try:
        from_date = datetime.strptime(from_date_str, "%Y-%m-%d")
        to_date = datetime.strptime(to_date_str, "%Y-%m-%d")
    except ValueError:
        return {'status': 'error', 'message': 'Invalid date format. Please use YYYY-MM-DD format.'}

    result = {}

    for recruiter_name in recruiter_usernames:
        # Initialize data structures
        monthly_data = {}
        total_months = (to_date.year - from_date.year) * 12 + (to_date.month - from_date.month) + 1

        for i in range(total_months):
            # Calculate start and end of current month
            start_date = from_date + relativedelta(months=i)
            end_date = start_date + relativedelta(months=1) - relativedelta(days=1)

            # Query candidates for the recruiter within the current month (mockup query)
            candidates = db.session.query(Candidate).filter(
                Candidate.recruiter == recruiter_name,
                Candidate.date_created >= start_date,
                Candidate.date_created <= end_date,
                Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
            ).all()

            candidates_data = []
            total_screening_candidates = 0
            total_days_to_close = 0
            count_of_onboarded_positions = 0
            total_candidates = len(candidates)
            unsuccessful_closures = 0

            for candidate in candidates:
                if candidate.status == 'SCREENING':
                    total_screening_candidates += 1
                elif candidate.status == 'ON-BOARDED':
                    count_of_onboarded_positions += 1

                    # Calculate days to close (mockup calculation)
                    if candidate.date_created and candidate.last_working_date:
                        days_to_close = (candidate.last_working_date - candidate.date_created).days
                        total_days_to_close += days_to_close

                    # Prepare candidate data (mockup preparation)
                    candidate_data = {
                        'candidate_name': candidate.name,
                        'job_id': candidate.job_id,  # Assuming job_id is a regular column
                        'client': candidate.client,
                        'recruiter': candidate.recruiter,
                        'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
                        'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
                        'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
                        'profile': candidate.profile,
                        'status': candidate.status
                    }
                    candidates_data.append(candidate_data)
                elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
                    unsuccessful_closures += 1

            # Calculate average days to close (mockup calculation)
            average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

            # Calculate percentage of onboarded candidates (mockup calculation)
            percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

            # Store monthly data
            monthly_data[start_date.strftime('%Y-%m')] = {
                'candidates': candidates_data,
                'total_days_to_close': total_days_to_close,
                'count_of_screening_candidates': total_screening_candidates,
                'count_of_onboarded_positions': count_of_onboarded_positions,
                'unsuccessful_closures': unsuccessful_closures,
                'average_days_to_close': average_days_to_close,
                'percentage_onboarded': percentage_onboarded,
                'total_candidates_count': total_candidates
            }

        # Evaluate trend in closure rates (mockup evaluation)
        first_month_data = next(iter(monthly_data.values()), None)
        last_month_data = next(iter(list(monthly_data.values())[::-1]), None)  # Convert dict_values to list and then reverse

        if first_month_data and last_month_data:
            initial_onboarded_percentage = first_month_data.get('percentage_onboarded', 0)
            final_onboarded_percentage = last_month_data.get('percentage_onboarded', 0)

            if final_onboarded_percentage > initial_onboarded_percentage:
                trend = 'improving'
            elif final_onboarded_percentage < initial_onboarded_percentage:
                trend = 'declining'
            else:
                trend = 'stable'

            # Prepare line graph data
            line_graph_data = prepare_line_graph_data(recruiter_name, monthly_data)

            result[recruiter_name] = {
                'line_graph_data': line_graph_data,
                'monthly_data': monthly_data,
                'overall_summary': {
                    'total_months_analyzed': total_months,
                    'trend_in_closure_rates': trend
                }
            }
        else:
            result[recruiter_name] = {
                'monthly_data': monthly_data,
                'overall_summary': {
                    'total_months_analyzed': total_months,
                    'trend_in_closure_rates': 'insufficient data'  # Or handle as appropriate
                }
            }

    return result

# Function to prepare line graph data
def prepare_line_graph_data(recruiter_name, monthly_data):
    line_graph_data = {}

    for month, data in monthly_data.items():
        line_graph_data[month] = {
            'Total Days to Close': data['total_days_to_close'],
            'Screening Candidates': data['count_of_screening_candidates'],
            'Onboarded Positions': data['count_of_onboarded_positions'],
            'Unsuccessful Closures': data['unsuccessful_closures'],
            'Average Days to Close': data['average_days_to_close'],
            'Percentage Onboarded': data['percentage_onboarded'],
            'Total Candidates Count': data['total_candidates_count']
        }

    return line_graph_data


# def calculate_historical_performance_analysis(recruiter_usernames, from_date_str, to_date_str):
#     try:
#         from_date = datetime.strptime(from_date_str, "%Y-%m-%d")
#         to_date = datetime.strptime(to_date_str, "%Y-%m-%d")
#     except ValueError:
#         return {'status': 'error', 'message': 'Invalid date format. Please use YYYY-MM-DD format.'}

#     result = {}

#     for recruiter_name in recruiter_usernames:
#         # Initialize data structures
#         monthly_data = {}
#         total_months = (to_date.year - from_date.year) * 12 + (to_date.month - from_date.month) + 1

#         for i in range(total_months):
#             # Calculate start and end of current month
#             start_date = from_date + relativedelta(months=i)
#             end_date = start_date + relativedelta(months=1) - relativedelta(days=1)

#             # Query candidates for the recruiter within the current month
#             candidates = db.session.query(Candidate).filter(
#                 Candidate.recruiter == recruiter_name,
#                 Candidate.date_created >= start_date,
#                 Candidate.date_created <= end_date,
#                 Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
#             ).all()

#             candidates_data = []
#             total_screening_candidates = 0
#             total_days_to_close = 0
#             count_of_onboarded_positions = 0
#             total_candidates = len(candidates)
#             unsuccessful_closures = 0

#             for candidate in candidates:
#                 if candidate.status == 'SCREENING':
#                     total_screening_candidates += 1
#                 elif candidate.status == 'ON-BOARDED':
#                     count_of_onboarded_positions += 1

#                     # Calculate days to close
#                     if candidate.date_created and candidate.data_updated_date:
#                         days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                         total_days_to_close += days_to_close

#                     # Prepare candidate data
#                     candidate_data = {
#                         'candidate_name': candidate.name,
#                         'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                         'client': candidate.client,
#                         'recruiter': candidate.recruiter,
#                         'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                         'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                         'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                         'profile': candidate.profile,
#                         'status': candidate.status
#                     }
#                     candidates_data.append(candidate_data)
#                 elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
#                     unsuccessful_closures += 1

#             # Calculate average days to close
#             average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#             # Calculate percentage of onboarded candidates
#             percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

#             # Store monthly data
#             monthly_data[start_date.strftime('%Y-%m')] = {
#                 'candidates': candidates_data,
#                 'total_days_to_close': total_days_to_close,
#                 'count_of_screening_candidates': total_screening_candidates,
#                 'count_of_onboarded_positions': count_of_onboarded_positions,
#                 'unsuccessful_closures': unsuccessful_closures,
#                 'average_days_to_close': average_days_to_close,
#                 'percentage_onboarded': percentage_onboarded,
#                 'total_candidates_count': total_candidates
#             }

#         # Evaluate trend in closure rates
#         first_month_data = next(iter(monthly_data.values()), None)
#         last_month_data = next(iter(list(monthly_data.values())[::-1]), None)  # Convert dict_values to list and then reverse

#         if first_month_data and last_month_data:
#             initial_onboarded_percentage = first_month_data.get('percentage_onboarded', 0)
#             final_onboarded_percentage = last_month_data.get('percentage_onboarded', 0)

#             if final_onboarded_percentage > initial_onboarded_percentage:
#                 trend = 'improving'
#             elif final_onboarded_percentage < initial_onboarded_percentage:
#                 trend = 'declining'
#             else:
#                 trend = 'stable'

#             result[recruiter_name] = {
#                 'monthly_data': monthly_data,
#                 'overall_summary': {
#                     'total_months_analyzed': total_months,
#                     'trend_in_closure_rates': trend
#                 }
#             }
#         else:
#             result[recruiter_name] = {
#                 'monthly_data': monthly_data,
#                 'overall_summary': {
#                     'total_months_analyzed': total_months,
#                     'trend_in_closure_rates': 'insufficient data'  # Or handle as appropriate
#                 }
#             }

#     return result


# def get_time_to_close_analysis(recruiter_usernames):
#     result = {}

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter
#         candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0
#         total_candidates = len(candidates)
#         unsuccessful_closures = 0

#         for candidate in candidates:
#             if candidate.status == 'SCREENING':
#                 total_screening_candidates += 1
#             elif candidate.status == 'ON-BOARDED':
#                 count_of_onboarded_positions += 1

#                 # Calculate days to close
#                 if candidate.date_created and candidate.data_updated_date:
#                     days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                     total_days_to_close += days_to_close

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': candidate.name,
#                     'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                     'client': candidate.client,
#                     'recruiter': candidate.recruiter,
#                     'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                     'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                     'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                     'profile': candidate.profile,
#                     'status': candidate.status
#                 }
#                 candidates_data.append(candidate_data)
#             elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
#                 unsuccessful_closures += 1

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Calculate percentage of onboarded candidates
#         percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'unsuccessful_closures': unsuccessful_closures,
#             'average_days_to_close': average_days_to_close,
#             'percentage_onboarded': percentage_onboarded,
#             'total_candidates_count': total_candidates
#         }

#         result[recruiter_name] = recruiter_data

#     # Sort recruiters by percentage of onboarded candidates
#     ranked_recruiters = sorted(result.values(), key=lambda x: x['percentage_onboarded'], reverse=True)

#     # Assign rankings
#     rank = 0
#     for data in ranked_recruiters:
#         rank += 1
#         if data['count_of_onboarded_positions'] > 0:
#             data['ranking'] = rank
#         else:
#             data['ranking'] = 0

#     return ranked_recruiters

# def get_time_to_close_analysis(recruiter_usernames):
#     result = {}

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter
#         candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0
#         total_candidates = len(candidates)
#         unsuccessful_closures = 0

#         for candidate in candidates:
#             if candidate.status == 'SCREENING':
#                 total_screening_candidates += 1
#             elif candidate.status == 'ON-BOARDED':
#                 count_of_onboarded_positions += 1

#                 # Calculate days to close
#                 if candidate.date_created and candidate.data_updated_date:
#                     days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                     total_days_to_close += days_to_close

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': candidate.name,
#                     'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                     'client': candidate.client,
#                     'recruiter': candidate.recruiter,
#                     'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                     'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                     'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                     'profile': candidate.profile,
#                     'status': candidate.status
#                 }
#                 candidates_data.append(candidate_data)
#             elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
#                 unsuccessful_closures += 1

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Calculate percentage of onboarded candidates
#         percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'unsuccessful_closures': unsuccessful_closures,
#             'average_days_to_close': average_days_to_close,
#             'percentage_onboarded': percentage_onboarded,
#             'total_candidates_count': total_candidates
#         }

#         result[recruiter_name] = recruiter_data

#     # Sort recruiters by percentage of onboarded candidates
#     ranked_recruiters = sorted(result.items(), key=lambda x: (x[1]['percentage_onboarded']), reverse=True)

#     # Assign rankings
#     rank = 0
#     for recruiter, data in ranked_recruiters:
#         if data['total_candidates_count'] > 0:
#             rank += 1
#             result[recruiter]['ranking'] = rank
#         else:
#             result[recruiter]['ranking'] = 0

#     return list(result.values())
    

# def get_time_to_close_analysis(recruiter_usernames):
#     result = []

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter
#         candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0
#         total_candidates = len(candidates)
#         unsuccessful_closures = 0

#         for candidate in candidates:
#             if candidate.status == 'SCREENING':
#                 total_screening_candidates += 1
#             elif candidate.status == 'ON-BOARDED':
#                 count_of_onboarded_positions += 1

#                 # Calculate days to close
#                 if candidate.date_created and candidate.data_updated_date:
#                     days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                     total_days_to_close += days_to_close

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': candidate.name,
#                     'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                     'client': candidate.client,
#                     'recruiter': candidate.recruiter,
#                     'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                     'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                     'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                     'profile': candidate.profile,
#                     'status': candidate.status
#                 }
#                 candidates_data.append(candidate_data)
#             elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
#                 unsuccessful_closures += 1

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Calculate percentage of onboarded candidates
#         percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'unsuccessful_closures': unsuccessful_closures,
#             'average_days_to_close': average_days_to_close,
#             'percentage_onboarded': percentage_onboarded,
#             'total_candidates_count': total_candidates
#         }

#         result.append(recruiter_data)

#     # Calculate ranking based on percentage of onboarded candidates (highest percentage gets higher rank)
#     result.sort(key=lambda x: x['percentage_onboarded'], reverse=True)  # Sort by percentage_onboarded descending
#     for i, recruiter_data in enumerate(result, start=1):
#         recruiter_data['ranking'] = i

#     return result


# def get_time_to_close_analysis(recruiter_usernames):
#     result = []

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter
#         candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0
#         total_candidates = len(candidates)
#         unsuccessful_closures = 0

#         for candidate in candidates:
#             if candidate.status == 'SCREENING':
#                 total_screening_candidates += 1
#             elif candidate.status == 'ON-BOARDED':
#                 count_of_onboarded_positions += 1

#                 # Calculate days to close
#                 if candidate.date_created and candidate.data_updated_date:
#                     days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                     total_days_to_close += days_to_close

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': candidate.name,
#                     'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                     'client': candidate.client,
#                     'recruiter': candidate.recruiter,
#                     'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                     'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                     'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                     'profile': candidate.profile,
#                     'status': candidate.status
#                 }
#                 candidates_data.append(candidate_data)
#             elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
#                 unsuccessful_closures += 1

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Calculate percentage of onboarded candidates
#         percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'unsuccessful_closures': unsuccessful_closures,
#             'average_days_to_close': average_days_to_close,
#             'percentage_onboarded': percentage_onboarded,
#             'total_candidates_count': total_candidates
#         }

#         result.append(recruiter_data)

#     # Calculate ranking based on count of onboarded positions (highest count gets higher rank)
#     result.sort(key=lambda x: x['count_of_onboarded_positions'], reverse=True)  # Sort by count_of_onboarded_positions descending
#     for i, recruiter_data in enumerate(result, start=1):
#         recruiter_data['ranking'] = i

#     return result


# def get_(recruiter_usernames):
#     result = []

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter where status is 'SCREENING' or 'ON-BOARDED'
#         candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status.in_(['SCREENING', 'ON-BOARDED'])
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0

#         for candidate in candidates:
#             if candidate.status == 'SCREENING':
#                 total_screening_candidates += 1
#             elif candidate.status == 'ON-BOARDED':
#                 count_of_onboarded_positions += 1

#                 # Calculate days to close
#                 if candidate.date_created and candidate.data_updated_date:
#                     days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                     total_days_to_close += days_to_close

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': candidate.name,
#                     'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                     'client': candidate.client,
#                     'recruiter': candidate.recruiter,
#                     'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                     'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                     'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                     'profile': candidate.profile,
#                     'status': candidate.status
#                 }
#                 candidates_data.append(candidate_data)

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Calculate percentage of onboarded candidates
#         total_candidates = len(candidates)
#         if total_candidates > 0:
#             percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100
#         else:
#             percentage_onboarded = 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'average_days_to_close': average_days_to_close,
#             'percentage_onboarded': percentage_onboarded,
#             'total_candidates_count': total_candidates
#         }

#         result.append(recruiter_data)

#     # Calculate ranking based on count of onboarded positions (highest count gets higher rank)
#     result.sort(key=lambda x: x['count_of_onboarded_positions'], reverse=True)  # Sort by count_of_onboarded_positions descending
#     for i, recruiter_data in enumerate(result, start=1):
#         recruiter_data['ranking'] = i

#     return result



# def get_(recruiter_usernames):
#     result = []

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter where status is 'SCREENING' or 'ON-BOARDED'
#         candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status.in_(['SCREENING', 'ON-BOARDED'])
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0

#         for candidate in candidates:
#             if candidate.status == 'SCREENING':
#                 total_screening_candidates += 1
#             elif candidate.status == 'ON-BOARDED':
#                 count_of_onboarded_positions += 1

#                 # Calculate days to close
#                 if candidate.date_created and candidate.data_updated_date:
#                     days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                     total_days_to_close += days_to_close

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': candidate.name,
#                     'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                     'client': candidate.client,
#                     'recruiter': candidate.recruiter,
#                     'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                     'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                     'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                     'profile': candidate.profile,
#                     'status': candidate.status
#                 }
#                 candidates_data.append(candidate_data)

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'average_days_to_close': average_days_to_close
#         }

#         result.append(recruiter_data)

#     return result



# def get_time_to_close_analysis(recruiter_usernames):
#     # data=request.json
#     # recruiter_usernames = data.get('recruiter_usernames', [])
#     result = []

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter where status is 'SCREENING'
#         screening_candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status == 'SCREENING'
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0

#         for candidate in screening_candidates:
#             total_screening_candidates += 1

#             # Check if the candidate eventually transitions to 'ON-BOARDED'
#             onboarded_candidate = db.session.query(Candidate).filter(
#                 Candidate.id == candidate.id,
#                 Candidate.recruiter == recruiter_name,
#                 Candidate.status == 'ON-BOARDED'
#             ).first()

#             if onboarded_candidate and onboarded_candidate.date_created and onboarded_candidate.data_updated_date:
#                 # Calculate days to close
#                 days_to_close = (onboarded_candidate.data_updated_date - onboarded_candidate.date_created).days

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': onboarded_candidate.name,
#                     'job_id': onboarded_candidate.job_id,
#                     'client': onboarded_candidate.client,
#                     'recruiter': onboarded_candidate.recruiter,
#                     'date_created': onboarded_candidate.date_created.strftime('%Y-%m-%d'),
#                     'date_updated': onboarded_candidate.data_updated_date.strftime('%Y-%m-%d'),
#                     'days_to_close': days_to_close,
#                     'profile': onboarded_candidate.profile,
#                     'status': onboarded_candidate.status
#                 }
#                 candidates_data.append(candidate_data)
#                 total_days_to_close += days_to_close
#                 count_of_onboarded_positions += 1

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'average_days_to_close': average_days_to_close
#         }

#         result.append(recruiter_data)

#     return result

def get_submission_counts(candidates_query, from_date, to_date, interval):
    if interval == 'daily':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.DATE(Candidate.date_created)).with_entities(
            func.DATE(Candidate.date_created).label('date_part'),
            func.count().label('count')
        )
    elif interval == 'weekly':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.DATE(Candidate.date_created)).with_entities(
            func.DATE(Candidate.date_created).label('date_part'),
            func.count().label('count')
        )
    elif interval == 'monthly':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.TO_CHAR(Candidate.date_created, 'YYYY-MM')).with_entities(
            func.TO_CHAR(Candidate.date_created, 'YYYY-MM').label('date_part'),
            func.count().label('count')
        )
    elif interval == 'yearly':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.TO_CHAR(Candidate.date_created, 'YYYY')).with_entities(
            func.TO_CHAR(Candidate.date_created, 'YYYY').label('date_part'),
            func.count().label('count')
        )
    else:
        return []

    submission_counts = grouped_query.all()
    return [{'date_part': str(item.date_part), 'count': item.count} for item in submission_counts]

def get_role_industry_location_analysis():
    # Get all distinct roles from JobPost table
    roles = db.session.query(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client
    ).distinct().all()

    # Query to get count of all candidates by role, location, job type, and client
    role_industry_location_analysis = db.session.query(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client,
        func.count(Candidate.id).label('total_count')
    ).join(
        Candidate,
        Candidate.job_id == JobPost.id  # Assuming job_id links Candidate to JobPost
    ).filter(
        JobPost.role.isnot(None)  # Ensure JobPost.role is not None (to filter out non-linked candidates)
    ).group_by(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client
    ).all()

    # Query to get the count of on-boarded candidates by role, location, job type, and client
    on_boarded_candidates_analysis = db.session.query(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client,
        func.count(Candidate.id).label('on_boarded_count')
    ).join(
        Candidate,
        Candidate.job_id == JobPost.id
    ).filter(
        Candidate.status == 'ON-BOARDED'
    ).group_by(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client
    ).all()

    # Combine results
    result = []
    on_boarded_dict = {(item.role, item.location, item.job_type, item.client): item.on_boarded_count for item in on_boarded_candidates_analysis}
    total_dict = {(item.role, item.location, item.job_type, item.client): item.total_count for item in role_industry_location_analysis}

    for role, location, job_type, client in roles:
        key = (role, location, job_type, client)
        total_count = total_dict.get(key, 0)
        on_boarded_count = on_boarded_dict.get(key, 0)
        on_boarded_percentage = (on_boarded_count / total_count) * 100 if total_count > 0 else 0

        result.append({
            'role': role,
            'location': location,
            'job_type': job_type,
            'client': client,
            'total_count': total_count,
            'on_boarded_count': on_boarded_count,
            'on_boarded_percentage': on_boarded_percentage
        })

    return result

# def get_role_industry_location_analysis():
#     # Query to get count of selected candidates by role, location, job type, and client
#     role_industry_location_analysis = db.session.query(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type,
#         JobPost.client,
#         func.count(Candidate.id).label('selected_count')
#     ).join(
#         Candidate,
#         Candidate.job_id == JobPost.id  # Assuming job_id links Candidate to JobPost
#     ).filter(
#         # Candidate.status == 'SELECTED',
#         JobPost.role.isnot(None)  # Ensure JobPost.role is not None (to filter out non-linked candidates)
#     ).group_by(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type,
#         JobPost.client
#     ).all()

#     # Query to get the count of on-boarded candidates by role, location, job type, and client
#     on_boarded_candidates_analysis = db.session.query(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type,
#         JobPost.client,
#         func.count(Candidate.id).label('on_boarded_count')
#     ).join(
#         Candidate,
#         Candidate.job_id == JobPost.id
#     ).filter(
#         Candidate.status == 'ON-BOARDED'
#     ).group_by(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type,
#         JobPost.client
#     ).all()

#     # Combine results
#     result = []
#     on_boarded_dict = {(item.role, item.location, item.job_type, item.client): item.on_boarded_count for item in on_boarded_candidates_analysis}

#     for item in role_industry_location_analysis:
#         key = (item.role, item.location, item.job_type, item.client)
#         selected_count = item.selected_count
#         on_boarded_count = on_boarded_dict.get(key, 0)
#         on_boarded_percentage = (on_boarded_count / selected_count) * 100 if selected_count > 0 else 0

#         result.append({
#             'role': item.role,
#             'location': item.location,
#             'job_type': item.job_type,
#             'client': item.client,
#             'selected_count': selected_count,
#             'on_boarded_count': on_boarded_count,
#             'on_boarded_percentage': on_boarded_percentage
#         })

#     return result

# def get_role_industry_location_analysis(recruiter_username, from_date, to_date):
#     role_industry_location_analysis = db.session.query(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type,
#         func.count(Candidate.id).label('count')
#     ).join(
#         Candidate,
#         and_(
#             Candidate.profile == JobPost.role,
#             Candidate.job_id == JobPost.id  # Assuming job_id links Candidate to JobPost
#         )
#     ).filter(
#         Candidate.recruiter == recruiter_username,
#         # Candidate.date_created >= from_date,
#         # Candidate.date_created <= to_date,
#         ~Candidate.status.in_(['ON-BOARDED',
#                 'SCREEN REJECTED', 'L1-REJECTED', 'L2-REJECTED', 'L3-REJECTED', 'OFFER-DECLINED', 
#                 'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO-SHOW']),
#         JobPost.role.isnot(None)  # Ensure JobPost.role is not None (to filter out non-linked candidates)
#     ).group_by(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type
#     ).all()

#     remaining_candidates = [{
#         'role': item.role,
#         'location': item.location,
#         'job_type': item.job_type,
#         'count': item.count
#     } for item in role_industry_location_analysis]

#     # Additional query to count linked candidates per role
#     linked_candidates_count = {}
#     for role_info in remaining_candidates:
#         role = role_info['role']
#         count = role_info['count']
#         linked_candidates_count[role] = count

#     total_remaining_candidates = sum(item['count'] for item in remaining_candidates)

#     rejected_candidates = db.session.query(
#         Candidate.id,
#         Candidate.name
#     ).filter(
#         Candidate.recruiter == recruiter_username,
#         # Candidate.date_created >= from_date,
#         # Candidate.date_created <= to_date,
#         ~Candidate.status.in_(['ON-BOARDED',
#                 'SCREEN REJECTED', 'L1-REJECTED', 'L2-REJECTED', 'L3-REJECTED', 'OFFER-DECLINED', 
#                 'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO-SHOW'])
#     ).all()

#     rejected_candidates_list = [{'id': candidate.id, 'name': candidate.name} for candidate in rejected_candidates]
#     total_rejected_candidates = len(rejected_candidates_list)

#     on_boarded_candidates = db.session.query(
#         Candidate.id,
#         Candidate.name,
#         Candidate.client,
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type
#     ).join(
#         JobPost,
#         Candidate.job_id == JobPost.id
#     ).filter(
#         Candidate.recruiter == recruiter_username,
#         # Candidate.date_created >= from_date,
#         # Candidate.date_created <= to_date,
#         Candidate.status == 'ON-BOARDED'
#     ).all()

#     on_boarded_candidates_list = [{
#         'candidate_id': candidate.id,
#         'candidate_name': candidate.name,
#         'client': candidate.client,
#         'role': candidate.role,
#         'location': candidate.location,
#         'job_type': candidate.job_type
#     } for candidate in on_boarded_candidates]

#     total_on_boarded_candidates = len(on_boarded_candidates_list)

#     on_boarded_percentage = (total_on_boarded_candidates / total_remaining_candidates) * 100 if total_remaining_candidates > 0 else 0

#     result = {
#         'remaining_candidates': remaining_candidates,
#         'total_remaining_candidates': total_remaining_candidates,
#         'linked_candidates_count': linked_candidates_count,  # Include linked candidates count here
#         'rejected_candidates': rejected_candidates_list,
#         'total_rejected_candidates': total_rejected_candidates,
#         'on_boarded_candidates': on_boarded_candidates_list,
#         'total_on_boarded_candidates': total_on_boarded_candidates,
#         'on_boarded_percentage': on_boarded_percentage
#     }

#     return result

# def get_role_industry_location_analysis(recruiter_username, from_date, to_date):
#     role_industry_location_analysis = db.session.query(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type,
#         func.count(Candidate.id).label('count')
#     ).join(
#         Candidate,
#         and_(
#             Candidate.profile == JobPost.role,
#             Candidate.job_id == JobPost.id  # Assuming job_id links Candidate to JobPost
#         )
#     ).filter(
#         Candidate.recruiter == recruiter_username,
#         Candidate.date_created >= from_date,
#         Candidate.date_created <= to_date,
#         ~Candidate.status.in_(['REJECTED', 'ON-BOARDED']),
#         JobPost.role.isnot(None)  # Ensure JobPost.role is not None (to filter out non-linked candidates)
#     ).group_by(
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type
#     ).all()

#     remaining_candidates = [{
#         'role': item.role,
#         'location': item.location,
#         'job_type': item.job_type,
#         'count': item.count
#     } for item in role_industry_location_analysis]

#     # Additional query to count linked candidates per role
#     linked_candidates_count = {}
#     for role_info in remaining_candidates:
#         role = role_info['role']
#         count = role_info['count']
#         linked_candidates_count[role] = count

#     total_remaining_candidates = sum(item['count'] for item in remaining_candidates)

#     rejected_candidates = db.session.query(
#         Candidate.id,
#         Candidate.name
#     ).filter(
#         Candidate.recruiter == recruiter_username,
#         Candidate.date_created >= from_date,
#         Candidate.date_created <= to_date,
#         Candidate.status == 'REJECTED'
#     ).all()

#     rejected_candidates_list = [{'id': candidate.id, 'name': candidate.name} for candidate in rejected_candidates]
#     total_rejected_candidates = len(rejected_candidates_list)

#     on_boarded_candidates = db.session.query(
#         Candidate.id,
#         Candidate.name,
#         Candidate.client,
#         JobPost.role,
#         JobPost.location,
#         JobPost.job_type
#     ).join(
#         JobPost,
#         Candidate.job_id == JobPost.id
#     ).filter(
#         Candidate.recruiter == recruiter_username,
#         Candidate.date_created >= from_date,
#         Candidate.date_created <= to_date,
#         Candidate.status == 'ON-BOARDED'
#     ).all()

#     on_boarded_candidates_list = [{
#         'candidate_id': candidate.id,
#         'candidate_name': candidate.name,
#         'client': candidate.client,
#         'role': candidate.role,
#         'location': candidate.location,
#         'job_type': candidate.job_type
#     } for candidate in on_boarded_candidates]

#     total_on_boarded_candidates = len(on_boarded_candidates_list)

#     on_boarded_percentage = (total_on_boarded_candidates / total_remaining_candidates) * 100 if total_remaining_candidates > 0 else 0

#     result = {
#         'remaining_candidates': remaining_candidates,
#         'total_remaining_candidates': total_remaining_candidates,
#         'linked_candidates_count': linked_candidates_count,  # Include linked candidates count here
#         'rejected_candidates': rejected_candidates_list,
#         'total_rejected_candidates': total_rejected_candidates,
#         'on_boarded_candidates': on_boarded_candidates_list,
#         'total_on_boarded_candidates': total_on_boarded_candidates,
#         'on_boarded_percentage': on_boarded_percentage
#     }

#     return result


def get_conversion_rate(query):
    total_submissions = query.count()
    if total_submissions > 0:
        successful_closures = query.filter(Candidate.status == 'ON-BOARDED').count()
        conversion_rate = successful_closures / total_submissions
    else:
        conversion_rate = 0.0

    return conversion_rate


def get_client_closure_rates(query):
    client_closure_counts = query.filter(Candidate.status.in_(['SELECTED', 'ON-BOARDED'])).group_by(Candidate.client).with_entities(
        Candidate.client,
        func.count().label('count')
    ).all()

    client_closure_rates = [{'client': item.client, 'count': item.count} for item in client_closure_counts]

    highest_closure_client = max(client_closure_rates, key=lambda x: x['count']) if client_closure_rates else None
    lowest_closure_client = min(client_closure_rates, key=lambda x: x['count']) if client_closure_rates else None

    highest_closure_candidates = []
    lowest_closure_candidates = []

    if highest_closure_client:
        highest_closure_candidates = query.filter(
            Candidate.client == highest_closure_client['client'],
            Candidate.status.in_(['SELECTED', 'ON-BOARDED'])
        ).all()
        highest_closure_candidates = [{
            'candidate_name': candidate.name,
            'job_id': candidate.job_id,
            'client': candidate.client,
            'recruiter': candidate.recruiter,
            'date_created': candidate.date_created.strftime('%Y-%m-%d'),
            'time_created': candidate.date_created.strftime('%H:%M:%S'),
            'profile': candidate.profile,
            'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
            'status': candidate.status
        } for candidate in highest_closure_candidates]

    if lowest_closure_client:
        lowest_closure_candidates = query.filter(
            Candidate.client == lowest_closure_client['client'],
            Candidate.status.in_(['SELECTED', 'ON-BOARDED'])
        ).all()
        lowest_closure_candidates = [{
            'candidate_name': candidate.name,
            'job_id': candidate.job_id,
            'client': candidate.client,
            'recruiter': candidate.recruiter,
            'date_created': candidate.date_created.strftime('%Y-%m-%d'),
            'time_created': candidate.date_created.strftime('%H:%M:%S'),
            'profile': candidate.profile,
            'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
            'status': candidate.status
        } for candidate in lowest_closure_candidates]

    return client_closure_rates, highest_closure_client, lowest_closure_client, highest_closure_candidates, lowest_closure_candidates


def generate_bar_graph_data(recruiter_data):
    bar_graph_data = {
        'recruiters': [],
        'selected_candidates_counts': []
    }

    for recruiter, data in recruiter_data.items():
        bar_graph_data['recruiters'].append(recruiter)
        bar_graph_data['selected_candidates_counts'].append(data['selected_candidates_count'])

    return bar_graph_data


@app.route('/time_to_close_position_for_recruiter', methods=['POST'])
def get_time_to_close_analysis():
    data = request.json
    recruiter_usernames = data.get('recruiter_names', [])
    result = {}

    for recruiter_name in recruiter_usernames:
        # Query candidates for the recruiter
        candidates = db.session.query(Candidate).filter(
            Candidate.recruiter == recruiter_name,
            Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
        ).all()

        candidates_data = []
        total_screening_candidates = 0
        total_days_to_close = 0
        count_of_onboarded_positions = 0
        total_candidates = len(candidates)
        unsuccessful_closures = 0

        for candidate in candidates:
            if candidate.status == 'SCREENING':
                total_screening_candidates += 1
            elif candidate.status == 'ON-BOARDED':
                count_of_onboarded_positions += 1

                # Calculate days to close
                if candidate.date_created and candidate.data_updated_date:
                    days_to_close = (candidate.data_updated_date - candidate.date_created).days
                    total_days_to_close += days_to_close

                # Prepare candidate data
                candidate_data = {
                    'candidate_name': candidate.name,
                    'job_id': candidate.job_id,  # Assuming job_id is a regular column
                    'client': candidate.client,
                    'recruiter': candidate.recruiter,
                    'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
                    'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
                    'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
                    'profile': candidate.profile,
                    'status': candidate.status
                }
                candidates_data.append(candidate_data)
            elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
                unsuccessful_closures += 1

        # Calculate average days to close
        average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

        # Calculate percentage of onboarded candidates
        percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

        # Append summary and candidates data for the recruiter
        recruiter_data = {
            'recruiter_name': recruiter_name,
            'candidates': candidates_data,
            'total_days_to_close': total_days_to_close,
            'count_of_screening_candidates': total_screening_candidates,
            'count_of_onboarded_positions': count_of_onboarded_positions,
            'unsuccessful_closures': unsuccessful_closures,
            'average_days_to_close': average_days_to_close,
            'percentage_onboarded': percentage_onboarded,
            'total_candidates_count': total_candidates
        }

        result[recruiter_name] = recruiter_data

    # Sort recruiters by average days to close (ascending order)
    ranked_recruiters = sorted(result.values(), key=lambda x: x['average_days_to_close'])

    # Assign rankings
    rank = 1
    for data in ranked_recruiters:
        if data['count_of_onboarded_positions'] > 0:
            data['ranking'] = rank
            rank += 1
        else:
            data['ranking'] = 0

    return jsonify(ranked_recruiters)


# @app.route('/time_to_close_position_for_recruiter', methods=['POST'])
# def get_time_to_close_analysis():
#     data=request.json
#     recruiter_usernames = data.get('recruiter_names', [])
#     result = {}

#     for recruiter_name in recruiter_usernames:
#         # Query candidates for the recruiter
#         candidates = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
#         ).all()

#         candidates_data = []
#         total_screening_candidates = 0
#         total_days_to_close = 0
#         count_of_onboarded_positions = 0
#         total_candidates = len(candidates)
#         unsuccessful_closures = 0

#         for candidate in candidates:
#             if candidate.status == 'SCREENING':
#                 total_screening_candidates += 1
#             elif candidate.status == 'ON-BOARDED':
#                 count_of_onboarded_positions += 1

#                 # Calculate days to close
#                 if candidate.date_created and candidate.data_updated_date:
#                     days_to_close = (candidate.data_updated_date - candidate.date_created).days
#                     total_days_to_close += days_to_close

#                 # Prepare candidate data
#                 candidate_data = {
#                     'candidate_name': candidate.name,
#                     'job_id': candidate.job_id,  # Assuming job_id is a regular column
#                     'client': candidate.client,
#                     'recruiter': candidate.recruiter,
#                     'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
#                     'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
#                     'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
#                     'profile': candidate.profile,
#                     'status': candidate.status
#                 }
#                 candidates_data.append(candidate_data)
#             elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
#                 unsuccessful_closures += 1

#         # Calculate average days to close
#         average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

#         # Calculate percentage of onboarded candidates
#         percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

#         # Append summary and candidates data for the recruiter
#         recruiter_data = {
#             'recruiter_name': recruiter_name,
#             'candidates': candidates_data,
#             'total_days_to_close': total_days_to_close,
#             'count_of_screening_candidates': total_screening_candidates,
#             'count_of_onboarded_positions': count_of_onboarded_positions,
#             'unsuccessful_closures': unsuccessful_closures,
#             'average_days_to_close': average_days_to_close,
#             'percentage_onboarded': percentage_onboarded,
#             'total_candidates_count': total_candidates
#         }

#         result[recruiter_name] = recruiter_data

#     # Sort recruiters by percentage of onboarded candidates
#     # ranked_recruiters = sorted(result.values(), key=lambda x: x['percentage_onboarded'], reverse=True)
#     ranked_recruiters = sorted(result.values(), key=lambda x: x['average_days_to_close'], reverse=True)

#     # Assign rankings
#     rank = 0
#     for data in ranked_recruiters:
#         rank += 1
#         if data['count_of_onboarded_positions'] > 0:
#             data['ranking'] = rank
#         else:
#             data['ranking'] = 0

#     return jsonify(ranked_recruiters)

    



# import itertools


# @app.route('/generate_excel', methods=['POST'])
# def generate_excel():
#     data = request.json

#     if not data:
#         return jsonify({'error': 'No JSON data provided'})

#     user_id = data.get('user_id')
#     from_date_str = data.get('from_date')
#     to_date_str = data.get('to_date')
#     recruiter_names = data.get('recruiter_names', [])

#     if not recruiter_names:
#         return jsonify({'error': 'Please select any Recruiter'})

#     try:
#         from_date = datetime.strptime(from_date_str, "%d-%m-%Y")
#         to_date = datetime.strptime(to_date_str, "%d-%m-%Y")
#     except ValueError:
#         return jsonify({'error': 'Invalid date format. Please use DD-MM-YYYY format.'})

#     session = Session()

#     # Generate all recruiter-date combinations within the specified range
#     all_recruiter_date_combinations = list(itertools.product(recruiter_names, pd.date_range(from_date, to_date, freq='D')))

#     # Sort combinations by date (earliest to latest)
#     all_recruiter_date_combinations.sort(key=lambda x: x[1])

#     # Fetch all candidates within the specified date range
#     candidates_query = session.query(
#         Candidate.recruiter,
#         Candidate.date_created,
#         func.count(Candidate.id).label('count')
#     ).filter(
#         Candidate.recruiter.in_(recruiter_names),
#         Candidate.date_created >= from_date,
#         Candidate.date_created <= to_date
#     ).group_by(Candidate.recruiter, Candidate.date_created).all()

#     # Convert candidate data to a dictionary for easy lookup
#     candidate_data_dict = {
#         (row.recruiter, row.date_created.strftime("%d-%m-%Y")): row.count
#         for row in candidates_query
#     }

#     # Prepare data for the report
#     data = []
#     for recruiter, date in all_recruiter_date_combinations:
#         count = candidate_data_dict.get((recruiter, date.strftime("%d-%m-%Y")), 0)
#         data.append({
#             "recruiter": recruiter,
#             "date_created": date.strftime("%d-%m-%Y"),
#             "count": count
#         })

#     # Convert data to DataFrame
#     data_df = pd.DataFrame(data)

#     # Creating a pivot table for the report
#     pivot_table = data_df.pivot_table(
#         index='recruiter',
#         columns='date_created',
#         values='count',
#         aggfunc='sum',
#         fill_value=0,
#         margins=True,
#         margins_name='Grand Total'
#     )

#     # Reordering columns to match from_date to to_date
#     date_columns = pd.date_range(from_date, to_date, freq='D').strftime('%d-%m-%Y')
#     pivot_table = pivot_table.reindex(columns=date_columns)

#     # Convert pivot table to JSON string
#     pivot_table_json = pivot_table.to_json(orient='index')

#     # Convert JSON string to dictionary for manipulation
#     pivot_table_dict = json.loads(pivot_table_json)

#     # Add 'from_date' and 'to_date' strings to the dictionary
#     pivot_table_dict['from_date_str'] = from_date_str
#     pivot_table_dict['to_date_str'] = to_date_str

#     # Convert dictionary back to JSON string
#     final_json_response = json.dumps(pivot_table_dict)

#     return jsonify(final_json_response)

import itertools

@app.route('/generate_excel', methods=['POST'])
def generate_excel():
    data = request.json

    if not data:
        return jsonify({'error': 'No JSON data provided'})

    user_id = data.get('user_id')
    from_date_str = data.get('from_date')
    to_date_str = data.get('to_date')
    recruiter_names = data.get('recruiter_names', [])

    if not recruiter_names:
        return jsonify({'error': 'Please select any Recruiter'})

    try:
        from_date = datetime.strptime(from_date_str, "%d-%m-%Y")
        to_date = datetime.strptime(to_date_str, "%d-%m-%Y")
    except ValueError:
        return jsonify({'error': 'Invalid date format. Please use DD-MM-YYYY format.'})

    session = Session()
    
    # Generate all recruiter-date combinations within the specified range
    all_recruiter_date_combinations = list(itertools.product(recruiter_names, pd.date_range(from_date, to_date, freq='D')))

    # Sort combinations by date (earliest to latest)
    all_recruiter_date_combinations.sort(key=lambda x: x[1])

    # Fetch all candidates within the specified date range
    candidates_query = session.query(Candidate.recruiter, Candidate.date_created, func.count(Candidate.id).label('count')).filter(
        Candidate.recruiter.in_(recruiter_names),
        Candidate.date_created >= from_date,
        Candidate.date_created <= to_date
    ).group_by(Candidate.recruiter, Candidate.date_created).all()

    # Convert candidate data to a dictionary for easy lookup
    candidate_data_dict = {(row.recruiter, row.date_created.strftime("%d-%m-%Y")): row.count for row in candidates_query}

    # Prepare data for the report
    data = []
    for recruiter, date in all_recruiter_date_combinations:
        count = candidate_data_dict.get((recruiter, date.strftime("%d-%m-%Y")), 0)
        data.append({
            "recruiter": recruiter,
            "date_created": date.strftime("%d-%m-%Y"),
            "count": count
        })

    # Convert data to DataFrame
    data_df = pd.DataFrame(data)

    # Creating a pivot table for the report
    pivot_table = data_df.pivot_table(index='recruiter', columns='date_created', values='count', aggfunc='sum',
                                      fill_value=0, margins=True, margins_name='Grand Total')

    # # Reordering columns to match from_date to to_date
    # date_columns = pd.date_range(from_date, to_date, freq='D').strftime('%d-%m-%Y')
    # pivot_table = pivot_table[date_columns]

    # Convert pivot table to JSON for response
    styled_pivot_table_json = pivot_table.to_json()

    return jsonify({
        'user_id': user_id,
        'from_date_str': from_date_str,
        'to_date_str': to_date_str,
        'pivot_table': styled_pivot_table_json
    })


# def re_send_notification(recruiter_email, job_id):
#     msg = Message('Job Update Notification', sender='ganesh.s@makonissoft.com', recipients=[recruiter_email])
#     msg.body = f'Hello,\n\nThe job post with ID {job_id} has been updated.\n\nPlease check your dashboard for more details.'
#     mail.send(msg)


def job_removed_send_notification(recruiter_email, new_recruiter_name, job_data, job_id):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                Job Removal Notification
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>The job post with ID <b>{job_id}</b> has been removed from your login.</p>
            <p>Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 20%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check your dashboard for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        f'Job Removal Notification: Job ID {job_id}',
        sender='kanuparthisaiganesh582@gmail.com',
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)



def job_updated_send_notification(recruiter_email, new_recruiter_name, job_data, job_id):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                Job Update Notification
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>The job post with ID <b>{job_id}</b> has been updated.</p>
            <p>Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 20%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check your dashboard for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        f'Job Update Notification: Job ID {job_id}',
        sender='kanuparthisaiganesh582@gmail.com',
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)


@app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
def edit_job_post(job_post_id):
    try:
        data = request.json
        user_id = data.get('user_id')
        
        # Retrieve the user
        user = User.query.filter_by(id=user_id).first()
        
        if user and user.user_type == 'management':
            # Retrieve the job post to be edited
            job_post = JobPost.query.get(job_post_id)
            
            if job_post:
                old_recruiter_usernames = set(job_post.recruiter.split(', ')) if isinstance(job_post.recruiter, str) else set()

                # Check if any field except 'recruiter' is updated
                fields_updated = set(data.keys()) - {'recruiter'}
                if fields_updated:
                    # Update job post fields
                    job_post.client = data.get('client', job_post.client)
                    job_post.experience_min = data.get('experience_min', job_post.experience_min)
                    job_post.experience_max = data.get('experience_max', job_post.experience_max)
                    job_post.budget_min = data.get('budget_min', job_post.budget_min)
                    job_post.budget_max = data.get('budget_max', job_post.budget_max)
                    job_post.location = data.get('location', job_post.location)
                    job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
                    job_post.notice_period = data.get('notice_period', job_post.notice_period)
                    job_post.role = data.get('role', job_post.role)
                    job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
                    job_post.mode = data.get('mode', job_post.mode)
                    job_post.job_status = data.get('job_status', job_post.job_status)
                    job_post.skills = data.get('skills', job_post.skills)
                    
                    recruiters = data.get('recruiter', job_post.recruiter)
                    recruiters = set(recruiters if isinstance(recruiters, list) else [recruiters])
                    job_post.recruiter = ', '.join(list(recruiters))
                    
                    job_type = data.get('Job_Type')
                    job_post.job_type = job_type
                    if job_type == 'Contract':
                        job_post.contract_in_months = data.get('Job_Type_details')
                    
                    # Handle jd_pdf field
                    jd_pdf = data.get('jd_pdf')
                    if jd_pdf is not None:
                        jd_binary = base64.b64decode(jd_pdf)
                        job_post.jd_pdf = jd_binary
                        job_post.jd_pdf_present = True

                    # Update job post in the database
                    db.session.commit()
                    
                    # Create notification records for each recruiter
                    if job_post.recruiter:
                        for recruiter in recruiters:
                            notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
                            if notification:
                                notification.num_notification += 1
                            else:
                                new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
                                db.session.add(new_notification)
                                new_notification.num_notification = 1
                    
                    # Update candidate details
                    candidates = Candidate.query.filter_by(job_id=job_post_id).all()
                    for candidate in candidates:
                        candidate.client = job_post.client
                        candidate.profile = job_post.role

                    db.session.commit()
                    
                    return jsonify({
                        'status': 'success',
                        "message": "Job post details updated successfully",
                        'job_post_id': job_post_id,
                        'old_recruiter_usernames': list(old_recruiter_usernames),
                        'new_recruiter_usernames': list(recruiters)
                    }), 200
                else:
                    return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
            else:
                return jsonify({'status': 'error', "message": "Job post not found"})
        else:
            return jsonify({'status': 'error', "message": "Unauthorized"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/send_edit_notifications', methods=['POST'])
def send_edit_notifications():
    try:
        data = request.json
        job_post_id = data.get('job_post_id')
        old_recruiter_usernames = set(data.get('old_recruiter_usernames', []))
        new_recruiter_usernames = set(data.get('new_recruiter_usernames', []))

        if not job_post_id:
            return jsonify({'status': 'error', 'message': 'job_post_id is required'}), 400

        job_post = JobPost.query.filter_by(id=job_post_id).first()
        if not job_post:
            return jsonify({'status': 'error', 'message': 'Job post not found'}), 404

        job_data = f"<tr><td>{job_post.id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

        unchanged_recruiters = set(old_recruiter_usernames) & set(new_recruiter_usernames)
        removed_recruiters = set(old_recruiter_usernames) - set(new_recruiter_usernames)
        added_recruiters = set(new_recruiter_usernames) - set(old_recruiter_usernames)

        all_recruiters = User.query.filter(
            User.username.in_(set(new_recruiter_usernames).union(set(old_recruiter_usernames))),
            User.user_type == 'recruiter',
            User.is_active == True,
            User.is_verified == True
        ).all()

        all_recruiter_emails = {recruiter.username: recruiter.email for recruiter in all_recruiters}

        for recruiter_name in unchanged_recruiters:
            email = all_recruiter_emails.get(recruiter_name)
            if email:
                job_updated_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post.id)

        for recruiter_name in removed_recruiters:
            email = all_recruiter_emails.get(recruiter_name)
            if email:
                job_removed_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post.id)

        for recruiter_name in added_recruiters:
            email = all_recruiter_emails.get(recruiter_name)
            if email:
                post_job_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data)
        
        return jsonify({'status': 'success', 'message': 'Notifications sent successfully'}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
        

# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 old_recruiter_usernames = set(job_post.recruiter.split(', ')) if isinstance(job_post.recruiter, str) else set()

#                 # Check if any field except 'recruiter' is updated
#                 fields_updated = set(data.keys()) - {'recruiter'}
#                 if fields_updated:
#                     # Update job post fields
#                     job_post.client = data.get('client', job_post.client)
#                     job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                     job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                     job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                     job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                     job_post.location = data.get('location', job_post.location)
#                     job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                     job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                     job_post.role = data.get('role', job_post.role)
#                     job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                     job_post.mode = data.get('mode', job_post.mode)
#                     job_post.job_status = data.get('job_status', job_post.job_status)
#                     job_post.skills = data.get('skills', job_post.skills)
                    
#                     recruiters = data.get('recruiter', job_post.recruiter)
#                     recruiters = set(recruiters if isinstance(recruiters, list) else [recruiters])
#                     job_post.recruiter = ', '.join(list(recruiters))
                    
#                     job_type = data.get('Job_Type')
#                     job_post.job_type = job_type
#                     if job_type == 'Contract':
#                         job_post.contract_in_months = data.get('Job_Type_details')
                    
#                     # Handle jd_pdf field
#                     jd_pdf = data.get('jd_pdf')
#                     if jd_pdf is not None:
#                         jd_binary = base64.b64decode(jd_pdf)
#                         job_post.jd_pdf = jd_binary
#                         job_post.jd_pdf_present = True

#                     # Update job post in the database
#                     db.session.commit()
                    
#                     # Create notification records for each recruiter
#                     if job_post.recruiter:
#                         for recruiter in recruiters:
#                             notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                             if notification:
#                                 notification.num_notification += 1
#                             else:
#                                 new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                                 db.session.add(new_notification)
#                                 new_notification.num_notification = 1
                    
#                     # Update candidate details
#                     candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                     for candidate in candidates:
#                         candidate.client = job_post.client
#                         candidate.profile = job_post.role

#                     db.session.commit()
                    
#                     return jsonify({'status': 'success', "message": "Job post details updated successfully", 'job_post_id': job_post_id}), 200
#                 else:
#                     return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
#             else:
#                 return jsonify({'status': 'error', "message": "Job post not found"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# @app.route('/send_edit_notifications', methods=['POST'])
# def send_edit_notifications():
#     try:
#         data = request.json
#         job_post_id = data.get('job_post_id')
#         old_recruiter_usernames= data.get('old_recruiter_usernames')

#         if not job_post_id:
#             return jsonify({'status': 'error', 'message': 'job_post_id is required'}), 400

#         job_post = JobPost.query.filter_by(id=job_post_id).first()
#         if not job_post:
#             return jsonify({'status': 'error', 'message': 'Job post not found'}), 404

#         job_data = f"<tr><td>{job_post.id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

#         # old_recruiter_usernames = set(job_post.recruiter.split(', ')) if isinstance(job_post.recruiter, str) else set()

#         new_recruiters = User.query.filter(User.username.in_(old_recruiter_usernames)).all()
#         new_recruiter_usernames = {recruiter.username for recruiter in new_recruiters}

#         unchanged_recruiters = old_recruiter_usernames & new_recruiter_usernames
#         removed_recruiters = old_recruiter_usernames - new_recruiter_usernames
#         added_recruiters = new_recruiter_usernames - old_recruiter_usernames

#         all_recruiters = User.query.filter(
#             User.username.in_(new_recruiter_usernames.union(old_recruiter_usernames)),
#             User.user_type == 'recruiter',
#             User.is_active == True,
#             User.is_verified == True
#         ).all()

#         all_recruiter_emails = {recruiter.username: recruiter.email for recruiter in all_recruiters}

#         for recruiter_name in unchanged_recruiters:
#             email = all_recruiter_emails.get(recruiter_name)
#             if email:
#                 job_updated_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post.id)

#         for recruiter_name in removed_recruiters:
#             email = all_recruiter_emails.get(recruiter_name)
#             if email:
#                 job_removed_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post.id)

#         for recruiter_name in added_recruiters:
#             email = all_recruiter_emails.get(recruiter_name)
#             if email:
#                 post_job_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data)
        
#         return jsonify({'status': 'success', 'message': 'Notifications sent successfully'}), 200
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 old_recruiter_usernames = set(job_post.recruiter.split(', ')) if isinstance(job_post.recruiter, str) else set()

#                 # Check if any field except 'recruiter' is updated
#                 fields_updated = set(data.keys()) - {'recruiter'}
#                 if fields_updated:
#                     # Update job post fields
#                     job_post.client = data.get('client', job_post.client)
#                     job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                     job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                     job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                     job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                     job_post.location = data.get('location', job_post.location)
#                     job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                     job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                     job_post.role = data.get('role', job_post.role)
#                     job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                     job_post.mode = data.get('mode', job_post.mode)
#                     job_post.job_status = data.get('job_status', job_post.job_status)
#                     job_post.skills = data.get('skills', job_post.skills)
                    
#                     recruiters = data.get('recruiter', job_post.recruiter)
#                     recruiters = set(recruiters if isinstance(recruiters, list) else [recruiters])
#                     job_post.recruiter = ', '.join(list(recruiters))
                    
#                     job_type = data.get('Job_Type')
#                     job_post.job_type = job_type
#                     if job_type == 'Contract':
#                         job_post.contract_in_months = data.get('Job_Type_details')
                    
#                     # Handle jd_pdf field
#                     jd_pdf = data.get('jd_pdf')
#                     if jd_pdf is not None:
#                         jd_binary = base64.b64decode(jd_pdf)
#                         job_post.jd_pdf = jd_binary
#                         job_post.jd_pdf_present = True

#                     # Update job post in the database
#                     db.session.commit()
                    
#                     # Create notification records for each recruiter
#                     if job_post.recruiter:
#                         for recruiter in recruiters:
#                             notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                             if notification:
#                                 notification.num_notification += 1
#                             else:
#                                 new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                                 db.session.add(new_notification)
#                                 new_notification.num_notification = 1
                    
#                     # Update candidate details
#                     candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                     for candidate in candidates:
#                         candidate.client = job_post.client
#                         candidate.profile = job_post.role

#                     db.session.commit()
                    
#                     # Retrieve the email addresses of the recruiters
#                     recruiter_emails = [recruiter.email for recruiter in User.query.filter(
#                         User.username.in_(recruiters),
#                         User.user_type == 'recruiter',
#                         User.is_active == True,
#                         User.is_verified == True
#                     )]
                    
#                     job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"
                    
#                     old_recruiters = User.query.filter(User.username.in_(old_recruiter_usernames)).all()

#                     new_recruiters = User.query.filter(User.username.in_(recruiters)).all()
#                     new_recruiter_usernames = {recruiter.username for recruiter in new_recruiters}

#                     unchanged_recruiters = old_recruiter_usernames & new_recruiter_usernames
#                     removed_recruiters = old_recruiter_usernames - new_recruiter_usernames
#                     added_recruiters = new_recruiter_usernames - old_recruiter_usernames

#                     all_recruiters = User.query.filter(
#                         User.username.in_(new_recruiter_usernames.union(old_recruiter_usernames)),
#                         User.user_type == 'recruiter',
#                         User.is_active == True,
#                         User.is_verified == True
#                     ).all()

#                     all_recruiter_emails = {recruiter.username: recruiter.email for recruiter in all_recruiters}

#                     for recruiter_name in unchanged_recruiters:
#                         email = all_recruiter_emails.get(recruiter_name)
#                         if email:
#                             job_updated_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post_id)

#                     for recruiter_name in removed_recruiters:
#                         email = all_recruiter_emails.get(recruiter_name)
#                         if email:
#                             job_removed_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post_id)

#                     for recruiter_name in added_recruiters:
#                         email = all_recruiter_emails.get(recruiter_name)
#                         if email:
#                             post_job_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data)
                    
#                     return jsonify({'status': 'success', "message": "Job post details updated successfully"})
#                 else:
#                     if old_recruiter_usernames:
#                         job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"
#                         for recruiter_name in old_recruiter_usernames:
#                             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#                             if recruiter:
#                                 job_updated_send_notification(recruiter.email, recruiter.username, user.username, job_data, job_post_id)
                    
#                     return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
#             else:
#                 return jsonify({'status': 'error', "message": "Job post not found"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500



# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 old_recruiter_usernames = job_post.recruiter.split(', ') if isinstance(job_post.recruiter, str) else []

#                 # Check if any field except 'recruiter' is updated
#                 fields_updated = set(data.keys()) - {'recruiter'}
#                 if fields_updated:
#                     # Update job post fields
#                     job_post.client = data.get('client', job_post.client)
#                     job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                     job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                     job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                     job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                     job_post.location = data.get('location', job_post.location)
#                     job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                     job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                     job_post.role = data.get('role', job_post.role)
#                     job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                     job_post.mode = data.get('mode', job_post.mode)
#                     job_post.job_status = data.get('job_status', job_post.job_status)
#                     job_post.skills = data.get('skills', job_post.skills)
                    
#                     recruiters = data.get('recruiter', job_post.recruiter)
#                     recruiters = recruiters if isinstance(recruiters, list) else [recruiters]
#                     job_post.recruiter = ', '.join(list(set(recruiters)))
                    
#                     job_type = data.get('Job_Type')
#                     if job_type == 'Contract':
#                         job_post.contract_in_months = data.get('Job_Type_details')
                    
#                     # Handle jd_pdf field
#                     jd_pdf = data.get('jd_pdf')
#                     if jd_pdf is not None:
#                         jd_binary = base64.b64decode(jd_pdf)
#                         job_post.jd_pdf = jd_binary
#                         job_post.jd_pdf_present = True
#                     else:
#                         job_post.jd_pdf_present = False

#                     # Update job post in the database
#                     db.session.commit()
                    
#                     # Create notification records for each recruiter
#                     if job_post.recruiter:
#                         recruiters = list(set(job_post.recruiter.split(', ')))
#                         for recruiter in recruiters:
#                             notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                             if notification:
#                                 notification.num_notification += 1
#                             else:
#                                 new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                                 db.session.add(new_notification)
#                                 new_notification.num_notification = 1
                    
#                     # Update candidate details
#                     candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                     for candidate in candidates:
#                         candidate.client = job_post.client
#                         candidate.profile = job_post.role

#                     db.session.commit()
                    
#                     # Retrieve the email addresses of the recruiters
#                     recruiter_emails = [recruiter.email for recruiter in User.query.filter(
#                         User.username.in_(recruiters),
#                         User.user_type == 'recruiter',
#                         User.is_active == True,
#                         User.is_verified == True
#                     )]
                    
#                     job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

#                     for email in recruiter_emails:
#                         if email in [r.email for r in User.query.filter(User.username.in_(old_recruiter_usernames))]:
#                             job_updated_send_notification(recruiter_email=email, new_recruiter_name=user.username, job_data=job_data, job_id=job_post_id)
#                         else:
#                             post_job_send_notification(recruiter_email=email, new_recruiter_name=user.username, job_data=job_data)
                    
#                     return jsonify({'status': 'success', "message": "Job post details updated successfully"})
#                 else:
#                     if old_recruiter_usernames:
#                         job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"
#                         for recruiter_name in old_recruiter_usernames:
#                             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#                             if recruiter:
#                                 re_send_notification(recruiter.email, recruiter.username, user.username, job_data, job_post_id)
                    
#                     return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
#             else:
#                 return jsonify({'status': 'error', "message": "Job post not found"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Store the old recruiters' usernames
#                 old_recruiter_usernames = job_post.recruiter.split(',') if job_post.recruiter else []

#                 # Check if any field except 'recruiter' is updated
#                 fields_updated = set(data.keys()) - {'recruiter'}
#                 if fields_updated:
#                     # Update job post fields
#                     job_post.client = data.get('client', job_post.client)
#                     job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                     job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                     job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                     job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                     job_post.location = data.get('location', job_post.location)
#                     job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                     job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                     job_post.role = data.get('role', job_post.role)
#                     job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                     job_post.mode = data.get('mode', job_post.mode)
#                     job_post.job_status = data.get('job_status', job_post.job_status)
#                     job_post.skills = data.get('skills', job_post.skills)
#                     recruiters = data.get('recruiter', job_post.recruiter)
#                     if recruiters:
#                         # Ensure recruiters are unique
#                         unique_recruiters = list(set(recruiters))
#                         job_post.recruiter = ', '.join(unique_recruiters)
                    
#                     job_type = data.get('Job_Type')
#                     if job_type == 'Contract':
#                         job_post.contract_in_months = data.get('Job_Type_details')
                    
#                     # Handle jd_pdf field
#                     jd_pdf = data.get('jd_pdf')
#                     if jd_pdf is not None:
#                         jd_binary = base64.b64decode(jd_pdf)
#                         job_post.jd_pdf = jd_binary
#                         job_post.jd_pdf_present = True
#                     else:
#                         job_post.jd_pdf_present = False

                    

#                     # Update job post in the database
#                     db.session.commit()
                    
#                     # Create notification records for each recruiter
#                     if job_post.recruiter:
#                         recruiters = list(set(job_post.recruiter.split(', ')))
#                         for recruiter in recruiters:
#                             notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                             if notification:
#                                 notification.num_notification += 1
#                             else:
#                                 new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                                 db.session.add(new_notification)
#                                 new_notification.num_notification = 1
                    
#                     # Update candidate details
#                     candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                     for candidate in candidates:
#                         candidate.client = job_post.client
#                         candidate.profile = job_post.role

#                     db.session.commit()
                    
#                     # Retrieve the email addresses of the recruiters
#                     recruiter_emails = [recruiter.email for recruiter in User.query.filter(
#                         User.username.in_(recruiters),
#                         User.user_type == 'recruiter',
#                         User.is_active == True,
#                         User.is_verified == True
#                     )]
                    
#                     job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

#                     for email in recruiter_emails:
#                         if email in [r.email for r in User.query.filter(User.username.in_(old_recruiter_usernames))]:
#                             job_updated_send_notification(recruiter_email=email, new_recruiter_name=user.username, job_data=job_data, job_id=job_post_id)
#                         else:
#                             post_job_send_notification(recruiter_email=email, new_recruiter_name=user.username, job_data=job_data, job_id=job_post_id)
                    
#                     return jsonify({'status': 'success', "message": "Job post details updated successfully"})
#                 else:
#                     # No fields updated other than recruiter, so only send notifications to old recruiters
#                     if old_recruiter_usernames:
#                         job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"
#                         for recruiter_name in old_recruiter_usernames:
#                             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#                             if recruiter:
#                                 re_send_notification(recruiter.email, user.username, job_data, job_post_id)
                    
#                     return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
#             else:
#                 return jsonify({'status': 'error', "message": "Job post not found"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500



# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Store the old recruiters' usernames
#                 old_recruiter_usernames = job_post.recruiter.split(',') if job_post.recruiter else []

#                 # Check if any field except 'recruiter' is updated
#                 fields_updated = set(data.keys()) - {'recruiter'}
#                 if fields_updated:
#                     # Update job post fields
#                     job_post.client = data.get('client', job_post.client)
#                     job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                     job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                     job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                     job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                     job_post.location = data.get('location', job_post.location)
#                     job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                     job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                     job_post.role = data.get('role', job_post.role)
#                     job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                     job_post.mode = data.get('mode', job_post.mode)
#                     job_post.job_status = data.get('job_status', job_post.job_status)
#                     job_post.skills = data.get('skills', job_post.skills)
            
#                     job_type = data.get('Job_Type')
#                     if job_type == 'Contract':
#                         job_post.contract_in_months = data.get('Job_Type_details')
                    
#                     # Handle jd_pdf field
#                     jd_pdf = data.get('jd_pdf')
#                     if jd_pdf is not None:
#                         jd_binary = base64.b64decode(jd_pdf)
#                         job_post.jd_pdf = jd_binary
#                         job_post.jd_pdf_present = True
#                     else:
#                         job_post.jd_pdf_present = False

#                     # Update job post in the database
#                     db.session.commit()
                    
#                     # Create notification records for each recruiter
#                     if job_post.recruiter:
#                         recruiters = list(set(job_post.recruiter.split(', ')))
#                         for recruiter in recruiters:
#                             notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                             if notification:
#                                 notification.num_notification += 1
#                             else:
#                                 new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                                 db.session.add(new_notification)
#                                 new_notification.num_notification = 1
                    
#                     # Update candidate details
#                     candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                     for candidate in candidates:
#                         candidate.client = job_post.client
#                         candidate.profile = job_post.role

#                     db.session.commit()
                    
#                     # Retrieve the email addresses of the recruiters
#                     recruiter_emails = [recruiter.email for recruiter in User.query.filter(
#                         User.username.in_(recruiters),
#                         User.user_type == 'recruiter',
#                         User.is_active == True,
#                         User.is_verified == True
#                     )]
                    
#                     job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

#                     for email in recruiter_emails:
#                         if email in [r.email for r in User.query.filter(User.username.in_(old_recruiter_usernames))]:
#                             re_send_notification(recruiter_email=email, new_recruiter_name=user.username, job_data=job_data, job_id=job_post_id)
#                         else:
#                             post_job_send_notification(recruiter_email=email, new_recruiter_name=user.username, job_data=job_data, job_id=job_post_id)
                    
#                     return jsonify({'status': 'success', "message": "Job post details updated successfully"})
#                 else:
#                     # No fields updated other than recruiter, so only send notifications to old recruiters
#                     if old_recruiter_usernames:
#                         job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"
#                         for recruiter_name in old_recruiter_usernames:
#                             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#                             if recruiter:
#                                 re_send_notification(recruiter.email, user.username, job_data, job_post_id)
                    
#                     return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
#             else:
#                 return jsonify({'status': 'error', "message": "Job post not found"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Update job post fields
#                 job_post.client = data.get('client', job_post.client)
#                 job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                 job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                 job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                 job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                 job_post.location = data.get('location', job_post.location)
#                 job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                 job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                 job_post.role = data.get('role', job_post.role)
#                 job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                 job_post.mode = data.get('mode', job_post.mode)
#                 job_post.job_status = data.get('job_status', job_post.job_status)
#                 # job_post.job_type = data.get('Job_Type', job_post.job_type)  # Updated key 'job_type' to 'Job_Type'
#                 job_post.skills = data.get('skills', job_post.skills)
#                 # job_post.jd_pdf = data.get('jd_pdf', job_post.jd_pdf)

#                 job_type = data.get('Job_Type')
#                 if job_type == 'Contract':
#                     job_post.contract_in_months = data.get('Job_Type_details')
#                     # job_type = job_type + '(' + Job_Type_details + ' Months )'
#                 else:
#                     pass
                
#                 # Handle jd_pdf field
#                 jd_pdf = data.get('jd_pdf')
#                 if jd_pdf is not None:
#                     # Decode the base64 encoded PDF file
#                     jd_binary = base64.b64decode(jd_pdf)
#                     # Update job_post with the decoded binary data
#                     job_post.jd_pdf = jd_binary
#                     # Set jd_pdf_present to True since PDF is present
#                     job_post.jd_pdf_present = True
#                 else:
#                     # If jd_pdf is None, set jd_pdf_present to False
#                     job_post.jd_pdf_present = False
                    
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     # Ensure recruiters are unique
#                     unique_recruiters = list(set(recruiters))
#                     job_post.recruiter = ', '.join(unique_recruiters)
                    
#                 # Update data_updated_date and data_updated_time
#                 # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
#                 # job_post.data_updated_date = current_datetime.date()
#                 # job_post.data_updated_time = current_datetime.time()
                
#                 # Update job post in the database
#                 db.session.commit()
                
#                 # Iterate over recruiters and create notification records for each
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     unique_recruiters = list(set(recruiters))
#                     for recruiter in unique_recruiters:
#                         # Check if a notification exists for the job post and recruiter combination
#                         notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                         if notification:
#                             # If notification exists, increment num_notification by 1
#                             notification.num_notification += 1
#                         else:
#                             # If notification does not exist, create a new record with num_notification set to 1
#                             new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                             db.session.add(new_notification)
#                             new_notification.num_notification = 1
                
#                 # Update candidate details if job_post_id is present in the Candidate table
#                 candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                 for candidate in candidates:
#                     candidate.client = job_post.client
#                     candidate.profile = job_post.role
#                     # Add more fields if necessary

#                 db.session.commit()
                
#                 # Retrieve the email addresses of the recruiters
#                 recruiter_emails = [recruiter.email for recruiter in User.query.filter(User.username.in_(unique_recruiters),
#                                                                                       User.user_type == 'recruiter',
#                                                                                       User.is_active == True,
#                                                                                       User.is_verified == True)]
#                 job_data = f"<tr><td>{job_post_id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"
                
#                 for email in recruiter_emails:
#                     if email in [r.email for r in User.query.filter(User.username.in_(job_post.recruiter.split(',')))]:
#                         # If recruiter already associated, send a job update notification
                        
#                         # re_send_notification(recruiter_email=email, job_id=job_post_id)
#                         job_updated_send_notification(recruiter_email, new_recruiter_name, job_data, job_post_id)
#                     else:
#                         # If new recruiter added, send a new job notification
                        
#                         # send_notification(recruiter_email=email)
#                         post_job_send_notification(recruiter_email, new_recruiter_name, job_data,job_post_id)
                        
                
#                 # Return success message
#                 return jsonify({'status': 'success', "message": "Job post details updated successfully"})
#             else:
#                 return jsonify({'status': 'error',"message": "Job post are not updated successfully"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Update job post fields
#                 job_post.client = data.get('client', job_post.client)
#                 job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                 job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                 job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                 job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                 job_post.location = data.get('location', job_post.location)
#                 job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                 job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                 job_post.role = data.get('role', job_post.role)
#                 job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                 job_post.mode = data.get('mode', job_post.mode)
#                 job_post.job_status = data.get('job_status', job_post.job_status)
#                 job_post.job_type = data.get('Job_Type', job_post.job_type)  # Updated key 'job_type' to 'Job_Type'
#                 job_post.skills = data.get('skills', job_post.skills)
#                 # job_post.jd_pdf = data.get('jd_pdf', job_post.jd_pdf)

#                 job_type = data.get('job_type')
#                 if job_type == 'Contract':
#                     job_post.contract_in_months = data.get('Job_Type_details')
#                     # job_type = job_type + '(' + Job_Type_details + ' Months )'
#                 # else:
#                 #     pass
                
#                 # Handle jd_pdf field
#                 jd_pdf = data.get('jd_pdf')
#                 if jd_pdf is not None:
#                     # Decode the base64 encoded PDF file
#                     jd_binary = base64.b64decode(jd_pdf)
#                     # Update job_post with the decoded binary data
#                     job_post.jd_pdf = jd_binary
#                     # Set jd_pdf_present to True since PDF is present
#                     job_post.jd_pdf_present = True
#                 else:
#                     # If jd_pdf is None, set jd_pdf_present to False
#                     job_post.jd_pdf_present = False
                    
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     # Ensure recruiters are unique
#                     unique_recruiters = list(set(recruiters))
#                     job_post.recruiter = ', '.join(unique_recruiters)
                    
#                 # Update data_updated_date and data_updated_time
#                 current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
#                 job_post.data_updated_date = current_datetime.date()
#                 job_post.data_updated_time = current_datetime.time()
                
#                 # Update job post in the database
#                 db.session.commit()
                
#                 # Iterate over recruiters and create notification records for each
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     unique_recruiters = list(set(recruiters))
#                     for recruiter in unique_recruiters:
#                         # Check if a notification exists for the job post and recruiter combination
#                         notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                         if notification:
#                             # If notification exists, increment num_notification by 1
#                             notification.num_notification += 1
#                         else:
#                             # If notification does not exist, create a new record with num_notification set to 1
#                             new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                             db.session.add(new_notification)
#                             new_notification.num_notification = 1
                
#                 # Update candidate details if job_post_id is present in the Candidate table
#                 candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                 for candidate in candidates:
#                     candidate.client = job_post.client
#                     candidate.profile = job_post.role
#                     # Add more fields if necessary

#                 db.session.commit()
                
#                 # Send notifications
#                 for recruiter in unique_recruiters:
#                     if recruiter in job_post.recruiter.split(','):
#                         # If recruiter already associated, send a job update notification
#                         re_send_notification(recruiter_email=recruiter, job_id=job_post_id)
#                     else:
#                         # If new recruiter added, send a new job notification
#                         send_notification(recruiter_email=recruiter)
                
#                 # Return success message
#                 return jsonify({'status': 'success', "message": "Job post details updated successfully"})
#             else:
#                 return jsonify({'status': 'error',"message": "Job post are not updated successfully"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Update job post fields
#                 job_post.client = data.get('client', job_post.client)
#                 job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                 job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                 job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                 job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                 job_post.location = data.get('location', job_post.location)
#                 job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                 job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                 job_post.role = data.get('role', job_post.role)
#                 job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                 job_post.mode = data.get('mode', job_post.mode)
#                 job_post.job_status = data.get('job_status', job_post.job_status)
#                 job_post.job_type = data.get('Job_Type', job_post.job_type)  # Updated key 'job_type' to 'Job_Type'
#                 job_post.skills = data.get('skills', job_post.skills)
#                 # job_post.jd_pdf = data.get('jd_pdf', job_post.jd_pdf)

#                 job_type = data.get('job_type')
#                 if job_type == 'Contract':
#                     job_post.contract_in_months = data.get('Job_Type_details')
#                     # job_type = job_type + '(' + Job_Type_details + ' Months )'
#                 # else:
#                 #     pass
                
#                 # Handle jd_pdf field
#                 jd_pdf = data.get('jd_pdf')
#                 if jd_pdf is not None:
#                     # Decode the base64 encoded PDF file
#                     jd_binary = base64.b64decode(jd_pdf)
#                     # Update job_post with the decoded binary data
#                     job_post.jd_pdf = jd_binary
#                     # Set jd_pdf_present to True since PDF is present
#                     job_post.jd_pdf_present = True
#                 else:
#                     # If jd_pdf is None, set jd_pdf_present to False
#                     job_post.jd_pdf_present = False
                    
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     # Ensure recruiters are unique
#                     unique_recruiters = list(set(recruiters))
#                     job_post.recruiter = ', '.join(unique_recruiters)
                    
#                 # Update data_updated_date and data_updated_time
#                 current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
#                 job_post.data_updated_date = current_datetime.date()
#                 job_post.data_updated_time = current_datetime.time()
                
#                 # Update job post in the database
#                 db.session.commit()
                
#                 # Iterate over recruiters and create notification records for each
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     unique_recruiters = list(set(recruiters))
#                     for recruiter in unique_recruiters:
#                         # Check if a notification exists for the job post and recruiter combination
#                         notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                         if notification:
#                             # If notification exists, increment num_notification by 1
#                             notification.num_notification += 1
#                         else:
#                             # If notification does not exist, create a new record with num_notification set to 1
#                             new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                             db.session.add(new_notification)
#                             new_notification.num_notification = 1
                
#                 # Update candidate details if job_post_id is present in the Candidate table
#                 candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                 for candidate in candidates:
#                     candidate.client = job_post.client
#                     candidate.profile = job_post.role
#                     # Add more fields if necessary

#                 db.session.commit()
                
#                 # Return success message
#                 return jsonify({'status': 'success', "message": "Job post details updated successfully"})
#             else:
#                 return jsonify({'status': 'error',"message": "Job post are not updated successfully"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Update job post fields
#                 job_post.client = data.get('client', job_post.client)
#                 job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                 job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                 job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                 job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                 job_post.location = data.get('location', job_post.location)
#                 job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                 job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                 job_post.role = data.get('role', job_post.role)
#                 job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                 job_post.mode = data.get('mode', job_post.mode)
#                 job_post.job_status = data.get('job_status', job_post.job_status)
#                 job_post.job_type = data.get('Job_Type', job_post.job_type)  # Updated key 'job_type' to 'Job_Type'
#                 job_post.skills = data.get('skills', job_post.skills)
#                 # job_post.jd_pdf = data.get('jd_pdf', job_post.jd_pdf)
                
#                 # Handle jd_pdf field
#                 jd_pdf = data.get('jd_pdf')
#                 if jd_pdf is not None:
#                     # Decode the base64 encoded PDF file
#                     jd_binary = base64.b64decode(jd_pdf)
#                     # Update job_post with the decoded binary data
#                     job_post.jd_pdf = jd_binary
#                     # Set jd_pdf_present to True since PDF is present
#                     job_post.jd_pdf_present = True
#                 else:
#                     # If jd_pdf is None, set jd_pdf_present to False
#                     job_post.jd_pdf_present = False
                    
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     # Ensure recruiters are unique
#                     unique_recruiters = list(set(recruiters))
#                     job_post.recruiter = ', '.join(unique_recruiters)
                    
#                 # Update data_updated_date and data_updated_time
#                 current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
#                 job_post.data_updated_date = current_datetime.date()
#                 job_post.data_updated_time = current_datetime.time()
                
#                 # Update job post in the database
#                 db.session.commit()
                
#                 # Iterate over recruiters and create notification records for each
#                 recruiters = data.get('recruiter', job_post.recruiter)
#                 if recruiters:
#                     unique_recruiters = list(set(recruiters))
#                     for recruiter in unique_recruiters:
#                         # Check if a notification exists for the job post and recruiter combination
#                         notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                         if notification:
#                             # If notification exists, increment num_notification by 1
#                             notification.num_notification += 1
#                         else:
#                             # If notification does not exist, create a new record with num_notification set to 1
#                             new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                             db.session.add(new_notification)
#                             new_notification.num_notification = 1
                
#                 db.session.commit()
                
#                 # Return success message
#                 return jsonify({'status': 'success',"message": "Job post updated successfully"})
#             else:
#                 return jsonify({'status': 'error',"message": "Job post are not updated successfully"})
#         else:
#             return jsonify({'status': 'error',"message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500



# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Update job post fields
#                 job_post.client = data.get('client', job_post.client)
#                 job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                 job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                 job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                 job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                 job_post.location = data.get('location', job_post.location)
#                 job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                 job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                 job_post.role = data.get('role', job_post.role)
#                 job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                 job_post.mode = data.get('mode', job_post.mode)
#                 job_post.job_status = data.get('job_status', job_post.job_status)
#                 job_post.job_type = data.get('Job_Type', job_post.job_type)  # Updated key 'job_type' to 'Job_Type'
#                 job_post.skills = data.get('skills', job_post.skills)
#                 job_post.jd_pdf = data.get('jd_pdf', job_post.jd_pdf)
#                 recruiters = data.get('recruiter', job_post.recruiter)

#                 # Update data_updated_date and data_updated_time
#                 current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
#                 job_post.data_updated_date = current_datetime.date()
#                 job_post.data_updated_time = current_datetime.time()
                
#                 # Update job post in the database
#                 db.session.commit()
                
#                 # Iterate over recruiters and create notification records for each
#                 for recruiter in recruiters:
#                     # Check if a notification exists for the job post and recruiter combination
#                     notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                     if notification:
#                         # If notification exists, increment num_notification by 1
#                         notification.num_notification += 1
#                     else:
#                         # If notification does not exist, create a new record with num_notification set to 1
#                         new_notification = Notification(job_post_id=job_post_id, recruiter_name=job_post.recruiter)
#                         db.session.add(new_notification)
#                         db.session.commit()
#                         new_notification.num_notification = 1
                
#                 db.session.commit()
                
#                 # Return success message
#                 return jsonify({"message": "Job post updated successfully"}), 200
#             else:
#                 return jsonify({"error": "Job post not found"}), 404
#         else:
#             return jsonify({"error": "Unauthorized"}), 401
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         # Check if the user exists and has the right permissions
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 # Update job post fields
#                 job_post.client = data.get('client', job_post.client)
#                 job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                 job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                 job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                 job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                 job_post.location = data.get('location', job_post.location)
#                 job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                 job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                 job_post.role = data.get('role', job_post.role)
#                 job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                 job_post.mode = data.get('mode', job_post.mode)
#                 job_post.job_status = data.get('job_status', job_post.job_status)
#                 job_post.job_type = data.get('Job_Type', job_post.job_type)  # Updated key 'job_type' to 'Job_Type'
#                 job_post.skills = data.get('skills', job_post.skills)
#                 job_post.jd_pdf = data.get('jd_pdf', job_post.jd_pdf)
#                 job_post.recruiter=data.get('recruiter',job_post.recruiter)

#                 # Update data_updated_date and data_updated_time
#                 # current_datetime = datetime.now()
#                 current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
#                 job_post.data_updated_date = current_datetime.date()
#                 job_post.data_updated_time = current_datetime.time()
                
#                 # Update job post in the database
#                 db.session.commit()
                
#                 # Increment num_notification count by 1 for each notification associated with the job post
#                 notifications = Notification.query.filter_by(job_post_id=job_post_id).all()
#                 for notification in notifications:
#                     notification.num_notification += 1
                
#                 db.session.commit()
                
#                 # Return success message
#                 return jsonify({"message": "Job post updated successfully"}), 200
#             else:
#                 return jsonify({"error": "Job post not found"}), 404
#         else:
#             return jsonify({"error": "Unauthorized"}), 401
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


@app.route('/jobs_notification/<int:user_id>', methods=['GET'])
def get_jobs_notification(user_id):
    # Retrieve the user
    user = User.query.filter_by(id=user_id).first()
    
    # Check if the user exists and has the right permissions
    if user and user.user_type == 'recruiter':
        recruiter_name = user.username
        
        # Retrieve the notifications for the recruiter where num_notification >= 1
        notifications = Notification.query.filter_by(recruiter_name=recruiter_name).filter(Notification.num_notification >= 1).all()
        
        # Format the notifications as a list of dictionaries
        notifications_list = [
            {
                # 'id': notification.id,
                'job_post_id': notification.job_post_id,
                'recruiter_name': notification.recruiter_name,
                'notification_status': notification.notification_status,
                'num_notification': notification.num_notification
            } for notification in notifications
        ]
        
        return jsonify(notifications_list), 200
    else:
        return jsonify({'error': 'User not found or does not have the right permissions'}), 404


@app.route('/checked_jobs_notification/<int:user_id>', methods=['POST'])
def checked_jobs_notification(user_id):
    data = request.json
    checked_notification_status = data.get('checked_notification_status')

    user = User.query.filter_by(id=user_id).first()
    
    # Check if the user exists and has the right permissions
    if user and user.user_type == 'recruiter':
        recruiter_name = user.username
        
        # Retrieve the notifications for the recruiter
        notifications = Notification.query.filter_by(recruiter_name=recruiter_name).all()
        
        if checked_notification_status:
            # Update the num_notification to 0 for each notification
            for notification in notifications:
                notification.notification_status = checked_notification_status
                notification.num_notification = 0
                db.session.commit()
                
            # Delete notifications where num_notification is 0
            notifications_to_delete = Notification.query.filter_by(recruiter_name=recruiter_name, num_notification=0).all()
            for notification in notifications_to_delete:
                db.session.delete(notification)
                db.session.commit()
        
        # Format the notifications as a list of dictionaries
        notifications_list = [
            {
                'id': notification.id,
                'job_post_id': notification.job_post_id,
                'recruiter_name': notification.recruiter_name,
                'notification_status': notification.notification_status,
                'num_notification': notification.num_notification
            } for notification in notifications
        ]
        
        return jsonify(notifications_list), 200
    else:
        return jsonify({'error': 'User not found or does not have the right permissions'}), 404


# @app.route('/checked_jobs_notification/<int:user_id>', methods=['POST'])
# def checked_jobs_notification(user_id):
#     data = request.json
#     checked_notification_status = data.get('checked_notification_status')

#     user = User.query.filter_by(id=user_id).first()
    
#     # Check if the user exists and has the right permissions
#     if user and user.user_type == 'recruiter':
#         recruiter_name = user.username
        
#         # Retrieve the notifications for the recruiter
#         notifications = Notification.query.filter_by(recruiter_name=recruiter_name).all()
        
#         if checked_notification_status:
#             # Update the num_notification to 0 for each notification
#             for notification in notifications:
#                 notification.notification_status = checked_notification_status
#                 notification.num_notification = 0
#                 db.session.commit()
        
#         # Format the notifications as a list of dictionaries
#         notifications_list = [
#             {
#                 'id': notification.id,
#                 'job_post_id': notification.job_post_id,
#                 'recruiter_name': notification.recruiter_name,
#                 'notification_status': notification.notification_status,
#                 'num_notification': notification.num_notification
#             } for notification in notifications
#         ]
        
#         return jsonify(notifications_list), 200
#     else:
#         return jsonify({'error': 'User not found or does not have the right permissions'}), 404
    
@app.route('/get_candidate_data')
def get_candidate_data():
    candidates = Candidate.query.all()
    candidate_data = []
    for candidate in candidates:
        candidate_data.append({
            'id': candidate.id,
            'name': candidate.name,
            'email': candidate.email,
            'client': candidate.client,
            'current_company':candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location':candidate.current_job_location,
            'preferred_job_location':candidate.preferred_job_location,
            'skills':candidate.skills,
            'status':candidate.status,
        })
    return jsonify(candidate_data)


@app.route('/send_email', methods=['POST'])
def send_email():
    recipient_email = request.form.get('recipient_email')

    if not recipient_email:
        flash('Recipient email is required.', 'error')
        return redirect(url_for('careers'))

    # Create a link to the page you want to send
    page_link = 'http://127.0.0.1:5001/careers'  # Replace with the actual link

    # Create the email content with a hyperlink
    email_content = f"Click the link below to view active job posts: <a href='{page_link}'>{page_link}</a>"

    # Create an email message
    message = Message('Active Job Posts', sender='kanuparthisaiganesh582@gmail.com', recipients=[recipient_email])
    message.html = email_content

    # Send the email
    mail.send(message)

    flash('Email sent successfully!', 'success')
    return redirect(url_for('careers'))

#new
@app.route('/careers', methods=['GET'])
def careers():
    user_type = session.get('user_type', None)
    is_logged_in = 'user_id' in session
    candidate_message = request.args.get('candidate_message')
    print(candidate_message)

    # Query the database to retrieve active job posts and sort them by date_created in descending order
    active_jobs = JobPost.query.filter_by(job_status='Active').order_by(JobPost.date_created.desc()).all()

    return render_template('careers.html', jobs=active_jobs, user_type=user_type, is_logged_in=is_logged_in,candidate_message=candidate_message)

#new
@app.route('/apply_careers', methods=['GET', 'POST'])
def apply_careers():
    user_id = session.get('user_id')
    if not user_id:
        # User is not authenticated, you can redirect them to a login page or take appropriate action
        return redirect(url_for('career_login'))
    user = Career_user.query.get(user_id)
    if request.method == 'GET':
        job_id = request.args.get('job_id')
        client = request.args.get('client')
        profile = request.args.get('profile')
        name = user.name
        email = user.email

        if job_id:
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Hold')).first()
            if matching_job_post:
                return render_template('job_on_hold.html')
        
        job_post = JobPost.query.get(job_id)
        experience_min = job_post.experience_min

        job_ids = db.session.query(JobPost.id).filter(JobPost.client == client, JobPost.job_status == 'Active').all()
        job_roles = db.session.query(JobPost.role).filter(JobPost.client == client).all()

        ids = [job_id[0] for job_id in job_ids]
        roles = [job_role[0] for job_role in job_roles]

        candidate_data = None
        if 'candidate_data' in request.args:
            candidate_data = ast.literal_eval(request.args['candidate_data'])

        return render_template('apply_careers.html', candidate_data=candidate_data, job_id=job_id,
                               client=client, profile=profile, ids=ids, roles=roles,
                               name=name, email=email,experience_min=experience_min)

    if request.method == 'POST':
        try:
            job_id = request.form['job_id']
            name = request.form['name']
            mobile = request.form['mobile']
            email = request.form['email']
            client = request.form['client']
            profile = request.form['profile']
            skills = request.form['skills']

            # Ensure client and job_id are integers
            job_id = int(job_id)

            # Check if the job post is active
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
            if not matching_job_post:
                return render_template('job_on_hold.html')

            # Handle other form fields...
            current_company = request.form['current_company']
            position = request.form['position']
            current_job_location = request.form['current_job_location']
            preferred_job_location = request.form['preferred_job_location']
            qualifications = request.form['qualifications']
            experience = request.form['experience']
            exp_months = request.form['exp_months']
            experience = experience + '.' + exp_months
            relevant_experience = request.form['relevant_experience']
            relevant_exp_months = request.form['relevant_exp_months']
            relevant_experience = relevant_experience + '.' + relevant_exp_months
            currency_type_current = request.form['currency_type_current']
            currency_type_except = request.form['currency_type_except']
            current_ctc = currency_type_current + " " + request.form['current_ctc']
            expected_ctc = currency_type_except + " " + request.form['expected_ctc']
            linkedin = request.form['linkedin']

            # Handle file upload
            filename = None
            resume_binary = None
            if 'resume' in request.files:
                resume_file = request.files['resume']
                if resume_file and allowed_file(resume_file.filename):
                    # Convert the resume file to binary data
                    resume_binary = resume_file.read()
                    filename = secure_filename(resume_file.filename)
                else:
                    return render_template('apply_careers.html', error_message='Invalid file extension')

            notice_period = request.form['notice_period']
            last_working_date = None
            buyout = False
            period_of_notice = None

            if notice_period == 'yes':
                last_working_date = request.form['last_working_date']
                buyout = 'buyout' in request.form
            elif notice_period == 'no':
                period_of_notice = request.form['months']
                buyout = 'buyout' in request.form
            elif notice_period == 'completed':
                last_working_date = request.form['last_working_date']

            holding_offer = request.form['holding_offer']

            if holding_offer == 'yes':
                total = request.form['total']
                if total == '':
                    total = 0
                else:
                    total = int(request.form['total'])
                package_in_lpa = request.form['package_in_lpa']
                if package_in_lpa == '':
                    package_in_lpa = 0
                else:
                    package_in_lpa = float(request.form['package_in_lpa'])
            else:
                total = None
                package_in_lpa = None

            reason_for_job_change = request.form.get('reason_for_job_change')
            remarks = request.form.get('remarks')

            reference = request.form['reference']
            reference_name = None
            reference_position = None
            reference_information = None

            if reference == 'yes':
                reference_name = request.form['reference_name']
                reference_position = request.form['reference_position']
                reference_information = request.form['reference_information']
            elif reference == 'no':
                reference_name = None
                reference_position = None
                reference_information = None

            existing_candidate = Candidate.query.filter(
                and_(Candidate.profile == profile, Candidate.client == client, Candidate.email == email,
                     Candidate.mobile == mobile)).first()
            if existing_candidate:
                return render_template('candidate_exists.html',
                                       error_message='Candidate with the same profile and client already exists')

            # Create a new Candidate object
            new_candidate = Candidate(
                job_id=job_id,
                name=name,
                mobile=mobile,
                email=email,
                client=client,
                current_company=current_company,
                position=position,
                profile=profile,
                resume=resume_binary,
                current_job_location=current_job_location,
                preferred_job_location=preferred_job_location,
                qualifications=qualifications,
                experience=experience,
                relevant_experience=relevant_experience,
                current_ctc=current_ctc,
                expected_ctc=expected_ctc,
                notice_period=notice_period,
                last_working_date=last_working_date if notice_period == 'yes' or notice_period == 'completed' else None,
                buyout=buyout,
                holding_offer=holding_offer,
                total=total,
                package_in_lpa=package_in_lpa,
                linkedin_url=linkedin,
                reason_for_job_change=reason_for_job_change,
                status='None',
                remarks=remarks,
                skills=skills,
                period_of_notice=period_of_notice,
                reference=reference,
                reference_name=reference_name,
                reference_position=reference_position,
                reference_information=reference_information
            )

            new_candidate.date_created = date.today()
            new_candidate.time_created = datetime.now().time()

            # Commit the new candidate to the database
            db.session.add(new_candidate)
            db.session.commit()

            try:
                msg = Message('Successful Submission of Your Job Application', sender='kanuparthisaiganesh582@gmail.com', recipients=[email])
                msg.body = f"Dear { name },\n Congratulations! Your job application has been successfully submitted for the position at {client} for the role of {profile}. We appreciate your interest in joining our team.\n\n  Our dedicated recruiter will review your application, and you can expect to hear from us within the next 24 hours.\n\nBest wishes for your application process!\n\n Regards, \n\nTeam\nMakonis Talent Track Pro\nrecruiterpro@makonissoft.com\n"
                mail.send(msg)
            except Exception as e:
                # Handle email sending errors, log the error
                return render_template('error.html', error_message=f"Error sending thank-you email: {str(e)}")

            return redirect(url_for('careers', candidate_message='Candidate Added Successfully'))

        except Exception as e:
            # Handle any exceptions here (e.g., log the error, return an error page)
            return render_template('error.html', error_message=str(e))

    return redirect(url_for('careers'))


#new
# User Login
@app.route('/career_login', methods=['GET', 'POST'])
def career_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = Career_user.query.filter_by(username=username, password=password).first()

        if user:
            # Store the user's session or token
            session['user_id'] = user.id
            return redirect(url_for('careers'))

    return render_template('career_login.html')

#new
@app.route('/career_logout')
def career_logout():
    # Clear the user's session
    session.pop('user_id', None)
    return redirect(url_for('careers'))

#new
@app.route('/career_register', methods=['GET', 'POST'])
def career_register():
    if request.method == 'POST':
        username = request.form['username']
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']

        # Create a new user and add it to the database
        new_user = Career_user(username=username, name=name, email=email, password=password)
        db.session.add(new_user)
        db.session.commit()

        return redirect(url_for('career_login'))

    return render_template('career_registration.html')

#new
@app.route('/career_dashboard')
def career_dashboard():
    edit_candidate_message = request.args.get('edit_candidate_message')
    if 'user_id' in session and 'user_type' in session:
        page_no = request.args.get('page_no')
        candidate_message = request.args.get('candidate_message')
        signup_message = request.args.get('signup_message')
        job_message = request.args.get('job_message')
        update_candidate_message = request.args.get('update_candidate_message')
        delete_message = request.args.get("delete_message")

        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']

        if user_type == 'management':
            users = User.query.all()
            candidates = Candidate.query.filter((Candidate.reference.is_not(None))).all()
            candidates = sorted(candidates, key=lambda candidate: candidate.id)
            JobsPosted = JobPost.query.all()
            # data = json.dumps(candidates, sort_keys=False)
            return render_template('career_dashboard.html', users=users, user_type=user_type, user_name=user_name,
                                   candidates=candidates, update_candidate_message=update_candidate_message,
                                   candidate_message=candidate_message, delete_message=delete_message,
                                   JobsPosted=JobsPosted, signup_message=signup_message, job_message=job_message,
                                   page_no=page_no, edit_candidate_message=edit_candidate_message)
        elif user_type == 'recruiter':
            recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
            recruiter_name = User.query.get(user_id).name
            if recruiter:
                candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name,
                                                         Candidate.reference.is_not(None))).all()
                candidates = sorted(candidates, key=lambda candidate: candidate.id)
                career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
                                                                  Career_notification.recruiter_name == user_name).count()
                career_notifications = Career_notification.query.filter(
                    Career_notification.recruiter_name.contains(recruiter_name)).all()

                for career_notification in career_notifications:
                    if career_notification.notification_status == False:
                        career_notification.notification_status = True
                        db.session.commit()
                return render_template('career_dashboard.html', user=recruiter, user_type=user_type, user_name=user_name,
                                       candidates=candidates, candidate_message=candidate_message,
                                       update_candidate_message=update_candidate_message,
                                       career_count_notification_no=career_count_notification_no,
                                       edit_candidate_message=edit_candidate_message, page_no=page_no)
        else:
            user = User.query.filter_by(id=user_id).first()
            if user:
                candidates = Candidate.query.filter_by(recruiter=user.name).all()  # Filter candidates by user's name
                return render_template('career_dashboard.html', user=user, user_type=user_type, candidates=candidates)

    return redirect(url_for('index'))

#new
@app.route('/website_candidate_assign', methods=['GET', 'POST'])
def website_candidate_assign():
    assignment_message = request.args.get('assignment_message')
    if 'user_id' in session and 'user_type' in session and session['user_type'] == 'management':
        user_name = session['user_name']
        recruiters = User.query.filter_by(user_type='recruiter').all()

        if request.method == 'POST':
            assign_recruiter_id = request.form.get('assign_recruiter_id')
            selected_candidate_ids = request.form.getlist('selected_candidate_ids')

            if assign_recruiter_id and selected_candidate_ids:
                assigned_recruiter = User.query.get(assign_recruiter_id)
                if assigned_recruiter:
                    # Fetch selected candidates by their IDs
                    candidates = Candidate.query.filter(
                        Candidate.id.in_(selected_candidate_ids),
                        Candidate.recruiter.is_(None),
                        Candidate.management.is_(None)
                    ).all()

                    for candidate in candidates:
                        # Assign the selected recruiter to the candidate
                        candidate.recruiter = assigned_recruiter.name
                        # Send an email to the assigned recruiter
                        send_career_email(assigned_recruiter.email, 'Alert! New Candidate Assignment ',
                                          f'Dear {assigned_recruiter.name}\n\n,A new candidate application has been assigned to you. Please access your dashboard to view the details.\n\nCandidate Name: {candidate.name}\n\nClient: {candidate.client}\n\nRole: {candidate.profile}\n\nAssigned by Manager: {user_name}\n\nFeel free to reach out if you have any questions during the recruitment process.\n\nRegards,\n\nTeam\nMakonis Talent Track Pro\nrecruiterpro@makonissoft.com')

                    db.session.commit()

                    # Create notifications for the assigned recruiter
                    notifications = []
                    for candidate in candidates:
                        notification = Career_notification(
                            recruiter_name=assigned_recruiter.name,
                            notification_status=False  # You may set this to True for unread notifications
                        )
                        notifications.append(notification)

                    db.session.add_all(notifications)
                    db.session.commit()

                    return redirect(
                        url_for('website_candidate_assign', assignment_message='Candidates Assigned Successfully'))

        candidates = Candidate.query.filter(
            Candidate.recruiter.is_(None),
            Candidate.management.is_(None)
        ).all()

        candidate_count = Candidate.query.filter(
            Candidate.recruiter.is_(None),
            Candidate.management.is_(None)
        ).count()

        return render_template(
            'website_candidate_assign.html',
            recruiters=recruiters,
            candidates=candidates,
            assignment_message=assignment_message,
            user_name=user_name,
            candidate_count=candidate_count
        )

    return redirect(url_for('index'))

#new
def send_career_email(to, subject, message):
    msg = Message(subject, sender='kanuparthisaiganesh582@gmail.com', recipients=[to])
    msg.body = message
    mail.send(msg)

####################################################################################################################################

import base64
import io
import re
from flask import Flask, request, jsonify
import fitz  # PyMuPDF
from docx import Document

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(file):
    """
    Extract text from PDF or DOCX files.
    
    Parameters:
        file (BytesIO): File-like object.
    
    Returns:
        str: Extracted text.
    """
    try:
        file.seek(0)
        header = file.read(4)
        file.seek(0)
        if header.startswith(b'%PDF'):
            return extract_text_from_pdf(file)
        elif header.startswith(b'PK\x03\x04'):
            return extract_text_from_docx(file)
        else:
            return ""  # Unsupported file format
    except Exception as e:
        print(f"Error determining file type: {e}")
        return ""

def extract_text_from_pdf(file):
    """
    Extract text from a PDF file.
    
    Parameters:
        file (BytesIO): PDF file-like object.
    
    Returns:
        str: Extracted text.
    """
    text = ""
    try:
        with fitz.open(stream=file, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file):
    """
    Extract text from a DOCX file.
    
    Parameters:
        file (BytesIO): DOCX file-like object.
    
    Returns:
        str: Extracted text.
    """
    text = ""
    try:
        doc = Document(file)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text

def extract_skills_from_resume(text, skills_list):
    found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
    return found_skills

def extract_email(text):
    email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_matches = re.findall(email_regex, text)
    return email_matches[-1].rstrip('.,') if email_matches else "No email found"
    
def extract_phone_number(text):
    phone_regex = r'\b\d{10}\b'
    phone_matches = re.findall(phone_regex, text)
    return phone_matches[-1] if phone_matches else "No phone number found"

# def extract_phone_number(text):
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[-1] if phone_matches else "No phone number found"

def extract_name(text):
    """
    Extract the name from the first few lines of the resume text.
    
    Parameters:
        text (str): Resume text.
    
    Returns:
        str: Extracted name.
    """
    lines = text.split('\n')
    name_words = []  # List to store the words of the name
    
    # Regular expressions to identify lines that are likely contact details
    phone_pattern = re.compile(r'\b(\+?\d[\d\-\.\s]+)?\d{10}\b')
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    
    for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
        # Skip lines that are likely to be contact details
        if phone_pattern.search(line) or email_pattern.search(line):
            continue
        
        # Remove common salutations and titles
        cleaned_line = re.sub(r'\b(Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam)\b', '', line, flags=re.IGNORECASE).strip()
        
        # Extract names with up to three words
        words = cleaned_line.split()
        name_words.extend(words)  # Add words from the current line to the list
        
        if len(name_words) <= 2:
            continue  # Continue accumulating words if we have less than or equal to three words
        else:
            # Stop accumulating if we exceed three words and return the concatenated name
            return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
    # Return the concatenated name if found within the first five lines
    if name_words:
        return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
    return "No name found"

# def extract_name(text):
#     """
#     Extract the name from the first few lines of the resume text.
    
#     Parameters:
#         text (str): Resume text.
    
#     Returns:
#         str: Extracted name.
#     """
#     lines = text.split('\n')
#     name_words = []  # List to store the words of the name
#     for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
#         # Remove common salutations and titles
#         cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE).strip()
#         # Extract names with up to three words
#         words = cleaned_line.split()
#         name_words.extend(words)  # Add words from the current line to the list
#         if len(name_words) <= 2:
#             continue  # Continue accumulating words if we have less than or equal to three words
#         else:
#             # Stop accumulating if we exceed three words and return the concatenated name
#             return ' '.join(word.capitalize() for word in name_words).rstrip('.,')
#     # Return the concatenated name if found within the first five lines
#     if name_words:
#         return ' '.join(word.capitalize() for word in name_words).rstrip('.,')
#     return "No name found"

# def extract_name(text):
#     """
#     Extract the name from the first few lines of the resume text.
    
#     Parameters:
#         text (str): Resume text.
    
#     Returns:
#         str: Extracted name.
#     """
#     lines = text.split('\n')
#     for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
#         # Remove common salutations and titles
#         cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE).strip()
#         # Extract names with at least two words
#         words = cleaned_line.split()
#         if len(words) >= 1:
#             # Capitalize the first letter of each word in the name
#             return ' '.join(word.capitalize() for word in words).rstrip('.,')
#     return "No name found"

# def extract_name(text):
#     """
#     Extract the name from the first few lines of the resume text.
    
#     Parameters:
#         text (str): Resume text.
    
#     Returns:
#         str: Extracted name.
#     """
#     lines = text.split('\n')
#     last_name = ""
#     for line in lines:
#         # Remove common salutations and titles
#         cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE)
#         # Extract names with at least two words
#         words = cleaned_line.split()
#         if len(words) >= 2:
#             # Capitalize the first letter of each word in the name
#             last_name = ' '.join(word.capitalize() for word in words).rstrip('.,')
#     return last_name if last_name else "No name found"

@app.route('/parse_resume', methods=['POST'])
def parse_resume():
    if 'resume' not in request.json:
        return jsonify({'status':'error',"message": "No resume data provided"})
    
    data = request.json
    resume_data = data['resume']
    
    try:
        decoded_resume = base64.b64decode(resume_data)
    except Exception as e:
        return jsonify({'status':'error',"message": "Invalid resume data"})
    
    resume_file = io.BytesIO(decoded_resume)
    resume_text = extract_text(resume_file)
    
    if not resume_text:
        return jsonify({'status':'error',"message": "No text found in the resume data"})

    it_skills = [ 
        'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
        'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 
        'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
        'Data Analysis and Visualization','Artificial Intelligence','Programming',
        'Database Management (SQL)','Web Development (HTML, CSS, JavaScript)',
        'Machine Learning and Artificial Intelligence','Network Administration',
        'Software Development and Testing','Embedded Systems','CAD and 3D Modeling',
        'HTML5', 'CSS3', 'Jquery', 'Bootstrap', 'XML', 'JSON', 'ABAP', 'SAPUI5',
        'Agile Methodology', 'Frontend Development', 'Jira', 'Odata', 'BTP', 'Fiori Launchpad', 
        'Python', 'JavaScript', 'HTML', 'CSS','React', 'Node.js', 'Django', 'Git', 'AWS',
        'Linux','DevOps','Linear Regression','Logistic Regression','Decision Tree',
        'SVM (Support Vector Machine)','Ensembles','Random Forest','Clustering',
        'PCA (Principal Component Analysis)','K-means','Recommendation System',
        'Market Basket Analysis','CNN','RNN','LSTM','Natural Language Processing',
        'NLTK','LGBM','XGBoost','Transformers','Siamese network','BTYD (Buy Till You Die)',
        'ML Ops Tools: Azure Synapse','Azure ML','Azure Databricks','ML flow','Airflow',
        'Kubernetes','Dockers','Data Streaming – Kafka','Flask','LT Spice','Wireshark',
        'Ansys Lumerical','Zemax OpticStudio','Xilinx Vivado','Google Collab','MATLAB'
    ]
    
    non_it_skills = [
        'Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership',
        'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail',
        'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 
        'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking',
        'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 
        'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 
        'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency',
        'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 
        'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 
        'Social Media Management', 'Content Creation', 'Graphic Design', 'Video Editing', 'Photography', 
        'Data Entry', 'Administrative Support', 'Customer Support'
    ]

    extracted_it_skills = extract_skills_from_resume(resume_text, it_skills)
    extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
    non_it_skills_final = list(set(extracted_nonit_skills) - set(extracted_it_skills))

    skills_it = ", ".join(extracted_it_skills) if extracted_it_skills else "No skills found"
    skills_non_it = ", ".join(non_it_skills_final) if non_it_skills_final else "No skills found"

    email_text = extract_email(resume_text)
    phone_text = extract_phone_number(resume_text)
    name_text = extract_name(resume_text)

    return jsonify({
        'status':'success',
        'message':'resume parsed successfully',
        "name": name_text,
        "mail": email_text,
        "phone": phone_text,
        "skill1": skills_it,
        "skill2": skills_non_it
    })


# def extract_text(file):
#     """
#     Extract text from PDF or DOCX files.
    
#     Parameters:
#         file (BytesIO): File-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     try:
#         file.seek(0)
#         header = file.read(4)
#         file.seek(0)
#         if header.startswith(b'%PDF'):
#             return extract_text_from_pdf(file)
#         elif header.startswith(b'PK\x03\x04'):
#             return extract_text_from_docx(file)
#         else:
#             return ""  # Unsupported file format
#     except Exception as e:
#         print(f"Error determining file type: {e}")
#         return ""

# def extract_text_from_pdf(file):
#     """
#     Extract text from a PDF file.
    
#     Parameters:
#         file (BytesIO): PDF file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         with fitz.open(stream=file, filetype="pdf") as doc:
#             for page in doc:
#                 text += page.get_text()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#     return text

# def extract_text_from_docx(file):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         file (BytesIO): DOCX file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(file)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

# def extract_skills_from_resume(text, skills_list):
#     found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
#     return found_skills

# def extract_email(text):
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[0].rstrip('.,') if email_matches else "No email found"

# def extract_phone_number(text):
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[0] if phone_matches else "No phone number found"

# def extract_name(text):
#   """
#   Extract the name from the first few lines of the resume text, handling common salutations,
#   titles, single-word names, and punctuation at the end.

#   Args:
#       text (str): Resume text.

#   Returns:
#       str: Extracted name (or "No name found" if no suitable name is found).
#   """

#   lines = text.split('\n')
#   for line in lines[:3]:  # Check only the first 3 lines
#     cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE)
#     words = cleaned_line.strip().split()

#     # Handle single-word names (e.g., initials)
#     if len(words) == 1:
#       # If it's all uppercase, consider it a name
#       if words[0].isupper():
#         return words[0]
#       else:
#         continue  # Skip single-word names that are not all uppercase

#     # Extract and capitalize name with two or more words
#     if len(words) >= 2:
#       # Remove trailing punctuation (., !) from the last word
#       last_word = words[-1].rstrip('.,!')
#       name = ' '.join(word.capitalize() for word in words[:-1]) + ' ' + last_word
#       return name

#   return "No name found"
    
# @app.route('/parse_resume', methods=['POST'])
# def parse_resume():
#     if 'resume' not in request.json:
#         return jsonify({"error": "No resume data provided"}), 400
    
#     data = request.json
#     resume_data = data['resume']
    
#     try:
#         decoded_resume = base64.b64decode(resume_data)
#     except Exception as e:
#         return jsonify({"error": "Invalid resume data"}), 400
    
#     resume_file = io.BytesIO(decoded_resume)
#     resume_text = extract_text(resume_file)
    
#     if not resume_text:
#         return jsonify({"error": "No text found in the resume data"}), 400

#     it_skills = [ 
#         'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
#         'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 
#         'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
#         'Data Analysis and Visualization','Artificial Intelligence','Programming',
#         'Database Management (SQL)','Web Development (HTML, CSS, JavaScript)',
#         'Machine Learning and Artificial Intelligence','Network Administration',
#         'Software Development and Testing','Embedded Systems','CAD and 3D Modeling',
#         'HTML5', 'CSS3', 'Jquery', 'Bootstrap', 'XML', 'JSON', 'ABAP', 'SAPUI5',
#         'Agile Methodology', 'Frontend Development', 'Jira', 'Odata', 'BTP', 'Fiori Launchpad', 
#         'Python', 'JavaScript', 'HTML', 'CSS','React', 'Node.js', 'Django', 'Git', 'AWS',
#         'Linux','DevOps','Linear Regression','Logistic Regression','Decision Tree',
#         'SVM (Support Vector Machine)','Ensembles','Random Forest','Clustering',
#         'PCA (Principal Component Analysis)','K-means','Recommendation System',
#         'Market Basket Analysis','CNN','RNN','LSTM','Natural Language Processing',
#         'NLTK','LGBM','XGBoost','Transformers','Siamese network','BTYD (Buy Till You Die)',
#         'ML Ops Tools: Azure Synapse','Azure ML','Azure Databricks','ML flow','Airflow',
#         'Kubernetes','Dockers','Data Streaming – Kafka','Flask','LT Spice','Wireshark',
#         'Ansys Lumerical','Zemax OpticStudio','Xilinx Vivado','Google Collab','MATLAB'
#     ]
    
#     non_it_skills = [
#         'Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership',
#         'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail',
#         'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 
#         'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking',
#         'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 
#         'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 
#         'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency',
#         'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 
#         'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 
#         'Social Media Management', 'Content Creation', 'Graphic Design', 'Video Editing', 'Photography', 
#         'Data Entry', 'Administrative Support', 'Customer Support', 'Teaching', 'Mentoring', 'Coaching', 
#         'Retail Management', 'Hospitality Management', 'Event Management', 'Creative Writing', 'Content Marketing', 
#         'Copywriting', 'Publications', 'Translation', 'Counseling', 'Fitness Instruction', 'Nutrition', 'Wellness', 
#         'Fashion Design', 'Interior Design', 'Artistic Skills', 'Music', 'Sports', 'Culinary Arts', 'Photography', 
#         'Videography', 'Project Coordination', 'Community Outreach', 'Volunteer Management', 'Fundraising', 
#         'Political Campaigning', 'Government Relations', 'Policy Analysis', 'Nonprofit Management', 'Grant Writing', 
#         'Fundraising', 'Event Planning', 'Real Estate', 'Property Management', 'Construction Management', 
#         'Facilities Management', 'Environmental Sustainability', 'Energy Management', 'Public Health', 
#         'Healthcare Administration', 'Nursing', 'Dental Hygiene', 'Pharmacy', 'Physical Therapy', 'Occupational Therapy', 
#         'Social Work', 'Child Care', 'Elderly Care', 'Counseling', 'Psychology', 'Sociology', 'Anthropology', 'Archaeology', 
#         'Geography', 'History', 'Political Science', 'Economics', 'Philosophy', 'Theology', 'Linguistics', 'Literature', 
#         'Creative Writing', 'Journalism', 'Broadcasting', 'Public Relations', 'Marketing', 'Advertising', 'Market Research', 
#         'Retail Sales', 'Wholesale Sales', 'Account Management', 'Client Relations', 'Customer Service', 'Conflict Resolution', 
#         'Presentation Skills', 'Public Speaking', 'Writing', 'Editing', 'Proofreading', 'Content Creation', 'Graphic Design', 
#         'Visual Merchandising', 'Retail Operations', 'Inventory Management', 'Supply Chain', 'Logistics', 'Quality Assurance', 
#         'Process Improvement', 'Project Management', 'Financial Planning', 'Budgeting', 'Financial Analysis', 'Bookkeeping', 
#         'Data Entry', 'Administrative Support', 'Executive Assistance', 'Time Management', 'Organizational Skills', 'Event Planning', 
#         'Event Coordination', 'Event Marketing', 'Catering', 'Venue Management', 'Wedding Planning', 'Trade Show Coordination', 
#         'Customer Service', 'Conflict Resolution', 'Problem-Solving', 'Decision Making', 'Team Collaboration', 'Leadership', 
#         'Supervision', 'Employee Training', 'Performance Management', 'Recruitment', 'Human Resources', 'Payroll Administration', 
#         'Employee Relations', 'Safety Compliance', 'Labor Relations', 'Legal Compliance', 'Contract Negotiation', 'Risk Management', 
#         'Policy Development', 'Quality Management', 'Process Improvement', 'Supply Chain Management', 'Logistics', 'Inventory Control', 
#         'Procurement', 'Distribution', 'Quality Assurance', 'Process Improvement', 'Product Development', 'Marketing Strategy', 
#         'Brand Management', 'Market Research', 'Public Relations', 'Social Media Management', 'Content Creation', 'Copywriting', 
#         'Email Marketing', 'Sales Strategy', 'Client Relationship Management', 'Sales Forecasting', 'Lead Generation', 
#         'Account Management', 'Customer Retention', 'Sales Presentations', 'Networking', 'Public Speaking', 'Team Collaboration', 
#         'Project Management', 'Client Communications', 'Technical Support', 'Troubleshooting', 'Network Administration', 
#         'Quality Assurance', 'Project Management', 'Technical Writing', 'Documentation', 'Research and Development', 'Innovation', 
#         'Problem-Solving', 'Critical Thinking', 'Attention to Detail', 'Collaboration', 'Time Management', 'Adaptability', 'Leadership', 
#         'Creativity', 'Analytical Skills', 'Data Analysis', 'Statistical Analysis', 'Mathematics', 'Physics', 'Chemistry', 'Biology', 
#         'Geology', 'Environmental Science', 'Meteorology', 'Agricultural Science', 'Animal Science', 'Food Science', 'Nutrition', 
#         'Dietetics', 'Physical Therapy', 'Occupational Therapy', 'Speech-Language Pathology', 'Nursing', 'Pharmacy', 'Dentistry', 
#         'Veterinary Medicine', 'Medical Research', 'Medical Writing', 'Clinical Trials', 'Epidemiology', 'Public Health', 
#         'Healthcare Administration', 'Health Informatics', 'Fitness Instruction', 'Nutrition Counseling', 'Wellness Coaching', 
#         'Yoga Instruction', 'Personal Training', 'Physical Education', 'Sports Coaching', 'Athletic Training', 'Recreation', 
#         'Dance Instruction', 'Music Instruction', 'Art Instruction', 'Photography', 'Video Editing', 'Graphic Design', 'Interior Design', 
#         'Fashion Design', 'Culinary Arts', 'Baking', 'Cooking', 'Restaurant Management', 'Hotel Management', 'Tourism', 'Event Planning', 
#         'Museum Management', 'Library Science', 'Archiving', 'Curatorial Work', 'Conservation', 'Environmental Science', 
#         'Sustainability', 'Renewable Energy', 'Climate Change', 'Environmental Policy', 'Wildlife Conservation', 'Forestry', 
#         'Natural Resource Management', 'Ecology', 'Geography', 'Urban Planning', 'Civil Engineering', 'Structural Engineering', 
#         'Transportation Engineering', 'Geotechnical Engineering', 'Environmental Engineering', 'Water Resources Engineering', 
#         'Surveying', 'Architecture', 'Landscape Architecture', 'Interior Design', 'Urban Design', 'Real Estate Development', 
#         'Property Management', 'Construction Management', 'Building Inspection', 'Facilities Management', 'Space Planning', 
#         'Urban Planning', 'Public Administration', 'Policy Analysis', 'Government Relations', 'Political Campaigning', 'Public Policy', 
#         'Economics', 'Finance', 'Accounting', 'Actuarial Science', 'MS Office','Powerpoint','ms word','ms excel' 
#     ]

#     extracted_skills = extract_skills_from_resume(resume_text, it_skills)
#     extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
#     non_it_skills = list(set(extracted_nonit_skills) - set(extracted_skills))

#     skills_it = ", ".join(extracted_skills) if extracted_skills else "No skills found"
#     skills_non_it = ", ".join(non_it_skills) if non_it_skills else "No skills found"

#     mail_text = extract_email(resume_text)
#     phone_text = extract_phone_number(resume_text)

#     first_line = resume_text.split('\n')[0]
#     words = first_line.split()
#     first_3_words = ' '.join(words[:3])

#     name_text = "No name found"
#     if "RESUME" in first_line or "Resume" in first_line or "BIODATA" in first_line or "BioData" in first_line or "biodata" in first_line:
#         second_line = resume_text.split('\n')[1] if len(resume_text.split('\n')) > 1 else ""
#         words1 = second_line.split()
#         first_5_words_in_2 = ' '.join(words1[:5])
#         if any(keyword in first_5_words_in_2 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words1[:4]).title().rstrip('.,')
#         else:
#             name_text = ' '.join(words1).title().rstrip('.,')
#     else:
#         first_5_words_in_1 = ' '.join(words[:5])
#         if any(keyword in first_5_words_in_1 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words[:4]).title().rstrip('.,')
#         else:
#             name_text = ' '.join(words).title().rstrip('.,')

#     return jsonify({
#         "name": name_text,
#         "mail": mail_text,
#         "phone": phone_text,
#         "skill1": skills_it,
#         "skill2": skills_non_it
#     })

###########################################################################################################


# def extract_text(file):
#     """
#     Extract text from PDF or DOCX files.
    
#     Parameters:
#         file (BytesIO): File-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     try:
#         file.seek(0)
#         header = file.read(4)
#         file.seek(0)
#         if header.startswith(b'%PDF'):
#             return extract_text_from_pdf(file)
#         elif header.startswith(b'PK\x03\x04'):
#             return extract_text_from_docx(file)
#         else:
#             return ""  # Unsupported file format
#     except Exception as e:
#         print(f"Error determining file type: {e}")
#         return ""

# def extract_text_from_pdf(file):
#     """
#     Extract text from a PDF file.
    
#     Parameters:
#         file (BytesIO): PDF file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         with fitz.open(stream=file, filetype="pdf") as doc:
#             for page in doc:
#                 text += page.get_text()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#     return text

# def extract_text_from_docx(file):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         file (BytesIO): DOCX file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(file)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

# def extract_skills_from_resume(text, skills_list):
#     found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
#     return found_skills

# def extract_email(text):
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[0] if email_matches else "No email found"

# def extract_phone_number(text):
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[0] if phone_matches else "No phone number found"

# @app.route('/parse_resume', methods=['POST'])
# def parse_resume():
#     if 'resume' not in request.json:
#         return jsonify({"error": "No resume data provided"}), 400
    
#     data = request.json
#     resume_data = data['resume']
    
#     try:
#         decoded_resume = base64.b64decode(resume_data)
#     except Exception as e:
#         return jsonify({"error": "Invalid resume data"}), 400
    
#     resume_file = io.BytesIO(decoded_resume)
#     resume_text = extract_text(resume_file)
    
#     if not resume_text:
#         return jsonify({"error": "No text found in the resume data"}), 400

#     it_skills = [ 
#         'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
#         'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 
#         'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
#         'Data Analysis and Visualization','Artificial Intelligence','Programming',
#         'Database Management (SQL)','Web Development (HTML, CSS, JavaScript)',
#         'Machine Learning and Artificial Intelligence','Network Administration',
#         'Software Development and Testing','Embedded Systems','CAD and 3D Modeling',
#         'HTML5', 'CSS3', 'Jquery', 'Bootstrap', 'XML', 'JSON', 'ABAP', 'SAPUI5',
#         'Agile Methodology', 'Frontend Development', 'Jira', 'Odata', 'BTP', 'Fiori Launchpad', 
#         'Python', 'JavaScript', 'HTML', 'CSS','React', 'Node.js', 'Django', 'Git', 'AWS',
#         'Linux','DevOps','Linear Regression','Logistic Regression','Decision Tree',
#         'SVM (Support Vector Machine)','Ensembles','Random Forest','Clustering',
#         'PCA (Principal Component Analysis)','K-means','Recommendation System',
#         'Market Basket Analysis','CNN','RNN','LSTM','Natural Language Processing',
#         'NLTK','LGBM','XGBoost','Transformers','Siamese network','BTYD (Buy Till You Die)',
#         'ML Ops Tools: Azure Synapse','Azure ML','Azure Databricks','ML flow','Airflow',
#         'Kubernetes','Dockers','Data Streaming – Kafka','Flask','LT Spice','Wireshark',
#         'Ansys Lumerical','Zemax OpticStudio','Xilinx Vivado','Google Collab','MATLAB'
#     ]
    
#     non_it_skills = [
#         'Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership',
#         'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail',
#         'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 
#         'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking',
#         'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 
#         'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 
#         'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency',
#         'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 
#         'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 
#         'Social Media Management', 'Content Creation', 'Graphic Design', 'Video Editing', 'Photography', 
#         'Data Entry', 'Administrative Support', 'Customer Support', 'Teaching', 'Mentoring', 'Coaching', 
#         'Retail Management', 'Hospitality Management', 'Event Management', 'Creative Writing', 'Content Marketing', 
#         'Copywriting', 'Publications', 'Translation', 'Counseling', 'Fitness Instruction', 'Nutrition', 'Wellness', 
#         'Fashion Design', 'Interior Design', 'Artistic Skills', 'Music', 'Sports', 'Culinary Arts', 'Photography', 
#         'Videography', 'Project Coordination', 'Community Outreach', 'Volunteer Management', 'Fundraising', 
#         'Political Campaigning', 'Government Relations', 'Policy Analysis', 'Nonprofit Management', 'Grant Writing', 
#         'Fundraising', 'Event Planning', 'Real Estate', 'Property Management', 'Construction Management', 
#         'Facilities Management', 'Environmental Sustainability', 'Energy Management', 'Public Health', 
#         'Healthcare Administration', 'Nursing', 'Dental Hygiene', 'Pharmacy', 'Physical Therapy', 'Occupational Therapy', 
#         'Social Work', 'Child Care', 'Elderly Care', 'Counseling', 'Psychology', 'Sociology', 'Anthropology', 'Archaeology', 
#         'Geography', 'History', 'Political Science', 'Economics', 'Philosophy', 'Theology', 'Linguistics', 'Literature', 
#         'Creative Writing', 'Journalism', 'Broadcasting', 'Public Relations', 'Marketing', 'Advertising', 'Market Research', 
#         'Retail Sales', 'Wholesale Sales', 'Account Management', 'Client Relations', 'Customer Service', 'Conflict Resolution', 
#         'Presentation Skills', 'Public Speaking', 'Writing', 'Editing', 'Proofreading', 'Content Creation', 'Graphic Design', 
#         'Visual Merchandising', 'Retail Operations', 'Inventory Management', 'Supply Chain', 'Logistics', 'Quality Assurance', 
#         'Process Improvement', 'Project Management', 'Financial Planning', 'Budgeting', 'Financial Analysis', 'Bookkeeping', 
#         'Data Entry', 'Administrative Support', 'Executive Assistance', 'Time Management', 'Organizational Skills', 'Event Planning', 
#         'Event Coordination', 'Event Marketing', 'Catering', 'Venue Management', 'Wedding Planning', 'Trade Show Coordination', 
#         'Customer Service', 'Conflict Resolution', 'Problem-Solving', 'Decision Making', 'Team Collaboration', 'Leadership', 
#         'Supervision', 'Employee Training', 'Performance Management', 'Recruitment', 'Human Resources', 'Payroll Administration', 
#         'Employee Relations', 'Safety Compliance', 'Labor Relations', 'Legal Compliance', 'Contract Negotiation', 'Risk Management', 
#         'Policy Development', 'Quality Management', 'Process Improvement', 'Supply Chain Management', 'Logistics', 'Inventory Control', 
#         'Procurement', 'Distribution', 'Quality Assurance', 'Process Improvement', 'Product Development', 'Marketing Strategy', 
#         'Brand Management', 'Market Research', 'Public Relations', 'Social Media Management', 'Content Creation', 'Copywriting', 
#         'Email Marketing', 'Sales Strategy', 'Client Relationship Management', 'Sales Forecasting', 'Lead Generation', 
#         'Account Management', 'Customer Retention', 'Sales Presentations', 'Networking', 'Public Speaking', 'Team Collaboration', 
#         'Project Management', 'Client Communications', 'Technical Support', 'Troubleshooting', 'Network Administration', 
#         'Quality Assurance', 'Project Management', 'Technical Writing', 'Documentation', 'Research and Development', 'Innovation', 
#         'Problem-Solving', 'Critical Thinking', 'Attention to Detail', 'Collaboration', 'Time Management', 'Adaptability', 'Leadership', 
#         'Creativity', 'Analytical Skills', 'Data Analysis', 'Statistical Analysis', 'Mathematics', 'Physics', 'Chemistry', 'Biology', 
#         'Geology', 'Environmental Science', 'Meteorology', 'Agricultural Science', 'Animal Science', 'Food Science', 'Nutrition', 
#         'Dietetics', 'Physical Therapy', 'Occupational Therapy', 'Speech-Language Pathology', 'Nursing', 'Pharmacy', 'Dentistry', 
#         'Veterinary Medicine', 'Medical Research', 'Medical Writing', 'Clinical Trials', 'Epidemiology', 'Public Health', 
#         'Healthcare Administration', 'Health Informatics', 'Fitness Instruction', 'Nutrition Counseling', 'Wellness Coaching', 
#         'Yoga Instruction', 'Personal Training', 'Physical Education', 'Sports Coaching', 'Athletic Training', 'Recreation', 
#         'Dance Instruction', 'Music Instruction', 'Art Instruction', 'Photography', 'Video Editing', 'Graphic Design', 'Interior Design', 
#         'Fashion Design', 'Culinary Arts', 'Baking', 'Cooking', 'Restaurant Management', 'Hotel Management', 'Tourism', 'Event Planning', 
#         'Museum Management', 'Library Science', 'Archiving', 'Curatorial Work', 'Conservation', 'Environmental Science', 
#         'Sustainability', 'Renewable Energy', 'Climate Change', 'Environmental Policy', 'Wildlife Conservation', 'Forestry', 
#         'Natural Resource Management', 'Ecology', 'Geography', 'Urban Planning', 'Civil Engineering', 'Structural Engineering', 
#         'Transportation Engineering', 'Geotechnical Engineering', 'Environmental Engineering', 'Water Resources Engineering', 
#         'Surveying', 'Architecture', 'Landscape Architecture', 'Interior Design', 'Urban Design', 'Real Estate Development', 
#         'Property Management', 'Construction Management', 'Building Inspection', 'Facilities Management', 'Space Planning', 
#         'Urban Planning', 'Public Administration', 'Policy Analysis', 'Government Relations', 'Political Campaigning', 'Public Policy', 
#         'Economics', 'Finance', 'Accounting', 'Actuarial Science', 'MS Office','Powerpoint','ms word','ms excel' 
#     ]

#     extracted_skills = extract_skills_from_resume(resume_text, it_skills)
#     extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
#     non_it_skills = list(set(extracted_nonit_skills) - set(extracted_skills))

#     skills_it = ", ".join(extracted_skills) if extracted_skills else "No skills found"
#     skills_non_it = ", ".join(non_it_skills) if non_it_skills else "No skills found"

#     mail_text = extract_email(resume_text)
#     phone_text = extract_phone_number(resume_text)

#     first_line = resume_text.split('\n')[0]
#     words = first_line.split()
#     first_3_words = ' '.join(words[:3])

#     name_text = "No name found"
#     if "RESUME" in first_line or "Resume" in first_line or "BIODATA" in first_line or "BioData" in first_line or "biodata" in first_line:
#         second_line = resume_text.split('\n')[1] if len(resume_text.split('\n')) > 1 else ""
#         words1 = second_line.split()
#         first_5_words_in_2 = ' '.join(words1[:5])
#         if any(keyword in first_5_words_in_2 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words1[:4]).title()
#         else:
#             name_text = ' '.join(words1).title()
#     else:
#         first_5_words_in_1 = ' '.join(words[:5])
#         if any(keyword in first_5_words_in_1 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words[:4]).title()
#         else:
#             name_text = ' '.join(words).title()

#     return jsonify({
#         "name": name_text,
#         "mail": mail_text,
#         "phone": phone_text,
#         "skill1": skills_it,
#         "skill2": skills_non_it
#     })





# def extract_text(file_path):
#     """
#     Extract text from PDF or DOCX files.
    
#     Parameters:
#         file_path (str): Path to the file.
    
#     Returns:
#         str: Extracted text.
#     """
#     if file_path.endswith('.pdf'):
#         return extract_text_from_pdf(file_path)
#     elif file_path.endswith('.docx'):
#         return extract_text_from_docx(file_path)
#     else:
#         return ""  # Unsupported file format

# def extract_text_from_pdf(pdf_path):
#     """
#     Extract text from a PDF file.
    
#     Parameters:
#         pdf_path (str): Path to the PDF file.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         with fitz.open(pdf_path) as doc:
#             for page in doc:
#                 text += page.get_text()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#     return text

# def extract_text_from_docx(docx_path):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         docx_path (str): Path to the DOCX file.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(docx_path)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

# def extract_skills_from_resume(text, skills_list):
#     # This function will search for skills in the text based on the given skills list
#     found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
#     return found_skills

# def extract_email(text):
#     # Regex to find email in text
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[0] if email_matches else "No email found"

# def extract_phone_number(text):
#     # Regex to find phone number in text
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[0] if phone_matches else "No phone number found"

# @app.route('/parse_resume', methods=['POST'])
# def parse_resume():
#     if 'resume' not in request.json:
#         return jsonify({"error": "No resume data provided"}), 400
    
#     data = request.json
#     resume_data = data['resume']
    
#     print("resume_data :",resume_data)

#     try:
#         decoded_resume = base64.b64decode(resume_data)
#     except Exception as e:
#         return jsonify({"error": "Invalid resume data"}), 400
    
#     resume_file = io.BytesIO(decoded_resume)
#     resume_text = extract_text(resume_file)
    
#     if not resume_text:
#         return jsonify({"error": "No text found in the resume data"}), 400

#     it_skills = [ 'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
#                     'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
#         'Data Analysis and Visualization','Artificial Intelligence','Programming'
#         'Database Management (SQL)',
#         'Web Development (HTML, CSS, JavaScript)',
#         'Machine Learning and Artificial Intelligence',
#         'Network Administration',
#         'Software Development and Testing',
#         'Embedded Systems',
#         'CAD and 3D Modeling',
#         'HTML5', 'CSS3'  ,'Jquery' ,'Bootstrap' ,'XML' ,'JSON' ,'ABAP'
#         ,'SAPUI5',
#         'Agile Methodology' ,'Frontend Development' ,'Jira' ,'Odata' ,'BTP' ,'Fiori Launchpad', 
#         'Python', 'JavaScript', 'HTML', 'CSS',
#         'React', 'Node.js', 'Django', 'Git', 'AWS', 'Linux','DevOps',
#         'Linear Regression',
#         'Logistic Regression',
#         'Decision Tree',
#         'SVM (Support Vector Machine)',
#         'Ensembles',
#         'Random Forest',
#         'Clustering',
#         'PCA (Principal Component Analysis)',
#         'K-means',
#         'Recommendation System',
#         'Market Basket Analysis',
#         'CNN',
#         'RNN',
#         'LSTM',
#         'Natural Language Processing',
#         'NLTK',
#         'LGBM',
#         'XGBoost',
#         'Transformers',
#         'Siamese network',
#         'BTYD (Buy Till You Die)',
#         'ML Ops Tools: Azure Synapse',
#         'Azure ML',
#         'Azure Databricks',
#         'ML flow',
#         'Airflow',
#         'Kubernetes',
#         'Dockers',
#         'Data Streaming – Kafka',
#         'Flask','LT Spice',
#         'Wireshark',
#         'Ansys Lumerical',
#         'Zemax OpticStudio',
#         'Xilinx Vivado',
#         'Google Collab',
#         'MATLAB'
#         ]
#     non_it_skills = ['Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership', 'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail', 'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking', 'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency', 'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 'Social Media Management', 'Content Creation',
#                         'Graphic Design', 'Video Editing', 'Photography', 'Data Entry', 'Administrative Support', 'Customer Support', 'Teaching', 'Mentoring', 'Coaching', 'Retail Management', 'Hospitality Management', 'Event Management', 'Creative Writing', 'Content Marketing', 'Copywriting', 'Publications', 'Translation', 'Counseling', 'Fitness Instruction', 'Nutrition', 'Wellness', 'Fashion Design', 'Interior Design', 'Artistic Skills', 'Music', 'Sports', 'Culinary Arts', 'Photography', 'Videography', 'Project Coordination', 'Community Outreach', 'Volunteer Management', 'Fundraising', 'Political Campaigning', 'Government Relations', 'Policy Analysis', 'Nonprofit Management', 'Grant Writing', 'Fundraising', 'Event Planning', 'Real Estate', 'Property Management', 'Construction Management', 'Facilities Management', 'Environmental Sustainability', 'Energy Management', 'Public Health', 'Healthcare Administration', 
#                         'Nursing', 'Dental Hygiene', 'Pharmacy', 'Physical Therapy', 'Occupational Therapy', 'Social Work', 'Child Care', 'Elderly Care', 'Counseling', 'Psychology', 'Sociology', 'Anthropology', 'Archaeology', 'Geography', 'History', 'Political Science', 'Economics', 'Philosophy', 'Theology', 'Linguistics', 'Literature', 'Creative Writing', 'Journalism', 'Broadcasting', 'Public Relations', 'Marketing', 'Advertising', 'Market Research', 'Retail Sales', 'Wholesale Sales', 'Account Management', 'Client Relations', 'Customer Service', 'Conflict Resolution', 'Presentation Skills', 'Public Speaking', 'Writing', 'Editing', 'Proofreading', 'Content Creation', 'Graphic Design', 'Visual Merchandising', 'Retail Operations', 'Inventory Management', 'Supply Chain', 'Logistics', 'Quality Assurance', 'Process Improvement', 'Project Management', 'Financial Planning', 'Budgeting', 'Financial Analysis',
#                         'Bookkeeping', 'Data Entry', 'Administrative Support', 'Executive Assistance', 'Time Management', 'Organizational Skills', 'Event Planning', 'Event Coordination', 'Event Marketing', 'Catering', 'Venue Management', 'Wedding Planning', 'Trade Show Coordination', 'Customer Service', 'Conflict Resolution', 'Problem-Solving', 'Decision Making', 'Team Collaboration', 'Leadership', 'Supervision', 'Employee Training', 'Performance Management', 'Recruitment', 'Human Resources', 'Payroll Administration', 'Employee Relations', 'Safety Compliance', 'Labor Relations', 'Legal Compliance', 'Contract Negotiation', 'Risk Management', 'Policy Development', 'Quality Management', 'Process Improvement', 'Supply Chain Management', 'Logistics', 'Inventory Control', 'Procurement', 'Distribution', 'Quality Assurance', 'Process Improvement', 'Product Development', 'Marketing Strategy', 'Brand Management', 'Market Research', 
#                         'Public Relations', 'Social Media Management', 'Content Creation', 'Copywriting', 'Email Marketing', 'Sales Strategy', 'Client Relationship Management', 'Sales Forecasting', 'Lead Generation', 'Account Management', 'Customer Retention', 'Sales Presentations', 'Networking', 'Public Speaking', 'Team Collaboration', 'Project Management', 'Client Communications', 'Technical Support', 'Troubleshooting', 'Network Administration',   'Quality Assurance', 'Project Management', 'Technical Writing', 'Documentation',
#                         'Research and Development', 'Innovation', 'Problem-Solving', 'Critical Thinking', 'Attention to Detail', 'Collaboration', 'Time Management', 'Adaptability', 'Leadership', 'Creativity', 'Analytical Skills', 'Data Analysis', 'Statistical Analysis', 'Mathematics', 'Physics', 'Chemistry', 'Biology', 'Geology', 'Environmental Science', 'Meteorology', 'Agricultural Science', 'Animal Science', 'Food Science', 'Nutrition', 'Dietetics', 'Physical Therapy', 'Occupational Therapy', 'Speech-Language Pathology', 'Nursing', 'Pharmacy', 'Dentistry', 'Veterinary Medicine', 'Medical Research', 'Medical Writing', 'Clinical Trials', 'Epidemiology', 'Public Health', 'Healthcare Administration', 'Health Informatics', 'Fitness Instruction', 'Nutrition Counseling', 'Wellness Coaching', 'Yoga Instruction', 'Personal Training', 'Physical Education', 'Sports Coaching', 'Athletic Training', 'Recreation', 'Dance Instruction', 'Music Instruction',
#                         'Art Instruction', 'Photography', 'Video Editing', 'Graphic Design', 'Interior Design', 'Fashion Design', 'Culinary Arts', 'Baking', 'Cooking', 'Restaurant Management', 'Hotel Management', 'Tourism', 'Event Planning', 'Museum Management', 'Library Science', 'Archiving', 'Curatorial Work', 'Conservation', 'Environmental Science', 'Sustainability', 'Renewable Energy', 'Climate Change', 'Environmental Policy', 'Wildlife Conservation', 'Forestry', 'Natural Resource Management', 'Ecology', 'Geography', 'Urban Planning', 'Civil Engineering', 'Structural Engineering', 'Transportation Engineering', 'Geotechnical Engineering', 'Environmental Engineering', 'Water Resources Engineering', 'Surveying', 'Architecture', 'Landscape Architecture', 'Interior Design', 'Urban Design', 'Real Estate Development', 'Property Management', 'Construction Management', 'Building Inspection', 'Facilities Management', 'Space Planning', 'Urban Planning', 
#                         'Public Administration', 'Policy Analysis', 'Government Relations', 'Political Campaigning', 'Public Policy', 'Economics', 'Finance', 'Accounting', 'Actuarial Science' ,'MS Office','Powerpoint','ms word','ms excel' ]

#     extracted_skills = extract_skills_from_resume(resume_text, it_skills)
#     extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
#     non_it_skills = list(set(extracted_nonit_skills) - set(extracted_skills))

#     skills_it = ", ".join(extracted_skills) if extracted_skills else "No skills found"
#     skills_non_it = ", ".join(non_it_skills) if non_it_skills else "No skills found"

#     mail_text = extract_email(resume_text)
#     phone_text = extract_phone_number(resume_text)

#     first_line = resume_text.split('\n')[0]  # Extract the first line
#     words = first_line.split()
#     first_3_words = ' '.join(words[:3])

#     name_text = "No name found"
#     if "RESUME" in first_line or "Resume" in first_line or "BIODATA" in first_line or "BioData" in first_line or "biodata" in first_line:

#         print("if : entering")
        
#         second_line = resume_text.split('\n')[1] if len(resume_text.split('\n')) > 1 else ""
#         words1 = second_line.split()
#         first_5_words_in_2 = ' '.join(words1[:5])
#         if any(keyword in first_5_words_in_2 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words1[:4]).title()
#         else:
#             name_text = ' '.join(words1).title()
#     else:
#         first_5_words_in_1 = ' '.join(words[:5])
#         if any(keyword in first_5_words_in_1 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words[:4]).title()
#         else:
#             name_text = ' '.join(words).title()

#     return jsonify({
#         "name": name_text,
#         "mail": mail_text,
#         "phone": phone_text,
#         "skill1": skills_it,
#         "skill2": skills_non_it
#     })


if __name__ == '__main__':
    app.run(debug=True, host="0.0.0.0",port=5000)
